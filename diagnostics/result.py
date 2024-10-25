# 30240903 1413

def isDigit(q):
    q = q.replace(',', '.')
    q = q.replace('^', '')
    q = q.replace('*', '')    
    q = q.replace('**', '')
    try:
        float(q)
        return True
    except ValueError:
        return False  

def send_oDLK(f_oDLK):

    import datetime
    cT = datetime.datetime.now().strftime('%Y.%m.%d.%H.%M.%S')

    import docx
    doc = docx.Document()
    NameFile = '# OD_'

    p = doc.add_paragraph()
    runner = p.add_run(NameFile)
    runner.bold = True

    t8o8r = cT # Время регистрации
    obDan1R = []

    o_mM = set()
    oLang = {'английский', 'арабский', 'испанский', 'итальянский', 'казахский', 'киргизский', 'китайский', 'корейский', 
            'монгольский', 'немецкий', 'персидский', 'португальский', 'русский', 'турецкий', 'французский', 'хинди', 'японский'}

    oValuta = {'AMD (Армянский драм)', 'AZN (Азербайджанский манат)', 'BYR (Белорусский рубль)', 'EUR (Евро)', 
               'GEL (Лари)', 'KGS (Сом)', 'KZT (Тенге)', 'RUB (Российский рубль)', 'USD (Доллар США)'}

    if f_oDLK[2][0] != 'Фамилия' and f_oDLK[2][0] != 'Название клиники':
        o_mM.add('mM')
        doc.add_paragraph("# оЗп_И: '" + str(f_oDLK[2][0]) + "', # Не 'Фамилия' и не 'Название клиники'")

    if f_oDLK[2][0] == 'Фамилия':
        if len(f_oDLK) != 17:
            o_mM.add('mM')
            doc.add_paragraph("print('f_oDLK имеет некорректную длинну: " + str(len(f_oDLK)) + ", вместо 17')")
            if 17 - len(f_oDLK) > 0:
                doc.add_paragraph("# f_oDLK - нехватка " + str(17 - len(f_oDLK)) + " элементов!")
            if 17 - len(f_oDLK) < 0:
                doc.add_paragraph("# f_oDLK - избыток " + str(len(f_oDLK) - 17) + " элементов!")
        else:

            if f_oDLK[0][0] != 'IdProfile':
                doc.add_paragraph("# оЗп_И: '" + str(f_oDLK[0][0]) + "', # IdProfile")
                o_mM.add('mM')
            elif type(f_oDLK[0][1]) != int:
                doc.add_paragraph("# оЗп_З: " + str(f_oDLK[0][1]) + " тип данных не int, # IdProfile")
                o_mM.add('mM')
            else:
                doc.add_paragraph(str(f_oDLK[0][1]) + ", # 0 IdProfile")
                obDan1R.append(f_oDLK[0][1])

            if f_oDLK[1][0] != 'Email':
                doc.add_paragraph("# оЗп_И: '" + str(f_oDLK[1][0]) + "', # Email")
                o_mM.add('mM')
            elif type(f_oDLK[1][1]) != str or f_oDLK[1][1] == '' or '@' not in f_oDLK[1][1] or '.' not in f_oDLK[1][1]:
                doc.add_paragraph("# оЗп_З: '" + str(f_oDLK[1][1]) + "', # Email")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + f_oDLK[1][1] + "', # 1 Email")
                obDan1R.append(f_oDLK[1][1])

            if f_oDLK[2][1] == '':
                doc.add_paragraph("# оЗп_З: '', # Фамилия")
                o_mM.add('mM')
            elif type(f_oDLK[2][1]) != str:
                doc.add_paragraph("# оЗп_З: " + str(f_oDLK[2][1]) + ", # Фамилия")
                o_mM.add('mM')
                doc.add_paragraph(str(f_oDLK[2][1]) + ", # Фамилия")
            else:
                doc.add_paragraph("'" + f_oDLK[2][1] + "', # 2 Фамилия")
                obDan1R.append(f_oDLK[2][1])

            if f_oDLK[3][0] != 'Имя':
                doc.add_paragraph("# оЗп_И: '" + str(f_oDLK[3][0]) + "', # Имя")
                o_mM.add('mM')
            elif f_oDLK[3][1] == '':
                doc.add_paragraph("# оЗп_З: '', # Имя")
                o_mM.add('mM')
            elif type(f_oDLK[3][1]) != str:
                doc.add_paragraph("# оЗп_З: " + str(f_oDLK[3][1]) + ", # Имя")
                o_mM.add('mM')
                doc.add_paragraph(str(f_oDLK[3][1]) + ", # Имя")          
            else:
                doc.add_paragraph("'" + f_oDLK[3][1] + "', # 3 Имя")       
                obDan1R.append(f_oDLK[3][1])

            if f_oDLK[4][0] != 'Отчество':
                doc.add_paragraph("# оЗп_И: '" + str(f_oDLK[4][0]) + "', # Отчество")
                o_mM.add('mM')
            elif type(f_oDLK[4][1]) != str:
                doc.add_paragraph("# оЗп_З: " + str(f_oDLK[4][1]) + ", # Отчество")
                o_mM.add('mM')
                doc.add_paragraph(str(f_oDLK[4][1]) + ", # Отчество")
            else:
                doc.add_paragraph("'" + f_oDLK[4][1] + "', # 4 Отчество")       
                obDan1R.append(f_oDLK[4][1])

            if f_oDLK[5][0] != 'Пол':
                doc.add_paragraph("# оЗп_И: '" + str(f_oDLK[5][0]) + "', # Пол")
                o_mM.add('mM')
            elif f_oDLK[5][1] != 'мужской' and f_oDLK[5][1] != 'женский':
                doc.add_paragraph("# оЗп_З: '" + str(f_oDLK[5][1]) + "', # Пол")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + f_oDLK[5][1] + "', # 5 Пол")       
                obDan1R.append(f_oDLK[5][1])
                
            if f_oDLK[6][0] != 'Дата рождения':
                doc.add_paragraph("# оЗп_И: '" + str(f_oDLK[6][0]) + "', # Дата рождения")
                o_mM.add('mM')
            elif type(f_oDLK[6][1]) == str:
                if isDigit(f_oDLK[6][1].replace('.', '')) != True or len(f_oDLK[6][1]) != 10 or len(f_oDLK[6][1].replace('.', '')) != 8:
                    doc.add_paragraph("# оЗп_З: '" + str(f_oDLK[6][1]) + "', # Дата рождения")
                    o_mM.add('mM')
                else:
                    doc.add_paragraph("'" + f_oDLK[6][1] + "', # 6 Дата рождения")
                    obDan1R.append(f_oDLK[6][1])

            o_mML = set()
            if f_oDLK[7][0][0] != 'Укажите языки, на которых могут быть предоставлены услуги':
                doc.add_paragraph("# оЗп_И: '" + str(f_oDLK[7][0][0]) + "', # Укажите языки, на которых могут быть предоставлены услуги")
                o_mML.add('mM')
            elif type(f_oDLK[7][0][1]) != list:
                doc.add_paragraph("# оЗп_З: '" + f_oDLK[7][0][1] + "' не является списком, # Укажите языки, на которых могут быть предоставлены услуги")
            else:
                for q in f_oDLK[7][0][1]:
                    if q not in oLang and q != '':
                        doc.add_paragraph("# оЗп_З: '" + str(q) + "', # Укажите языки, на которых могут быть предоставлены услуги, нет такого языка в списке")
                        o_mML.add('mM')    
            if type(f_oDLK[7][1]) == list:
                if f_oDLK[7][1][0] != 'добавить язык, отсутствующий в списке':
                    doc.add_paragraph("# оЗп_И: '" + str(f_oDLK[7][1][0]) + "', # добавить язык, отсутствующий в списке")
                    o_mML.add('mM')
                elif type(f_oDLK[7][1][1]) != str:
                    doc.add_paragraph("# оЗп_З: '" + str(f_oDLK[7][1][1]) + "', # добавить язык, отсутствующий в списке")
                    o_mML.add('mM')
            if o_mML == set():
                f_alL = []
                for q in f_oDLK[7][0][1]:
                    if q != '':
                        f_alL.append(q)
                if f_oDLK[7][1][1] != '':
                    f_alL.append(f_oDLK[7][1][1])
                if len(f_alL) == 0:
                        doc.add_paragraph("# оЗп_З: не добавлен ни один язык (поле 'Укажите языки, на которых могут быть предоставлены услуги' и 'добавить язык, отсутствующий в списке') ")
                        o_mML.add('mM')
                str_alL = ''
                for q in f_alL:
                    if f_alL.index(q) != len(f_alL)-1:
                        str_alL = str_alL+q+', '
                    else:
                        str_alL = str_alL+q
                doc.add_paragraph("'" + str_alL + "', # 7 Языки")       
                obDan1R.append(str_alL)
            o_mM |= o_mML

            if f_oDLK[8][0] != 'WhatsApp':
                doc.add_paragraph("# оЗп_И: '" + str(f_oDLK[8][0]) + "', # WhatsApp")
                o_mM.add('mM')
            elif type(f_oDLK[8][1]) != int or isDigit(str(f_oDLK[8][1])) == False:
                doc.add_paragraph("# оЗп_З: '" + str(f_oDLK[8][1]) + "', # WhatsApp")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + str(f_oDLK[8][1]) + "', # 8 WhatsApp") 
                obDan1R.append(f_oDLK[8][1])

            if f_oDLK[9][0] != 'Telegram':
                doc.add_paragraph("# оЗп_И: '" + str(f_oDLK[9][0]) + "', # Telegram")
                o_mM.add('mM')
            elif type(f_oDLK[9][1]) != str:
                doc.add_paragraph("# оЗп_З: " + str(f_oDLK[9][1]) + ", # Telegram")
                o_mM.add('mM')
                doc.add_paragraph(str(f_oDLK[9][1]) + ", # Telegram")
            else:
                doc.add_paragraph("'" + f_oDLK[9][1] + "', # 9 Telegram") 
                obDan1R.append(f_oDLK[9][1])
                
            if f_oDLK[10][0] != 'Телефон':
                doc.add_paragraph("# оЗп_И: '" + str(f_oDLK[10][0]) + "', # Телефон")
                o_mM.add('mM')
            elif type(f_oDLK[10][1]) != int and f_oDLK[10][1] != '' or isDigit(str(f_oDLK[10][1])) == False and f_oDLK[10][1] != '':
                doc.add_paragraph("# оЗп_З: '" + str(f_oDLK[10][1]) + "', # Телефон")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + str(f_oDLK[10][1]) + "', # 10 Телефон") 
                obDan1R.append(f_oDLK[10][1])
                
            if f_oDLK[11][0] != 'Страна':
                doc.add_paragraph("# оЗп_И: '" + str(f_oDLK[11][0]) + "', # Страна")
                o_mM.add('mM')
            elif f_oDLK[11][1] == '':
                doc.add_paragraph("# оЗп_З: '', # Страна")
                o_mM.add('mM')
            elif type(f_oDLK[11][1]) != str:
                doc.add_paragraph("# оЗп_З: " + str(f_oDLK[11][1]) + ", # Страна")
                o_mM.add('mM')
                doc.add_paragraph(str(f_oDLK[11][1]) + ", # Страна")
            else:
                doc.add_paragraph("'" + f_oDLK[11][1] + "', # 11 Страна") 
                obDan1R.append(f_oDLK[11][1])
                
            if f_oDLK[12][0] != 'Город, населенный пункт':
                doc.add_paragraph("# оЗп_И: '" + str(f_oDLK[12][0]) + "', # Город, населенный пункт")
                o_mM.add('mM')
            elif f_oDLK[12][1] == '':
                doc.add_paragraph("# оЗп_З: '', # Город, населенный пункт")
                o_mM.add('mM')
            elif type(f_oDLK[12][1]) != str:
                doc.add_paragraph("# оЗп_З: " + str(f_oDLK[12][1]) + ", # Город, населенный пункт")
                o_mM.add('mM')
                doc.add_paragraph(str(f_oDLK[12][1]) + ", # 12 Город, населенный пункт")
            else:
                doc.add_paragraph("'" + f_oDLK[12][1] + "', # 12 Город, населенный пункт") 
                obDan1R.append(f_oDLK[12][1])
                
            if f_oDLK[13][0] != 'Компьютерный анализ симптомов пациента':
                doc.add_paragraph("# оЗп_И: '" + str(f_oDLK[13][0]) + "', # Компьютерный анализ")
                o_mM.add('mM')
            if f_oDLK[13][1] != 'хочу получать только симптомы пациента без их компьютерного анализа (бесплатно)' and f_oDLK[13][1] != 'хочу получать симптомы пациента и три наиболее вероятных диагноза или состояния, выявленных искусственным интеллектом сайта (бесплатно)' and f_oDLK[13][1] != 'хочу получать симптомы пациента и их подробный анализ искусственным интеллектом сайта, оплата через ежемесячную подписку':
                doc.add_paragraph("# оЗп_З: '" + str(f_oDLK[13][1]) + "', # Компьютерный анализ")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + f_oDLK[13][1] + "', # 13 Компьютерный анализ") 
                obDan1R.append(f_oDLK[13][1])

            if f_oDLK[14][0] != 'Валюта при указании стоимости услуг':
                doc.add_paragraph("# оЗп_И: '" + str(f_oDLK[14][0]) + "', # Валюта при указании стоимости услуг")
                o_mM.add('mM')
            elif f_oDLK[14][1] == '':
                doc.add_paragraph("# оЗп_З: '', # Валюта при указании стоимости услуг")
                o_mM.add('mM')
            elif type(f_oDLK[14][1]) != str:
                doc.add_paragraph("# оЗп_З: " + str(f_oDLK[14][1]) + ", # Валюта при указании стоимости услуг")
                o_mM.add('mM')
            elif f_oDLK[14][1] not in oValuta:
                doc.add_paragraph("# оЗп_З: '" + str(f_oDLK[14][1]) + "', # Валюта при указании стоимости услуг, нет такой валюты в списке")
                o_mM.add('mM')  
            else:
                doc.add_paragraph("'" + f_oDLK[14][1] + "', # 14 Валюта при указании стоимости услуг") 
                obDan1R.append(f_oDLK[14][1])
                
            if f_oDLK[16][0] != 'Ученая степень':
                doc.add_paragraph("# оЗп_И: '" + str(f_oDLK[16][0]) + "', # Ученая степень")
                o_mM.add('mM')
            elif type(f_oDLK[16][1]) != str:
                doc.add_paragraph("# оЗп_З: '" + str(f_oDLK[16][1]) + "', # Ученая степень")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + f_oDLK[16][1] + "', # 15 Ученая степень") 
                obDan1R.append(f_oDLK[16][1])

#---------------------
        # edu__1reg
#---------------------

    if f_oDLK[2][0] == 'Название клиники':
        if len(f_oDLK) != 14:
            o_mM.add('mM')
            doc.add_paragraph("print('f_oDLK имеет некорректную длинну: " + str(len(f_oDLK)) + ", вместо 14')")
            if 14 - len(f_oDLK) > 0:
                doc.add_paragraph("# f_oDLK - нехватка " + str(14 - len(f_oDLK)) + " элементов!")
            if 14 - len(f_oDLK) < 0:
                doc.add_paragraph("# f_oDLK - избыток " + str(len(f_oDLK) - 14) + " элементов!")

        else:
            if f_oDLK[0][0] != 'IdProfile':
                doc.add_paragraph("# оЗп_И: '" + str(f_oDLK[0][0]) + "', # IdProfile")
                o_mM.add('mM')
            elif type(f_oDLK[0][1]) != int:
                doc.add_paragraph("# оЗп_З: " + str(f_oDLK[0][1]) + " тип данных не int, # IdProfile")
                o_mM.add('mM')
            else:
                doc.add_paragraph(str(f_oDLK[0][1]) + ", # 0 IdProfile")
                obDan1R.append(f_oDLK[0][1])

            if f_oDLK[1][0] != 'Email':
                doc.add_paragraph("# оЗп_И: '" + str(f_oDLK[1][0]) + "', # Email")
                o_mM.add('mM')
            elif type(f_oDLK[1][1]) != str or f_oDLK[1][1] == '' or '@' not in f_oDLK[1][1] or '.' not in f_oDLK[1][1]:
                doc.add_paragraph("# оЗп_З: '" + str(f_oDLK[1][1]) + "', # Email")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + f_oDLK[1][1] + "', # 1 Email")
                obDan1R.append(f_oDLK[1][1])

            if f_oDLK[2][1] == '':
                doc.add_paragraph("# оЗп_З: '', # Название клиники")
                o_mM.add('mM')
            elif type(f_oDLK[2][1]) != str:
                doc.add_paragraph("# оЗп_З: " + str(f_oDLK[2][1]) + ", # Название клиники")
                o_mM.add('mM')
                doc.add_paragraph(str(f_oDLK[2][1]) + ", # Название клиники")
            else:
                doc.add_paragraph("'" + f_oDLK[2][1] + "', # 2 Название клиники") 
                obDan1R.append(f_oDLK[2][1])
                
            if f_oDLK[3][0] != 'Адрес клиники':
                doc.add_paragraph("# оЗп_И: '" + str(f_oDLK[3][0]) + "', # Адрес клиники")
                o_mM.add('mM')
            elif f_oDLK[3][1] == '':
                doc.add_paragraph("# оЗп_З: '', # Адрес клиники")
                o_mM.add('mM')
            elif type(f_oDLK[3][1]) != str:
                doc.add_paragraph("# оЗп_З: " + str(f_oDLK[3][1]) + ", # Адрес клиники")
                o_mM.add('mM')
                doc.add_paragraph(str(f_oDLK[3][1]) + ", # Адрес клиники")           
            else:
                doc.add_paragraph("'" + f_oDLK[3][1] + "', # 3 Адрес клиники") 
                obDan1R.append(f_oDLK[3][1])
                
            if f_oDLK[4][0] != 'Время работы клиники':
                doc.add_paragraph("# оЗп_И: '" + str(f_oDLK[4][0]) + "', # Время работы клиники")
                o_mM.add('mM')
            elif f_oDLK[4][1] == '':
                doc.add_paragraph("# оЗп_З: '', # Время работы клиники")
                o_mM.add('mM')
            elif type(f_oDLK[4][1]) != str:
                doc.add_paragraph("# оЗп_З: " + str(f_oDLK[4][1]) + ", # Время работы клиники")
                o_mM.add('mM')
                doc.add_paragraph(str(f_oDLK[4][1]) + ", # Время работы клиники")
            else:
                doc.add_paragraph("'" + f_oDLK[4][1] + "', # 4 Время работы клиники") 
                obDan1R.append(f_oDLK[4][1])
                
            obDan1R.append('')
            obDan1R.append('')

            o_mML = set()
            if f_oDLK[5][0][0] != 'Укажите языки, на которых могут быть предоставлены услуги':
                doc.add_paragraph("# оЗп_И: '" + str(f_oDLK[5][0][0]) + "', # Укажите языки, на которых могут быть предоставлены услуги")
                o_mML.add('mM')
            elif type(f_oDLK[5][0][1]) != list:
                doc.add_paragraph("# оЗп_З: '" + f_oDLK[5][0][1] + "' не является списком, # Укажите языки, на которых могут быть предоставлены услуги")
            else:
                for q in f_oDLK[5][0][1]:
                    if q not in oLang and q != '':
                        doc.add_paragraph("# оЗп_З: '" + str(q) + "', # Укажите языки, на которых могут быть предоставлены услуги, нет такого языка в списке")
                        o_mML.add('mM')    
            if type(f_oDLK[5][1]) == list:
                if f_oDLK[5][1][0] != 'добавить язык, отсутствующий в списке':
                    doc.add_paragraph("# оЗп_И: '" + str(f_oDLK[5][1][0]) + "', # добавить язык, отсутствующий в списке")
                    o_mML.add('mM')
                elif type(f_oDLK[5][1][1]) != str:
                    doc.add_paragraph("# оЗп_З: '" + str(f_oDLK[5][1][1]) + "', # добавить язык, отсутствующий в списке")
                    o_mML.add('mM')
            if o_mML == set():
                f_alL = []
                for q in f_oDLK[5][0][1]:
                    if q != '':
                        f_alL.append(q)
                if f_oDLK[5][1][1] != '':
                    f_alL.append(f_oDLK[5][1][1])
                if len(f_alL) == 0:
                        doc.add_paragraph("# оЗп_З: не добавлен ни один язык (поле 'Укажите языки, на которых могут быть предоставлены услуги' и 'добавить язык, отсутствующий в списке') ")
                        o_mML.add('mM')
                str_alL = ''
                for q in f_alL:
                    if f_alL.index(q) != len(f_alL)-1:
                        str_alL = str_alL+q+', '
                    else:
                        str_alL = str_alL+q
                doc.add_paragraph("'" + str_alL + "', # 5 Языки")       
                obDan1R.append(str_alL)
            o_mM |= o_mML

            if f_oDLK[6][0] != 'WhatsApp':
                doc.add_paragraph("# оЗп_И: '" + str(f_oDLK[6][0]) + "', # WhatsApp")
                o_mM.add('mM')
            elif type(f_oDLK[6][1]) != int or isDigit(str(f_oDLK[6][1])) == False:
                doc.add_paragraph("# оЗп_З: '" + str(f_oDLK[6][1]) + "', # WhatsApp")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + str(f_oDLK[6][1]) + "', # 6 WhatsApp") 
                obDan1R.append(f_oDLK[6][1])
                
            if f_oDLK[7][0] != 'Telegram':
                doc.add_paragraph("# оЗп_И: '" + str(f_oDLK[7][0]) + "', # Telegram")
                o_mM.add('mM')
            elif type(f_oDLK[7][1]) != str:
                doc.add_paragraph("# оЗп_З: " + str(f_oDLK[7][1]) + ", # Telegram")
                o_mM.add('mM')
                doc.add_paragraph(str(f_oDLK[7][1]) + ", # Telegram")
            else:
                doc.add_paragraph("'" + f_oDLK[7][1] + "', # 7 Telegram") 
                obDan1R.append(f_oDLK[7][1])
                
            if f_oDLK[8][0] != 'Телефон':
                doc.add_paragraph("# оЗп_И: '" + str(f_oDLK[8][0]) + "', # Телефон")
                o_mM.add('mM')
            elif type(f_oDLK[8][1]) != int and f_oDLK[8][1] != '' or isDigit(str(f_oDLK[8][1])) == False and f_oDLK[8][1] != '':
                doc.add_paragraph("# оЗп_З: '" + str(f_oDLK[8][1]) + "', # Телефон")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + str(f_oDLK[8][1]) + "', # 8 Телефон") 
                obDan1R.append(f_oDLK[8][1])
                
            if f_oDLK[9][0] != 'Страна':
                doc.add_paragraph("# оЗп_И: '" + str(f_oDLK[9][0]) + "', # Страна")
                o_mM.add('mM')
            elif f_oDLK[9][1] == '':
                doc.add_paragraph("# оЗп_З: '', # Страна")
                o_mM.add('mM')
            elif type(f_oDLK[9][1]) != str:
                doc.add_paragraph("# оЗп_З: " + str(f_oDLK[9][1]) + ", # Страна")
                o_mM.add('mM')
                doc.add_paragraph(str(f_oDLK[9][1]) + ", # Страна")
            else:
                doc.add_paragraph("'" + f_oDLK[9][1] + "', # 9 Страна") 
                obDan1R.append(f_oDLK[9][1])
                
            if f_oDLK[10][0] != 'Город, населенный пункт':
                doc.add_paragraph("# оЗп_И: '" + str(f_oDLK[10][0]) + "', # Город, населенный пункт")
                o_mM.add('mM')
            elif f_oDLK[10][1] == '':
                doc.add_paragraph("# оЗп_З: '', # Город, населенный пункт")
                o_mM.add('mM')
            elif type(f_oDLK[10][1]) != str:
                doc.add_paragraph("# оЗп_З: " + str(f_oDLK[10][1]) + ", # Город, населенный пункт")
                o_mM.add('mM')
                doc.add_paragraph(str(f_oDLK[10][1]) + ", # Город, населенный пункт")
            else:
                doc.add_paragraph("'" + f_oDLK[10][1] + "', # 10 Город, населенный пункт") 
                obDan1R.append(f_oDLK[10][1])
                
            if f_oDLK[11][0] != 'Компьютерный анализ симптомов пациента':
                doc.add_paragraph("# оЗп_И: '" + str(f_oDLK[11][0]) + "', # Компьютерный анализ")
                o_mM.add('mM')
            if f_oDLK[11][1] != 'хочу получать только симптомы пациента без их компьютерного анализа (бесплатно)' and f_oDLK[11][1] != 'хочу получать симптомы пациента и три наиболее вероятных диагноза или состояния, выявленных искусственным интеллектом сайта (бесплатно)' and f_oDLK[11][1] != 'хочу получать симптомы пациента и их подробный анализ искусственным интеллектом сайта, оплата через ежемесячную подписку':
                doc.add_paragraph("# оЗп_З: '" + str(f_oDLK[11][1]) + "', # Компьютерный анализ")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + f_oDLK[11][1] + "', # 11 Компьютерный анализ") 
                obDan1R.append(f_oDLK[11][1])

            if f_oDLK[12][0] != 'Валюта при указании стоимости услуг':
                doc.add_paragraph("# оЗп_И: '" + str(f_oDLK[12][0]) + "', # Валюта при указании стоимости услуг")
                o_mM.add('mM')
            elif f_oDLK[12][1] == '':
                doc.add_paragraph("# оЗп_З: '', # Валюта при указании стоимости услуг")
                o_mM.add('mM')
            elif type(f_oDLK[12][1]) != str:
                doc.add_paragraph("# оЗп_З: " + str(f_oDLK[12][1]) + ", # Валюта при указании стоимости услуг")
                o_mM.add('mM')
            elif f_oDLK[12][1] not in oValuta:
                doc.add_paragraph("# оЗп_З: '" + str(f_oDLK[12][1]) + "', # Валюта при указании стоимости услуг, нет такой валюты в списке")
                o_mM.add('mM')  
            else:
                doc.add_paragraph("'" + f_oDLK[12][1] + "', # 12 Валюта при указании стоимости услуг") 
                obDan1R.append(f_oDLK[12][1])

            if f_oDLK[13][0] != 'Вызов на дом':
                doc.add_paragraph("# оЗп_И: '" + str(f_oDLK[13][0]) + "', # Вызов на дом")
                o_mM.add('mM')
            if f_oDLK[13][1] != 'возможен вызов на дом' and f_oDLK[13][1] != 'вызов на дом не принимается':
                doc.add_paragraph("# оЗп_З: '" + str(f_oDLK[13][1]) + "', # Вызов на дом")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + f_oDLK[13][1] + "', # 13 Вызов на дом") 
                obDan1R.append(f_oDLK[13][1])
                
            # doc.add_paragraph("'" + t8o8r + "', # 16 Время регистрации")
            obDan1R.append(t8o8r)

#---------------------

    if f_oDLK[2][0] == 'Фамилия':
        doc.add_paragraph("print('    Специалист, " + str(f_oDLK[2][1]) + " " + str(f_oDLK[3][1]) + ", " + str(f_oDLK[5][1]) + ", " + str(f_oDLK[6][1]) + ", " + str(f_oDLK[12][1]) + "'), ")
    if f_oDLK[2][0] == 'Название клиники':
        doc.add_paragraph("print('    Клиника, " + str(f_oDLK[2][1]) + ", " + str(f_oDLK[3][1]) + ", " + str(f_oDLK[10][1]) + "'), ")
    doc.add_paragraph("print('---------------------------')\n")

    if o_mM != set():
        print(f"oDLK_Ошибка_{cT}")
        doc.add_paragraph("print('" + NameFile + " отклонено оЗп_. ----------------------------------------------------------------')")
        doc.save(NameFile+cT+'_M.docx')
        return NameFile+cT+'_M.docx'
    else:
        doc.save(NameFile+'V.docx')
        print(f"oDLK_Верно_{cT}\n")
        print(f"obDan1R {obDan1R}\n")
        
        return NameFile+'V.docx'    
    
#-------------------------------------------------------------------------------------------------------------------------------------

# 20240909 1621

def isDigit(q):
    q = q.replace(',', '.')
    q = q.replace('^', '')
    q = q.replace('*', '')    
    q = q.replace('**', '')
    try:
        float(q)
        return True
    except ValueError:
        return False  

def send_servLK (f_servLK):

    import datetime
    cT = datetime.datetime.now().strftime('%Y.%m.%d.%H.%M.%S') # Время регистрации

    import docx
    doc = docx.Document()
    NameFile = '# Serv_'

    p = doc.add_paragraph()
    runner = p.add_run(NameFile)
    runner.bold = True

    serv1R = []
    o_mM = set()
    oLang = {'английский', 'арабский', 'испанский', 'итальянский', 'казахский', 'киргизский', 'китайский', 'корейский', 
            'монгольский', 'немецкий', 'персидский', 'португальский', 'русский', 'турецкий', 'французский', 'хинди', 'японский'}

    if f_servLK[2][0] != 'Персонализированность услуги' and f_servLK[2][0] != 'Вид услуги':
        doc.add_paragraph("# оЗп_И: '" + f_servLK[2][0] + "', # Имя поля [1][0] не 'Персонализированность услуги' и не 'Вид услуги'")
        o_mM.add('mM')
        doc.add_paragraph("print('" + NameFile + " не добавлена из-за оЗп_И Имя поля [1][0] не 'Персонализированность услуги' и не 'Вид услуги'.')")
        doc.add_paragraph("isprav.add('serv1R')") 

#---------------------

    if f_servLK[2][0] == 'Вид услуги':
        if len(f_servLK) != 14:
            o_mM.add('mM')
            doc.add_paragraph("print('f_servLKSpec имеет некорректную длинну: " + str(len(f_servLK)) + ", вместо 14')")
            if 14 - len(f_servLK) > 0:
                doc.add_paragraph("# f_servLKSpec - нехватка " + str(14 - len(f_servLK)) + " элементов!")
            if 14 - len(f_servLK) < 0:
                doc.add_paragraph("# f_servLKSpec - избыток " + str(len(f_servLK) - 14) + " элементов!")
        else:
            doc.add_paragraph("\nserv1R.add(") 
            doc.add_paragraph("(")

            if f_servLK[0][0] != 'Id':
                doc.add_paragraph("# оЗп_И: '" + str(f_servLK[0][0]) + "', # Id")
                o_mM.add('mM')
            elif type(f_servLK[0][1]) != int:
                doc.add_paragraph("# оЗп_З: " + str(f_servLK[0][1]) + "тип данных не int, # Id")
                o_mM.add('mM')
            else:
                doc.add_paragraph(str(f_servLK[0][1]) + ", # 0 Id")
                serv1R.append(f_servLK[0][1])

            if f_servLK[1][0] != 'IdProfile':
                doc.add_paragraph("# оЗп_И: '" + str(f_servLK[1][0]) + "', # IdProfile")
                o_mM.add('mM')
            elif type(f_servLK[1][1]) != int:
                doc.add_paragraph("# оЗп_З: " + str(f_servLK[1][1]) + " тип данных не int, # IdProfile")
                o_mM.add('mM')
            else:
                doc.add_paragraph("" + str(f_servLK[1][1]) + ", # 1 IdProfile")
                serv1R.append(f_servLK[1][1])

            doc.add_paragraph("'', # 2 '' / Персонализированность услуги")
            serv1R.append('')
            doc.add_paragraph("'', # 3 '' / Фамилия")
            serv1R.append('')
            doc.add_paragraph("'', # 4 '' / Имя")
            serv1R.append('')
            doc.add_paragraph("'', # 5 '' / Отчество")
            serv1R.append('')
            doc.add_paragraph("'', # 6 '' / Пол")
            serv1R.append('')
            doc.add_paragraph("'', # 7 '' / Дата рождения")
            serv1R.append('')

            if f_servLK[2][0] != 'Вид услуги':
                doc.add_paragraph("# оЗп_И: '" + str(f_servLK[2][0]) + "', # Вид услуги")
                o_mM.add('mM')
            elif f_servLK[2][1] != 'дистанционная консультация' and f_servLK[2][1] != 'консультация на очном приеме' and f_servLK[2][1] != 'консультация, возможна дистанционная или на очном приеме' and f_servLK[2][1] != 'процедура (манипуляция, операция, …)':
                doc.add_paragraph("# оЗп_З: '" + str(f_servLK[2][1]) + "', # Вид услуги")
                o_mM.add('mM')
            elif f_servLK[2][1] == 'процедура (манипуляция, операция, …)' and len(f_servLK[2]) != 3:
                doc.add_paragraph("# оЗп_З: поле Название процедуры - не обязательное # Вид услуги")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + f_servLK[2][1] + "', # 8 Конс. дист. / Конс. оч. / Дист. или оч. конс. / Проц.")
                serv1R.append(f_servLK[2][1])
                if len(f_servLK[2]) == 2:
                    doc.add_paragraph("'', # 9 Название процедуры")
                    serv1R.append('')
                if len(f_servLK[2]) == 3:
                    doc.add_paragraph("'" + f_servLK[2][2] + "', # 9 Название процедуры")
                    serv1R.append(f_servLK[2][2])

            if f_servLK[3][0] != 'Тип услуги':
                doc.add_paragraph("# оЗп_И: '" + str(f_servLK[3][0]) + "', # Тип услуги")
                o_mM.add('mM')
            elif f_servLK[3][1] != 'врачебная услуга' and f_servLK[3][1] != 'не врачебная услуга' and f_servLK[3][1] != 'не медицинская услуга':
                doc.add_paragraph("# оЗп_З: '" + str(f_servLK[3][1]) + "', # Тип услуги")
                o_mM.add('mM')
            elif f_servLK[3][1] == 'не медицинская услуга' and len(f_servLK[3]) != 3:
                doc.add_paragraph("# оЗп_З: поле Тип услуги по не медицинской услуге - len != 3 # Тип услуги")
                o_mM.add('mM')
            elif f_servLK[3][1] == 'врачебная услуга' and len(f_servLK[3]) != 6:
                doc.add_paragraph("# оЗп_З: поле Тип услуги по врачебной услуге - len != 6 # Тип услуги")
                o_mM.add('mM')
            elif f_servLK[3][1] == 'не врачебная услуга' and len(f_servLK[3]) != 6:
                doc.add_paragraph("# оЗп_З: поле Тип услуги по не врачебной услуге - len != 6 # Тип услуги")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + f_servLK[3][1] + "', # 10 ВрУсл / НеВрУсл / НеМедУсл")
                serv1R.append(f_servLK[3][1])
                doc.add_paragraph("'" + f_servLK[3][2] + "', # 11 Название специальности")
                serv1R.append(f_servLK[3][2])
                if len(f_servLK[3]) == 3:
                    doc.add_paragraph("'', # 12 Сертификат по специальности, начало (чч.мм.гггг)")
                    serv1R.append('')
                    doc.add_paragraph("'', # 13 Сертификат по специальности, окончание (чч.мм.гггг)")
                    serv1R.append('')
                    doc.add_paragraph("'', # 14 Категория по специальности")
                    serv1R.append('')                    
                if len(f_servLK[3]) == 6:
                    doc.add_paragraph("'" + f_servLK[3][3] + "', # 12 Сертификат по специальности, начало (чч.мм.гггг)")
                    serv1R.append(f_servLK[3][3])
                    doc.add_paragraph("'" + f_servLK[3][4] + "', # 13 Сертификат по специальности, окончание (чч.мм.гггг)")
                    serv1R.append(f_servLK[3][4])
                    doc.add_paragraph("'" + f_servLK[3][5] + "', # 14 Категория по специальности")
                    serv1R.append(f_servLK[3][5])

            if f_servLK[5][0] != 'Возраст пациентов, лет':
                doc.add_paragraph("# оЗп_И: '" + str(f_servLK[5][0]) + "', # Возраст пациентов, лет")
                o_mM.add('mM')
            elif type(f_servLK[5][1]) != int and type(f_servLK[5][1]) != float and f_servLK[5][1] != '':
                doc.add_paragraph("# оЗп_З: '" + str(f_servLK[5][1]) + "', # Возраст пациентов, лет, от")
            else:
                serv1R.append(f_servLK[5][1])
                if type(f_servLK[5][1]) == int or type(f_servLK[5][1]) == float:
                    doc.add_paragraph(str(f_servLK[5][1]) + ", # 15 Возраст пациентов, лет, от")
                if type(f_servLK[5][1]) == str:
                    doc.add_paragraph("'" + str(f_servLK[5][1]) + "', # 15 Возраст пациентов, лет, от")

            if type(f_servLK[5][2]) != int and type(f_servLK[5][2]) != float and f_servLK[5][2] != '':
                doc.add_paragraph("# оЗп_З: '" + str(f_servLK[5][2]) + "', # Возраст пациентов, лет, до")
            else:
                serv1R.append(f_servLK[5][2])
                if type(f_servLK[5][2]) == int or type(f_servLK[5][2]) == float:
                    doc.add_paragraph(str(f_servLK[5][2]) + ", # 16 Возраст пациентов, лет, до")
                if type(f_servLK[5][2]) == str:
                    doc.add_paragraph("'" + str(f_servLK[5][2]) + "', # 16 Возраст пациентов, лет, до")

            if f_servLK[6][0] != 'Стаж по предлагаемой услуге, лет':
                doc.add_paragraph("# оЗп_И: '" + str(f_servLK[6][0]) + "', # Стаж по предлагаемой услуге, лет")
                o_mM.add('mM')
            elif type(f_servLK[6][1]) != int and type(f_servLK[6][1]) != float:
                doc.add_paragraph("# оЗп_З: " + str(f_servLK[6][1]) + ", # Стаж по предлагаемой услуге, лет")
                o_mM.add('mM')
            else:
                doc.add_paragraph(str(f_servLK[6][1]) + ", # 17 Стаж по предлагаемой услуге, лет")
                serv1R.append(f_servLK[6][1])

            if f_servLK[7][0] != 'Стоимость':
                doc.add_paragraph("# оЗп_И: '" + str(f_servLK[7][0]) + "', # Стоимость")
                o_mM.add('mM')
            elif type(f_servLK[7][1]) != int and type(f_servLK[7][1]) != float:
                doc.add_paragraph("# оЗп_З: " + str(f_servLK[7][1]) + ", # Стоимость")
                o_mM.add('mM')
            else:
                doc.add_paragraph(str(f_servLK[7][1]) + ", # 18 Стоимость")
                serv1R.append(f_servLK[7][1])

            if f_servLK[8][0] != 'Возможность пациента оплатить услугу через сайт':
                doc.add_paragraph("# оЗп_И: '" + str(f_servLK[8][0]) + "', # Возможность пациента оплатить услугу через сайт")
                o_mM.add('mM')
            elif f_servLK[8][1] != 'не предоставлять пациенту возможность оплатить услугу через сайт' and f_servLK[8][1] != 'предоставить пациенту возможность оплатить услугу через сайт (взимается % от стоимости услуги за транзакционные издержки)' and f_servLK[8][1] != '':
                doc.add_paragraph("# оЗп_З: '" + str(f_servLK[8][1]) + "', # Возможность пациента оплатить услугу через сайт")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + f_servLK[8][1] + "', # 19 Возможность пациента оплатить услугу через сайт")
                serv1R.append(f_servLK[8][1])

            if f_servLK[9][0] != 'Ключевые поисковые слова предлагаемой услуги':
                doc.add_paragraph("# оЗп_И: '" + str(f_servLK[9][0]) + "', # Ключевые поисковые слова предлагаемой услуги")
                o_mM.add('mM')
            elif type(f_servLK[9][1]) != str:
                doc.add_paragraph("# оЗп_З: '" + str(f_servLK[9][1]) + "', # Ключевые поисковые слова предлагаемой услуги")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + str(f_servLK[9][1]) + "', # 20 Ключевые поисковые слова предлагаемой услуги")
                serv1R.append(f_servLK[9][1])

            if f_servLK[10][0] != 'Название клиники, где оказывается услуга':
                doc.add_paragraph("# оЗп_И: '" + str(f_servLK[10][0]) + "', # Название клиники, где оказывается услуга")
                o_mM.add('mM')
            elif type(f_servLK[10][1]) != str:
                doc.add_paragraph("# оЗп_З: '" + str(f_servLK[10][1]) + "', # Название клиники, где оказывается услуга")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + str(f_servLK[10][1]) + "', # 21 Название клиники, где оказывается услуга")
                serv1R.append(f_servLK[10][1])

            if f_servLK[11][0] != 'Адрес, где оказывается услуга':
                doc.add_paragraph("# оЗп_И: '" + str(f_servLK[11][0]) + "', # Адрес, где оказывается услуга")
                o_mM.add('mM')
            elif type(f_servLK[11][1]) != str:
                doc.add_paragraph("# оЗп_З: '" + str(f_servLK[11][1]) + "', # Адрес, где оказывается услуга")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + str(f_servLK[11][1]) + "', # 22 Адрес, где оказывается услуга")
                serv1R.append(f_servLK[11][1])

            if f_servLK[12][0] != 'Время приема':
                doc.add_paragraph("# оЗп_И: '" + str(f_servLK[12][0]) + "', # Время приема")
                o_mM.add('mM')
            elif type(f_servLK[12][1]) != str:
                doc.add_paragraph("# оЗп_З: " + str(f_servLK[12][1]) + ", # Время приема")
                o_mM.add('mM')
            else:                   
                doc.add_paragraph("'" + str(f_servLK[12][1]) + "', # 23 Время приема")  
                serv1R.append(f_servLK[12][1])

            if f_servLK[4][0] != 'Возможность оказать данную услугу на дому':
                doc.add_paragraph("# оЗп_И: '" + str(f_servLK[4][0]) + "', # Возможность оказать данную услугу на дому")
                o_mM.add('mM')
            elif f_servLK[4][1] != 'возможен вызов на дом' and f_servLK[4][1] != 'вызов на дом не принимается':
                doc.add_paragraph("# оЗп_З: '" + str(f_servLK[4][1]) + "', # Возможность оказать данную услугу на дому")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + f_servLK[4][1] + "', # 24 Возможность оказать данную услугу на дому")
                serv1R.append(f_servLK[4][1])

            if f_servLK[13][0] != 'Дополнительная информация о предлагаемой услуге':
                doc.add_paragraph("# оЗп_И: '" + str(f_servLK[13][0]) + "', # Дополнительная информация о предлагаемой услуге")
                o_mM.add('mM')
            elif type(f_servLK[13][1]) != str:
                doc.add_paragraph("# оЗп_З: " + str(f_servLK[13][1]) + ", # Дополнительная информация о предлагаемой услуге")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + str(f_servLK[13][1]) + "', # 25 Дополнительная информация о предлагаемой услуге") 
                serv1R.append(f_servLK[13][1])

            doc.add_paragraph("),")
            doc.add_paragraph(")")

#---------------------
    if f_servLK[2][0] == 'Персонализированность услуги':
        if f_servLK[2][1] != 'персонализированная услуга (указываются данные специалиста)' and f_servLK[2][1] != 'не персонализированная услуга (не указываются данные специалиста)':
            o_mM.add('mM')
            doc.add_paragraph("# оЗп_З: '" + str(f_servLK[2][1]) + " не шаблонный варинат ('персонализированная услуга (указываются данные специалиста)', 'не персонализированная услуга (не указываются данные специалиста)') ', # Персонализированность услуги")
        doc.add_paragraph("'" + str(f_servLK[2][1]) + "', # Персонализированность услуги")
#---------------------

    if f_servLK[2][1] == 'персонализированная услуга (указываются данные специалиста)':
        if len(f_servLK) != 19:
            o_mM.add('mM')
            doc.add_paragraph("print('f_servLKPerson имеет некорректную длинну: " + str(len(f_servLK)) + ", вместо 19')")
            if 19 - len(f_servLK) > 0:
                doc.add_paragraph("# f_servLKPerson - нехватка " + str(19 - len(f_servLK)) + " элементов!")
            if 19 - len(f_servLK) < 0:
                doc.add_paragraph("# f_servLKPerson - избыток " + str(len(f_servLK) - 19) + " элементов!")

        else:

            doc.add_paragraph("\nserv1R.add(") 
            doc.add_paragraph("(")

            if f_servLK[0][0] != 'Id':
                doc.add_paragraph("# оЗп_И: '" + str(f_servLK[0][0]) + "', # Id")
                o_mM.add('mM')
            elif type(f_servLK[0][1]) != int:
                doc.add_paragraph("# оЗп_З: '" + str(f_servLK[0][1]) + "', # Id")
                o_mM.add('mM')
            else:
                doc.add_paragraph("" + str(f_servLK[0][1]) + ", # 0 Id")
                serv1R.append(f_servLK[0][1])

            if f_servLK[1][0] != 'IdProfile':
                doc.add_paragraph("# оЗп_И: '" + str(f_servLK[1][0]) + "', # IdProfile")
                o_mM.add('mM')
            elif type(f_servLK[1][1]) != int:
                doc.add_paragraph("# оЗп_З: '" + str(f_servLK[1][1]) + "', # IdProfile")
                o_mM.add('mM')
            else:
                doc.add_paragraph("" + str(f_servLK[1][1]) + ", # 1 IdProfile")
                serv1R.append(f_servLK[1][1])

            if f_servLK[2][0] != 'Персонализированность услуги':
                doc.add_paragraph("# оЗп_И: '" + str(f_servLK[2][0]) + "', # Персонализированность услуги")
                o_mM.add('mM')
            elif f_servLK[2][1] != 'персонализированная услуга (указываются данные специалиста)' and f_servLK[2][1] != 'не персонализированная услуга (не указываются данные специалиста)' and f_servLK[2][1] != '':
                doc.add_paragraph("# оЗп_З: '" + str(f_servLK[2][1]) + "', # Персонализированность услуги")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + f_servLK[2][1] + "', # 2 Персонализированность услуги")
                serv1R.append(f_servLK[2][1])

            if f_servLK[3][0] != 'Фамилия':
                doc.add_paragraph("# оЗп_И: '" + str(f_servLK[3][0]) + "', # Фамилия")
                o_mM.add('mM')
            elif f_servLK[3][1] == '':
                doc.add_paragraph("# оЗп_З: '', # Фамилия")
                o_mM.add('mM')
            elif type(f_servLK[3][1]) != str:
                doc.add_paragraph("# оЗп_З: " + str(f_servLK[3][1]) + ", # Фамилия")
                o_mM.add('mM')
                doc.add_paragraph(str(f_servLK[3][1]) + ", # Фамилия")          
            else:
                doc.add_paragraph("'" + f_servLK[3][1] + "', # 3 Фамилия")
                serv1R.append(f_servLK[3][1])
          
            if f_servLK[4][0] != 'Имя':
                doc.add_paragraph("# оЗп_И: '" + str(f_servLK[4][0]) + "', # Имя")
                o_mM.add('mM')
            elif f_servLK[4][1] == '':
                doc.add_paragraph("# оЗп_З: '', # Имя")
                o_mM.add('mM')
            elif type(f_servLK[4][1]) != str:
                doc.add_paragraph("# оЗп_З: " + str(f_servLK[4][1]) + ", # Имя")
                o_mM.add('mM')
                doc.add_paragraph(str(f_servLK[4][1]) + ", # Имя")          
            else:
                doc.add_paragraph("'" + f_servLK[4][1] + "', # 4 Имя")
                serv1R.append(f_servLK[4][1]) 

            if f_servLK[5][0] != 'Отчество':
                doc.add_paragraph("# оЗп_И: '" + str(f_servLK[5][0]) + "', # Отчество")
                o_mM.add('mM')
            elif type(f_servLK[5][1]) != str:
                doc.add_paragraph("# оЗп_З: " + str(f_servLK[5][1]) + ", # Отчество")
                o_mM.add('mM')
                doc.add_paragraph(str(f_servLK[5][1]) + ", # Отчество")
            else:
                doc.add_paragraph("'" + f_servLK[5][1] + "', # 5 Отчество")       
                serv1R.append(f_servLK[5][1])

            if f_servLK[6][0] != 'Пол':
                doc.add_paragraph("# оЗп_И: '" + str(f_servLK[6][0]) + "', # Пол")
                o_mM.add('mM')
            elif f_servLK[6][1] != 'мужской' and f_servLK[6][1] != 'женский':
                doc.add_paragraph("# оЗп_З: '" + str(f_servLK[6][1]) + "', # Пол")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + f_servLK[6][1] + "', # 6 Пол")       
                serv1R.append(f_servLK[6][1])

            if f_servLK[7][0] != 'Дата рождения':
                doc.add_paragraph("# оЗп_И: '" + str(f_servLK[7][0]) + "', # Дата рождения")
                o_mM.add('mM')
            elif type(f_servLK[7][1]) == str:
                if isDigit(f_servLK[7][1].replace('.', '')) != True or len(f_servLK[7][1]) != 10 or len(f_servLK[7][1].replace('.', '')) != 8:
                    doc.add_paragraph("# оЗп_З: '" + str(f_servLK[7][1]) + "', # Дата рождения")
                    o_mM.add('mM')
                else:
                    doc.add_paragraph("'" + f_servLK[7][1] + "', # 7 '' / Дата рождения")
                    serv1R.append(f_servLK[7][1])

            if f_servLK[8][0] != 'Вид услуги':
                doc.add_paragraph("# оЗп_И: '" + str(f_servLK[8][0]) + "', # Вид услуги")
                o_mM.add('mM')
            elif f_servLK[8][1] != 'дистанционная консультация' and f_servLK[8][1] != 'консультация на очном приеме' and f_servLK[8][1] != 'консультация, возможна дистанционная или на очном приеме' and f_servLK[8][1] != 'процедура (манипуляция, операция, …)':
                doc.add_paragraph("# оЗп_З: '" + str(f_servLK[8][1]) + "', # Вид услуги")
                o_mM.add('mM')
            elif f_servLK[8][1] == 'процедура (манипуляция, операция, …)' and len(f_servLK[8]) != 3:
                doc.add_paragraph("# оЗп_З: поле Название процедуры - не обязательное # Вид услуги")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + f_servLK[8][1] + "', # 8 Конс. дист. / Конс. оч. / Дист. или оч. конс. / Проц.")
                serv1R.append(f_servLK[8][1])
                if len(f_servLK[8]) == 2:
                    doc.add_paragraph("'', # 9 Название процедуры")
                    serv1R.append('')
                if len(f_servLK[8]) == 3:
                    doc.add_paragraph("'" + f_servLK[8][2] + "', # 9 Название процедуры")
                    serv1R.append(f_servLK[8][2])

            if f_servLK[9][0] != 'Тип услуги':
                doc.add_paragraph("# оЗп_И: '" + str(f_servLK[9][0]) + "', # Тип услуги")
                o_mM.add('mM')
            elif f_servLK[9][1] != 'врачебная услуга' and f_servLK[9][1] != 'не врачебная услуга' and f_servLK[9][1] != 'не медицинская услуга':
                doc.add_paragraph("# оЗп_З: '" + str(f_servLK[9][1]) + "', # Тип услуги")
                o_mM.add('mM')
            elif f_servLK[9][1] == 'не медицинская услуга' and len(f_servLK[9]) != 3:
                doc.add_paragraph("# оЗп_З: поле Тип услуги по не медицинской услуге - len != 3 # Тип услуги")
                o_mM.add('mM')
            elif f_servLK[9][1] == 'врачебная услуга' and len(f_servLK[9]) != 6:
                doc.add_paragraph("# оЗп_З: поле Тип услуги по врачебной услуге - len != 6 # Тип услуги")
                o_mM.add('mM')
            elif f_servLK[9][1] == 'не врачебная услуга' and len(f_servLK[9]) != 6:
                doc.add_paragraph("# оЗп_З: поле Тип услуги по не врачебной услуге - len != 6 # Тип услуги")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + f_servLK[9][1] + "', # 10 ВрУсл / НеВрУсл / НеМедУсл")
                serv1R.append(f_servLK[9][1])
                doc.add_paragraph("'" + f_servLK[9][2] + "', # 11 Название специальности")
                serv1R.append(f_servLK[9][2])
                if len(f_servLK[9]) == 3:
                    doc.add_paragraph("'', # 12 Сертификат по специальности, начало (чч.мм.гггг)")
                    serv1R.append('')
                    doc.add_paragraph("'', # 13 Сертификат по специальности, окончание (чч.мм.гггг)")
                    serv1R.append('')
                    doc.add_paragraph("'', # 14 Категория по специальности")
                    serv1R.append('')                    
                if len(f_servLK[9]) == 6:
                    doc.add_paragraph("'" + f_servLK[9][3] + "', # 12 Сертификат по специальности, начало (чч.мм.гггг)")
                    serv1R.append(f_servLK[9][3])
                    doc.add_paragraph("'" + f_servLK[9][4] + "', # 13 Сертификат по специальности, окончание (чч.мм.гггг)")
                    serv1R.append(f_servLK[9][4])
                    doc.add_paragraph("'" + f_servLK[9][5] + "', # 14 Категория по специальности")
                    serv1R.append(f_servLK[9][5])

            if f_servLK[10][0] != 'Возраст пациентов, лет':
                doc.add_paragraph("# оЗп_И: '" + str(f_servLK[10][0]) + "', # Возраст пациентов, лет")
                o_mM.add('mM')
            elif type(f_servLK[10][1]) != int and type(f_servLK[10][1]) != float and f_servLK[10][1] != '':
                doc.add_paragraph("# оЗп_З: '" + str(f_servLK[10][1]) + "', # Возраст пациентов, лет, от")
            else:
                serv1R.append(f_servLK[10][1])
                if type(f_servLK[10][1]) == int or type(f_servLK[10][1]) == float:
                    doc.add_paragraph(str(f_servLK[10][1]) + ", # 15 Возраст пациентов, лет, от")
                if type(f_servLK[10][1]) == str:
                    doc.add_paragraph("'" + str(f_servLK[10][1]) + "', # 15 Возраст пациентов, лет, от")

            if type(f_servLK[10][2]) != int and type(f_servLK[10][2]) != float and f_servLK[10][2] != '':
                doc.add_paragraph("# оЗп_З: '" + str(f_servLK[10][2]) + "', # Возраст пациентов, лет, до")
            else:
                serv1R.append(f_servLK[10][2])
                if type(f_servLK[10][2]) == int or type(f_servLK[10][2]) == float:
                    doc.add_paragraph(str(f_servLK[10][2]) + ", # 16 Возраст пациентов, лет, до")
                if type(f_servLK[10][2]) == str:
                    doc.add_paragraph("'" + str(f_servLK[10][2]) + "', # 16 Возраст пациентов, лет, до")

            if f_servLK[11][0] != 'Стаж по предлагаемой услуге, лет':
                doc.add_paragraph("# оЗп_И: '" + str(f_servLK[11][0]) + "', # Стаж по предлагаемой услуге, лет")
                o_mM.add('mM')
            elif type(f_servLK[11][1]) != int and type(f_servLK[11][1]) != float:
                doc.add_paragraph("# оЗп_З: " + str(f_servLK[11][1]) + ", # Стаж по предлагаемой услуге, лет")
                o_mM.add('mM')
            else:
                doc.add_paragraph(str(f_servLK[11][1]) + ", # 17 Стаж по предлагаемой услуге, лет")
                serv1R.append(f_servLK[11][1])

            if f_servLK[12][0] != 'Стоимость':
                doc.add_paragraph("# оЗп_И: '" + str(f_servLK[12][0]) + "', # Стоимость")
                o_mM.add('mM')
            elif type(f_servLK[12][1]) != int and type(f_servLK[12][1]) != float:
                doc.add_paragraph("# оЗп_З: " + str(f_servLK[12][1]) + ", # Стоимость")
                o_mM.add('mM')
            else:
                doc.add_paragraph(str(f_servLK[12][1]) + ", # 18 Стоимость")
                serv1R.append(f_servLK[12][1])

            if f_servLK[13][0] != 'Возможность пациента оплатить услугу через сайт':
                doc.add_paragraph("# оЗп_И: '" + str(f_servLK[13][0]) + "', # Возможность пациента оплатить услугу через сайт")
                o_mM.add('mM')
            elif f_servLK[13][1] != 'не предоставлять пациенту возможность оплатить услугу через сайт' and f_servLK[13][1] != 'предоставить пациенту возможность оплатить услугу через сайт (взимается % от стоимости услуги за транзакционные издержки)' and f_servLK[13][1] != '':
                doc.add_paragraph("# оЗп_З: '" + str(f_servLK[13][1]) + "', # Возможность пациента оплатить услугу через сайт")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + f_servLK[13][1] + "', # 19 Возможность пациента оплатить услугу через сайт")
                serv1R.append(f_servLK[13][1])

            if f_servLK[14][0] != 'Ключевые поисковые слова предлагаемой услуги':
                doc.add_paragraph("# оЗп_И: '" + str(f_servLK[14][0]) + "', # Ключевые поисковые слова предлагаемой услуги")
                o_mM.add('mM')
            elif type(f_servLK[14][1]) != str:
                doc.add_paragraph("# оЗп_З: '" + str(f_servLK[14][1]) + "', # Ключевые поисковые слова предлагаемой услуги")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + str(f_servLK[14][1]) + "', # 20 Ключевые поисковые слова предлагаемой услуги")
                serv1R.append(f_servLK[14][1])

            o_mME = set()
            if f_servLK[15][0] != 'Email услуги':
                doc.add_paragraph("# оЗп_И: '" + str(f_servLK[15][0]) + "', # Email услуги")
                o_mME.add('mM')
            elif f_servLK[15][1] != '':
                if type(f_servLK[15][1]) != str or '@' not in f_servLK[15][1] or '.' not in f_servLK[15][1]:
                    doc.add_paragraph("# оЗп_З: '" + str(f_servLK[15][1]) + "', # Email услуги")
                    o_mME.add('mM')
            if o_mME == set():
                doc.add_paragraph("'" + f_servLK[15][1] + "', # 21 Email услуги")
                serv1R.append(f_servLK[15][1])
            o_mM |= o_mME

            doc.add_paragraph("'', # 22 '' / (Адрес услуги)")
            serv1R.append('')

            if f_servLK[16][0] != 'Время приема':
                doc.add_paragraph("# оЗп_И: '" + str(f_servLK[16][0]) + "', # Время приема")
                o_mM.add('mM')
            elif type(f_servLK[16][1]) != str:
                doc.add_paragraph("# оЗп_З: " + str(f_servLK[16][1]) + ", # Время приема")
                o_mM.add('mM')
            else:                   
                doc.add_paragraph("'" + str(f_servLK[16][1]) + "', # 23 Время приема")  
                serv1R.append(f_servLK[16][1])

            if f_servLK[17][0] != 'Возможность оказать данную услугу на дому':
                doc.add_paragraph("# оЗп_И: '" + str(f_servLK[17][0]) + "', # Возможность оказать данную услугу на дому")
                o_mM.add('mM')
            elif f_servLK[17][1] != '' and f_servLK[17][1] != 'возможен вызов на дом' and f_servLK[17][1] != 'вызов на дом не принимается':
                doc.add_paragraph("# оЗп_З: '" + str(f_servLK[17][1]) + "', # Возможность оказать данную услугу на дому")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + f_servLK[17][1] + "', # 24 Возможность оказать данную услугу на дому")
                serv1R.append(f_servLK[17][1])

            if f_servLK[18][0] != 'Дополнительная информация о предлагаемой услуге':
                doc.add_paragraph("# оЗп_И: '" + str(f_servLK[18][0]) + "', # Дополнительная информация о предлагаемой услуге")
                o_mM.add('mM')
            elif type(f_servLK[18][1]) != str:
                doc.add_paragraph("# оЗп_З: " + str(f_servLK[18][1]) + ", # Дополнительная информация о предлагаемой услуге")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + str(f_servLK[18][1]) + "', # 25 Дополнительная информация о предлагаемой услуге") 
                serv1R.append(f_servLK[18][1])

            doc.add_paragraph(")")
            doc.add_paragraph(")")

#---------------------

    if f_servLK[2][1] == 'не персонализированная услуга (не указываются данные специалиста)':
        if len(f_servLK) != 14:
            o_mM.add('mM')
            doc.add_paragraph("print('f_servLKNotPerson имеет некорректную длинну: " + str(len(f_servLK)) + ", вместо 14')")
            if 14 - len(f_servLK) > 0:
                doc.add_paragraph("# f_servLKNotPerson - нехватка " + str(14 - len(f_servLK)) + " элементов!")
            if 14 - len(f_servLK) < 0:
                doc.add_paragraph("# f_servLKNotPerson - избыток " + str(len(f_servLK) - 14) + " элементов!")
        else:

            doc.add_paragraph("\nserv1R.add(") 
            doc.add_paragraph("(")

            if f_servLK[0][0] != 'Id':
                doc.add_paragraph("# оЗп_И: '" + str(f_servLK[0][0]) + "', # Id")
                o_mM.add('mM')
            elif type(f_servLK[0][1]) != int:
                doc.add_paragraph("# оЗп_З: '" + str(f_servLK[0][1]) + "', # Id")
                o_mM.add('mM')
            else:
                doc.add_paragraph("" + str(f_servLK[0][1]) + ", # 0 Id")
                serv1R.append(f_servLK[0][1])

            if f_servLK[1][0] != 'IdProfile':
                doc.add_paragraph("# оЗп_И: '" + str(f_servLK[1][0]) + "', # IdProfile")
                o_mM.add('mM')
            elif type(f_servLK[1][1]) != int:
                doc.add_paragraph("# оЗп_З: '" + str(f_servLK[1][1]) + "', # IdProfile")
                o_mM.add('mM')
            else:
                doc.add_paragraph("" + str(f_servLK[1][1]) + ", # 1 IdProfile")
                serv1R.append(f_servLK[1][1])

            if f_servLK[2][0] != 'Персонализированность услуги':
                doc.add_paragraph("# оЗп_И: '" + str(f_servLK[2][0]) + "', # Персонализированность услуги")
                o_mM.add('mM')
            elif f_servLK[2][1] != 'персонализированная услуга (указываются данные специалиста)' and f_servLK[2][1] != 'не персонализированная услуга (не указываются данные специалиста)' and f_servLK[2][1] != '':
                doc.add_paragraph("# оЗп_З: '" + str(f_servLK[2][1]) + "', # Персонализированность услуги")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + f_servLK[2][1] + "', # 2 Персонализированность услуги")
                serv1R.append(f_servLK[2][1])

            doc.add_paragraph("'', # 3 '' / Фамилия")
            serv1R.append('')
            doc.add_paragraph("'', # 4 '' / Имя")
            serv1R.append('')
            doc.add_paragraph("'', # 5 '' / Отчество")
            serv1R.append('')
            doc.add_paragraph("'', # 6 '' / Пол")
            serv1R.append('')
            doc.add_paragraph("'', # 7 '' / Дата рождения")
            serv1R.append('')

            if f_servLK[3][0] != 'Вид услуги':
                doc.add_paragraph("# оЗп_И: '" + str(f_servLK[3][0]) + "', # Вид услуги")
                o_mM.add('mM')
            elif f_servLK[3][1] != 'дистанционная консультация' and f_servLK[3][1] != 'консультация на очном приеме' and f_servLK[3][1] != 'консультация, возможна дистанционная или на очном приеме' and f_servLK[3][1] != 'процедура (манипуляция, операция, …)':
                doc.add_paragraph("# оЗп_З: '" + str(f_servLK[3][1]) + "', # Вид услуги")
                o_mM.add('mM')
            elif f_servLK[3][1] == 'процедура (манипуляция, операция, …)' and len(f_servLK[3]) != 3:
                doc.add_paragraph("# оЗп_З: поле Название процедуры - не обязательное # Вид услуги")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + f_servLK[3][1] + "', # 8 Конс. дист. / Конс. оч. / Дист. или оч. конс. / Проц.")
                serv1R.append(f_servLK[3][1])
                if len(f_servLK[3]) == 2:
                    doc.add_paragraph("'', # 9 Название процедуры")
                    serv1R.append('')
                if len(f_servLK[3]) == 3:
                    doc.add_paragraph("'" + f_servLK[3][2] + "', # 9 Название процедуры")
                    serv1R.append(f_servLK[3][2])

            if f_servLK[4][0] != 'Тип услуги':
                doc.add_paragraph("# оЗп_И: '" + str(f_servLK[4][0]) + "', # Тип услуги")
                o_mM.add('mM')
            elif f_servLK[4][1] != 'врачебная услуга' and f_servLK[4][1] != 'не врачебная услуга' and f_servLK[4][1] != 'не медицинская услуга':
                doc.add_paragraph("# оЗп_З: '" + str(f_servLK[4][1]) + "', # Тип услуги")
                o_mM.add('mM')
            elif f_servLK[4][1] == 'не медицинская услуга' and len(f_servLK[4]) != 3:
                doc.add_paragraph("# оЗп_З: поле Тип услуги по не медицинской услуге - len != 3 # Тип услуги")
                o_mM.add('mM')
            elif f_servLK[4][1] == 'врачебная услуга' and len(f_servLK[4]) != 3:
                doc.add_paragraph("# оЗп_З: поле Тип услуги по врачебной услуге - len != 3 # Тип услуги")
                o_mM.add('mM')
            elif f_servLK[4][1] == 'не врачебная услуга' and len(f_servLK[4]) != 3:
                doc.add_paragraph("# оЗп_З: поле Тип услуги по не врачебной услуге - len != 3 # Тип услуги")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + f_servLK[4][1] + "', # 10 ВрУсл / НеВрУсл / НеМедУсл")
                serv1R.append(f_servLK[4][1])
                doc.add_paragraph("'" + f_servLK[4][2] + "', # 11 Название специальности")
                serv1R.append(f_servLK[4][2])

            doc.add_paragraph("'', # 12 '' (Сертификат по специальности, начало (чч.мм.гггг))")
            serv1R.append('')
            doc.add_paragraph("'', # 13 '' (Сертификат по специальности, окончание (чч.мм.гггг))")
            serv1R.append('')
            doc.add_paragraph("'', # 14 '' (Категория по специальности)")
            serv1R.append('')

            if f_servLK[5][0] != 'Возраст пациентов, лет':
                doc.add_paragraph("# оЗп_И: '" + str(f_servLK[5][0]) + "', # Возраст пациентов, лет")
                o_mM.add('mM')
            elif type(f_servLK[5][1]) != int and type(f_servLK[5][1]) != float and f_servLK[5][1] != '':
                doc.add_paragraph("# оЗп_З: '" + str(f_servLK[5][1]) + "', # Возраст пациентов, лет, от")
            else:
                serv1R.append(f_servLK[5][1])
                if type(f_servLK[5][1]) == int or type(f_servLK[5][1]) == float:
                    doc.add_paragraph(str(f_servLK[5][1]) + ", # 15 Возраст пациентов, лет, от")
                if type(f_servLK[5][1]) == str:
                    doc.add_paragraph("'" + str(f_servLK[5][1]) + "', # 15 Возраст пациентов, лет, от")

            if type(f_servLK[5][2]) != int and type(f_servLK[5][2]) != float and f_servLK[5][2] != '':
                doc.add_paragraph("# оЗп_З: '" + str(f_servLK[5][2]) + "', # Возраст пациентов, лет, до")
            else:
                serv1R.append(f_servLK[5][2])
                if type(f_servLK[5][2]) == int or type(f_servLK[5][2]) == float:
                    doc.add_paragraph(str(f_servLK[5][2]) + ", # 16 Возраст пациентов, лет, до")
                if type(f_servLK[5][2]) == str:
                    doc.add_paragraph("'" + str(f_servLK[5][2]) + "', # 16 Возраст пациентов, лет, до")

            doc.add_paragraph("'', # 17 '' (Стаж по предлагаемой услуге, лет)")
            serv1R.append('')

            if f_servLK[6][0] != 'Стоимость':
                doc.add_paragraph("# оЗп_И: '" + str(f_servLK[6][0]) + "', # Стоимость")
                o_mM.add('mM')
            elif type(f_servLK[6][1]) != int and type(f_servLK[6][1]) != float:
                doc.add_paragraph("# оЗп_З: " + str(f_servLK[6][1]) + ", # Стоимость")
                o_mM.add('mM')
            else:
                doc.add_paragraph(str(f_servLK[6][1]) + ", # 18 Стоимость")
                serv1R.append(f_servLK[6][1])

            if f_servLK[7][0] != 'Возможность пациента оплатить услугу через сайт':
                doc.add_paragraph("# оЗп_И: '" + str(f_servLK[7][0]) + "', # Возможность пациента оплатить услугу через сайт")
                o_mM.add('mM')
            elif f_servLK[7][1] != 'не предоставлять пациенту возможность оплатить услугу через сайт' and f_servLK[7][1] != 'предоставить пациенту возможность оплатить услугу через сайт (взимается % от стоимости услуги за транзакционные издержки)' and f_servLK[7][1] != '':
                doc.add_paragraph("# оЗп_З: '" + str(f_servLK[7][1]) + "', # Возможность пациента оплатить услугу через сайт")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + f_servLK[7][1] + "', # 19 Возможность пациента оплатить услугу через сайт")
                serv1R.append(f_servLK[7][1])

            if f_servLK[8][0] != 'Ключевые поисковые слова предлагаемой услуги':
                doc.add_paragraph("# оЗп_И: '" + str(f_servLK[8][0]) + "', # Ключевые поисковые слова предлагаемой услуги")
                o_mM.add('mM')
            elif type(f_servLK[8][1]) != str:
                doc.add_paragraph("# оЗп_З: '" + str(f_servLK[8][1]) + "', # Ключевые поисковые слова предлагаемой услуги")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + str(f_servLK[8][1]) + "', # 20 Ключевые поисковые слова предлагаемой услуги")
                serv1R.append(f_servLK[8][1])

            o_mME = set()
            if f_servLK[9][0] != 'Email услуги':
                doc.add_paragraph("# оЗп_И: '" + str(f_servLK[9][0]) + "', # Email услуги")
                o_mME.add('mM')
            elif f_servLK[9][1] != '':
                if type(f_servLK[9][1]) != str or '@' not in f_servLK[9][1] or '.' not in f_servLK[9][1]:
                    doc.add_paragraph("# оЗп_З: '" + str(f_servLK[9][1]) + "', # Email услуги")
                    o_mME.add('mM')
            if o_mME == set():
                doc.add_paragraph("'" + f_servLK[9][1] + "', # 21 Email услуги")
                serv1R.append(f_servLK[9][1])
            o_mM |= o_mME

            o_mML = set()
            if f_servLK[10][0][0] != 'Укажите языки, на которых может быть предоставлена данная услуга':
                doc.add_paragraph("# оЗп_И: '" + str(f_servLK[10][0][0]) + "', # Укажите языки, на которых могут быть предоставлены услуги")
                o_mML.add('mM')
            elif type(f_servLK[10][0][1]) != list:
                doc.add_paragraph("# оЗп_З: '" + f_servLK[10][0][1] + "' не является списком, # Укажите языки, на которых могут быть предоставлены услуги")
            else:
                for q in f_servLK[10][0][1]:
                    if q not in oLang and q != '':
                        doc.add_paragraph("# оЗп_З: '" + str(q) + "', # Укажите языки, на которых могут быть предоставлены услуги, нет такого языка в списке")
                        o_mML.add('mM')    
            if type(f_servLK[10][1]) == list:
                if f_servLK[10][1][0] != 'добавить язык, отсутствующий в списке':
                    doc.add_paragraph("# оЗп_И: '" + str(f_servLK[10][1][0]) + "', # добавить язык, отсутствующий в списке")
                    o_mML.add('mM')
                elif type(f_servLK[10][1][1]) != str:
                    doc.add_paragraph("# оЗп_З: '" + str(f_servLK[10][1][1]) + "', # добавить язык, отсутствующий в списке")
                    o_mML.add('mM')
            if o_mML == set():
                f_alL = []
                for q in f_servLK[10][0][1]:
                    if q != '':
                        f_alL.append(q)
                if f_servLK[10][1][1] != '':
                    f_alL.append(f_servLK[10][1][1])
                str_alL = ''
                for q in f_alL:
                    if f_alL.index(q) != len(f_alL)-1:
                        str_alL = str_alL+q+', '
                    else:
                        str_alL = str_alL+q
                doc.add_paragraph("'" + str_alL + "', # 22 Адрес услуги / '' / Языки")       
                serv1R.append(str_alL)
            o_mM |= o_mML

            if f_servLK[11][0] != 'Время приема':
                doc.add_paragraph("# оЗп_И: '" + str(f_servLK[11][0]) + "', # Время приема")
                o_mM.add('mM')
            elif type(f_servLK[11][1]) != str:
                doc.add_paragraph("# оЗп_З: " + str(f_servLK[11][1]) + ", # Время приема")
                o_mM.add('mM')
            else:                   
                doc.add_paragraph("'" + str(f_servLK[11][1]) + "', # 23 Время приема")  
                serv1R.append(f_servLK[11][1])

            if f_servLK[12][0] != 'Возможность оказать данную услугу на дому':
                doc.add_paragraph("# оЗп_И: '" + str(f_servLK[12][0]) + "', # Возможность оказать данную услугу на дому")
                o_mM.add('mM')
            elif f_servLK[12][1] != '' and f_servLK[12][1] != 'возможен вызов на дом' and f_servLK[12][1] != 'вызов на дом не принимается':
                doc.add_paragraph("# оЗп_З: '" + str(f_servLK[12][1]) + "', # Возможность оказать данную услугу на дому")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + f_servLK[12][1] + "', # 24 Возможность оказать данную услугу на дому")
                serv1R.append(f_servLK[12][1])

            if f_servLK[13][0] != 'Дополнительная информация о предлагаемой услуге':
                doc.add_paragraph("# оЗп_И: '" + str(f_servLK[13][0]) + "', # Дополнительная информация о предлагаемой услуге")
                o_mM.add('mM')
            elif type(f_servLK[13][1]) != str:
                doc.add_paragraph("# оЗп_З: " + str(f_servLK[13][1]) + ", # Дополнительная информация о предлагаемой услуге")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + str(f_servLK[13][1]) + "', # 25 Дополнительная информация о предлагаемой услуге") 
                serv1R.append(f_servLK[13][1])

            doc.add_paragraph(")")
            doc.add_paragraph(")")

#---------------------
    if f_servLK[2][0] == 'Вид услуги':
        if 'консультация' in f_servLK[2][1]:
                doc.add_paragraph("print('    Специалист, " + str(f_servLK[2][1]) + ", " + str(f_servLK[3][2]) + ", " + str(f_servLK[7][1]) + "'), ")
        if 'процедура' in f_servLK[2][1]:
                doc.add_paragraph("print('    Специалист, " + "процедура: " + str(f_servLK[2][2]) + ", " + str(f_servLK[3][2]) + ", " + str(f_servLK[7][1]) + "'), ")

    if f_servLK[2][0] == 'Персонализированность услуги':
        if f_servLK[2][1] == 'персонализированная услуга (указываются данные специалиста)':
            if 'консультация' in f_servLK[8][1]:
                    doc.add_paragraph("print('    Клиника, персонализированная, " + str(f_servLK[3][1]) + " " + str(f_servLK[4][1]) + ", " + str(f_servLK[8][1]) + ", " + str(f_servLK[9][2]) + ", " + str(f_servLK[12][1]) + "'), ")
            if 'процедура' in f_servLK[8][1]:
                    doc.add_paragraph("print('    Клиника, персонализированная, " + "процедура: " + str(f_servLK[3][1]) + " " + str(f_servLK[4][1]) + ", " + str(f_servLK[8][2]) + ", " + str(f_servLK[12][1]) + ", " + "'), ")

        if f_servLK[2][1] == 'не персонализированная услуга (не указываются данные специалиста)':
            if 'консультация' in f_servLK[3][1]:
                    doc.add_paragraph("print('    Клиника, не персонализированная, " + str(f_servLK[3][1]) + ", " + str(f_servLK[4][2]) + ", " + str(f_servLK[6][1]) + "'), ")
            if 'процедура' in f_servLK[3][1]:
                    doc.add_paragraph("print('    Клиника, не персонализированная, " + "процедура: " + str(f_servLK[3][2]) + ", " + str(f_servLK[4][2]) + ", " + str(f_servLK[6][1]) + "'), ")
    doc.add_paragraph("print('---------------------------')\n")
    if o_mM != set():
        print(f"servLK_Ошибка_{cT}")
        doc.add_paragraph("print('" + NameFile + " отклонено оЗп_. ----------------------------------------------------------------')")
        doc.save(NameFile+cT+'_M.docx')
        return NameFile+cT+'_M.docx'
    else:
        doc.save(NameFile+'V.docx')
        print(f"servLK_Верно_{cT}\n")
        print(f"serv1R {serv1R}\n")

        return NameFile+'V.docx'
    
#-------------------------------------------------------------------------------------------------------------------------------------

# 20240909 1619

def isDigit(q):
    q = q.replace(',', '.')
    q = q.replace('^', '')
    q = q.replace('*', '')    
    q = q.replace('**', '')
    try:
        float(q)
        return True
    except ValueError:
        return False  

def send_servDel(f_servDel):

    import datetime
    cT = datetime.datetime.now().strftime('%Y.%m.%d.%H.%M.%S') # Время регистрации

    import docx
    doc = docx.Document()
    NameFile = '# Serv_'

    p = doc.add_paragraph()
    runner = p.add_run(NameFile)
    runner.bold = True

    serv1D = []
    o_mM = set()
    oLang = {'английский', 'арабский', 'испанский', 'итальянский', 'казахский', 'киргизский', 'китайский', 'корейский', 
            'монгольский', 'немецкий', 'персидский', 'португальский', 'русский', 'турецкий', 'французский', 'хинди', 'японский'}

    if f_servDel[2][0] != 'Персонализированность услуги' and f_servDel[2][0] != 'Вид услуги':
        doc.add_paragraph("# оЗп_И: '" + f_servDel[2][0] + "', # Имя поля [1][0] не 'Персонализированность услуги' и не 'Вид услуги'")
        o_mM.add('mM')
        doc.add_paragraph("print('" + NameFile + " не добавлена из-за оЗп_И Имя поля [1][0] не 'Персонализированность услуги' и не 'Вид услуги'.')")
        doc.add_paragraph("isprav.add('serv1D')") 

#---------------------

    if f_servDel[2][0] == 'Вид услуги':
        if len(f_servDel) != 15:
            o_mM.add('mM')
            doc.add_paragraph("print('f_servDelSpec имеет некорректную длинну: " + str(len(f_servDel)) + ", вместо 15')")
            if 15 - len(f_servDel) > 0:
                doc.add_paragraph("# f_servDelSpec - нехватка " + str(15 - len(f_servDel)) + " элементов!")
            if 15 - len(f_servDel) < 0:
                doc.add_paragraph("# f_servDelSpec - избыток " + str(len(f_servDel) - 15) + " элементов!")

        else:
            doc.add_paragraph("\nserv1D.add(") 
            doc.add_paragraph("(")

            if f_servDel[0][0] != 'Id':
                doc.add_paragraph("# оЗп_И: '" + str(f_servDel[0][0]) + "', # Id")
                o_mM.add('mM')
            elif type(f_servDel[0][1]) != int:
                doc.add_paragraph("# оЗп_З: " + str(f_servDel[0][1]) + "тип данных не int, # Id")
                o_mM.add('mM')
            else:
                doc.add_paragraph(str(f_servDel[0][1]) + ", # 0 Id")
                serv1D.append(f_servDel[0][1])

            if f_servDel[1][0] != 'IdProfile':
                doc.add_paragraph("# оЗп_И: '" + str(f_servDel[1][0]) + "', # IdProfile")
                o_mM.add('mM')
            elif type(f_servDel[1][1]) != int:
                doc.add_paragraph("# оЗп_З: " + str(f_servDel[1][1]) + " тип данных не int, # IdProfile")
                o_mM.add('mM')
            else:
                doc.add_paragraph("" + str(f_servDel[1][1]) + ", # 1 IdProfile")
                serv1D.append(f_servDel[1][1])

            doc.add_paragraph("'', # 2 '' / Персонализированность услуги")
            serv1D.append('')
            doc.add_paragraph("'', # 3 '' / Фамилия")
            serv1D.append('')
            doc.add_paragraph("'', # 4 '' / Имя")
            serv1D.append('')
            doc.add_paragraph("'', # 5 '' / Отчество")
            serv1D.append('')
            doc.add_paragraph("'', # 6 '' / Пол")
            serv1D.append('')
            doc.add_paragraph("'', # 7 '' / Дата рождения")
            serv1D.append('')

            if f_servDel[2][0] != 'Вид услуги':
                doc.add_paragraph("# оЗп_И: '" + str(f_servDel[2][0]) + "', # Вид услуги")
                o_mM.add('mM')
            elif f_servDel[2][1] != 'дистанционная консультация' and f_servDel[2][1] != 'консультация на очном приеме' and f_servDel[2][1] != 'консультация, возможна дистанционная или на очном приеме' and f_servDel[2][1] != 'процедура (манипуляция, операция, …)':
                doc.add_paragraph("# оЗп_З: '" + str(f_servDel[2][1]) + "', # Вид услуги")
                o_mM.add('mM')
            elif f_servDel[2][1] == 'процедура (манипуляция, операция, …)' and len(f_servDel[2]) != 3:
                doc.add_paragraph("# оЗп_З: поле Название процедуры - не обязательное # Вид услуги")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + f_servDel[2][1] + "', # 8 Конс. дист. / Конс. оч. / Дист. или оч. конс. / Проц.")
                serv1D.append(f_servDel[2][1])
                if len(f_servDel[2]) == 2:
                    doc.add_paragraph("'', # 9 Название процедуры")
                    serv1D.append('')
                if len(f_servDel[2]) == 3:
                    doc.add_paragraph("'" + f_servDel[2][2] + "', # 9 Название процедуры")
                    serv1D.append(f_servDel[2][2])

            if f_servDel[3][0] != 'Тип услуги':
                doc.add_paragraph("# оЗп_И: '" + str(f_servDel[3][0]) + "', # Тип услуги")
                o_mM.add('mM')
            elif f_servDel[3][1] != 'врачебная услуга' and f_servDel[3][1] != 'не врачебная услуга' and f_servDel[3][1] != 'не медицинская услуга':
                doc.add_paragraph("# оЗп_З: '" + str(f_servDel[3][1]) + "', # Тип услуги")
                o_mM.add('mM')
            elif f_servDel[3][1] == 'не медицинская услуга' and len(f_servDel[3]) != 3:
                doc.add_paragraph("# оЗп_З: поле Тип услуги по не медицинской услуге - len != 3 # Тип услуги")
                o_mM.add('mM')
            elif f_servDel[3][1] == 'врачебная услуга' and len(f_servDel[3]) != 6:
                doc.add_paragraph("# оЗп_З: поле Тип услуги по врачебной услуге - len != 6 # Тип услуги")
                o_mM.add('mM')
            elif f_servDel[3][1] == 'не врачебная услуга' and len(f_servDel[3]) != 6:
                doc.add_paragraph("# оЗп_З: поле Тип услуги по не врачебной услуге - len != 6 # Тип услуги")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + f_servDel[3][1] + "', # 10 ВрУсл / НеВрУсл / НеМедУсл")
                serv1D.append(f_servDel[3][1])
                doc.add_paragraph("'" + f_servDel[3][2] + "', # 11 Название специальности")
                serv1D.append(f_servDel[3][2])
                if len(f_servDel[3]) == 3:
                    doc.add_paragraph("'', # 12 Сертификат по специальности, начало (чч.мм.гггг)")
                    serv1D.append('')
                    doc.add_paragraph("'', # 13 Сертификат по специальности, окончание (чч.мм.гггг)")
                    serv1D.append('')
                    doc.add_paragraph("'', # 14 Категория по специальности")
                    serv1D.append('')                    
                if len(f_servDel[3]) == 6:
                    doc.add_paragraph("'" + f_servDel[3][3] + "', # 12 Сертификат по специальности, начало (чч.мм.гггг)")
                    serv1D.append(f_servDel[3][3])
                    doc.add_paragraph("'" + f_servDel[3][4] + "', # 13 Сертификат по специальности, окончание (чч.мм.гггг)")
                    serv1D.append(f_servDel[3][4])
                    doc.add_paragraph("'" + f_servDel[3][5] + "', # 14 Категория по специальности")
                    serv1D.append(f_servDel[3][5])

            if f_servDel[5][0] != 'Возраст пациентов, лет':
                doc.add_paragraph("# оЗп_И: '" + str(f_servDel[5][0]) + "', # Возраст пациентов, лет")
                o_mM.add('mM')
            elif type(f_servDel[5][1]) != int and type(f_servDel[5][1]) != float and f_servDel[5][1] != '':
                doc.add_paragraph("# оЗп_З: '" + str(f_servDel[5][1]) + "', # Возраст пациентов, лет, от")
            else:
                serv1D.append(f_servDel[5][1])
                if type(f_servDel[5][1]) == int or type(f_servDel[5][1]) == float:
                    doc.add_paragraph(str(f_servDel[5][1]) + ", # 15 Возраст пациентов, лет, от")
                if type(f_servDel[5][1]) == str:
                    doc.add_paragraph("'" + str(f_servDel[5][1]) + "', # 15 Возраст пациентов, лет, от")

            if type(f_servDel[5][2]) != int and type(f_servDel[5][2]) != float and f_servDel[5][2] != '':
                doc.add_paragraph("# оЗп_З: '" + str(f_servDel[5][2]) + "', # Возраст пациентов, лет, до")
            else:
                serv1D.append(f_servDel[5][2])
                if type(f_servDel[5][2]) == int or type(f_servDel[5][2]) == float:
                    doc.add_paragraph(str(f_servDel[5][2]) + ", # 16 Возраст пациентов, лет, до")
                if type(f_servDel[5][2]) == str:
                    doc.add_paragraph("'" + str(f_servDel[5][2]) + "', # 16 Возраст пациентов, лет, до")

            if f_servDel[6][0] != 'Стаж по предлагаемой услуге, лет':
                doc.add_paragraph("# оЗп_И: '" + str(f_servDel[6][0]) + "', # Стаж по предлагаемой услуге, лет")
                o_mM.add('mM')
            elif type(f_servDel[6][1]) != int and type(f_servDel[6][1]) != float:
                doc.add_paragraph("# оЗп_З: " + str(f_servDel[6][1]) + ", # Стаж по предлагаемой услуге, лет")
                o_mM.add('mM')
            else:
                doc.add_paragraph(str(f_servDel[6][1]) + ", # 17 Стаж по предлагаемой услуге, лет")
                serv1D.append(f_servDel[6][1])

            if f_servDel[7][0] != 'Стоимость':
                doc.add_paragraph("# оЗп_И: '" + str(f_servDel[7][0]) + "', # Стоимость")
                o_mM.add('mM')
            elif type(f_servDel[7][1]) != int and type(f_servDel[7][1]) != float:
                doc.add_paragraph("# оЗп_З: " + str(f_servDel[7][1]) + ", # Стоимость")
                o_mM.add('mM')
            else:
                doc.add_paragraph(str(f_servDel[7][1]) + ", # 18 Стоимость")
                serv1D.append(f_servDel[7][1])

            if f_servDel[8][0] != 'Возможность пациента оплатить услугу через сайт':
                doc.add_paragraph("# оЗп_И: '" + str(f_servDel[8][0]) + "', # Возможность пациента оплатить услугу через сайт")
                o_mM.add('mM')
            elif f_servDel[8][1] != 'не предоставлять пациенту возможность оплатить услугу через сайт' and f_servDel[8][1] != 'предоставить пациенту возможность оплатить услугу через сайт (взимается % от стоимости услуги за транзакционные издержки)' and f_servDel[8][1] != '':
                doc.add_paragraph("# оЗп_З: '" + str(f_servDel[8][1]) + "', # Возможность пациента оплатить услугу через сайт")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + f_servDel[8][1] + "', # 19 Возможность пациента оплатить услугу через сайт")
                serv1D.append(f_servDel[8][1])

            if f_servDel[9][0] != 'Ключевые поисковые слова предлагаемой услуги':
                doc.add_paragraph("# оЗп_И: '" + str(f_servDel[9][0]) + "', # Ключевые поисковые слова предлагаемой услуги")
                o_mM.add('mM')
            elif type(f_servDel[9][1]) != str:
                doc.add_paragraph("# оЗп_З: '" + str(f_servDel[9][1]) + "', # Ключевые поисковые слова предлагаемой услуги")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + str(f_servDel[9][1]) + "', # 20 Ключевые поисковые слова предлагаемой услуги")
                serv1D.append(f_servDel[9][1])

            if f_servDel[10][0] != 'Название клиники, где оказывается услуга':
                doc.add_paragraph("# оЗп_И: '" + str(f_servDel[10][0]) + "', # Название клиники, где оказывается услуга")
                o_mM.add('mM')
            elif type(f_servDel[10][1]) != str:
                doc.add_paragraph("# оЗп_З: '" + str(f_servDel[10][1]) + "', # Название клиники, где оказывается услуга")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + str(f_servDel[10][1]) + "', # 21 Название клиники, где оказывается услуга")
                serv1D.append(f_servDel[10][1])

            if f_servDel[11][0] != 'Адрес, где оказывается услуга':
                doc.add_paragraph("# оЗп_И: '" + str(f_servDel[11][0]) + "', # Адрес, где оказывается услуга")
                o_mM.add('mM')
            elif type(f_servDel[11][1]) != str:
                doc.add_paragraph("# оЗп_З: '" + str(f_servDel[11][1]) + "', # Адрес, где оказывается услуга")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + str(f_servDel[11][1]) + "', # 22 Адрес, где оказывается услуга")
                serv1D.append(f_servDel[11][1])

            if f_servDel[12][0] != 'Время приема':
                doc.add_paragraph("# оЗп_И: '" + str(f_servDel[12][0]) + "', # Время приема")
                o_mM.add('mM')
            elif type(f_servDel[12][1]) != str:
                doc.add_paragraph("# оЗп_З: " + str(f_servDel[12][1]) + ", # Время приема")
                o_mM.add('mM')
            else:                   
                doc.add_paragraph("'" + str(f_servDel[12][1]) + "', # 23 Время приема")  
                serv1D.append(f_servDel[12][1])

            if f_servDel[4][0] != 'Возможность оказать данную услугу на дому':
                doc.add_paragraph("# оЗп_И: '" + str(f_servDel[4][0]) + "', # Возможность оказать данную услугу на дому")
                o_mM.add('mM')
            elif f_servDel[4][1] != 'возможен вызов на дом' and f_servDel[4][1] != 'вызов на дом не принимается':
                doc.add_paragraph("# оЗп_З: '" + str(f_servDel[4][1]) + "', # Возможность оказать данную услугу на дому")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + f_servDel[4][1] + "', # 24 Возможность оказать данную услугу на дому")
                serv1D.append(f_servDel[4][1])

            if f_servDel[13][0] != 'Дополнительная информация о предлагаемой услуге':
                doc.add_paragraph("# оЗп_И: '" + str(f_servDel[13][0]) + "', # Дополнительная информация о предлагаемой услуге")
                o_mM.add('mM')
            elif type(f_servDel[13][1]) != str:
                doc.add_paragraph("# оЗп_З: " + str(f_servDel[13][1]) + ", # Дополнительная информация о предлагаемой услуге")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + str(f_servDel[13][1]) + "', # 25 Дополнительная информация о предлагаемой услуге") 
                serv1D.append(f_servDel[13][1])

            if f_servDel[14] != 'servDel':
                doc.add_paragraph("# оИдент: '" + str(f_servDel[14]) + "', # servDel")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + str(f_servDel[14]) + "', # 26 servDel") 
                serv1D.append(f_servDel[14])

            doc.add_paragraph("),")
            doc.add_paragraph(")")

#---------------------
    if f_servDel[2][0] == 'Персонализированность услуги':
        if f_servDel[2][1] != 'персонализированная услуга (указываются данные специалиста)' and f_servDel[2][1] != 'не персонализированная услуга (не указываются данные специалиста)':
            o_mM.add('mM')
            doc.add_paragraph("# оЗп_З: '" + str(f_servDel[2][1]) + " не шаблонный варинат ('персонализированная услуга (указываются данные специалиста)', 'не персонализированная услуга (не указываются данные специалиста)') ', # Персонализированность услуги")
        doc.add_paragraph("'" + str(f_servDel[2][1]) + "', # Персонализированность услуги")
#---------------------

    if f_servDel[2][1] == 'персонализированная услуга (указываются данные специалиста)':
        if len(f_servDel) != 20:
            o_mM.add('mM')
            doc.add_paragraph("print('f_servDelPerson имеет некорректную длинну: " + str(len(f_servDel)) + ", вместо 20')")
            if 20 - len(f_servDel) > 0:
                doc.add_paragraph("# f_servDelPerson - нехватка " + str(20 - len(f_servDel)) + " элементов!")
            if 20 - len(f_servDel) < 0:
                doc.add_paragraph("# f_servDelPerson - избыток " + str(len(f_servDel) - 20) + " элементов!")

        else:

            doc.add_paragraph("\nserv1D.add(") 
            doc.add_paragraph("(")

            if f_servDel[0][0] != 'Id':
                doc.add_paragraph("# оЗп_И: '" + str(f_servDel[0][0]) + "', # Id")
                o_mM.add('mM')
            elif type(f_servDel[0][1]) != int:
                doc.add_paragraph("# оЗп_З: '" + str(f_servDel[0][1]) + "', # Id")
                o_mM.add('mM')
            else:
                doc.add_paragraph("" + str(f_servDel[0][1]) + ", # 0 Id")
                serv1D.append(f_servDel[0][1])

            if f_servDel[1][0] != 'IdProfile':
                doc.add_paragraph("# оЗп_И: '" + str(f_servDel[1][0]) + "', # IdProfile")
                o_mM.add('mM')
            elif type(f_servDel[1][1]) != int:
                doc.add_paragraph("# оЗп_З: '" + str(f_servDel[1][1]) + "', # IdProfile")
                o_mM.add('mM')
            else:
                doc.add_paragraph("" + str(f_servDel[1][1]) + ", # 1 IdProfile")
                serv1D.append(f_servDel[1][1])

            if f_servDel[2][0] != 'Персонализированность услуги':
                doc.add_paragraph("# оЗп_И: '" + str(f_servDel[2][0]) + "', # Персонализированность услуги")
                o_mM.add('mM')
            elif f_servDel[2][1] != 'персонализированная услуга (указываются данные специалиста)' and f_servDel[2][1] != 'не персонализированная услуга (не указываются данные специалиста)' and f_servDel[2][1] != '':
                doc.add_paragraph("# оЗп_З: '" + str(f_servDel[2][1]) + "', # Персонализированность услуги")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + f_servDel[2][1] + "', # 2 Персонализированность услуги")
                serv1D.append(f_servDel[2][1])

            if f_servDel[3][0] != 'Фамилия':
                doc.add_paragraph("# оЗп_И: '" + str(f_servDel[3][0]) + "', # Фамилия")
                o_mM.add('mM')
            elif f_servDel[3][1] == '':
                doc.add_paragraph("# оЗп_З: '', # Фамилия")
                o_mM.add('mM')
            elif type(f_servDel[3][1]) != str:
                doc.add_paragraph("# оЗп_З: " + str(f_servDel[3][1]) + ", # Фамилия")
                o_mM.add('mM')
                doc.add_paragraph(str(f_servDel[3][1]) + ", # Фамилия")          
            else:
                doc.add_paragraph("'" + f_servDel[3][1] + "', # 3 Фамилия")
                serv1D.append(f_servDel[3][1])
          
            if f_servDel[4][0] != 'Имя':
                doc.add_paragraph("# оЗп_И: '" + str(f_servDel[4][0]) + "', # Имя")
                o_mM.add('mM')
            elif f_servDel[4][1] == '':
                doc.add_paragraph("# оЗп_З: '', # Имя")
                o_mM.add('mM')
            elif type(f_servDel[4][1]) != str:
                doc.add_paragraph("# оЗп_З: " + str(f_servDel[4][1]) + ", # Имя")
                o_mM.add('mM')
                doc.add_paragraph(str(f_servDel[4][1]) + ", # Имя")          
            else:
                doc.add_paragraph("'" + f_servDel[4][1] + "', # 4 Имя")
                serv1D.append(f_servDel[4][1]) 

            if f_servDel[5][0] != 'Отчество':
                doc.add_paragraph("# оЗп_И: '" + str(f_servDel[5][0]) + "', # Отчество")
                o_mM.add('mM')
            elif type(f_servDel[5][1]) != str:
                doc.add_paragraph("# оЗп_З: " + str(f_servDel[5][1]) + ", # Отчество")
                o_mM.add('mM')
                doc.add_paragraph(str(f_servDel[5][1]) + ", # Отчество")
            else:
                doc.add_paragraph("'" + f_servDel[5][1] + "', # 5 Отчество")       
                serv1D.append(f_servDel[5][1])

            if f_servDel[6][0] != 'Пол':
                doc.add_paragraph("# оЗп_И: '" + str(f_servDel[6][0]) + "', # Пол")
                o_mM.add('mM')
            elif f_servDel[6][1] != 'мужской' and f_servDel[6][1] != 'женский':
                doc.add_paragraph("# оЗп_З: '" + str(f_servDel[6][1]) + "', # Пол")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + f_servDel[6][1] + "', # 6 Пол")       
                serv1D.append(f_servDel[6][1])

            if f_servDel[7][0] != 'Дата рождения':
                doc.add_paragraph("# оЗп_И: '" + str(f_servDel[7][0]) + "', # Дата рождения")
                o_mM.add('mM')
            elif type(f_servDel[7][1]) == str:
                if isDigit(f_servDel[7][1].replace('.', '')) != True or len(f_servDel[7][1]) != 10 or len(f_servDel[7][1].replace('.', '')) != 8:
                    doc.add_paragraph("# оЗп_З: '" + str(f_servDel[7][1]) + "', # Дата рождения")
                    o_mM.add('mM')
                else:
                    doc.add_paragraph("'" + f_servDel[7][1] + "', # 7 '' / Дата рождения")
                    serv1D.append(f_servDel[7][1])

            if f_servDel[8][0] != 'Вид услуги':
                doc.add_paragraph("# оЗп_И: '" + str(f_servDel[8][0]) + "', # Вид услуги")
                o_mM.add('mM')
            elif f_servDel[8][1] != 'дистанционная консультация' and f_servDel[8][1] != 'консультация на очном приеме' and f_servDel[8][1] != 'консультация, возможна дистанционная или на очном приеме' and f_servDel[8][1] != 'процедура (манипуляция, операция, …)':
                doc.add_paragraph("# оЗп_З: '" + str(f_servDel[8][1]) + "', # Вид услуги")
                o_mM.add('mM')
            elif f_servDel[8][1] == 'процедура (манипуляция, операция, …)' and len(f_servDel[8]) != 3:
                doc.add_paragraph("# оЗп_З: поле Название процедуры - не обязательное # Вид услуги")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + f_servDel[8][1] + "', # 8 Конс. дист. / Конс. оч. / Дист. или оч. конс. / Проц.")
                serv1D.append(f_servDel[8][1])
                if len(f_servDel[8]) == 2:
                    doc.add_paragraph("'', # 9 Название процедуры")
                    serv1D.append('')
                if len(f_servDel[8]) == 3:
                    doc.add_paragraph("'" + f_servDel[8][2] + "', # 9 Название процедуры")
                    serv1D.append(f_servDel[8][2])

            if f_servDel[9][0] != 'Тип услуги':
                doc.add_paragraph("# оЗп_И: '" + str(f_servDel[9][0]) + "', # Тип услуги")
                o_mM.add('mM')
            elif f_servDel[9][1] != 'врачебная услуга' and f_servDel[9][1] != 'не врачебная услуга' and f_servDel[9][1] != 'не медицинская услуга':
                doc.add_paragraph("# оЗп_З: '" + str(f_servDel[9][1]) + "', # Тип услуги")
                o_mM.add('mM')
            elif f_servDel[9][1] == 'не медицинская услуга' and len(f_servDel[9]) != 3:
                doc.add_paragraph("# оЗп_З: поле Тип услуги по не медицинской услуге - len != 3 # Тип услуги")
                o_mM.add('mM')
            elif f_servDel[9][1] == 'врачебная услуга' and len(f_servDel[9]) != 6:
                doc.add_paragraph("# оЗп_З: поле Тип услуги по врачебной услуге - len != 6 # Тип услуги")
                o_mM.add('mM')
            elif f_servDel[9][1] == 'не врачебная услуга' and len(f_servDel[9]) != 6:
                doc.add_paragraph("# оЗп_З: поле Тип услуги по не врачебной услуге - len != 6 # Тип услуги")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + f_servDel[9][1] + "', # 10 ВрУсл / НеВрУсл / НеМедУсл")
                serv1D.append(f_servDel[9][1])
                doc.add_paragraph("'" + f_servDel[9][2] + "', # 11 Название специальности")
                serv1D.append(f_servDel[9][2])
                if len(f_servDel[9]) == 3:
                    doc.add_paragraph("'', # 12 Сертификат по специальности, начало (чч.мм.гггг)")
                    serv1D.append('')
                    doc.add_paragraph("'', # 13 Сертификат по специальности, окончание (чч.мм.гггг)")
                    serv1D.append('')
                    doc.add_paragraph("'', # 14 Категория по специальности")
                    serv1D.append('')                    
                if len(f_servDel[9]) == 6:
                    doc.add_paragraph("'" + f_servDel[9][3] + "', # 12 Сертификат по специальности, начало (чч.мм.гггг)")
                    serv1D.append(f_servDel[9][3])
                    doc.add_paragraph("'" + f_servDel[9][4] + "', # 13 Сертификат по специальности, окончание (чч.мм.гггг)")
                    serv1D.append(f_servDel[9][4])
                    doc.add_paragraph("'" + f_servDel[9][5] + "', # 14 Категория по специальности")
                    serv1D.append(f_servDel[9][5])

            if f_servDel[10][0] != 'Возраст пациентов, лет':
                doc.add_paragraph("# оЗп_И: '" + str(f_servDel[10][0]) + "', # Возраст пациентов, лет")
                o_mM.add('mM')
            elif type(f_servDel[10][1]) != int and type(f_servDel[10][1]) != float and f_servDel[10][1] != '':
                doc.add_paragraph("# оЗп_З: '" + str(f_servDel[10][1]) + "', # Возраст пациентов, лет, от")
            else:
                serv1D.append(f_servDel[10][1])
                if type(f_servDel[10][1]) == int or type(f_servDel[10][1]) == float:
                    doc.add_paragraph(str(f_servDel[10][1]) + ", # 15 Возраст пациентов, лет, от")
                if type(f_servDel[10][1]) == str:
                    doc.add_paragraph("'" + str(f_servDel[10][1]) + "', # 15 Возраст пациентов, лет, от")

            if type(f_servDel[10][2]) != int and type(f_servDel[10][2]) != float and f_servDel[10][2] != '':
                doc.add_paragraph("# оЗп_З: '" + str(f_servDel[10][2]) + "', # Возраст пациентов, лет, до")
            else:
                serv1D.append(f_servDel[10][2])
                if type(f_servDel[10][2]) == int or type(f_servDel[10][2]) == float:
                    doc.add_paragraph(str(f_servDel[10][2]) + ", # 16 Возраст пациентов, лет, до")
                if type(f_servDel[10][2]) == str:
                    doc.add_paragraph("'" + str(f_servDel[10][2]) + "', # 16 Возраст пациентов, лет, до")

            if f_servDel[11][0] != 'Стаж по предлагаемой услуге, лет':
                doc.add_paragraph("# оЗп_И: '" + str(f_servDel[11][0]) + "', # Стаж по предлагаемой услуге, лет")
                o_mM.add('mM')
            elif type(f_servDel[11][1]) != int and type(f_servDel[11][1]) != float:
                doc.add_paragraph("# оЗп_З: " + str(f_servDel[11][1]) + ", # Стаж по предлагаемой услуге, лет")
                o_mM.add('mM')
            else:
                doc.add_paragraph(str(f_servDel[11][1]) + ", # 17 Стаж по предлагаемой услуге, лет")
                serv1D.append(f_servDel[11][1])

            if f_servDel[12][0] != 'Стоимость':
                doc.add_paragraph("# оЗп_И: '" + str(f_servDel[12][0]) + "', # Стоимость")
                o_mM.add('mM')
            elif type(f_servDel[12][1]) != int and type(f_servDel[12][1]) != float:
                doc.add_paragraph("# оЗп_З: " + str(f_servDel[12][1]) + ", # Стоимость")
                o_mM.add('mM')
            else:
                doc.add_paragraph(str(f_servDel[12][1]) + ", # 18 Стоимость")
                serv1D.append(f_servDel[12][1])

            if f_servDel[13][0] != 'Возможность пациента оплатить услугу через сайт':
                doc.add_paragraph("# оЗп_И: '" + str(f_servDel[13][0]) + "', # Возможность пациента оплатить услугу через сайт")
                o_mM.add('mM')
            elif f_servDel[13][1] != 'не предоставлять пациенту возможность оплатить услугу через сайт' and f_servDel[13][1] != 'предоставить пациенту возможность оплатить услугу через сайт (взимается % от стоимости услуги за транзакционные издержки)' and f_servDel[13][1] != '':
                doc.add_paragraph("# оЗп_З: '" + str(f_servDel[13][1]) + "', # Возможность пациента оплатить услугу через сайт")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + f_servDel[13][1] + "', # 19 Возможность пациента оплатить услугу через сайт")
                serv1D.append(f_servDel[13][1])

            if f_servDel[14][0] != 'Ключевые поисковые слова предлагаемой услуги':
                doc.add_paragraph("# оЗп_И: '" + str(f_servDel[14][0]) + "', # Ключевые поисковые слова предлагаемой услуги")
                o_mM.add('mM')
            elif type(f_servDel[14][1]) != str:
                doc.add_paragraph("# оЗп_З: '" + str(f_servDel[14][1]) + "', # Ключевые поисковые слова предлагаемой услуги")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + str(f_servDel[14][1]) + "', # 20 Ключевые поисковые слова предлагаемой услуги")
                serv1D.append(f_servDel[14][1])

            o_mME = set()
            if f_servDel[15][0] != 'Email услуги':
                doc.add_paragraph("# оЗп_И: '" + str(f_servDel[15][0]) + "', # Email услуги")
                o_mME.add('mM')
            elif f_servDel[15][1] != '':
                if type(f_servDel[15][1]) != str or '@' not in f_servDel[15][1] or '.' not in f_servDel[15][1]:
                    doc.add_paragraph("# оЗп_З: '" + str(f_servDel[15][1]) + "', # Email услуги")
                    o_mME.add('mM')
            if o_mME == set():
                doc.add_paragraph("'" + f_servDel[15][1] + "', # 21 Email услуги")
                serv1D.append(f_servDel[15][1])
            o_mM |= o_mME

            doc.add_paragraph("'', # 22 '' / (Адрес услуги)")
            serv1D.append('')

            if f_servDel[16][0] != 'Время приема':
                doc.add_paragraph("# оЗп_И: '" + str(f_servDel[16][0]) + "', # Время приема")
                o_mM.add('mM')
            elif type(f_servDel[16][1]) != str:
                doc.add_paragraph("# оЗп_З: " + str(f_servDel[16][1]) + ", # Время приема")
                o_mM.add('mM')
            else:                   
                doc.add_paragraph("'" + str(f_servDel[16][1]) + "', # 23 Время приема")  
                serv1D.append(f_servDel[16][1])

            if f_servDel[17][0] != 'Возможность оказать данную услугу на дому':
                doc.add_paragraph("# оЗп_И: '" + str(f_servDel[17][0]) + "', # Возможность оказать данную услугу на дому")
                o_mM.add('mM')
            elif f_servDel[17][1] != '' and f_servDel[17][1] != 'возможен вызов на дом' and f_servDel[17][1] != 'вызов на дом не принимается':
                doc.add_paragraph("# оЗп_З: '" + str(f_servDel[17][1]) + "', # Возможность оказать данную услугу на дому")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + f_servDel[17][1] + "', # 24 Возможность оказать данную услугу на дому")
                serv1D.append(f_servDel[17][1])

            if f_servDel[18][0] != 'Дополнительная информация о предлагаемой услуге':
                doc.add_paragraph("# оЗп_И: '" + str(f_servDel[18][0]) + "', # Дополнительная информация о предлагаемой услуге")
                o_mM.add('mM')
            elif type(f_servDel[18][1]) != str:
                doc.add_paragraph("# оЗп_З: " + str(f_servDel[18][1]) + ", # Дополнительная информация о предлагаемой услуге")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + str(f_servDel[18][1]) + "', # 25 Дополнительная информация о предлагаемой услуге") 
                serv1D.append(f_servDel[18][1])

            if f_servDel[19] != 'servDel':
                doc.add_paragraph("# оИдент: '" + str(f_servDel[19]) + "', # servDel")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + str(f_servDel[19]) + "', # 26 servDel") 
                serv1D.append(f_servDel[19])

            doc.add_paragraph(")")
            doc.add_paragraph(")")

#---------------------

    if f_servDel[2][1] == 'не персонализированная услуга (не указываются данные специалиста)':
        if len(f_servDel) != 15:
            o_mM.add('mM')
            doc.add_paragraph("print('f_servDelNotPerson имеет некорректную длинну: " + str(len(f_servDel)) + ", вместо 15')")
            if 15 - len(f_servDel) > 0:
                doc.add_paragraph("# f_servDelNotPerson - нехватка " + str(15 - len(f_servDel)) + " элементов!")
            if 15 - len(f_servDel) < 0:
                doc.add_paragraph("# f_servDelNotPerson - избыток " + str(len(f_servDel) - 15) + " элементов!")

        else:

            doc.add_paragraph("\nserv1D.add(") 
            doc.add_paragraph("(")

            if f_servDel[0][0] != 'Id':
                doc.add_paragraph("# оЗп_И: '" + str(f_servDel[0][0]) + "', # Id")
                o_mM.add('mM')
            elif type(f_servDel[0][1]) != int:
                doc.add_paragraph("# оЗп_З: '" + str(f_servDel[0][1]) + "', # Id")
                o_mM.add('mM')
            else:
                doc.add_paragraph("" + str(f_servDel[0][1]) + ", # 0 Id")
                serv1D.append(f_servDel[0][1])

            if f_servDel[1][0] != 'IdProfile':
                doc.add_paragraph("# оЗп_И: '" + str(f_servDel[1][0]) + "', # IdProfile")
                o_mM.add('mM')
            elif type(f_servDel[1][1]) != int:
                doc.add_paragraph("# оЗп_З: '" + str(f_servDel[1][1]) + "', # IdProfile")
                o_mM.add('mM')
            else:
                doc.add_paragraph("" + str(f_servDel[1][1]) + ", # 1 IdProfile")
                serv1D.append(f_servDel[1][1])

            if f_servDel[2][0] != 'Персонализированность услуги':
                doc.add_paragraph("# оЗп_И: '" + str(f_servDel[2][0]) + "', # Персонализированность услуги")
                o_mM.add('mM')
            elif f_servDel[2][1] != 'персонализированная услуга (указываются данные специалиста)' and f_servDel[2][1] != 'не персонализированная услуга (не указываются данные специалиста)' and f_servDel[2][1] != '':
                doc.add_paragraph("# оЗп_З: '" + str(f_servDel[2][1]) + "', # Персонализированность услуги")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + f_servDel[2][1] + "', # 2 Персонализированность услуги")
                serv1D.append(f_servDel[2][1])

            doc.add_paragraph("'', # 3 '' / Фамилия")
            serv1D.append('')
            doc.add_paragraph("'', # 4 '' / Имя")
            serv1D.append('')
            doc.add_paragraph("'', # 5 '' / Отчество")
            serv1D.append('')
            doc.add_paragraph("'', # 6 '' / Пол")
            serv1D.append('')
            doc.add_paragraph("'', # 7 '' / Дата рождения")
            serv1D.append('')

            if f_servDel[3][0] != 'Вид услуги':
                doc.add_paragraph("# оЗп_И: '" + str(f_servDel[3][0]) + "', # Вид услуги")
                o_mM.add('mM')
            elif f_servDel[3][1] != 'дистанционная консультация' and f_servDel[3][1] != 'консультация на очном приеме' and f_servDel[3][1] != 'консультация, возможна дистанционная или на очном приеме' and f_servDel[3][1] != 'процедура (манипуляция, операция, …)':
                doc.add_paragraph("# оЗп_З: '" + str(f_servDel[3][1]) + "', # Вид услуги")
                o_mM.add('mM')
            elif f_servDel[3][1] == 'процедура (манипуляция, операция, …)' and len(f_servDel[3]) != 3:
                doc.add_paragraph("# оЗп_З: поле Название процедуры - не обязательное # Вид услуги")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + f_servDel[3][1] + "', # 8 Конс. дист. / Конс. оч. / Дист. или оч. конс. / Проц.")
                serv1D.append(f_servDel[3][1])
                if len(f_servDel[3]) == 2:
                    doc.add_paragraph("'', # 9 Название процедуры")
                    serv1D.append('')
                if len(f_servDel[3]) == 3:
                    doc.add_paragraph("'" + f_servDel[3][2] + "', # 9 Название процедуры")
                    serv1D.append(f_servDel[3][2])

            if f_servDel[4][0] != 'Тип услуги':
                doc.add_paragraph("# оЗп_И: '" + str(f_servDel[4][0]) + "', # Тип услуги")
                o_mM.add('mM')
            elif f_servDel[4][1] != 'врачебная услуга' and f_servDel[4][1] != 'не врачебная услуга' and f_servDel[4][1] != 'не медицинская услуга':
                doc.add_paragraph("# оЗп_З: '" + str(f_servDel[4][1]) + "', # Тип услуги")
                o_mM.add('mM')
            elif f_servDel[4][1] == 'не медицинская услуга' and len(f_servDel[4]) != 3:
                doc.add_paragraph("# оЗп_З: поле Тип услуги по не медицинской услуге - len != 3 # Тип услуги")
                o_mM.add('mM')
            elif f_servDel[4][1] == 'врачебная услуга' and len(f_servDel[4]) != 3:
                doc.add_paragraph("# оЗп_З: поле Тип услуги по врачебной услуге - len != 3 # Тип услуги")
                o_mM.add('mM')
            elif f_servDel[4][1] == 'не врачебная услуга' and len(f_servDel[4]) != 3:
                doc.add_paragraph("# оЗп_З: поле Тип услуги по не врачебной услуге - len != 3 # Тип услуги")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + f_servDel[4][1] + "', # 10 ВрУсл / НеВрУсл / НеМедУсл")
                serv1D.append(f_servDel[4][1])
                doc.add_paragraph("'" + f_servDel[4][2] + "', # 11 Название специальности")
                serv1D.append(f_servDel[4][2])

            doc.add_paragraph("'', # 12 '' (Сертификат по специальности, начало (чч.мм.гггг))")
            serv1D.append('')
            doc.add_paragraph("'', # 13 '' (Сертификат по специальности, окончание (чч.мм.гггг))")
            serv1D.append('')
            doc.add_paragraph("'', # 14 '' (Категория по специальности)")
            serv1D.append('')

            if f_servDel[5][0] != 'Возраст пациентов, лет':
                doc.add_paragraph("# оЗп_И: '" + str(f_servDel[5][0]) + "', # Возраст пациентов, лет")
                o_mM.add('mM')
            elif type(f_servDel[5][1]) != int and type(f_servDel[5][1]) != float and f_servDel[5][1] != '':
                doc.add_paragraph("# оЗп_З: '" + str(f_servDel[5][1]) + "', # Возраст пациентов, лет, от")
            else:
                serv1D.append(f_servDel[5][1])
                if type(f_servDel[5][1]) == int or type(f_servDel[5][1]) == float:
                    doc.add_paragraph(str(f_servDel[5][1]) + ", # 15 Возраст пациентов, лет, от")
                if type(f_servDel[5][1]) == str:
                    doc.add_paragraph("'" + str(f_servDel[5][1]) + "', # 15 Возраст пациентов, лет, от")

            if type(f_servDel[5][2]) != int and type(f_servDel[5][2]) != float and f_servDel[5][2] != '':
                doc.add_paragraph("# оЗп_З: '" + str(f_servDel[5][2]) + "', # Возраст пациентов, лет, до")
            else:
                serv1D.append(f_servDel[5][2])
                if type(f_servDel[5][2]) == int or type(f_servDel[5][2]) == float:
                    doc.add_paragraph(str(f_servDel[5][2]) + ", # 16 Возраст пациентов, лет, до")
                if type(f_servDel[5][2]) == str:
                    doc.add_paragraph("'" + str(f_servDel[5][2]) + "', # 16 Возраст пациентов, лет, до")

            doc.add_paragraph("'', # 17 '' (Стаж по предлагаемой услуге, лет)")
            serv1D.append('')

            if f_servDel[6][0] != 'Стоимость':
                doc.add_paragraph("# оЗп_И: '" + str(f_servDel[6][0]) + "', # Стоимость")
                o_mM.add('mM')
            elif type(f_servDel[6][1]) != int and type(f_servDel[6][1]) != float:
                doc.add_paragraph("# оЗп_З: " + str(f_servDel[6][1]) + ", # Стоимость")
                o_mM.add('mM')
            else:
                doc.add_paragraph(str(f_servDel[6][1]) + ", # 18 Стоимость")
                serv1D.append(f_servDel[6][1])

            if f_servDel[7][0] != 'Возможность пациента оплатить услугу через сайт':
                doc.add_paragraph("# оЗп_И: '" + str(f_servDel[7][0]) + "', # Возможность пациента оплатить услугу через сайт")
                o_mM.add('mM')
            elif f_servDel[7][1] != 'не предоставлять пациенту возможность оплатить услугу через сайт' and f_servDel[7][1] != 'предоставить пациенту возможность оплатить услугу через сайт (взимается % от стоимости услуги за транзакционные издержки)' and f_servDel[7][1] != '':
                doc.add_paragraph("# оЗп_З: '" + str(f_servDel[7][1]) + "', # Возможность пациента оплатить услугу через сайт")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + f_servDel[7][1] + "', # 19 Возможность пациента оплатить услугу через сайт")
                serv1D.append(f_servDel[7][1])

            if f_servDel[8][0] != 'Ключевые поисковые слова предлагаемой услуги':
                doc.add_paragraph("# оЗп_И: '" + str(f_servDel[8][0]) + "', # Ключевые поисковые слова предлагаемой услуги")
                o_mM.add('mM')
            elif type(f_servDel[8][1]) != str:
                doc.add_paragraph("# оЗп_З: '" + str(f_servDel[8][1]) + "', # Ключевые поисковые слова предлагаемой услуги")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + str(f_servDel[8][1]) + "', # 20 Ключевые поисковые слова предлагаемой услуги")
                serv1D.append(f_servDel[8][1])

            o_mME = set()
            if f_servDel[9][0] != 'Email услуги':
                doc.add_paragraph("# оЗп_И: '" + str(f_servDel[9][0]) + "', # Email услуги")
                o_mME.add('mM')
            elif f_servDel[9][1] != '':
                if type(f_servDel[9][1]) != str or '@' not in f_servDel[9][1] or '.' not in f_servDel[9][1]:
                    doc.add_paragraph("# оЗп_З: '" + str(f_servDel[9][1]) + "', # Email услуги")
                    o_mME.add('mM')
            if o_mME == set():
                doc.add_paragraph("'" + f_servDel[9][1] + "', # 21 Email услуги")
                serv1D.append(f_servDel[9][1])
            o_mM |= o_mME

            o_mML = set()
            if f_servDel[10][0][0] != 'Укажите языки, на которых может быть предоставлена данная услуга':
                doc.add_paragraph("# оЗп_И: '" + str(f_servDel[10][0][0]) + "', # Укажите языки, на которых могут быть предоставлены услуги")
                o_mML.add('mM')
            elif type(f_servDel[10][0][1]) != list:
                doc.add_paragraph("# оЗп_З: '" + f_servDel[10][0][1] + "' не является списком, # Укажите языки, на которых могут быть предоставлены услуги")
            else:
                for q in f_servDel[10][0][1]:
                    if q not in oLang and q != '':
                        doc.add_paragraph("# оЗп_З: '" + str(q) + "', # Укажите языки, на которых могут быть предоставлены услуги, нет такого языка в списке")
                        o_mML.add('mM')    
            if type(f_servDel[10][1]) == list:
                if f_servDel[10][1][0] != 'добавить язык, отсутствующий в списке':
                    doc.add_paragraph("# оЗп_И: '" + str(f_servDel[10][1][0]) + "', # добавить язык, отсутствующий в списке")
                    o_mML.add('mM')
                elif type(f_servDel[10][1][1]) != str:
                    doc.add_paragraph("# оЗп_З: '" + str(f_servDel[10][1][1]) + "', # добавить язык, отсутствующий в списке")
                    o_mML.add('mM')
            if o_mML == set():
                f_alL = []
                for q in f_servDel[10][0][1]:
                    if q != '':
                        f_alL.append(q)
                if f_servDel[10][1][1] != '':
                    f_alL.append(f_servDel[10][1][1])
                str_alL = ''
                for q in f_alL:
                    if f_alL.index(q) != len(f_alL)-1:
                        str_alL = str_alL+q+', '
                    else:
                        str_alL = str_alL+q
                doc.add_paragraph("'" + str_alL + "', # 22 Адрес услуги / '' / Языки")       
                serv1D.append(str_alL)
            o_mM |= o_mML

            if f_servDel[11][0] != 'Время приема':
                doc.add_paragraph("# оЗп_И: '" + str(f_servDel[11][0]) + "', # Время приема")
                o_mM.add('mM')
            elif type(f_servDel[11][1]) != str:
                doc.add_paragraph("# оЗп_З: " + str(f_servDel[11][1]) + ", # Время приема")
                o_mM.add('mM')
            else:                   
                doc.add_paragraph("'" + str(f_servDel[11][1]) + "', # 23 Время приема")  
                serv1D.append(f_servDel[11][1])

            if f_servDel[12][0] != 'Возможность оказать данную услугу на дому':
                doc.add_paragraph("# оЗп_И: '" + str(f_servDel[12][0]) + "', # Возможность оказать данную услугу на дому")
                o_mM.add('mM')
            elif f_servDel[12][1] != '' and f_servDel[12][1] != 'возможен вызов на дом' and f_servDel[12][1] != 'вызов на дом не принимается':
                doc.add_paragraph("# оЗп_З: '" + str(f_servDel[12][1]) + "', # Возможность оказать данную услугу на дому")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + f_servDel[12][1] + "', # 24 Возможность оказать данную услугу на дому")
                serv1D.append(f_servDel[12][1])

            if f_servDel[13][0] != 'Дополнительная информация о предлагаемой услуге':
                doc.add_paragraph("# оЗп_И: '" + str(f_servDel[13][0]) + "', # Дополнительная информация о предлагаемой услуге")
                o_mM.add('mM')
            elif type(f_servDel[13][1]) != str:
                doc.add_paragraph("# оЗп_З: " + str(f_servDel[13][1]) + ", # Дополнительная информация о предлагаемой услуге")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + str(f_servDel[13][1]) + "', # 25 Дополнительная информация о предлагаемой услуге") 
                serv1D.append(f_servDel[13][1])

            if f_servDel[14] != 'servDel':
                doc.add_paragraph("# оИдент: '" + str(f_servDel[14]) + "', # servDel")
                o_mM.add('mM')
            else:
                doc.add_paragraph("'" + str(f_servDel[14]) + "', # 26 servDel") 
                serv1D.append(f_servDel[14])                

            doc.add_paragraph(")")
            doc.add_paragraph(")")

#---------------------
    if f_servDel[2][0] == 'Вид услуги':
        if 'консультация' in f_servDel[2][1]:
                doc.add_paragraph("print('    Специалист, " + str(f_servDel[2][1]) + ", " + str(f_servDel[3][2]) + ", " + str(f_servDel[7][1]) + "'), ")
        if 'процедура' in f_servDel[2][1]:
                doc.add_paragraph("print('    Специалист, " + "процедура: " + str(f_servDel[2][2]) + ", " + str(f_servDel[3][2]) + ", " + str(f_servDel[7][1]) + "'), ")

    if f_servDel[2][0] == 'Персонализированность услуги':
        if f_servDel[2][1] == 'персонализированная услуга (указываются данные специалиста)':
            if 'консультация' in f_servDel[8][1]:
                    doc.add_paragraph("print('    Клиника, персонализированная, " + str(f_servDel[3][1]) + " " + str(f_servDel[4][1]) + ", " + str(f_servDel[8][1]) + ", " + str(f_servDel[9][2]) + ", " + str(f_servDel[12][1]) + "'), ")
            if 'процедура' in f_servDel[8][1]:
                    doc.add_paragraph("print('    Клиника, персонализированная, " + "процедура: " + str(f_servDel[3][1]) + " " + str(f_servDel[4][1]) + ", " + str(f_servDel[8][2]) + ", " + str(f_servDel[12][1]) + ", " + "'), ")

        if f_servDel[2][1] == 'не персонализированная услуга (не указываются данные специалиста)':
            if 'консультация' in f_servDel[3][1]:
                    doc.add_paragraph("print('    Клиника, не персонализированная, " + str(f_servDel[3][1]) + ", " + str(f_servDel[4][2]) + ", " + str(f_servDel[6][1]) + "'), ")
            if 'процедура' in f_servDel[3][1]:
                    doc.add_paragraph("print('    Клиника, не персонализированная, " + "процедура: " + str(f_servDel[3][2]) + ", " + str(f_servDel[4][2]) + ", " + str(f_servDel[6][1]) + "'), ")
    doc.add_paragraph("print('---------------------------')\n")
    if o_mM != set():
        print(f"servDel_Ошибка_{cT}")
        doc.add_paragraph("print('" + NameFile + " отклонено оЗп_. ----------------------------------------------------------------')")
        doc.save(NameFile+cT+'_M.docx')
        return NameFile+cT+'_M.docx'
    else:
        doc.save(NameFile+'V.docx')
        print(f"servDel_Верно_{cT}\n")
        print(f"serv1D {serv1D}\n")

        return NameFile+'V.docx'
        
#-------------------------------------------------------------------------------------------------------------------------------------
    
# 20240903 1439

def isDigit(q):
    q = q.replace(',', '.')
    q = q.replace('^', '')
    q = q.replace('*', '')    
    q = q.replace('**', '')
    try:
        float(q)
        return True
    except ValueError:
        return False  

def send_oDSpClLK (f_oDSpClLK):

    import datetime
    cT = datetime.datetime.now().strftime('%Y.%m.%d.%H.%M.%S')

    import docx
    doc = docx.Document()
    # NameFile = '# Serv_'
    NameFile = '# ODSpCl_'

    p = doc.add_paragraph()
    runner = p.add_run(NameFile)
    runner.bold = True

    t8o8r = cT # Время регистрации
    oDSpCl1R = []

    o_mM = set()
    oLang = {'английский', 'арабский', 'испанский', 'итальянский', 'казахский', 'киргизский', 'китайский', 'корейский', 
            'монгольский', 'немецкий', 'персидский', 'португальский', 'русский', 'турецкий', 'французский', 'хинди', 'японский'}

    if len(f_oDSpClLK) != 10:
        o_mM.add('mM')
        doc.add_paragraph("print('f_oDSpClLKPerson имеет некорректную длинну: " + str(len(f_oDSpClLK)) + ", вместо 10')")
        if 10 - len(f_oDSpClLK) > 0:
            doc.add_paragraph("# f_oDSpClLKPerson - нехватка " + str(10 - len(f_oDSpClLK)) + " элементов!")
        if 10 - len(f_oDSpClLK) < 0:
            doc.add_paragraph("# f_oDSpClLKPerson - избыток " + str(len(f_oDSpClLK) - 10) + " элементов!")

    else:

        doc.add_paragraph("\noDSpCl1R.add(") 
        doc.add_paragraph("(")

        if f_oDSpClLK[0][0] != 'Id':
            doc.add_paragraph("# оЗп_И: '" + str(f_oDSpClLK[0][0]) + "', # Id")
            o_mM.add('mM')
        elif type(f_oDSpClLK[0][1]) != int:
            doc.add_paragraph("# оЗп_З: '" + str(f_oDSpClLK[0][1]) + "', # Id")
            o_mM.add('mM')
        else:
            doc.add_paragraph("" + str(f_oDSpClLK[0][1]) + ", # 0 Id")
            oDSpCl1R.append(f_oDSpClLK[0][1])

        if f_oDSpClLK[1][0] != 'IdProfile':
            doc.add_paragraph("# оЗп_И: '" + str(f_oDSpClLK[1][0]) + "', # IdProfile")
            o_mM.add('mM')
        elif type(f_oDSpClLK[1][1]) != int:
            doc.add_paragraph("# оЗп_З: '" + str(f_oDSpClLK[1][1]) + "', # IdProfile")
            o_mM.add('mM')
        else:
            doc.add_paragraph("" + str(f_oDSpClLK[1][1]) + ", # 1 IdProfile")
            oDSpCl1R.append(f_oDSpClLK[1][1])

        if f_oDSpClLK[2][0] != 'Фамилия':
            doc.add_paragraph("# оЗп_И: '" + str(f_oDSpClLK[2][0]) + "', # Фамилия")
            o_mM.add('mM')
        elif f_oDSpClLK[2][1] == '':
            doc.add_paragraph("# оЗп_З: '', # Фамилия")
            o_mM.add('mM')
        elif type(f_oDSpClLK[2][1]) != str:
            doc.add_paragraph("# оЗп_З: " + str(f_oDSpClLK[2][1]) + ", # Фамилия")
            o_mM.add('mM')
            doc.add_paragraph(str(f_oDSpClLK[2][1]) + ", # Фамилия")          
        else:
            doc.add_paragraph("'" + f_oDSpClLK[2][1] + "', # 3 Фамилия")
            oDSpCl1R.append(f_oDSpClLK[2][1])
        
        if f_oDSpClLK[3][0] != 'Имя':
            doc.add_paragraph("# оЗп_И: '" + str(f_oDSpClLK[3][0]) + "', # Имя")
            o_mM.add('mM')
        elif f_oDSpClLK[3][1] == '':
            doc.add_paragraph("# оЗп_З: '', # Имя")
            o_mM.add('mM')
        elif type(f_oDSpClLK[3][1]) != str:
            doc.add_paragraph("# оЗп_З: " + str(f_oDSpClLK[3][1]) + ", # Имя")
            o_mM.add('mM')
            doc.add_paragraph(str(f_oDSpClLK[3][1]) + ", # Имя")          
        else:
            doc.add_paragraph("'" + f_oDSpClLK[3][1] + "', # 4 Имя")
            oDSpCl1R.append(f_oDSpClLK[3][1]) 

        if f_oDSpClLK[4][0] != 'Отчество':
            doc.add_paragraph("# оЗп_И: '" + str(f_oDSpClLK[4][0]) + "', # Отчество")
            o_mM.add('mM')
        elif type(f_oDSpClLK[4][1]) != str:
            doc.add_paragraph("# оЗп_З: " + str(f_oDSpClLK[4][1]) + ", # Отчество")
            o_mM.add('mM')
            doc.add_paragraph(str(f_oDSpClLK[4][1]) + ", # Отчество")
        else:
            doc.add_paragraph("'" + f_oDSpClLK[4][1] + "', # 5 Отчество")       
            oDSpCl1R.append(f_oDSpClLK[4][1])

        if f_oDSpClLK[5][0] != 'Пол':
            doc.add_paragraph("# оЗп_И: '" + str(f_oDSpClLK[5][0]) + "', # Пол")
            o_mM.add('mM')
        elif f_oDSpClLK[5][1] != 'мужской' and f_oDSpClLK[5][1] != 'женский':
            doc.add_paragraph("# оЗп_З: '" + str(f_oDSpClLK[5][1]) + "', # Пол")
            o_mM.add('mM')
        else:
            doc.add_paragraph("'" + f_oDSpClLK[5][1] + "', # 6 Пол")       
            oDSpCl1R.append(f_oDSpClLK[5][1])

        if f_oDSpClLK[6][0] != 'Дата рождения':
            doc.add_paragraph("# оЗп_И: '" + str(f_oDSpClLK[6][0]) + "', # Дата рождения")
            o_mM.add('mM')
        if type(f_oDSpClLK[6][1]) == str:
            if isDigit(f_oDSpClLK[6][1].replace('.', '')) != True or len(f_oDSpClLK[6][1]) != 10 or len(f_oDSpClLK[6][1].replace('.', '')) != 8:
                doc.add_paragraph("# оЗп_З: '" + str(f_oDSpClLK[6][1]) + "', # Дата рождения")
                o_mM.add('mM')
        doc.add_paragraph("'" + str(f_oDSpClLK[6][1]) + "', # 7 '' / Дата рождения")
        oDSpCl1R.append(f_oDSpClLK[6][1])

        o_mML = set()
        if f_oDSpClLK[7][0][0] != 'Укажите языки, на которых могут быть предоставлены услуги специалиста':
            doc.add_paragraph("# оЗп_И: '" + str(f_oDSpClLK[7][0][0]) + "', # Укажите языки, на которых могут быть предоставлены услуги специалиста")
            o_mML.add('mM')
        elif type(f_oDSpClLK[7][0][1]) != list:
            doc.add_paragraph("# оЗп_З: '" + f_oDSpClLK[7][0][1] + "' не является списком, # Укажите языки, на которых могут быть предоставлены услуги специалиста")
        else:
            for q in f_oDSpClLK[7][0][1]:
                if q not in oLang and q != '':
                    doc.add_paragraph("# оЗп_З: '" + str(q) + "', # Укажите языки, на которых могут быть предоставлены услуги специалиста, нет такого языка в списке")
                    o_mML.add('mM')    
        if type(f_oDSpClLK[7][1]) == list:
            if f_oDSpClLK[7][1][0] != 'добавить язык, отсутствующий в списке':
                doc.add_paragraph("# оЗп_И: '" + str(f_oDSpClLK[7][1][0]) + "', # добавить язык, отсутствующий в списке")
                o_mML.add('mM')
            elif type(f_oDSpClLK[7][1][1]) != str:
                doc.add_paragraph("# оЗп_З: '" + str(f_oDSpClLK[7][1][1]) + "', # добавить язык, отсутствующий в списке")
                o_mML.add('mM')
        if o_mML == set():
            f_alL = []
            for q in f_oDSpClLK[7][0][1]:
                if q != '':
                    f_alL.append(q)
            if f_oDSpClLK[7][1][1] != '':
                f_alL.append(f_oDSpClLK[7][1][1])
            str_alL = ''
            for q in f_alL:
                if f_alL.index(q) != len(f_alL)-1:
                    str_alL = str_alL+q+', '
                else:
                    str_alL = str_alL+q
            doc.add_paragraph("'" + str_alL + "', # 8 Языки")       
            oDSpCl1R.append(str_alL)
        o_mM |= o_mML

        for q in f_oDSpClLK[8]:
            if q[0][0] != 'Начало':
                doc.add_paragraph("# оЗп_И: '" + str(q[0][0]) + "', # Обучение: Начало")
                o_mM.add('mM')
            if type(q[0][1]) != str:
                doc.add_paragraph("# оЗп_З: '" + str(q[0][1]) + "', # Обучение: Начало")
                o_mM.add('mM')
            if q[0][1] != '':
                if isDigit(q[0][1].replace('.', '')) != True or len(q[0][1]) != 10 or len(q[0][1].replace('.', '')) != 8:
                    doc.add_paragraph("# оЗп_З: '" + str(q[0][1]) + "', # Обучение: Начало")
                    o_mM.add('mM')
            doc.add_paragraph("'" + str(q[0][1]) + "', # 6 Обучение: Начало")

            if q[1][0] != 'Окончание':
                doc.add_paragraph("# оЗп_И: '" + str(q[1][0]) + "', # Обучение: Окончание")
                o_mM.add('mM')
            if type(q[1][1]) != str:
                doc.add_paragraph("# оЗп_З: '" + str(q[1][1]) + "', # Обучение: Окончание")
                o_mM.add('mM')
            if q[1][1] != '':
                if isDigit(q[1][1].replace('.', '')) != True or len(q[1][1]) != 10 or len(q[1][1].replace('.', '')) != 8:
                    doc.add_paragraph("# оЗп_З: '" + str(q[1][1]) + "', # Обучение: Окончание")
                    o_mM.add('mM')
            doc.add_paragraph("'" + str(q[1][1]) + "', # 7 Обучение: Окончание")

            if q[2][0] != 'Квалификация, специальность, тема':
                doc.add_paragraph("# оЗп_И: '" + str(q[2][0]) + "', # Обучение: Тематика")
                o_mM.add('mM')
            if type(q[2][1]) != str:
                doc.add_paragraph("# оЗп_З: '" + str(q[2][1]) + "', # Обучение: Тематика")
                o_mM.add('mM')
            doc.add_paragraph("'" + str(q[2][1]) + "', # 8 Обучение: Тематика")

            if q[3][0] != 'Учебное заведение':
                doc.add_paragraph("# оЗп_И: '" + str(q[3][0]) + "', # Обучение: Учебное заведение")
                o_mM.add('mM')
            if type(q[3][1]) != str:
                doc.add_paragraph("# оЗп_З: '" + str(q[3][1]) + "', # Обучение: Учебное заведение")
                o_mM.add('mM')
            doc.add_paragraph("'" + str(q[3][1]) + "', # 9 Обучение: Учебное заведение")


#---------------------

    doc.add_paragraph("print('    Языки специалиста клиники, : " + str(f_oDSpClLK[3][1]) + " " + str(f_oDSpClLK[4][1]) + ", " + str(f_oDSpClLK[7][1]) + "'), ")
    doc.add_paragraph("print('---------------------------')\n")
    if o_mM != set():
        print(f"oDSpClLK_Ошибка_{cT}")
        doc.add_paragraph("print('" + NameFile + " отклонено оЗп_. ----------------------------------------------------------------')")
        doc.save(NameFile+cT+'_M.docx')
        return NameFile+cT+'_M.docx'
    else:
        doc.save(NameFile+'V.docx')
        print(f"oDSpClLK_Верно_{cT}\n")

        print(f"oDSpCl1R {oDSpCl1R}\n")
        print(f"'{oDSpCl1R[0]}', # 0 Id")       
        print(f"'{oDSpCl1R[1]}', # 1 IdProfile") 
        print(f"'{oDSpCl1R[2]}', # 2 Фамилия") 
        print(f"'{oDSpCl1R[3]}', # 3 Имя") 
        print(f"'{oDSpCl1R[4]}', # 4 Отчество") 
        print(f"'{oDSpCl1R[5]}', # 5 Пол")  
        print(f"'{oDSpCl1R[6]}', # 6 Дата рождения") 
        print(f"'{oDSpCl1R[7]}', # 7 Языки") 

        return NameFile+'V.docx'


# 2024-9-3 13:46:5

zhR = {
('жалоба/осмотр: анальный зуд', 'анальный зуд'), 
('жалоба/осмотр: аномалия развития лицевого черепа', 'аномалия развития лицевого черепа'), 
('жалоба/осмотр: атаксия', 'не скоординированые движения'), 
('жалоба/осмотр: атаксия', 'нарушение согласованности движений различных мышц при условии отсутствия мышечной слабости'), 
('жалоба/осмотр: атаксия', 'неуклюжие движения'), 
('жалоба/осмотр: атаксия', 'атаксия'), 
('жалоба/осмотр: атаксия', 'нарушение согласованности движений и моторики'), 
('жалоба/осмотр: атаксия', 'неуклюжие не скоординированые движения, часто нарушается равновесие'), 
('жалоба/осмотр: атрофические изменения слизистой оболочки желудочно-кишечного тракта', 'атрофические изменения слизистой оболочки желудочно-кишечного тракта'), 
('жалоба/осмотр: атрофия мышц', 'атрофия мышц'), 
('жалоба/осмотр: атрофия сосочков языка', 'атрофия сосочков языка'), 
('жалоба/осмотр: атрофия сосочков языка', 'атрофический глоссит'), 
('жалоба/осмотр: атрофия сосочков языка', 'полированный язык'), 
('жалоба/осмотр: атрофия сосочков языка', 'хантеровский глоссит'), 
('жалоба/осмотр: атрофия сосочков языка', 'atrophic glossitis'), 
('жалоба/осмотр: атрофия сосочков языка', 'moeller glossitis'), 
('жалоба/осмотр: атрофия сосочков языка', 'лакированный язык'), 
('жалоба/осмотр: атрофия сосочков языка', 'moeller-hunter glossitis'), 
('жалоба/осмотр: атрофия сосочков языка', 'атрофичный язык'), 
('жалоба/осмотр: атрофия сосочков языка', 'сглаженность сосочков языка'), 
('жалоба/осмотр: атрофия сосочков языка', 'гунтеровский глоссит'), 
('жалоба/осмотр: атрофия сосочков языка', 'hunter glossitis'), 
('жалоба/осмотр: атрофия сосочков языка', 'atrophy of tongue papillae'), 
('жалоба/осмотр: ахоличный кал', 'белый кал'), 
('жалоба/осмотр: ахоличный кал', 'обесцвеченный кал'), 
('жалоба/осмотр: ахоличный кал', 'обесцвечивание кала'), 
('жалоба/осмотр: ахоличный кал', 'ахоличный кал'), 
('жалоба/осмотр: белый густой налет на языке', 'белый густой налет на языке'), 
('жалоба/осмотр: белый густой налет на языке', 'язык обложен густым белым налетом'), 
('жалоба/осмотр: беспокойное поведение', 'беспокойное поведение'), 
('жалоба/осмотр: беспокойное поведение', 'ажитация'), 
('жалоба/осмотр: бессонница', 'бессонница'), 
('жалоба/осмотр: бледность', 'бледность'), 
('жалоба/осмотр: бледность', 'бледный оттенок кожных покровов и слизистых'), 
('жалоба/осмотр: бледность кожи', 'бледность кожи'), 
('жалоба/осмотр: бледность кожи', 'бледность кожных покровов'), 
('жалоба/осмотр: бледность слизистых оболочек', 'бледность слизистых оболочек'), 
('жалоба/осмотр: вогнутость на поверхности ногтя', 'койлонихии'), 
('жалоба/осмотр: вогнутость на поверхности ногтя', 'койлонихия'), 
('жалоба/осмотр: вогнутость на поверхности ногтя', 'ложкообразная вогнутость ногтя'), 
('жалоба/осмотр: вогнутость на поверхности ногтя', 'блюдцеобразная вогнутость ногтя'), 
('жалоба/осмотр: вогнутость на поверхности ногтя', 'вдавление на поверхности ногтя'), 
('жалоба/осмотр: вогнутость на поверхности ногтя', 'вогнутость на поверхности ногтя'), 
('жалоба/осмотр: выпадение волос', 'выпадение волос'), 
('жалоба/осмотр: высыпания на коже', 'высыпания на коже'), 
('жалоба/осмотр: вязкий кал', 'вязкий кал'), 
('жалоба/осмотр: грязно-серый оттенок кожных покровов', 'грязно-серый оттенок кожных покровов'), 
('жалоба/осмотр: деменция', 'старческое слабоумие'), 
('жалоба/осмотр: деменция', 'старческий маразм'), 
('жалоба/осмотр: деменция', 'деградация памяти, мышления, поведения и способности выполнять ежедневные действия'), 
('жалоба/осмотр: деменция', 'деградация памяти, мышления, поведения'), 
('жалоба/осмотр: деменция', 'деменция'), 
('жалоба/осмотр: деменция', 'старческий склероз'), 
('жалоба/осмотр: деменция', 'приобретенное слабоумие'), 
('жалоба/осмотр: депрессия', 'депрессия'), 
('жалоба/осмотр: деформация ногтевых пластин', 'деформация ногтевых пластин'), 
('жалоба/осмотр: диарея', 'поносы'), 
('жалоба/осмотр: диарея', 'диарея'), 
('жалоба/осмотр: диарея', 'жидкий стул'), 
('жалоба/осмотр: диарея', 'диареи'), 
('жалоба/осмотр: диарея', 'понос'), 
('жалоба/осмотр: динамическая атаксия', 'динамическая атаксия'), 
('жалоба/осмотр: динамическая атаксия', 'нарушение координации при движении'), 
('жалоба/осмотр: желтушность', 'желтоватость кожи'), 
('жалоба/осмотр: желтушность', 'желтушность слизистых оболочек'), 
('жалоба/осмотр: желтушность', 'небольшая желтушность'), 
('жалоба/осмотр: желтушность', 'желтуха'), 
('жалоба/осмотр: желтушность', 'желтоватость слизистых оболочек'), 
('жалоба/осмотр: желтушность', 'желтушность кожи'), 
('жалоба/осмотр: желтушность', 'субиктеричность кожи'), 
('жалоба/осмотр: желтушность', 'желтушность'), 
('жалоба/осмотр: желтушность', 'иктеричность склер'), 
('жалоба/осмотр: желтушность', 'иктеричность'), 
('жалоба/осмотр: желтушность', 'желтушная окраска'), 
('жалоба/осмотр: желтушность', 'желтоватость'), 
('жалоба/осмотр: желтушность', 'субиктеричность'), 
('жалоба/осмотр: желтушность', 'желтоватость склер'), 
('жалоба/осмотр: желтушность', 'желтушность склер'), 
('жалоба/осмотр: желтые испражнения с незначительной примесью слизи', 'желтые испражнения с незначительной примесью слизи'), 
('жалоба/осмотр: желтые испражнения с незначительной примесью слизи', 'желтый кал с незначительной примесью слизи'), 
('жалоба/осмотр: желтые испражнения с незначительной примесью слизи', 'желтый стул с незначительной примесью слизи'), 
('жалоба/осмотр: задержка мочи', 'задержка мочи'), 
('жалоба/осмотр: изменение голоса', 'изменения голоса'), 
('жалоба/осмотр: изменение голоса', 'изменение голоса'), 
('жалоба/осмотр: изменение запаха кала', 'изменение запаха кала'), 
('жалоба/осмотр: изменение консистенции кала', 'изменение консистенции кала'), 
('жалоба/осмотр: изменение цвета кала', 'изменение цвета кала'), 
('жалоба/осмотр: икота', 'икота'), 
('жалоба/осмотр: истончение ногтей', 'истончение ногтей'), 
('жалоба/осмотр: истощение', 'истощение'), 
('жалоба/осмотр: кариес', 'развитие кариеса'), 
('жалоба/осмотр: кариес', 'кариес'), 
('жалоба/осмотр: кашель', 'кашель'), 
('жалоба/осмотр: кашель с кровью', 'кашель с кровью'), 
('жалоба/осмотр: кашицеобразный кал', 'кашицеобразный кал'), 
('жалоба/осмотр: кончики ногтей изогнуты вниз больше обычного', 'кончики ногтей изогнуты вниз больше обычного'), 
('жалоба/осмотр: кончики ногтей изогнуты вниз больше обычного', 'кончики ногтей начинают изгибаться вниз больше обычного'), 
('жалоба/осмотр: крапивница', 'хроническая крапивница'), 
('жалоба/осмотр: крапивница', 'крапивница'), 
('жалоба/осмотр: крапивница', 'крапивная лихорадка'), 
('жалоба/осмотр: крапивница', 'уртикария'), 
('жалоба/осмотр: крапивница', 'крапивная сыпь'), 
('жалоба/осмотр: крапивница', 'urticaria'), 
('жалоба/осмотр: куриная слепота', 'ухудшение ночного зрения'), 
('жалоба/осмотр: куриная слепота', 'куриная (ночная) слепота'), 
('жалоба/осмотр: куриная слепота', 'ухудшение сумеречного зрения'), 
('жалоба/осмотр: куриная слепота', 'гемералопия'), 
('жалоба/осмотр: куриная слепота', 'куриная слепота'), 
('жалоба/осмотр: куриная слепота', 'ночная слепота'), 
('жалоба/осмотр: ломкость волос', 'ломкость волос'), 
('жалоба/осмотр: ломкость ногтей', 'ломкость ногтей'), 
('жалоба/осмотр: метеоризм', 'метеоризм'), 
('жалоба/осмотр: метеоризм', 'избыточное газообразование'), 
('жалоба/осмотр: метеоризм', 'вздутие живота'), 
('жалоба/осмотр: метеоризм', 'вздутие'), 
('жалоба/осмотр: метеоризм', 'вспучивание живота'), 
('жалоба/осмотр: мышечные спазмы', 'мышечный спазм'), 
('жалоба/осмотр: мышечные спазмы', 'мышечные спазмы'), 
('жалоба/осмотр: налет на языке', 'налет на языке'), 
('жалоба/осмотр: налет на языке', 'обложенность налетом языка'), 
('жалоба/осмотр: наличие в кале непереваренной пищи', 'наличие в кале непереваренной пищи'), 
('жалоба/осмотр: нарушение глотания', 'нарушение глотания'), 
('жалоба/осмотр: нарушение глотания', 'застревание пищи в горле'), 
('жалоба/осмотр: нарушение глотания', 'застревание пищи в глотке'), 
('жалоба/осмотр: нарушение глотания', 'дисфагия'), 
('жалоба/осмотр: нарушение глотания', 'нарушение акта глотания'), 
('жалоба/осмотр: нарушение зрения', 'нарушение зрения'), 
('жалоба/осмотр: нарушение зрения', 'нарушения зрения'), 
('жалоба/осмотр: нарушение координации движений', 'нарушение координации движений'), 
('жалоба/осмотр: нарушение осанки', 'нарушение осанки'), 
('жалоба/осмотр: нарушение сна', 'плохой сон'), 
('жалоба/осмотр: нарушение сна', 'нарушение сна'), 
('жалоба/осмотр: нарушения слуха', 'нарушения слуха'), 
('жалоба/осмотр: невозможность удержать мочу при смехе, кашле, чихании', 'невозможность удержать мочу при смехе, кашле, чихании'), 
('жалоба/осмотр: недержание мочи', 'недержание мочи'), 
('жалоба/осмотр: недержание мочи', 'инконтиненция'), 
('жалоба/осмотр: недержание мочи', 'слабость сфинктеров мочевого пузыря'), 
('жалоба/осмотр: недержание мочи', 'неудержание мочи'), 
('жалоба/осмотр: неустойчивый стул', 'неустойчивый стул'), 
('жалоба/осмотр: нистагм', 'дрожание глаз'), 
('жалоба/осмотр: нистагм', 'колебание глаз'), 
('жалоба/осмотр: нистагм', 'нистагм'), 
('жалоба/осмотр: нистагм', 'тремор глаз'), 
('жалоба/осмотр: нистагм', 'бегание глаз'), 
('жалоба/осмотр: обморок', 'синкопе'), 
('жалоба/осмотр: обморок', 'обморок'), 
('жалоба/осмотр: образование на коже трещин', 'образование на коже трещин'), 
('жалоба/осмотр: обратимые нарушения зрения', 'обратимые нарушения зрения'), 
('жалоба/осмотр: обширные отеки', 'обширные отеки'), 
('жалоба/осмотр: обширные отеки', 'генерализованные отеки'), 
('жалоба/осмотр: обширные травмы', 'обширные травмы'), 
('жалоба/осмотр: одышка', 'одышка'), 
('жалоба/осмотр: ожог', 'ожог'), 
('жалоба/осмотр: ожог', 'ожоги'), 
('жалоба/осмотр: опущенное веко', 'опущение века'), 
('жалоба/осмотр: опущенное веко', 'опущенное веко'), 
('жалоба/осмотр: осиплость голоса', 'осиплость голоса'), 
('жалоба/осмотр: отрыжка', 'отрыжка воздухом или пищей'), 
('жалоба/осмотр: отрыжка', 'отрыжка'), 
('жалоба/осмотр: пенистость мочи', 'пенистость мочи'), 
('жалоба/осмотр: поза зародыша', 'поза зародыша'), 
('жалоба/осмотр: поза зародыша', 'вынужденное положение поза зародыша'), 
('жалоба/осмотр: поза зародыша', 'вынужденное положение тела поза зародыша'), 
('жалоба/осмотр: покраснение кончика языка', 'покраснение кончика языка'), 
('жалоба/осмотр: покраснение лица', 'покраснение лица'), 
('жалоба/осмотр: покраснение языка', 'эритема языка'), 
('жалоба/осмотр: покраснение языка', 'покраснение языка'), 
('жалоба/осмотр: покраснение языка', 'гиперемия языка'), 
('жалоба/осмотр: полиурия', 'полиурия'), 
('жалоба/осмотр: полиурия', 'выделение мочи в объеме более 3 л в день'), 
('жалоба/осмотр: полифекалия', 'полифекалия'), 
('жалоба/осмотр: поперечная исчерченность ногтей', 'поперечная исчерченность ногтей'), 
('жалоба/осмотр: постоянное покашливание', 'постоянное покашливание'), 
('жалоба/осмотр: потемнение мочи', 'потемнение мочи'), 
('жалоба/осмотр: потеря веса', 'снижение веса'), 
('жалоба/осмотр: потеря веса', 'похудение'), 
('жалоба/осмотр: потеря веса', 'похудание'), 
('жалоба/осмотр: потеря веса', 'снижение массы тела'), 
('жалоба/осмотр: потеря веса', 'потеря веса'), 
('жалоба/осмотр: потливость', 'потливость'), 
('жалоба/осмотр: потливость', 'повышенная потливость'), 
('жалоба/осмотр: потливость', 'повышенное потоотделение'), 
('жалоба/осмотр: приступы кашля и/или удушья преимущественно в ночное время, после обильной еды', 'приступы кашля и/или удушья преимущественно в ночное время, после обильной еды'), 
('жалоба/осмотр: психоз', 'психоз'), 
('жалоба/осмотр: раннее поседение волос', 'раннее поседение волос'), 
('жалоба/осмотр: расчесы на коже', 'расчесы на коже'), 
('жалоба/осмотр: рвота', 'рвота'), 
('жалоба/осмотр: рвота с кровью', 'рвота с кровью'), 
('жалоба/осмотр: рвота, не приносящая облегчения, вначале пищей, затем содержимым с примесью желчи, сопровождается тошнотой', 'рвота, не приносящая облегчения, вначале пищей, затем содержимым с примесью желчи, сопровождается тошнотой'), 
('жалоба/осмотр: редкий сухой кашель', 'редкий сухой кашель'), 
('жалоба/осмотр: редкий сухой кашель', 'редкий сухой кашель или с небольшим количеством трудноотделяемой мокроты, может быть мучительным, приступообразным'), 
('жалоба/осмотр: серый цвет кала', 'серый цвет кала'), 
('жалоба/осмотр: симптом "красных капелек"', 'четко отграниченные ярко-красные пятна на коже живота, груди, иногда в области спины'), 
('жалоба/осмотр: симптом "красных капелек"', 'симптом "красных капелек"'), 
('жалоба/осмотр: симптом "синих склер"', 'синеватость или голубизна склер'), 
('жалоба/осмотр: симптом "синих склер"', 'синеватость склер'), 
('жалоба/осмотр: симптом "синих склер"', 'симптом "синих склер"'), 
('жалоба/осмотр: симптом "синих склер"', 'голубизна склер'), 
('жалоба/осмотр: симптом Айзенберга I', 'удар ребром ладони ниже угла правой лопатки вызывает умеренную и "пронизывающую" локальную боль в области желчного пузыря'), 
('жалоба/осмотр: симптом Айзенберга I', 'симптом Айзенберга I'), 
('жалоба/осмотр: симптом Айзенберга I', 'удар ребром ладони ниже угла правой лопатки вызывает умеренную локальную боль и "пронизывающую" - в области желчного пузыря'), 
('жалоба/осмотр: симптом Алиева', 'появление наряду с локальной болезненностью при пальпации в точках Боаса или Маккензи иррадирующей боли по направлению к желчному пузырю'), 
('жалоба/осмотр: симптом Алиева', 'симптом Алиева'), 
('жалоба/осмотр: симптом Гротта', 'симптом Гротта'), 
('жалоба/осмотр: симптом Гротта', 'атрофия подкожно-жировой клетчатки в зоне, соответствующей проекции поджелудочной железы на переднюю брюшную стенку'), 
('жалоба/осмотр: симптом Кера', 'боль в правом подреберье во время глубокого вдоха'), 
('жалоба/осмотр: симптом Кера', 'симптом Кера'), 
('жалоба/осмотр: симптом Мерфи', 'симптом Мерфи'), 
('жалоба/осмотр: симптом Мерфи', 'боль при пальпации в правом подреберье на вдохе'), 
('жалоба/осмотр: симптом Щеткина-Блюмберга', 'симптом Щеткина-Блюмберга'), 
('жалоба/осмотр: симптом Щеткина-Блюмберга', 'резкое усилениеболи в животепри быстром снятиипальпирующейруки с передней брюшной стенки после надавливания'), 
('жалоба/осмотр: симптомы поражения черепных нервов', 'поражение черепных нервов'), 
('жалоба/осмотр: симптомы поражения черепных нервов', 'симптомы поражения черепных нервов'), 
('жалоба/осмотр: снижение объема рабочей памяти', 'снижение объема рабочей памяти'), 
('жалоба/осмотр: снижение скорости мышления', 'снижение скорости мышления'), 
('жалоба/осмотр: снижение скорости мышления', 'уменьшение умственной функции'), 
('жалоба/осмотр: снижение скорости реакции на неожиданные ситуации', 'снижение скорости реакции на неожиданные ситуации'), 
('жалоба/осмотр: снижение способности к концентрации внимания', 'снижение способности к концентрации внимания'), 
('жалоба/осмотр: снижение способности контролировать импульсивное поведение', 'снижение способности контролировать импульсивное поведение'), 
('жалоба/осмотр: снижение тургора кожи', 'снижение тургора кожи'), 
('жалоба/осмотр: статическая атаксия', 'нарушение равновесия при стоянии'), 
('жалоба/осмотр: статическая атаксия', 'статическая атаксия'), 
('жалоба/осмотр: стеаторея', 'наличие избыточного жира в кале'), 
('жалоба/осмотр: стеаторея', 'стеаторея'), 
('жалоба/осмотр: стеаторея', 'сальный вид кала'), 
('жалоба/осмотр: судороги', 'судороги'), 
('жалоба/осмотр: сухость глаз', 'сухость глаз'), 
('жалоба/осмотр: сухость глаз', 'сухость роговицы и конъюнктивы глаза'), 
('жалоба/осмотр: сухость глаз', 'синдром сухого глаза'), 
('жалоба/осмотр: сухость глаз', 'ксерофтальмия'), 
('жалоба/осмотр: сухость кожи', 'ксеродермия'), 
('жалоба/осмотр: сухость кожи', 'ксероз'), 
('жалоба/осмотр: сухость кожи', 'снижение влажности кожи'), 
('жалоба/осмотр: сухость кожи', 'сухая кожа'), 
('жалоба/осмотр: сухость кожи', 'ксероз кожи'), 
('жалоба/осмотр: сухость кожи', 'сухость кожи'), 
('жалоба/осмотр: сухость языка', 'сухость языка'), 
('жалоба/осмотр: сыпь', 'кожная сыпь'), 
('жалоба/осмотр: сыпь', 'сыпь'), 
('жалоба/осмотр: тошнота', 'тошнота'), 
('жалоба/осмотр: травма уха', 'травма уха'), 
('жалоба/осмотр: травма уха', 'травмы уха'), 
('жалоба/осмотр: трещины в углу рта', 'трещины, заеды в углах рта'), 
('жалоба/осмотр: трещины в углу рта', 'ангулярный стоматит'), 
('жалоба/осмотр: трещины в углу рта', 'трещины в уголках рта'), 
('жалоба/осмотр: трещины в углу рта', 'трещины в углах рта'), 
('жалоба/осмотр: трещины в углу рта', 'трещины в углу рта'), 
('жалоба/осмотр: трещины в углу рта', 'заеды в углах рта'), 
('жалоба/осмотр: тусклость волос', 'тусклость волос'), 
('жалоба/осмотр: тусклость ногтей', 'тусклость ногтей'), 
('жалоба/осмотр: увеличение лимфатических узлов', 'увеличение лимфатических узлов'), 
('жалоба/осмотр: увеличение миндалин', 'увеличение миндалин'), 
('жалоба/осмотр: увеличение миндалин', 'отек миндалин'), 
('жалоба/осмотр: увеличение миндалин', 'отечность миндалин'), 
('жалоба/осмотр: урчание живота', 'урчание живота'), 
('жалоба/осмотр: урчание живота', 'урчание'), 
('жалоба/осмотр: участки пигментации на лице и конечностях', 'участки пигментации на лице и конечностях'), 
('жалоба/осмотр: частое мочеиспускание', 'частое мочеиспускание'), 
('жалоба/осмотр: черный стул', 'черный стул'), 
('жалоба/осмотр: черный стул', 'мелена'), 
('жалоба/осмотр: чихание', 'чихание'), 
('жалоба/осмотр: шелушение кожи', 'шелушение кожи'), 
('жалоба/осмотр: экхимозы', 'синяки'), 
('жалоба/осмотр: экхимозы', 'экхимозы'), 
('жалоба/осмотр: экхимозы', 'кровоизлияния в кожу или слизистую оболочку'), 
('жалоба/осмотр: экхимозы', 'синяк'), 
('жалоба/осмотр: энурез', 'энурез'), 
('жалоба/осмотр: энурез', 'ночное недержание мочи'), 
('жалоба/осмотр: эрозия эмали зубов', 'эрозия эмали зубов'), 
('жалоба/пальпация: боль при пальпации живота в эпигастральной области', 'боль при пальпации живота в верхней части'), 
('жалоба/пальпация: боль при пальпации живота в эпигастральной области', 'боль при пальпации живота в эпигастральной области'), 
('жалоба/пальпация: боль при пальпации живота в эпигастральной области', 'боль при пальпации живота в верхней области'), 
('жалоба/пальпация: боль при пальпации живота в эпигастральной области', 'болезненность при пальпации живота в эпигастральной области'), 
('жалоба/пальпация: боль при пальпации живота в эпигастральной области', 'умеренная болезненность в эпигастральной области при пальпации'), 
('жалоба/пальпация: боль при пальпации живота в эпигастральной области', 'при пальпации живота умеренная болезненность в эпигастральной области'), 
('жалоба/пальпация: боль при пальпация живота с иррадиацией в спину, позвоночник, левое подреберье, надплечье', 'боль при пальпация живота с иррадиацией в спину, позвоночник, левое подреберье, надплечье'), 
('жалоба/пальпация: пальпаторно болезненность в правом подреберье', 'боль в правом подреберье при пальпации'), 
('жалоба/пальпация: пальпаторно болезненность в правом подреберье', 'пальпаторно болезненность в правом подреберье'), 
('жалоба/пальпация: снижение тургора кожи', 'снижение тургора кожи'), 
('жалоба/пальпация: умеренная болезненность в пилоро-дуоденальной области при пальпации', 'умеренная болезненность в пилоро-дуоденальной области при пальпации'), 
('жалоба: болезненность в левом реберно-позвоночном углу', 'болезненность в зоне Мэйо-Робсона'), 
('жалоба: болезненность в левом реберно-позвоночном углу', 'болезненность в левом реберно-позвоночном углу (зона Мэйо-Робсона)'), 
('жалоба: болезненность в левом реберно-позвоночном углу', 'болезненность в левом реберно-позвоночном углу'), 
('жалоба: болезненность языка', 'болезненность языка'), 
('жалоба: боль', 'боль'), 
('жалоба: боль', 'dolor'), 
('жалоба: боль', 'алгия'), 
('жалоба: боль в горле', 'боль в горле'), 
('жалоба: боль в горле при глотании', 'боль в горле при глотании'), 
('жалоба: боль в груди', 'боль в груди'), 
('жалоба: боль в животе', 'абдоминальная боль'), 
('жалоба: боль в животе', 'боль в животе'), 
('жалоба: боль в костях', 'боль в костях'), 
('жалоба: боль в мышцах', 'боль в мышцах'), 
('жалоба: боль в области пупка', 'боль в области пупка'), 
('жалоба: боль в плече', 'боль в плече'), 
('жалоба: боль в правом подреберье', 'боль и дискомфорт в правом подреберье'), 
('жалоба: боль в правом подреберье', 'дискомфорт в правом подреберье'), 
('жалоба: боль в правом подреберье', 'боль в правом подреберье'), 
('жалоба: боль в руке или плече', 'боль в руке или плече'), 
('жалоба: боль в суставах', 'боль в суставах'), 
('жалоба: боль в ухе', 'боль в ухе'), 
('жалоба: боль в эпигастральной области с иррадиацией в спину', 'боли в эпигастральной области с иррадиацией в спину'), 
('жалоба: боль в эпигастральной области с иррадиацией в спину', 'боли в верхней части живота отдающие в спину'), 
('жалоба: боль в эпигастральной области с иррадиацией в спину', 'боль в эпигастральной области с иррадиацией в спину'), 
('жалоба: боль в эпигастральной области с иррадиацией в спину', 'боль в верхней части живота отдающая в спину'), 
('жалоба: боль в эпигастрии', 'умеренная болезненность в эпигастральной области живота'), 
('жалоба: боль в эпигастрии', 'боли в эпигастрии'), 
('жалоба: боль в эпигастрии', 'боль в эпигастрии'), 
('жалоба: боль в эпигастрии', 'боль в верхнем отделе живота'), 
('жалоба: боль в эпигастрии', 'умеренная болезненность в эпигастральной области'), 
('жалоба: боль в эпигастрии', 'боль в верхней части живота'), 
('жалоба: боль в эпигастрии', 'дискомфорт в верхнем отделе живота'), 
('жалоба: боль в эпигастрии', 'боль в эпигастральной области'), 
('жалоба: боль и жжение в языке', 'боль и жжение в языке'), 
('жалоба: боль и жжение во рту', 'боль и жжение во рту'), 
('жалоба: боль и распирание в области языка', 'боль и распирание в области языка'), 
('жалоба: боль при глотании', 'боль при глотании'), 
('жалоба: боль при глотании', 'боль при глотании пищи'), 
('жалоба: боль при глотании', 'одинофагия'), 
('жалоба: головная боль', 'головная боль'), 
('жалоба: головокружение', 'головокружение'), 
('жалоба: горечь во рту', 'ощущение горечи во рту'), 
('жалоба: горечь во рту', 'горечь во рту'), 
('жалоба: жажда', 'постоянная жажда'), 
('жалоба: жажда', 'жажда'), 
('жалоба: желание употреблять в пищу что-то необычное и малосъедобное (мел, зубной порошок, уголь, глину, песок, лед, сырое тесто, фарш, крупу, …)', 'желание употреблять в пищу что-то необычное и малосъедобное (мел, зубной порошок, уголь, глину, песок, лед, сырое тесто, фарш, крупу, …)'), 
('жалоба: желчная колика', 'острая, интенсивная, нетерпимая, жгучая, распирающая, сжимающая, схваткообразная боль в эпигастрии и правом подреберье, заставляющая больного метаться в кровати, продолжительностью от 15 мин до 5 ч, достигая максимума в течение 20-30 мин'), 
('жалоба: желчная колика', 'желчная колика'), 
('жалоба: желчная колика', 'приступ желчной колики'), 
('жалоба: заложенность носа', 'заложенность носа'), 
('жалоба: запор', 'запоры'), 
('жалоба: запор', 'запор'), 
('жалоба: затруднения при глотании пищи', 'затруднения при глотании пищи'), 
('жалоба: звон в ушах', 'звон в ушах'), 
('жалоба: извращение вкуса', 'изменение вкуса'), 
('жалоба: извращение вкуса', 'извращение вкуса'), 
('жалоба: извращение обоняния', 'извращение обоняния'), 
('жалоба: изжога', 'изжога'), 
('жалоба: кожный зуд', 'кожный зуд'), 
('жалоба: кожный зуд', 'упорный кожный зуд'), 
('жалоба: ломота и боли в теле', 'ломота и боли в теле'), 
('жалоба: ломота и боли в теле', 'ломота в теле'), 
('жалоба: мышечная боль', 'миалгия'), 
('жалоба: мышечная боль', 'мышечная боль'), 
('жалоба: мышечная слабость', 'мышечная слабость'), 
('жалоба: непереносимость алкоголя', 'непереносимость алкоголя'), 
('жалоба: непереносимость газированных напитков', 'непереносимость газированных напитков'), 
('жалоба: непереносимость глютена', 'непереносимость глютена'), 
('жалоба: непереносимость глютена', 'целиакия'), 
('жалоба: непереносимость жареной пищи', 'непереносимость жареной пищи'), 
('жалоба: непереносимость жирной пищи', 'непереносимость жирной пищи'), 
('жалоба: непереносимость лактозы', 'непереваривание молока'), 
('жалоба: непереносимость лактозы', 'непереносимость лактозы'), 
('жалоба: непереносимость лактозы', 'неусвоение молока'), 
('жалоба: непереносимость острой пищи', 'непереносимость острой пищи'), 
('жалоба: непереносимость яиц', 'непереносимость яиц'), 
('жалоба: непреодолимые позывы на мочеиспускание', 'императивные позывы на мочеиспускание'), 
('жалоба: непреодолимые позывы на мочеиспускание', 'непреодолимые позывы на мочеиспускание'), 
('жалоба: непрерывная рвота не приносящая облегчение', 'непрерывная рвота не приносящая облегчение'), 
('жалоба: онемение конечностей', 'онемение и парестезии в конечностях'), 
('жалоба: онемение конечностей', 'онемение конечностей'), 
('жалоба: першение в горле', 'першение в горле'), 
('жалоба: полидипсия', 'неестественно сильная жажда'), 
('жалоба: полидипсия', 'полидипсия'), 
('жалоба: потеря вкуса', 'потеря вкуса'), 
('жалоба: потеря вкуса', 'агевзия'), 
('жалоба: потеря обоняния', 'потеря обоняния'), 
('жалоба: пристрастие к необычным запахам (бензин, ацетон, лаки, краски, гуталин, сигаретный дым, …)', 'пристрастие к необычным запахам (бензин, ацетон, лаки, краски, гуталин, сигаретный дым, …)'), 
('жалоба: раздражительность', 'раздражительность'), 
('жалоба: слабость', 'слабость'), 
('жалоба: слабость', 'общая слабость'), 
('жалоба: слабость', 'недомогание'), 
('жалоба: слабость сфинктеров', 'слабость сфинктеров'), 
('жалоба: слепое пятно в поле зрения', 'слепое пятно в поле зрения'), 
('жалоба: слепое пятно в поле зрения', 'слепое пятно в зрительном поле'), 
('жалоба: слепое пятно в поле зрения', 'скотома'), 
('жалоба: стресс', 'стресс'), 
('жалоба: стресс', 'стрессовые ситуации'), 
('жалоба: сухость во рту', 'сухость во рту'), 
('жалоба: тяжесть в эпигастрии', 'тяжесть в эпигастральной области'), 
('жалоба: тяжесть в эпигастрии', 'тяжесть в эпигастрии'), 
('жалоба: тяжесть в эпигастрии', 'чувство тяжести в животе после приема пищи'), 
('жалоба: тяжесть в эпигастрии', 'тяжесть в верхней части живота'), 
('жалоба: умеренный озноб', 'умеренный озноб'), 
('жалоба: усиление симптомов хронической сердечной недостаточности', 'усиление симптомов хронической сердечной недостаточности'), 
('жалоба: утомляемость', 'быстрая утомляемость'), 
('жалоба: утомляемость', 'утомляемость'), 
('жалоба: утомляемость', 'повышенная утомляемость'), 
('жалоба: чувство "комка" в горле', 'чувство "комка" в горле'), 
('жалоба: чувство раннего насыщения', 'чувство переполнения после еды'), 
('жалоба: чувство раннего насыщения', 'чувство быстрого насыщения'), 
('жалоба: чувство раннего насыщения', 'чувство раннего насыщения'), 
('жалоба: шум в ушах', 'шум в ушах'), 
}

vzNM = {
('C-реактивный белок', '', '', '', '', '', '', 100.0, 'мг/л', '', '', 'C-реактивный белок: от и более 100 мг/л', ''), 
('C-реактивный белок', '', '', '', '', '', 1.0, 20.0, 'мг/л', '', 'C-реактивный белок: от 1 до 20 мг/л', '', ''), 
('C-реактивный белок', '', '', '', '', '', 1.0, 3.0, 'мг/л', 'C-реактивный белок: до 1 мг/л', 'C-реактивный белок: от 1 до 3 мг/л', 'C-реактивный белок: от и более 3 мг/л', ''), 
('C-реактивный белок', '', '', '', '', '', '', 1.0, 'мг/л', '', '', 'C-реактивный белок: от и более 1 мг/л', ''), 
('IgA в сыворотке крови', '', 19, '', '', '', 0.7, 4.0, 'г/л', 'IgA в сыворотке крови: снижение', 'IgA в сыворотке крови: норма', 'IgA в сыворотке крови: повышение', ''), 
('Pb71', 'пол: мужской', 18, 45, '', '', 3.8, 5.1, '', '', 'Pb71: green', '', ''), 
('Pc23', '', 18, 45, '', '', 1.5, 3.7, '', '', 'Pc23: green', '', ''), 
('Tc', '', 50, 100, '', '', 9.23, 12.71, '', '', 'Tc: green', '', 401643), 
('Tc', '', 50, 100, '', '', 9.23, 12.71, '', '', 'Tc: green', '', 614273), 
('Tc', '', 50, 100, '', '', 9.23, 12.71, '', '', 'Tc: green', '', 278751), 
('Tc', '', 50, 100, '', '', 9.23, 12.71, '', '', 'Tc: green', '', 215721), 
('агрегация тромбоцитов', '', '', '', '', '', 25, 75, '%', 'агрегация тромбоцитов: снижение', 'агрегация тромбоцитов: норма', 'агрегация тромбоцитов: повышение', ''), 
('АЛТ', 'пол: мужской', 18, '', '', '', 0, 41.0, 'Ед/л', 'АЛТ: снижение', 'АЛТ: норма', 'АЛТ: повышение', ''), 
('АЛТ', '', 1, 7, '', '', 0, 29.0, 'Ед/л', 'АЛТ: снижение', 'АЛТ: норма', 'АЛТ: повышение', ''), 
('АЛТ', 'пол: женский', 18, '', '', '', 0, 33.0, 'Ед/л', 'АЛТ: снижение', 'АЛТ: норма', 'АЛТ: повышение', ''), 
('АЛТ', '', 7, 18, '', '', 0, 37.0, 'Ед/л', 'АЛТ: снижение', 'АЛТ: норма', 'АЛТ: повышение', ''), 
('АЛТ', '', 0, 1, '', '', 0, 56.0, 'Ед/л', 'АЛТ: снижение', 'АЛТ: норма', 'АЛТ: повышение', ''), 
('альбумин', '', 14, 18, '', '', 32.0, 45.0, 'г/л', 'альбумин: снижение', 'альбумин: норма', 'альбумин: повышение', ''), 
('альбумин', '', 1, 14, '', '', 38.0, 54.0, 'г/л', 'альбумин: снижение', 'альбумин: норма', 'альбумин: повышение', ''), 
('альбумин', '', 18, '', '', '', 35.0, 52.0, 'г/л', 'альбумин: снижение', 'альбумин: норма', 'альбумин: повышение', ''), 
('амилаза общая', '', '', '', '', '', 28, 100, 'Ед/л', 'амилаза общая: снижение', 'амилаза общая: норма', 'амилаза общая: повышение', ''), 
('амилаза панкреатическая', '', 0, 1, '', '', 0, 8.0, 'Ед/л', 'амилаза панкреатическая: снижение', 'амилаза панкреатическая: норма', 'амилаза панкреатическая: повышение', ''), 
('амилаза панкреатическая', '', 1, 10, '', '', 0, 31.0, 'Ед/л', 'амилаза панкреатическая: снижение', 'амилаза панкреатическая: норма', 'амилаза панкреатическая: повышение', ''), 
('амилаза панкреатическая', '', 18, '', '', '', 0, 53, 'Ед/л', 'амилаза панкреатическая: снижение', 'амилаза панкреатическая: норма', 'амилаза панкреатическая: повышение', ''), 
('амилаза панкреатическая', '', 10, 18, '', '', 0, 39.0, 'Ед/л', 'амилаза панкреатическая: снижение', 'амилаза панкреатическая: норма', 'амилаза панкреатическая: повышение', ''), 
('АСТ', 'пол: мужской', 18, '', '', '', '', 40, 'Ед/л', 'АСТ: снижение', 'АСТ: норма', 'АСТ: повышение', ''), 
('АСТ', 'пол: женский', 0, 1, '', '', '', 58, 'Ед/л', 'АСТ: снижение', 'АСТ: норма', 'АСТ: повышение', ''), 
('АСТ', 'пол: женский', 7, 13, '', '', '', 44, 'Ед/л', 'АСТ: снижение', 'АСТ: норма', 'АСТ: повышение', ''), 
('АСТ', 'пол: женский', 18, '', '', '', '', 32, 'Ед/л', 'АСТ: снижение', 'АСТ: норма', 'АСТ: повышение', ''), 
('АСТ', 'пол: женский', 4, 7, '', '', '', 48, 'Ед/л', 'АСТ: снижение', 'АСТ: норма', 'АСТ: повышение', ''), 
('АСТ', 'пол: женский', 1, 4, '', '', '', 59, 'Ед/л', 'АСТ: снижение', 'АСТ: норма', 'АСТ: повышение', ''), 
('АСТ', 'пол: женский', 13, 18, '', '', '', 39, 'Ед/л', 'АСТ: снижение', 'АСТ: норма', 'АСТ: повышение', ''), 
('АФП', 'пол: женский', 1, '', 12.0, 15.0, 15.0, 60.0, 'МЕ/мл', 'АФП: снижение при беременности', 'АФП: норма', 'АФП: повышение при беременности', ''), 
('АФП', 'пол: женский', 1, '', '', 0.0, 5.8, 500.0, 'МЕ/мл', '', 'АФП: от 5.8 до 500 МЕ/мл', '', ''), 
('АФП', 'пол: женский', 1, '', 30.0, 32.0, 100.0, 250.0, 'МЕ/мл', '', '', 'АФП: повышение', ''), 
('АФП', 'пол: женский', 1, '', 24.0, 28.0, 52.0, 140.0, 'МЕ/мл', 'АФП: снижение при беременности', 'АФП: норма', 'АФП: повышение при беременности', ''), 
('АФП', 'пол: женский', 1, '', 28.0, 30.0, 67.0, 150.0, 'МЕ/мл', 'АФП: снижение при беременности', 'АФП: норма', 'АФП: повышение при беременности', ''), 
('АФП', 'пол: женский', 1, '', 19.0, 24.0, 27.0, 125.0, 'МЕ/мл', 'АФП: снижение при беременности', 'АФП: норма', 'АФП: повышение при беременности', ''), 
('АФП', 'пол: женский', 1, '', 19.0, 24.0, 27.0, 125.0, 'МЕ/мл', '', '', 'АФП: повышение', ''), 
('АФП', 'пол: женский', 1, '', 24.0, 28.0, 52.0, 140.0, 'МЕ/мл', '', '', 'АФП: повышение', ''), 
('АФП', 'пол: женский', 1, '', 0.0, 12.0, 0.0, 15.0, 'МЕ/мл', '', '', 'АФП: повышение', ''), 
('АФП', 'пол: женский', 1, '', 15.0, 19.0, 15.0, 95.0, 'МЕ/мл', 'АФП: снижение при беременности', 'АФП: норма', 'АФП: повышение при беременности', ''), 
('АФП', 'пол: женский', 1, '', 12.0, 15.0, 15.0, 60.0, 'МЕ/мл', '', '', 'АФП: повышение', ''), 
('АФП', 'пол: мужской', 1, '', '', '', 0.0, 5.8, 'МЕ/мл', '', 'АФП: норма', 'АФП: повышение', ''), 
('АФП', 'пол: женский', 1, '', 0.0, 12.0, 0.0, 15.0, 'МЕ/мл', 'АФП: снижение при беременности', 'АФП: норма', 'АФП: повышение при беременности', ''), 
('АФП', 'пол: женский', 1, '', 30.0, 32.0, 100.0, 250.0, 'МЕ/мл', 'АФП: снижение при беременности', 'АФП: норма', 'АФП: повышение при беременности', ''), 
('АФП', 'пол: мужской', 1, '', '', '', 5.8, 500.0, 'МЕ/мл', '', 'АФП: от 5.8 до 500 МЕ/мл', '', ''), 
('АФП', 'пол: женский', 1, '', 15.0, 19.0, 15.0, 95.0, 'МЕ/мл', '', '', 'АФП: повышение', ''), 
('АФП', 'пол: женский', 1, '', '', 0.0, 5.8, 100.0, 'МЕ/мл', '', 'АФП: от 5.8 до 100 МЕ/мл', '', ''), 
('АФП', 'пол: мужской', 1, '', '', '', 5.8, 100.0, 'МЕ/мл', '', 'АФП: от 5.8 до 100 МЕ/мл', '', ''), 
('АФП', 'пол: женский', 1, '', '', 0.0, 0.0, 5.8, 'МЕ/мл', '', 'АФП: норма', 'АФП: повышение', ''), 
('АФП', 'пол: женский', 1, '', 28.0, 30.0, 67.0, 150.0, 'МЕ/мл', '', '', 'АФП: повышение', ''), 
('базофилы', '', '', '', '', '', 0.0, 0.08, '10^9/л', 'базофилы: снижение', 'базофилы: норма', 'базофилы: повышение', ''), 
('базофилы, %', '', '', '', '', '', 0.0, 1.2, '%', 'базофилы, %: снижение', 'базофилы, %: норма', 'базофилы, %: повышение', ''), 
('белок общий', '', 3, 18, '', '', 60.0, 80.0, 'г/л', 'белок общий: снижение', 'белок общий: норма', 'белок общий: повышение', ''), 
('белок общий', '', 18, '', '', '', 64.0, 83.0, 'г/л', 'белок общий: снижение', 'белок общий: норма', 'белок общий: повышение', ''), 
('белок общий', '', 1, 3, '', '', 56.0, 75.0, 'г/л', 'белок общий: снижение', 'белок общий: норма', 'белок общий: повышение', ''), 
('белок общий в моче за сутки, г/сутки', '', '', '', '', '', 0.0, 0.3, 'г/сутки', '', 'белок общий в моче за сутки: норма', 'белок общий в моче за сутки: повышение', ''), 
('белок общий в моче за сутки, мг/сутки', '', '', '', '', '', 0.0, 300.0, 'мг/сутки', '', 'белок общий в моче за сутки: норма', 'белок общий в моче за сутки: повышение', ''), 
('белок общий в моче после интенсивной физической нагрузки, г/сутки', '', '', '', '', '', 0.0, 0.3, 'г/сутки', '', 'белок общий в моче после интенсивной физической нагрузки: норма', 'белок общий в моче после интенсивной физической нагрузки: повышение', ''), 
('белок общий в моче после интенсивной физической нагрузки, мг/сутки', '', '', '', '', '', 0.0, 300.0, 'мг/сутки', '', 'белок общий в моче после интенсивной физической нагрузки: норма', 'белок общий в моче после интенсивной физической нагрузки: повышение', ''), 
('белок общий в моче, г/л', '', '', '', '', '', 0.0, 0.15, 'г/л', '', 'белок общий в моче: норма', 'белок общий в моче: повышение', ''), 
('белок общий в моче, мг/л', '', '', '', '', '', 0.0, 150.0, 'мг/л', '', 'белок общий в моче: норма', 'белок общий в моче: повышение', ''), 
('бета-глобулин', '', '', '', '', '', 7.9, 13.7, '%', 'бета-глобулин: снижение', 'бета-глобулин: норма', 'бета-глобулин: повышение', ''), 
('билирубин общий', '', 1, '', '', '', 0, 21.0, 'мкмоль/л', 'билирубин общий: снижение', 'билирубин общий: норма', 'билирубин общий: повышение', ''), 
('билирубин прямой', '', '', '', '', '', '', 5.0, 'мкмоль/л', '', 'билирубин прямой: норма', 'билирубин прямой: повышение', ''), 
('витамин B12', '', '', '', '', '', 197.0, 771.0, 'пг/мл', 'витамин B12: снижение', 'витамин B12: норма', 'витамин B12: повышение', ''), 
('гамма-глобулин', '', '', '', '', '', 11.1, 18.8, '%', 'гамма-глобулин: снижение', 'гамма-глобулин: норма', 'гамма-глобулин: повышение', ''), 
('гамма-ГТП', '', 3, 6, '', '', 0.0, 23.0, 'Ед/л', 'гамма-ГТП: снижение', 'гамма-ГТП: норма', 'гамма-ГТП: повышение', ''), 
('гамма-ГТП', 'пол: женский', 12, 17, '', '', 0, 33.0, 'Ед/л', '', 'гамма-ГТП: норма', 'гамма-ГТП: повышение', ''), 
('гамма-ГТП', '', 1, 3, '', '', 0.0, 18.0, 'Ед/л', 'гамма-ГТП: снижение', 'гамма-ГТП: норма', 'гамма-ГТП: повышение', ''), 
('гамма-ГТП', 'пол: мужской', 12, 17, '', '', 0.0, 45.0, 'Ед/л', '', 'гамма-ГТП: норма', 'гамма-ГТП: повышение', ''), 
('гамма-ГТП', 'пол: женский', 17, '', '', '', 6.0, 42.0, 'Ед/л', 'гамма-ГТП: снижение', 'гамма-ГТП: норма', 'гамма-ГТП: повышение', ''), 
('гамма-ГТП', 'пол: мужской', 17, '', '', '', 10.0, 71.0, 'Ед/л', 'гамма-ГТП: снижение', 'гамма-ГТП: норма', 'гамма-ГТП: повышение', ''), 
('гамма-ГТП', '', 6, 12, '', '', 0.0, 17.0, 'Ед/л', 'гамма-ГТП: снижение', 'гамма-ГТП: норма', 'гамма-ГТП: повышение', ''), 
('гастрин 17 (базальный)', '', '', '', '', '', '', 5.0, 'пмоль/л', '', '', 'гастрин 17 (базальный): от и более 5 пмоль/л', ''), 
('гастрин 17 (базальный)', '', '', '', '', '', 1.0, 7.0, 'пмоль/л', 'гастрин 17 (базальный): снижение', 'гастрин 17 (базальный): норма', 'гастрин 17 (базальный): повышение', ''), 
('гастрин 17 (базальный)', '', '', '', '', '', 10.0, 20.0, 'пмоль/л', '', 'гастрин 17 (базальный): от 10 до 20 пмоль/л', '', ''), 
('гастрин 17 (базальный)', '', '', '', '', '', 7.0, 10.0, 'пмоль/л', '', 'гастрин 17 (базальный): от 7 до 10 пмоль/л', '', ''), 
('гастрин 17 (базальный)', '', '', '', '', '', '', 20.0, 'пмоль/л', '', '', 'гастрин 17 (базальный): от и более 20 пмоль/л', ''), 
('гематокрит', 'пол: женский', '', '', '', '', 35.0, 45.0, '%', 'гематокрит: снижение', 'гематокрит: норма', 'гематокрит: повышение', 11.278), 
('гематокрит', 'пол: мужской', '', '', '', '', 3.9, 4.9, 'л/л', 'гематокрит: снижение', 'гематокрит: норма', 'гематокрит: повышение', 12.478), 
('гематокрит', 'пол: женский', '', '', '', '', 3.5, 4.5, 'л/л', 'гематокрит: снижение', 'гематокрит: норма', 'гематокрит: повышение', 11.278), 
('гематокрит', 'пол: мужской', '', '', '', '', 39.0, 49.0, '%', 'гематокрит: снижение', 'гематокрит: норма', 'гематокрит: повышение', 12.478), 
('гемоглобин', '', 1, 5, '', '', 11.0, 14.0, 'г/дл', 'гемоглобин: снижение', 'гемоглобин: норма', 'гемоглобин: повышение', 35.333), 
('гемоглобин', 'пол: мужской', 18, 45, '', '', 9, 13.2, 'г/дл', '', 'гемоглобин: от 9 г/дл и ниже нормы', '', 'neS4'), 
('гемоглобин', 'пол: женский', 12, 15, '', '', 11.5, 15, 'г/дл', 'гемоглобин: снижение', 'гемоглобин: норма', 'гемоглобин: повышение', 37.222), 
('гемоглобин', 'пол: мужской', 65, '', '', '', 90, 126, 'г/л', '', 'гемоглобин: от 90 г/л и ниже нормы', '', 'neS4'), 
('гемоглобин', 'пол: женский', 65, '', '', '', 117, 161, 'г/л', 'гемоглобин: снижение', 'гемоглобин: норма', 'гемоглобин: повышение', 38.522), 
('гемоглобин', 'пол: женский', 12, 15, '', '', 9, 11.5, 'г/дл', '', 'гемоглобин: от 9 г/дл и ниже нормы', '', 'neS4'), 
('гемоглобин', 'пол: женский', 12, 15, '', '', 90, 115, 'г/л', '', 'гемоглобин: от 90 г/л и ниже нормы', '', 'neS4'), 
('гемоглобин', '', 1, 5, '', '', 110.0, 140.0, 'г/л', 'гемоглобин: снижение', 'гемоглобин: норма', 'гемоглобин: повышение', 35.333), 
('гемоглобин', 'пол: мужской', 65, '', '', '', 9, 12.6, 'г/дл', '', 'гемоглобин: от 9 г/дл и ниже нормы', '', 'neS4'), 
('гемоглобин', 'пол: мужской', 12, 15, '', '', 120.0, 160.0, 'г/л', 'гемоглобин: снижение', 'гемоглобин: норма', 'гемоглобин: повышение', 39.111), 
('гемоглобин', 'пол: женский', 45, 65, '', '', 117, 160, 'г/л', 'гемоглобин: снижение', 'гемоглобин: норма', 'гемоглобин: повышение', 38.444), 
('гемоглобин', 'пол: мужской', 45, 65, '', '', 9, 13.1, 'г/дл', '', 'гемоглобин: от 9 г/дл и ниже нормы', '', 'neS4'), 
('гемоглобин', '', 10, 12, '', '', 12, 15, 'г/дл', 'гемоглобин: снижение', 'гемоглобин: норма', 'гемоглобин: повышение', 38.333), 
('гемоглобин', 'пол: женский', 15, 18, '', '', 117, 153, 'г/л', 'гемоглобин: снижение', 'гемоглобин: норма', 'гемоглобин: повышение', 37.9), 
('гемоглобин', 'пол: мужской', 12, 15, '', '', 12.0, 16.0, 'г/дл', 'гемоглобин: снижение', 'гемоглобин: норма', 'гемоглобин: повышение', 39.111), 
('гемоглобин', '', 10, 12, '', '', 90, 120, 'г/л', '', 'гемоглобин: от 90 г/л и ниже нормы', '', 'neS4'), 
('гемоглобин', 'пол: женский', 15, 18, '', '', 11.7, 15.3, 'г/дл', 'гемоглобин: снижение', 'гемоглобин: норма', 'гемоглобин: повышение', 37.9), 
('гемоглобин', 'пол: мужской', 18, 45, '', '', 90, 132, 'г/л', '', 'гемоглобин: от 90 г/л и ниже нормы', '', 'neS4'), 
('гемоглобин', 'пол: женский', 12, 15, '', '', 115.0, 150.0, 'г/л', 'гемоглобин: снижение', 'гемоглобин: норма', 'гемоглобин: повышение', 37.222), 
('гемоглобин', 'пол: мужской', 45, 65, '', '', 131, 172, 'г/л', 'гемоглобин: снижение', 'гемоглобин: норма', 'гемоглобин: повышение', 42.489), 
('гемоглобин', '', '', '', '', '', 70, 90, 'г/л', 'гемоглобин: до 70 г/л', 'гемоглобин: от 70 до 90 г/л', '', 'neS4'), 
('гемоглобин', 'пол: женский', 45, 65, '', '', 11.7, 16, 'г/дл', 'гемоглобин: снижение', 'гемоглобин: норма', 'гемоглобин: повышение', 38.444), 
('гемоглобин', 'пол: мужской', 65, '', '', '', 126, 174, 'г/л', 'гемоглобин: снижение', 'гемоглобин: норма', 'гемоглобин: повышение', 41.533), 
('гемоглобин', '', 5, 10, '', '', 115.0, 145.0, 'г/л', 'гемоглобин: снижение', 'гемоглобин: норма', 'гемоглобин: повышение', 36.833), 
('гемоглобин', 'пол: мужской', 15, 18, '', '', 117.0, 166.0, 'г/л', 'гемоглобин: снижение', 'гемоглобин: норма', 'гемоглобин: повышение', 38.911), 
('гемоглобин', 'пол: женский', 18, 45, '', '', 117.0, 155.0, 'г/л', 'гемоглобин: снижение', 'гемоглобин: норма', 'гемоглобин: повышение', 38.056), 
('гемоглобин', 'пол: женский', 45, 65, '', '', 9, 11.7, 'г/дл', '', 'гемоглобин: от 9 г/дл и ниже нормы', '', 'neS4'), 
('гемоглобин', 'пол: мужской', 45, 65, '', '', 13.1, 17.2, 'г/дл', 'гемоглобин: снижение', 'гемоглобин: норма', 'гемоглобин: повышение', 42.489), 
('гемоглобин', '', '', '', '', '', 7, 9, 'г/дл', 'гемоглобин: до 7 г/дл', 'гемоглобин: от 7 до 9 г/дл', '', 'neS4'), 
('гемоглобин', '', 5, 10, '', '', 11.5, 14.5, 'г/дл', 'гемоглобин: снижение', 'гемоглобин: норма', 'гемоглобин: повышение', 36.833), 
('гемоглобин', 'пол: мужской', 45, 65, '', '', 90, 131, 'г/л', '', 'гемоглобин: от 90 г/л и ниже нормы', '', 'neS4'), 
('гемоглобин', 'пол: мужской', 18, 45, '', '', 132, 173, 'г/л', 'гемоглобин: снижение', 'гемоглобин: норма', 'гемоглобин: повышение', 42.789), 
('гемоглобин', 'пол: мужской', 65, '', '', '', 12.6, 17.4, 'г/дл', 'гемоглобин: снижение', 'гемоглобин: норма', 'гемоглобин: повышение', 41.533), 
('гемоглобин', '', 10, 12, '', '', 9, 12, 'г/дл', '', 'гемоглобин: от 9 г/дл и ниже нормы', '', 'neS4'), 
('гемоглобин', '', 10, 12, '', '', 120, 150, 'г/л', 'гемоглобин: снижение', 'гемоглобин: норма', 'гемоглобин: повышение', 38.333), 
('гемоглобин', 'пол: женский', 65, '', '', '', 9, 11.7, 'г/дл', '', 'гемоглобин: от 9 г/дл и ниже нормы', '', 'neS4'), 
('гемоглобин', 'пол: женский', 45, 65, '', '', 90, 117, 'г/л', '', 'гемоглобин: от 90 г/л и ниже нормы', '', 'neS4'), 
('гемоглобин', 'пол: женский', 15, 18, '', '', 90, 117, 'г/л', '', 'гемоглобин: от 90 г/л и ниже нормы', '', 'neS4'), 
('гемоглобин', 'пол: женский', 65, '', '', '', 90, 117, 'г/л', '', 'гемоглобин: от 90 г/л и ниже нормы', '', 'neS4'), 
('гемоглобин', 'пол: мужской', 18, 45, '', '', 13.2, 17.3, 'г/дл', 'гемоглобин: снижение', 'гемоглобин: норма', 'гемоглобин: повышение', 42.789), 
('гемоглобин', 'пол: женский', 15, 18, '', '', 9, 11.7, 'г/дл', '', 'гемоглобин: от 9 г/дл и ниже нормы', '', 'neS4'), 
('гемоглобин', 'пол: мужской', 15, 18, '', '', 11.7, 16.6, 'г/дл', 'гемоглобин: снижение', 'гемоглобин: норма', 'гемоглобин: повышение', 38.911), 
('гемоглобин', 'пол: женский', 65, '', '', '', 11.7, 16.1, 'г/дл', 'гемоглобин: снижение', 'гемоглобин: норма', 'гемоглобин: повышение', 38.522), 
('гемоглобин', 'пол: женский', 18, 45, '', '', 11.7, 15.5, 'г/дл', 'гемоглобин: снижение', 'гемоглобин: норма', 'гемоглобин: повышение', 38.056), 
('гликированный гемоглобин', '', '', '', '', '', 6.0, 6.5, '%', 'гликированный гемоглобин: до 6%', 'гликированный гемоглобин: от 6 до 6,5%', 'гликированный гемоглобин: от 6,5%', ''), 
('глюкоза', '', '', 14, '', '', 3.3, 5.6, 'ммоль/л', 'глюкоза: снижение', 'глюкоза: норма', 'глюкоза: повышение', ''), 
('глюкоза', '', 14, '', '', 0, 4.1, 6.1, 'ммоль/л', 'глюкоза: снижение', 'глюкоза: норма', 'глюкоза: повышение', ''), 
('глюкоза', 'пол: женский', '', '', 0, '', 4.1, 5.1, 'ммоль/л', 'глюкоза: снижение', 'глюкоза: норма', 'глюкоза: повышение', ''), 
('гомоцистеин', 'пол: женский', '', '', '', '', 4.44, 13.56, 'мкмоль/л', 'гомоцистеин: снижение', 'гомоцистеин: норма', 'гомоцистеин: повышение', ''), 
('гомоцистеин', 'пол: мужской', '', '', '', '', 5.46, 16.2, 'мкмоль/л', 'гомоцистеин: снижение', 'гомоцистеин: норма', 'гомоцистеин: повышение', ''), 
('железо', 'пол: мужской', 1, 4, '', '', 5.2, 16.3, 'мкмоль/л', 'железо: снижение', 'железо: норма', 'железо: повышение', ''), 
('железо', 'пол: женский', 18, '', '', '', 6.6, 26, 'мкмоль/л', 'железо: снижение', 'железо: норма', 'железо: повышение', ''), 
('железо', 'пол: женский', 4, 7, '', '', 5, 16.7, 'мкмоль/л', 'железо: снижение', 'железо: норма', 'железо: повышение', ''), 
('железо', 'пол: женский', 16, 18, '', '', 5.9, 18.3, 'мкмоль/л', 'железо: снижение', 'железо: норма', 'железо: повышение', ''), 
('железо', 'пол: мужской', 4, 7, '', '', 4.5, 20.6, 'мкмоль/л', 'железо: снижение', 'железо: норма', 'железо: повышение', ''), 
('железо', 'пол: мужской', 13, 16, '', '', 4.7, 19.7, 'мкмоль/л', 'железо: снижение', 'железо: норма', 'железо: повышение', ''), 
('железо', 'пол: мужской', 10, 13, '', '', 5, 20, 'мкмоль/л', 'железо: снижение', 'железо: норма', 'железо: повышение', ''), 
('железо', 'пол: женский', 7, 10, '', '', 5.4, 18.6, 'мкмоль/л', 'железо: снижение', 'железо: норма', 'железо: повышение', ''), 
('железо', 'пол: женский', 13, 16, '', '', 5.4, 19.5, 'мкмоль/л', 'железо: снижение', 'железо: норма', 'железо: повышение', ''), 
('железо', 'пол: мужской', 18, '', '', '', 11, 28, 'мкмоль/л', 'железо: снижение', 'железо: норма', 'железо: повышение', ''), 
('железо', 'пол: мужской', 16, 18, '', '', 4.8, 24.7, 'мкмоль/л', 'железо: снижение', 'железо: норма', 'железо: повышение', ''), 
('железо', 'пол: мужской', 7, 10, '', '', 4.8, 17.2, 'мкмоль/л', 'железо: снижение', 'железо: норма', 'железо: повышение', ''), 
('железо', 'пол: женский', 10, 13, '', '', 5.7, 18.6, 'мкмоль/л', 'железо: снижение', 'железо: норма', 'железо: повышение', ''), 
('железо', 'пол: женский', 1, 4, '', '', 4.5, 18.1, 'мкмоль/л', 'железо: снижение', 'железо: норма', 'железо: повышение', ''), 
('желчные кислоты', '', '', '', '', '', 0, 10.0, 'мкмоль/л', 'желчные кислоты: снижение', 'желчные кислоты: норма', 'желчные кислоты: повышение', ''), 
('калий', '', '', '', '', '', 3.5, 5.1, 'ммоль/л', 'калий: снижение', 'калий: норма', 'калий: повышение', ''), 
('кальций общий', '', 60, 90, '', '', 2.22, 2.55, 'ммоль/л', 'кальций общий: снижение', 'кальций общий: норма', 'кальций общий: повышение', ''), 
('кальций общий', '', 1, 2, '', '', 2.25, 2.75, 'ммоль/л', 'кальций общий: снижение', 'кальций общий: норма', 'кальций общий: повышение', ''), 
('кальций общий', '', 2, 12, '', '', 2.2, 2.7, 'ммоль/л', 'кальций общий: снижение', 'кальций общий: норма', 'кальций общий: повышение', ''), 
('кальций общий', '', 18, 60, '', '', 2.15, 2.5, 'ммоль/л', 'кальций общий: снижение', 'кальций общий: норма', 'кальций общий: повышение', ''), 
('кальций общий', '', 90, '', '', '', 2.05, 2.4, 'ммоль/л', 'кальций общий: снижение', 'кальций общий: норма', 'кальций общий: повышение', ''), 
('кальций общий', '', 12, 18, '', '', 2.1, 2.55, 'ммоль/л', 'кальций общий: снижение', 'кальций общий: норма', 'кальций общий: повышение', ''), 
('коэффициент атерогенности', '', '', '', '', '', 0, 3.5, '', 'коэффициент атерогенности: снижение', 'коэффициент атерогенности: норма', 'коэффициент атерогенности: повышение', ''), 
('коэффициент больших тромбоцитов (P-LCR)', '', '', '', '', '', 13.0, 43.0, '%', 'коэффициент больших тромбоцитов (P-LCR): снижение', 'коэффициент больших тромбоцитов (P-LCR): норма', 'коэффициент больших тромбоцитов (P-LCR): повышение', ''), 
('креатинин', '', 9, 11, '', '', 34, 65, 'мкмоль/л', 'креатинин: снижение', 'креатинин: норма', 'креатинин: повышение', ''), 
('креатинин', '', 1, 3, '', '', 21, 36, 'мкмоль/л', 'креатинин: снижение', 'креатинин: норма', 'креатинин: повышение', ''), 
('креатинин', '', 7, 9, '', '', 35, 53, 'мкмоль/л', 'креатинин: снижение', 'креатинин: норма', 'креатинин: повышение', ''), 
('креатинин', 'пол: женский', 15, '', '', '', 44, 80, 'мкмоль/л', 'креатинин: снижение', 'креатинин: норма', 'креатинин: повышение', ''), 
('креатинин', 'пол: мужской', 15, '', '', '', 62, 106, 'мкмоль/л', 'креатинин: снижение', 'креатинин: норма', 'креатинин: повышение', ''), 
('креатинин', '', 5, 7, '', '', 28, 52, 'мкмоль/л', 'креатинин: снижение', 'креатинин: норма', 'креатинин: повышение', ''), 
('креатинин', '', 11, 13, '', '', 46, 70, 'мкмоль/л', 'креатинин: снижение', 'креатинин: норма', 'креатинин: повышение', ''), 
('креатинин', '', 3, 5, '', '', 27, 42, 'мкмоль/л', 'креатинин: снижение', 'креатинин: норма', 'креатинин: повышение', ''), 
('креатинин', '', 13, 15, '', '', 50, 77, 'мкмоль/л', 'креатинин: снижение', 'креатинин: норма', 'креатинин: повышение', ''), 
('креатинкиназа', 'пол: женский', 12, 17, '', '', 0, 123, 'Ед/л', 'креатинкиназа: снижение', 'креатинкиназа: норма', 'креатинкиназа: повышение', ''), 
('креатинкиназа', 'пол: мужской', 6, 12, '', '', 0, 247, 'Ед/л', 'креатинкиназа: снижение', 'креатинкиназа: норма', 'креатинкиназа: повышение', ''), 
('креатинкиназа', '', 1, 3, '', '', 0, 228, 'Ед/л', 'креатинкиназа: снижение', 'креатинкиназа: норма', 'креатинкиназа: повышение', ''), 
('креатинкиназа', '', 3, 6, '', '', 0, 149, 'Ед/л', 'креатинкиназа: снижение', 'креатинкиназа: норма', 'креатинкиназа: повышение', ''), 
('креатинкиназа', 'пол: женский', 17, '', '', '', 0, 167, 'Ед/л', 'креатинкиназа: снижение', 'креатинкиназа: норма', 'креатинкиназа: повышение', ''), 
('креатинкиназа', 'пол: мужской', 17, '', '', '', 0, 190, 'Ед/л', 'креатинкиназа: снижение', 'креатинкиназа: норма', 'креатинкиназа: повышение', ''), 
('креатинкиназа', 'пол: мужской', 12, 17, '', '', 0, 270, 'Ед/л', 'креатинкиназа: снижение', 'креатинкиназа: норма', 'креатинкиназа: повышение', ''), 
('креатинкиназа', 'пол: женский', 6, 12, '', '', 0, 154, 'Ед/л', 'креатинкиназа: снижение', 'креатинкиназа: норма', 'креатинкиназа: повышение', ''), 
('лактат', '', '', '', '', '', 0.5, 22, 'ммоль/л', 'лактат: снижение', 'лактат: норма', 'лактат: повышение', ''), 
('лактатдегидрогеназа', '', 1, 3, '', '', 0, 344, 'Ед/л', 'лактатдегидрогеназа: снижение', 'лактатдегидрогеназа: норма', 'лактатдегидрогеназа: повышение', ''), 
('лактатдегидрогеназа', 'пол: женский', 17, '', '', '', 135, 214, 'Ед/л', 'лактатдегидрогеназа: снижение', 'лактатдегидрогеназа: норма', 'лактатдегидрогеназа: повышение', ''), 
('лактатдегидрогеназа', '', 3, 6, '', '', 0, 314, 'Ед/л', 'лактатдегидрогеназа: снижение', 'лактатдегидрогеназа: норма', 'лактатдегидрогеназа: повышение', ''), 
('лактатдегидрогеназа', '', 0, 1, '', '', 0, 451, 'Ед/л', 'лактатдегидрогеназа: снижение', 'лактатдегидрогеназа: норма', 'лактатдегидрогеназа: повышение', ''), 
('лактатдегидрогеназа', '', 12, 17, '', '', 0, 279, 'Ед/л', 'лактатдегидрогеназа: снижение', 'лактатдегидрогеназа: норма', 'лактатдегидрогеназа: повышение', ''), 
('лактатдегидрогеназа', '', 6, 12, '', '', 0, 332, 'Ед/л', 'лактатдегидрогеназа: снижение', 'лактатдегидрогеназа: норма', 'лактатдегидрогеназа: повышение', ''), 
('лактатдегидрогеназа', 'пол: мужской', 17, '', '', '', 135, 225, 'Ед/л', 'лактатдегидрогеназа: снижение', 'лактатдегидрогеназа: норма', 'лактатдегидрогеназа: повышение', ''), 
('латентная железосвязывающая способность сыворотки крови', '', '', '', '', '', 20, 62, 'мкмоль/л', 'латентная железосвязывающая способность сыворотки крови: снижение', 'латентная железосвязывающая способность сыворотки крови: норма', 'латентная железосвязывающая способность сыворотки крови: повышение', ''), 
('лейкоциты', '', 4, 6, '', '', 5000, 14500, 'клеток/мкл', 'лейкоциты: снижение', 'лейкоциты: норма', 'лейкоциты: повышение', 1122.389), 
('лейкоциты', '', 6, 10, '', '', 4.5, 13.5, '10^9/л', 'лейкоциты: снижение', 'лейкоциты: норма', 'лейкоциты: повышение', 1010.5), 
('лейкоциты', '', 6, 10, '', '', 4500, 13500, 'клеток/мкл', 'лейкоциты: снижение', 'лейкоциты: норма', 'лейкоциты: повышение', 1010.5), 
('лейкоциты', '', 4, 6, '', '', 5, 14.5, '10^9/л', 'лейкоциты: снижение', 'лейкоциты: норма', 'лейкоциты: повышение', 1122.389), 
('лейкоциты', '', 1, 2, '', '', 6, 17, '10^9/л', 'лейкоциты: снижение', 'лейкоциты: норма', 'лейкоциты: повышение', 1346.556), 
('лейкоциты', '', 16, '', '', '', 4000, 10000, 'клеток/мкл', 'лейкоциты: снижение', 'лейкоциты: норма', 'лейкоциты: повышение', 896.667), 
('лейкоциты', '', 16, '', '', '', 4.0, 10.0, '10^9/л', 'лейкоциты: снижение', 'лейкоциты: норма', 'лейкоциты: повышение', 896.667), 
('лейкоциты', '', 1, 2, '', '', 6000, 17000, 'клеток/мкл', 'лейкоциты: снижение', 'лейкоциты: норма', 'лейкоциты: повышение', 1346.556), 
('лейкоциты', '', 0, 1, '', '', 6, 17.5, '10^9/л', 'лейкоциты: снижение', 'лейкоциты: норма', 'лейкоциты: повышение', 1346.944), 
('лейкоциты', '', 2, 4, '', '', 5.5, 15.5, '10^9/л', 'лейкоциты: снижение', 'лейкоциты: норма', 'лейкоциты: повышение', 1234.278), 
('лейкоциты', '', 0, 1, '', '', 6000, 17500, 'клеток/мкл', 'лейкоциты: снижение', 'лейкоциты: норма', 'лейкоциты: повышение', 1346.944), 
('лейкоциты', '', 2, 4, '', '', 5500, 15500, 'клеток/мкл', 'лейкоциты: снижение', 'лейкоциты: норма', 'лейкоциты: повышение', 1234.278), 
('лейкоциты', '', 10, 16, '', '', 4500, 13000, 'клеток/мкл', 'лейкоциты: снижение', 'лейкоциты: норма', 'лейкоциты: повышение', 1010.111), 
('лейкоциты', '', 10, 16, '', '', 4.5, 13, '10^9/л', 'лейкоциты: снижение', 'лейкоциты: норма', 'лейкоциты: повышение', 1010.111), 
('лимфоциты', '', 8, 10, '', '', 1.5, 6.5, '10^9/л', 'лимфоциты: снижение', 'лимфоциты: норма', 'лимфоциты: повышение', 338.389), 
('лимфоциты', '', 10, 16, '', '', 1.2, 5.2, '10^9/л', 'лимфоциты: снижение', 'лимфоциты: норма', 'лимфоциты: повышение', 270.711), 
('лимфоциты', '', 1, 2, '', '', 3000, 9500, 'клеток/мкл', 'лимфоциты: снижение', 'лимфоциты: норма', 'лимфоциты: повышение', 674.056), 
('лимфоциты', '', 2, 4, '', '', 2, 8, '10^9/л', 'лимфоциты: снижение', 'лимфоциты: норма', 'лимфоциты: повышение', 450.667), 
('лимфоциты', '', 0, 1, '', '', 2, 11, '10^9/л', 'лимфоциты: снижение', 'лимфоциты: норма', 'лимфоциты: повышение', 453.0), 
('лимфоциты', '', 4, 6, '', '', 1.5, 7, '10^9/л', 'лимфоциты: снижение', 'лимфоциты: норма', 'лимфоциты: повышение', 338.778), 
('лимфоциты', '', 16, '', '', '', 1.0, 4.8, '10^9/л', 'лимфоциты: снижение', 'лимфоциты: норма', 'лимфоциты: повышение', 225.956), 
('лимфоциты', '', 10, 16, '', '', 1200, 5200, 'клеток/мкл', 'лимфоциты: снижение', 'лимфоциты: норма', 'лимфоциты: повышение', 270.711), 
('лимфоциты', '', 6, 8, '', '', 1.5, 6.8, '10^9/л', 'лимфоциты: снижение', 'лимфоциты: норма', 'лимфоциты: повышение', 338.622), 
('лимфоциты', '', 8, 10, '', '', 1500, 6500, 'клеток/мкл', 'лимфоциты: снижение', 'лимфоциты: норма', 'лимфоциты: повышение', 338.389), 
('лимфоциты', '', 0, 1, '', '', 2000, 11000, 'клеток/мкл', 'лимфоциты: снижение', 'лимфоциты: норма', 'лимфоциты: повышение', 453.0), 
('лимфоциты', '', 1, 2, '', '', 3, 9.5, '10^9/л', 'лимфоциты: снижение', 'лимфоциты: норма', 'лимфоциты: повышение', 674.056), 
('лимфоциты', '', 2, 4, '', '', 2000, 8000, 'клеток/мкл', 'лимфоциты: снижение', 'лимфоциты: норма', 'лимфоциты: повышение', 450.667), 
('лимфоциты', '', 16, '', '', '', 1000.0, 4800.0, 'клеток/мкл', 'лимфоциты: снижение', 'лимфоциты: норма', 'лимфоциты: повышение', 225.956), 
('лимфоциты', '', 4, 6, '', '', 1500, 7000, 'клеток/мкл', 'лимфоциты: снижение', 'лимфоциты: норма', 'лимфоциты: повышение', 338.778), 
('лимфоциты', '', 6, 8, '', '', 1500, 6800, 'клеток/мкл', 'лимфоциты: снижение', 'лимфоциты: норма', 'лимфоциты: повышение', 338.622), 
('лимфоциты, %', '', 16, '', '', '', 19.0, 37.0, '%', 'лимфоциты, %: снижение', 'лимфоциты, %: норма', 'лимфоциты, %: повышение', 4.51), 
('лимфоциты, %', '', 16, '', '', '', 0.19, 0.37, 'доля', 'лимфоциты, %: снижение', 'лимфоциты, %: норма', 'лимфоциты, %: повышение', 4.51), 
('липаза', '', 18, '', '', '', 13, 60, 'Ед/л', 'липаза: снижение', 'липаза: норма', 'липаза: повышение', ''), 
('липаза', '', 13, 18, '', '', 0, 55, 'Ед/л', 'липаза: снижение', 'липаза: норма', 'липаза: повышение', ''), 
('липаза', '', 1, 13, '', '', 0, 31, 'Ед/л', 'липаза: снижение', 'липаза: норма', 'липаза: повышение', ''), 
('ЛПВП', 'пол: мужской', '', '', '', '', 1.0, '', 'ммоль/л', 'ЛПВП: снижение', 'ЛПВП: норма', 'ЛПВП: повышение', ''), 
('ЛПВП', 'пол: женский', '', '', '', '', 1.2, '', 'ммоль/л', 'ЛПВП: снижение', 'ЛПВП: норма', 'ЛПВП: повышение', ''), 
('ЛПНП', '', '', '', '', '', 0, 3.0, 'ммоль/л', 'ЛПНП: снижение', 'ЛПНП: норма', 'ЛПНП: повышение', ''), 
('магний', '', 1, 6, '', '', 0.7, 0.95, 'ммоль/л', 'магний: снижение', 'магний: норма', 'магний: повышение', ''), 
('магний', '', 12, 20, '', '', 0.7, 0.91, 'ммоль/л', 'магний: снижение', 'магний: норма', 'магний: повышение', ''), 
('магний', '', 20, '', '', '', 0.66, 1.07, 'ммоль/л', 'магний: снижение', 'магний: норма', 'магний: повышение', ''), 
('магний', '', 6, 12, '', '', 0.7, 0.86, 'ммоль/л', 'магний: снижение', 'магний: норма', 'магний: повышение', ''), 
('миоглобин', 'пол: мужской', '', '', '', '', 23, 72, 'мкг/л', 'миоглобин: снижение', 'миоглобин: норма', 'миоглобин: повышение', ''), 
('миоглобин', 'пол: женский', '', '', '', '', 19, 51, 'мкг/л', 'миоглобин: снижение', 'миоглобин: норма', 'миоглобин: повышение', ''), 
('моноциты', '', 16, '', '', '', 0.05, 0.82, '10^9/л', 'моноциты: снижение', 'моноциты: норма', 'моноциты: повышение', 11.749), 
('моноциты', '', 16, '', '', '', 50.0, 820.0, 'клеток/мкл', 'моноциты: снижение', 'моноциты: норма', 'моноциты: повышение', 11.749), 
('моноциты, %', '', 2, '', '', '', 0.03, 0.11, 'доля', 'моноциты, %: снижение', 'моноциты, %: норма', 'моноциты, %: повышение', 0.752), 
('моноциты, %', '', 2, '', '', '', 3.0, 11.0, '%', 'моноциты, %: снижение', 'моноциты, %: норма', 'моноциты, %: повышение', 0.752), 
('мочевая кислота', 'пол: женский', '', '', '', '', 142.8, 339.2, 'мкмоль/л', 'мочевая кислота: снижение', 'мочевая кислота: норма', 'мочевая кислота: повышение', ''), 
('мочевая кислота', 'пол: мужской', '', '', '', '', 202.3, 416.5, 'мкмоль/л', 'мочевая кислота: снижение', 'мочевая кислота: норма', 'мочевая кислота: повышение', ''), 
('мочевина', '', 4, 14, '', '', 2.5, 6, 'ммоль/л', 'мочевина: снижение', 'мочевина: норма', 'мочевина: повышение', ''), 
('мочевина', 'пол: мужской', 20, 50, '', '', 3.2, 7.3, 'ммоль/л', 'мочевина: снижение', 'мочевина: норма', 'мочевина: повышение', ''), 
('мочевина', '', 14, 20, '', '', 2.9, 7.5, 'ммоль/л', 'мочевина: снижение', 'мочевина: норма', 'мочевина: повышение', ''), 
('мочевина', 'пол: женский', 50, '', '', '', 3.5, 7.2, 'ммоль/л', 'мочевина: снижение', 'мочевина: норма', 'мочевина: повышение', ''), 
('мочевина', 'пол: женский', 20, 50, '', '', 2.6, 6.7, 'ммоль/л', 'мочевина: снижение', 'мочевина: норма', 'мочевина: повышение', ''), 
('мочевина', 'пол: мужской', 50, '', '', '', 3, 9.2, 'ммоль/л', 'мочевина: снижение', 'мочевина: норма', 'мочевина: повышение', ''), 
('мочевина', '', 0, 4, '', '', 1.8, 6, 'ммоль/л', 'мочевина: снижение', 'мочевина: норма', 'мочевина: повышение', ''), 
('натрий', '', '', '', '', '', 136, 145, 'ммоль/л', 'натрий: снижение', 'натрий: норма', 'натрий: повышение', ''), 
('нейтрофилы', '', 4, 8, '', '', 1.5, 8, '10^9/л', 'нейтрофилы: снижение', 'нейтрофилы: норма', 'нейтрофилы: повышение', 339.556), 
('нейтрофилы', '', 8, 16, '', '', 1800, 8000, 'клеток/мкл', 'нейтрофилы: снижение', 'нейтрофилы: норма', 'нейтрофилы: повышение', 406.222), 
('нейтрофилы', '', 1, 4, '', '', 1500, 8500, 'клеток/мкл', 'нейтрофилы: снижение', 'нейтрофилы: норма', 'нейтрофилы: повышение', 339.944), 
('нейтрофилы', '', 16, '', '', '', 1800.0, 7700.0, 'клеток/мкл', 'нейтрофилы: снижение', 'нейтрофилы: норма', 'нейтрофилы: повышение', 405.989), 
('нейтрофилы', '', 16, '', '', '', 1.8, 7.7, '10^9/л', 'нейтрофилы: снижение', 'нейтрофилы: норма', 'нейтрофилы: повышение', 405.989), 
('нейтрофилы', '', 1, 4, '', '', 1.5, 8.5, '10^9/л', 'нейтрофилы: снижение', 'нейтрофилы: норма', 'нейтрофилы: повышение', 339.944), 
('нейтрофилы', '', 8, 16, '', '', 1.8, 8, '10^9/л', 'нейтрофилы: снижение', 'нейтрофилы: норма', 'нейтрофилы: повышение', 406.222), 
('нейтрофилы', '', 4, 8, '', '', 1500, 8000, 'клеток/мкл', 'нейтрофилы: снижение', 'нейтрофилы: норма', 'нейтрофилы: повышение', 339.556), 
('нейтрофилы палочкоядерные, %', '', '', '', '', '', 1, 5, '%', 'нейтрофилы палочкоядерные, %: снижение', 'нейтрофилы палочкоядерные, %: норма', 'нейтрофилы палочкоядерные, %: повышение', ''), 
('нейтрофилы сегментоядерные, %', '', '', '', '', '', 40, 72, '%', 'нейтрофилы сегментоядерные, %: снижение', 'нейтрофилы сегментоядерные, %: норма', 'нейтрофилы сегментоядерные, %: повышение', ''), 
('нейтрофилы, %', '', 16, '', '', '', 0.47, 0.72, 'доля', 'нейтрофилы, %: снижение', 'нейтрофилы, %: норма', 'нейтрофилы, %: повышение', 11.004), 
('нейтрофилы, %', '', 4, 6, '', '', 32, 58, '%', 'нейтрофилы, %: снижение', 'нейтрофилы, %: норма', 'нейтрофилы, %: повышение', 7.562), 
('нейтрофилы, %', '', 0, 1, '', '', 16, 45, '%', 'нейтрофилы, %: снижение', 'нейтрофилы, %: норма', 'нейтрофилы, %: повышение', 3.906), 
('нейтрофилы, %', '', 2, 4, '', '', 32, 55, '%', 'нейтрофилы, %: снижение', 'нейтрофилы, %: норма', 'нейтрофилы, %: повышение', 7.539), 
('нейтрофилы, %', '', 10, 16, '', '', 0.43, 0.6, 'доля', 'нейтрофилы, %: снижение', 'нейтрофилы, %: норма', 'нейтрофилы, %: повышение', 10.022), 
('нейтрофилы, %', '', 1, 2, '', '', 0.28, 0.48, 'доля', 'нейтрофилы, %: снижение', 'нейтрофилы, %: норма', 'нейтрофилы, %: повышение', 6.596), 
('нейтрофилы, %', '', 0, 1, '', '', 0.16, 0.45, 'доля', 'нейтрофилы, %: снижение', 'нейтрофилы, %: норма', 'нейтрофилы, %: повышение', 3.906), 
('нейтрофилы, %', '', 8, 10, '', '', 0.41, 0.6, 'доля', 'нейтрофилы, %: снижение', 'нейтрофилы, %: норма', 'нейтрофилы, %: повышение', 9.578), 
('нейтрофилы, %', '', 10, 16, '', '', 43, 60, '%', 'нейтрофилы, %: снижение', 'нейтрофилы, %: норма', 'нейтрофилы, %: повышение', 10.022), 
('нейтрофилы, %', '', 8, 10, '', '', 41, 60, '%', 'нейтрофилы, %: снижение', 'нейтрофилы, %: норма', 'нейтрофилы, %: повышение', 9.578), 
('нейтрофилы, %', '', 1, 2, '', '', 28, 48, '%', 'нейтрофилы, %: снижение', 'нейтрофилы, %: норма', 'нейтрофилы, %: повышение', 6.596), 
('нейтрофилы, %', '', 6, 8, '', '', 38, 60, '%', 'нейтрофилы, %: снижение', 'нейтрофилы, %: норма', 'нейтрофилы, %: повышение', 8.911), 
('нейтрофилы, %', '', 4, 6, '', '', 0.32, 0.58, 'доля', 'нейтрофилы, %: снижение', 'нейтрофилы, %: норма', 'нейтрофилы, %: повышение', 7.562), 
('нейтрофилы, %', '', 2, 4, '', '', 0.32, 0.55, 'доля', 'нейтрофилы, %: снижение', 'нейтрофилы, %: норма', 'нейтрофилы, %: повышение', 7.539), 
('нейтрофилы, %', '', 6, 8, '', '', 0.38, 0.6, 'доля', 'нейтрофилы, %: снижение', 'нейтрофилы, %: норма', 'нейтрофилы, %: повышение', 8.911), 
('нейтрофилы, %', '', 16, '', '', '', 47, 72, '%', 'нейтрофилы, %: снижение', 'нейтрофилы, %: норма', 'нейтрофилы, %: повышение', 11.004), 
('пиридоксаль-5-фосфат в плазме', '', '', '', '', '', 5.0, 50.0, 'нг/мл', 'пиридоксаль-5-фосфат в плазме: снижение', 'пиридоксаль-5-фосфат в плазме: норма', 'пиридоксаль-5-фосфат в плазме: повышение', ''), 
('пиридоксаль-5-фосфат в цельной крови', '', '', '', '', '', 8.6, 27.2, 'нг/мл', 'пиридоксаль-5-фосфат в цельной крови: снижение', 'пиридоксаль-5-фосфат в цельной крови: норма', 'пиридоксаль-5-фосфат в цельной крови: повышение', ''), 
('прокальцитонин', '', '', '', '', '', 0.5, 2.0, 'нг/мл', '', 'прокальцитонин: от 0,5 до 2,0 нг/мл', '', ''), 
('прокальцитонин', '', '', '', '', '', 2.0, 10.0, 'нг/мл', '', 'прокальцитонин: от 2 до 10 нг/мл', 'прокальцитонин: более 10 нг/мл', ''), 
('прокальцитонин', '', '', '', '', '', 0.046, 0.5, 'нг/мл', '', 'прокальцитонин: от 0,046 до 0,5 нг/мл', '', ''), 
('прокальцитонин', '', '', '', '', '', '', 0.046, 'нг/мл', '', 'прокальцитонин: норма', 'прокальцитонин: повышение', ''), 
('распределение тромбоцитов по объему (PDW)', '', '', '', '', '', 10, 20, '%', 'распределение тромбоцитов по объему (PDW): снижение', 'распределение тромбоцитов по объему (PDW): норма', 'распределение тромбоцитов по объему (PDW): повышение', ''), 
('распределение эритроцитов по объему, коэффициент вариации (RDW-CV)', '', 1, '', '', '', 11.6, 14.8, '%', 'распределение эритроцитов по объему, коэффициент вариации (RDW-CV): снижение', 'распределение эритроцитов по объему, коэффициент вариации (RDW-CV): норма', 'распределение эритроцитов по объему, коэффициент вариации (RDW-CV): повышение', ''), 
('распределение эритроцитов по объему, стандартное отклонение (RDW-SD)', '', '', '', '', '', 37.0, 54.0, 'фл', 'распределение эритроцитов по объему, стандартное отклонение (RDW-SD): снижение', 'распределение эритроцитов по объему, стандартное отклонение (RDW-SD): норма', 'распределение эритроцитов по объему, стандартное отклонение (RDW-SD): повышение', ''), 
('СОЭ', '', '', 15, '', '', 2.0, 20.0, 'мм/ч', 'СОЭ: снижение', 'СОЭ: норма', 'СОЭ: повышение', ''), 
('СОЭ', 'пол: мужской', 15, 50, '', '', 2.0, 15.0, 'мм/ч', 'СОЭ: снижение', 'СОЭ: норма', 'СОЭ: повышение', ''), 
('СОЭ', 'пол: мужской', 50, '', '', '', 2.0, 20.0, 'мм/ч', 'СОЭ: снижение', 'СОЭ: норма', 'СОЭ: повышение', ''), 
('СОЭ', 'пол: женский', '', 50, '', '', 2.0, 20.0, 'мм/ч', 'СОЭ: снижение', 'СОЭ: норма', 'СОЭ: повышение', ''), 
('СОЭ', 'пол: женский', 50, '', '', '', 2.0, 30.0, 'мм/ч', 'СОЭ: снижение', 'СОЭ: норма', 'СОЭ: повышение', ''), 
('среднее содержание гемоглобина в эритроците (MCH)', '', 18, 45, '', '', 27.0, 34.0, 'пг', 'среднее содержание гемоглобина в эритроците (MCH): снижение', 'среднее содержание гемоглобина в эритроците (MCH): норма', 'среднее содержание гемоглобина в эритроците (MCH): повышение', ''), 
('средний объем тромбоцита (MPV)', '', '', '', '', '', 9.4, 12.4, 'фл', 'средний объем тромбоцита (MPV): снижение', 'средний объем тромбоцита (MPV): норма', 'средний объем тромбоцита (MPV): повышение', ''), 
('средний объем эритроцита (MCV)', '', 45, 65, '', '', 81, 101, 'фл', 'средний объем эритроцита (MCV): снижение', 'средний объем эритроцита (MCV): норма', 'средний объем эритроцита (MCV): повышение', ''), 
('средний объем эритроцита (MCV)', 'пол: женский', 15, 18, '', '', 78, 98, 'фл', 'средний объем эритроцита (MCV): снижение', 'средний объем эритроцита (MCV): норма', 'средний объем эритроцита (MCV): повышение', ''), 
('средний объем эритроцита (MCV)', 'пол: мужской', 12, 15, '', '', 77, 94, 'фл', 'средний объем эритроцита (MCV): снижение', 'средний объем эритроцита (MCV): норма', 'средний объем эритроцита (MCV): повышение', ''), 
('средний объем эритроцита (MCV)', 'пол: мужской', 18, 45, '', '', 80, 99, 'фл', 'средний объем эритроцита (MCV): снижение', 'средний объем эритроцита (MCV): норма', 'средний объем эритроцита (MCV): повышение', ''), 
('средний объем эритроцита (MCV)', '', 65, '', '', '', 81, 102, 'фл', 'средний объем эритроцита (MCV): снижение', 'средний объем эритроцита (MCV): норма', 'средний объем эритроцита (MCV): повышение', ''), 
('средний объем эритроцита (MCV)', 'пол: женский', 12, 15, '', '', 73, 95, 'фл', 'средний объем эритроцита (MCV): снижение', 'средний объем эритроцита (MCV): норма', 'средний объем эритроцита (MCV): повышение', ''), 
('средний объем эритроцита (MCV)', 'пол: женский', 18, 45, '', '', 81, 100, 'фл', 'средний объем эритроцита (MCV): снижение', 'средний объем эритроцита (MCV): норма', 'средний объем эритроцита (MCV): повышение', ''), 
('средний объем эритроцита (MCV)', 'пол: мужской', 15, 18, '', '', 79, 95, 'фл', 'средний объем эритроцита (MCV): снижение', 'средний объем эритроцита (MCV): норма', 'средний объем эритроцита (MCV): повышение', ''), 
('средняя концентрация гемоглобина в эритроците (MCHC)', '', 3, 12, '', '', 28, 36, 'г/дл', 'средняя концентрация гемоглобина в эритроците (MCHC): снижение', 'средняя концентрация гемоглобина в эритроците (MCHC): норма', 'средняя концентрация гемоглобина в эритроците (MCHC): повышение', 90.222), 
('средняя концентрация гемоглобина в эритроците (MCHC)', '', 3, 12, '', '', 280, 360, 'г/л', 'средняя концентрация гемоглобина в эритроците (MCHC): снижение', 'средняя концентрация гемоглобина в эритроците (MCHC): норма', 'средняя концентрация гемоглобина в эритроците (MCHC): повышение', 90.222), 
('средняя концентрация гемоглобина в эритроците (MCHC)', '', 19, '', '', '', 30, 38, 'г/дл', 'средняя концентрация гемоглобина в эритроците (MCHC): снижение', 'средняя концентрация гемоглобина в эритроците (MCHC): норма', 'средняя концентрация гемоглобина в эритроците (MCHC): повышение', 96.222), 
('средняя концентрация гемоглобина в эритроците (MCHC)', '', 12, 19, '', '', 330, 340, 'г/л', 'средняя концентрация гемоглобина в эритроците (MCHC): снижение', 'средняя концентрация гемоглобина в эритроците (MCHC): норма', 'средняя концентрация гемоглобина в эритроците (MCHC): повышение', 99.778), 
('средняя концентрация гемоглобина в эритроците (MCHC)', '', 0, 1, '', '', 290, 370, 'г/л', 'средняя концентрация гемоглобина в эритроците (MCHC): снижение', 'средняя концентрация гемоглобина в эритроците (MCHC): норма', 'средняя концентрация гемоглобина в эритроците (MCHC): повышение', 93.222), 
('средняя концентрация гемоглобина в эритроците (MCHC)', '', 0, 1, '', '', 29, 37, 'г/дл', 'средняя концентрация гемоглобина в эритроците (MCHC): снижение', 'средняя концентрация гемоглобина в эритроците (MCHC): норма', 'средняя концентрация гемоглобина в эритроците (MCHC): повышение', 93.222), 
('средняя концентрация гемоглобина в эритроците (MCHC)', '', 19, '', '', '', 300, 380, 'г/л', 'средняя концентрация гемоглобина в эритроците (MCHC): снижение', 'средняя концентрация гемоглобина в эритроците (MCHC): норма', 'средняя концентрация гемоглобина в эритроците (MCHC): повышение', 96.222), 
('средняя концентрация гемоглобина в эритроците (MCHC)', '', 1, 3, '', '', 28, 38, 'г/дл', 'средняя концентрация гемоглобина в эритроците (MCHC): снижение', 'средняя концентрация гемоглобина в эритроците (MCHC): норма', 'средняя концентрация гемоглобина в эритроците (MCHC): повышение', 91.778), 
('средняя концентрация гемоглобина в эритроците (MCHC)', '', 12, 19, '', '', 33, 34, 'г/дл', 'средняя концентрация гемоглобина в эритроците (MCHC): снижение', 'средняя концентрация гемоглобина в эритроците (MCHC): норма', 'средняя концентрация гемоглобина в эритроците (MCHC): повышение', 99.778), 
('средняя концентрация гемоглобина в эритроците (MCHC)', '', 1, 3, '', '', 280, 380, 'г/л', 'средняя концентрация гемоглобина в эритроците (MCHC): снижение', 'средняя концентрация гемоглобина в эритроците (MCHC): норма', 'средняя концентрация гемоглобина в эритроците (MCHC): повышение', 91.778), 
('температура тела', '', '', '', '', '', 95.0, 98.6, 'градусы Фаренгейта', 'температура тела: снижение', 'температура тела: норма', 'температура тела: повышение', 49.889), 
('температура тела', '', '', '', '', '', 35.0, 37.0, 'градусы Цельсия', 'температура тела: снижение', 'температура тела: норма', 'температура тела: повышение', 49.889), 
('трансферрин', '', '', '', '', '', 2, 3.6, 'г/л', 'трансферрин: снижение', 'трансферрин: норма', 'трансферрин: повышение', ''), 
('триглицериды', '', '', '', '', '', 0, 1.7, 'ммоль/л', 'триглицериды: снижение', 'триглицериды: норма', 'триглицериды: повышение', ''), 
('тромбокрит', '', '', '', '', '', 0.11, 0.36, '%', 'тромбокрит: снижение', 'тромбокрит: норма', 'тромбокрит: повышение', ''), 
('тромбоциты', '', 15, '', '', '', 150000, 400000, 'клеток/мкл', 'тромбоциты: снижение', 'тромбоциты: норма', 'тромбоциты: повышение', 33644.444), 
('тромбоциты', '', 15, '', '', '', 150.0, 400.0, '10^9/л', 'тромбоциты: снижение', 'тромбоциты: норма', 'тромбоциты: повышение', 33644.444), 
('ферритин', 'пол: мужской', 15, '', '', '', 20, 250, 'мкг/л', 'ферритин: снижение', 'ферритин: норма', 'ферритин: повышение', ''), 
('ферритин', '', 1, 15, '', '', 7, 140, 'мкг/л', 'ферритин: снижение', 'ферритин: норма', 'ферритин: повышение', ''), 
('ферритин', 'пол: женский', 15, '', '', '', 10, 120, 'мкг/л', 'ферритин: снижение', 'ферритин: норма', 'ферритин: повышение', ''), 
('фолиевая кислота', '', '', '', '', '', 3.1, 20.5, 'нг/мл', 'фолиевая кислота: снижение', 'фолиевая кислота: норма', 'фолиевая кислота: повышение', ''), 
('фосфор', '', 12, '', '', '', 0.81, 1.45, 'ммоль/л', 'фосфор: снижение', 'фосфор: норма', 'фосфор: повышение', ''), 
('фосфор', '', 2, 12, '', '', 1.45, 1.78, 'ммоль/л', 'фосфор: снижение', 'фосфор: норма', 'фосфор: повышение', ''), 
('фосфор', '', 0, 2, '', '', 1.45, 2.16, 'ммоль/л', 'фосфор: снижение', 'фосфор: норма', 'фосфор: повышение', ''), 
('фруктозамин', '', '', '', '', '', 170, 285, 'мкмоль/л', 'фруктозамин: снижение', 'фруктозамин: норма', 'фруктозамин: повышение', ''), 
('хлор', '', '', '', '', '', 98, 107, 'ммоль/л', 'хлор: снижение', 'хлор: норма', 'хлор: повышение', ''), 
('холестерин общий', '', '', '', '', '', 0, 5.2, 'ммоль/л', 'холестерин общий: снижение', 'холестерин общий: норма', 'холестерин общий: повышение', ''), 
('холинэстераза', 'пол: женский', '', '', '', '', 5860, 11800, 'Ед/л', 'холинэстераза: снижение', 'холинэстераза: норма', 'холинэстераза: повышение', ''), 
('холинэстераза', 'пол: мужской', '', '', '', '', 5800, 14600, 'Ед/л', 'холинэстераза: снижение', 'холинэстераза: норма', 'холинэстераза: повышение', ''), 
('цветовой показатель', '', '', '', '', '', 0.8, 1.0, '', 'цветовой показатель: снижение', 'цветовой показатель: норма', 'цветовой показатель: повышение', ''), 
('щелочная фосфатаза', 'пол: мужской', 15, 17, '', '', 89, 365, 'Ед/л', 'щелочная фосфатаза: снижение', 'щелочная фосфатаза: норма', 'щелочная фосфатаза: повышение', ''), 
('щелочная фосфатаза', 'пол: мужской', 19, '', '', '', 40, 150, 'Ед/л', 'щелочная фосфатаза: снижение', 'щелочная фосфатаза: норма', 'щелочная фосфатаза: повышение', ''), 
('щелочная фосфатаза', '', 1, 10, '', '', 156, 369, 'Ед/л', 'щелочная фосфатаза: снижение', 'щелочная фосфатаза: норма', 'щелочная фосфатаза: повышение', ''), 
('щелочная фосфатаза', 'пол: мужской', 13, 15, '', '', 127, 517, 'Ед/л', 'щелочная фосфатаза: снижение', 'щелочная фосфатаза: норма', 'щелочная фосфатаза: повышение', ''), 
('щелочная фосфатаза', 'пол: женский', 13, 15, '', '', 62, 280, 'Ед/л', 'щелочная фосфатаза: снижение', 'щелочная фосфатаза: норма', 'щелочная фосфатаза: повышение', ''), 
('щелочная фосфатаза', 'пол: мужской', 17, 19, '', '', 59, 164, 'Ед/л', 'щелочная фосфатаза: снижение', 'щелочная фосфатаза: норма', 'щелочная фосфатаза: повышение', ''), 
('щелочная фосфатаза', 'пол: женский', 15, '', '', '', 40, 150, 'Ед/л', 'щелочная фосфатаза: снижение', 'щелочная фосфатаза: норма', 'щелочная фосфатаза: повышение', ''), 
('щелочная фосфатаза', '', 10, 13, '', '', 141, 460, 'Ед/л', 'щелочная фосфатаза: снижение', 'щелочная фосфатаза: норма', 'щелочная фосфатаза: повышение', ''), 
('эластаза 1 в кале', '', '', '', '', '', 100.0, 200.0, 'мкг/г', 'эластаза 1 в кале: до 100 мкг/г', 'эластаза 1 в кале: от 100 до 200 мкг/г', 'эластаза 1 в кале: более 200 мкг/г', ''), 
('эозинофилы', '', 1, '', '', '', 0.02, 0.5, '10^9/л', 'эозинофилы: снижение', 'эозинофилы: норма', 'эозинофилы: повышение', 4.833), 
('эозинофилы', '', 1, '', '', '', 20.0, 500.0, 'клеток/мкл', 'эозинофилы: снижение', 'эозинофилы: норма', 'эозинофилы: повышение', 4.833), 
('эозинофилы, %', '', 1, '', '', '', 0.01, 0.05, 'доля', 'эозинофилы, %: снижение', 'эозинофилы, %: норма', 'эозинофилы, %: повышение', 0.261), 
('эозинофилы, %', '', 1, '', '', '', 1.0, 5.0, '%', 'эозинофилы, %: снижение', 'эозинофилы, %: норма', 'эозинофилы, %: повышение', 0.261), 
('эритроциты', 'пол: женский', 18, 45, '', '', 3800000.0, 5100000.0, 'клеток/мкл', 'эритроциты: снижение', 'эритроциты: норма', 'эритроциты: повышение', 844448.411), 
('эритроциты', 'пол: женский', 18, 45, '', '', 3.8, 5.1, '10^12/л', 'эритроциты: снижение', 'эритроциты: норма', 'эритроциты: повышение', 844448.411), 
('эритроциты', 'пол: мужской', 18, 45, '', '', 4.3, 5.7, '10^12/л', 'эритроциты: снижение', 'эритроциты: норма', 'эритроциты: повышение', 955559.989), 
('эритроциты', 'пол: мужской', 18, 45, '', '', 4300000.0, 5700000.0, 'клеток/мкл', 'эритроциты: снижение', 'эритроциты: норма', 'эритроциты: повышение', 955559.989), 
}

sm8uR = {
('Rg: язва 12 п.к.', 'язва двенадцатиперстной кишки'), 
('Rg: язва 12 п.к.', 'язва 12 п.к.'), 
('Rg: язва 12 п.к.', 'язва 12 перстной кишки'), 
('Rg: язва 12 п.к.', 'язва двенадцатиперсной кишки'), 
('Rg: язва желудка', 'язва желудка'), 
('гемоглобин: до 7 г/дл', 'гемоглобин менее 7 г/дл'), 
('гемоглобин: до 7 г/дл', 'гемоглобин ниже 7 г/дл'), 
('гемоглобин: до 70 г/л', 'гемоглобин менее 70 г/л'), 
('гемоглобин: до 70 г/л', 'гемоглобин ниже 70 г/л'), 
('гемоглобин: от 7 до 9 г/дл', 'гемоглобин 7-8.9 г/дл'), 
('гемоглобин: от 70 до 90 г/л', 'гемоглобин 70-89 г/л'), 
('жалоба/осмотр: обморок', 'синкопе'), 
('жалоба/осмотр: обморок', 'обморок'), 
('жалоба/осмотр: обширные отеки', 'генерализованные отеки'), 
('жалоба/осмотр: обширные отеки', 'обширные отеки'), 
('жалоба/осмотр: обширные травмы', 'обширные травмы'), 
('жалоба/осмотр: рвота с кровью', 'рвота с кровью'), 
('жалоба/осмотр: судороги', 'судороги'), 
('жалоба/осмотр: черный стул', 'мелена'), 
('жалоба/осмотр: черный стул', 'черный стул'), 
('ЭГДС: язва 12 п.к.', 'язва двенадцатиперстной кишки'), 
('ЭГДС: язва 12 п.к.', 'язва 12 п.к.'), 
('ЭГДС: язва 12 п.к.', 'язва 12 перстной кишки'), 
('ЭГДС: язва 12 п.к.', 'язва двенадцатиперсной кишки'), 
('ЭГДС: язва желудка', 'язва желудка'), 
}

sm_ds = {
('C-реактивный белок: до 1 мг/л', 'низкий сердечно-сосудистый риск', 'ar'), 
('C-реактивный белок: от 1 до 20 мг/л', 'COVID-19', 'ar'), 
('C-реактивный белок: от 1 до 20 мг/л', 'вирусная инфекция', 'ar'), 
('C-реактивный белок: от 1 до 3 мг/л', 'средний сердечно-сосудистый риск', 'ar'), 
('C-реактивный белок: от и более 1 мг/л', 'влияние курения', 'ar'), 
('C-реактивный белок: от и более 1 мг/л', 'ожирение', 'ar'), 
('C-реактивный белок: от и более 1 мг/л', 'сахарный диабет 2 типа', 'ar'), 
('C-реактивный белок: от и более 1 мг/л', 'влияние орального контрацептива', 'ar'), 
('C-реактивный белок: от и более 1 мг/л', 'сепсис', 'ar'), 
('C-реактивный белок: от и более 1 мг/л', 'нарушение гормонального фона', 'ar'), 
('C-реактивный белок: от и более 1 мг/л', 'влияние бета-блокатора', 'ar'), 
('C-реактивный белок: от и более 1 мг/л', 'опухолевый некроз', 'ar'), 
('C-реактивный белок: от и более 1 мг/л', 'злокачественное новообразование', 'ar'), 
('C-реактивный белок: от и более 1 мг/л', 'последствие операции', 'ar'), 
('C-реактивный белок: от и более 1 мг/л', 'влияние кортикостероида', 'ar'), 
('C-реактивный белок: от и более 1 мг/л', 'ожог', 'ar'), 
('C-реактивный белок: от и более 1 мг/л', 'атерогенная дислипидемия', 'ar'), 
('C-реактивный белок: от и более 1 мг/л', 'хронический вялотекущий воспалительный процесс', 'ar'), 
('C-реактивный белок: от и более 1 мг/л', 'беременность', 'ogr'), 
('C-реактивный белок: от и более 1 мг/л', 'атеросклероз сосудов', 'ar'), 
('C-реактивный белок: от и более 1 мг/л', 'влияние гормонального препарата', 'ar'), 
('C-реактивный белок: от и более 1 мг/л', 'инфаркт миокарда', 'ar'), 
('C-реактивный белок: от и более 1 мг/л', 'метастазы', 'ar'), 
('C-реактивный белок: от и более 1 мг/л', 'острое воспаление', 'ar'), 
('C-реактивный белок: от и более 1 мг/л', 'инсульт', 'ar'), 
('C-реактивный белок: от и более 1 мг/л', 'артериальная гипертензия', 'ar'), 
('C-реактивный белок: от и более 1 мг/л', 'остановка сердца', 'ar'), 
('C-реактивный белок: от и более 1 мг/л', 'влияние недавней физической нагрузки', 'ar'), 
('C-реактивный белок: от и более 1 мг/л', 'влияние НПВС', 'ar'), 
('C-реактивный белок: от и более 1 мг/л', 'травма', 'ar'), 
('C-реактивный белок: от и более 1 мг/л', 'влияние статина', 'ar'), 
('C-реактивный белок: от и более 100 мг/л', 'бактериальная инфекция', 'ar'), 
('C-реактивный белок: от и более 3 мг/л', 'высокий сердечно-сосудистый риск', 'ar'), 
('CA 125: повышение', 'серозная карцинома яичника', 'ar'), 
('CA 125: повышение', 'гепатоцеллюлярная карцинома', 'ar'), 
('CA 15-3: повышение', 'рак матки', 'ar'), 
('CA 15-3: повышение', 'рак молочной железы', 'ar'), 
('CA 19-9: повышение', 'колоректальный рак', 'ar'), 
('CA 19-9: повышение', 'холангиокарцинома', 'ar'), 
('CA 19-9: повышение', 'рак поджелудочной железы', 'ar'), 
('CA 19-9: повышение', 'гепатоцеллюлярная карцинома', 'ar'), 
('CA 19-9: повышение', 'рак желчного пузыря', 'ar'), 
('CA 19-9: повышение', 'рак желудка', 'ar'), 
('CA 50: повышение', 'колоректальный рак', 'ar'), 
('CA 50: повышение', 'рак предстательной железы', 'ar'), 
('CA 50: повышение', 'рак поджелудочной железы', 'ar'), 
('CA 50: повышение', 'рак молочной железы', 'ar'), 
('CA 50: повышение', 'рак матки', 'ar'), 
('CA 50: повышение', 'гепатоцеллюлярная карцинома', 'ar'), 
('Cyfra 21-1: повышение', 'рак мочевого пузыря', 'ar'), 
('Cyfra 21-1: повышение', 'немелкоклеточный рак легкого', 'ar'), 
('Cyfra 21-1: повышение', 'плоскоклеточный рак легкого', 'ar'), 
('FFMI: снижение', 'хронический панкреатит', 'ar'), 
('IgA в сыворотке крови: снижение', 'дефицит IgA', 'ar'), 
('Pb71: green', 'Pbo23', 'ar'), 
('Pc23: green', 'Pco23', 'ae'), 
('ph метрия желудка: ахлоргидрия желудочного сока', 'ахлоргидрия желудочного сока', 'ar'), 
('ph метрия желудка: гиперхлоргидрия желудочного сока', 'гиперхлоргидрия желудочного сока', 'ar'), 
('ph метрия желудка: гипохлоргидрия желудочного сока', 'гипохлоргидрия желудочного сока', 'ar'), 
('RDW-CV: повышение и MCV: повышение', 'влияние химиотерапии', 'ar'), 
('RDW-CV: повышение и MCV: повышение', 'железо-дефицитная анемия', 'ar'), 
('RDW-CV: повышение и MCV: повышение', 'сидеробластная анемия', 'ar'), 
('RDW-CV: повышение и MCV: повышение', 'аутоиммунная гемолитическая анемия', 'ar'), 
('RDW-CV: повышение и MCV: снижение', 'железо-дефицитная анемия', 'ar'), 
('RDW-CV: повышение и MCV: снижение', 'гемоглобинопатия', 'ar'), 
('RDW-CV: повышение и MCV: снижение', 'талассемия', 'ar'), 
('RDW-CV: повышение и MCV: снижение', 'витамин B12-дефицитная анемия', 'ar'), 
('RDW-CV: повышение и MCV: снижение', 'серповидно-клеточная анемия', 'ar'), 
('RDW-CV: повышение и MCV: снижение', 'гиповитаминоз B12', 'ar'), 
('RDW-CV: повышение и MCV: снижение', 'нарушение функции печени', 'ar'), 
('RDW-CV: повышение и MCV: снижение', 'гиповитаминоз B9', 'ar'), 
('RDW-SD: повышение и MCV: повышение', 'влияние химиотерапии', 'ar'), 
('RDW-SD: повышение и MCV: повышение', 'железо-дефицитная анемия', 'ar'), 
('RDW-SD: повышение и MCV: повышение', 'аутоиммунная гемолитическая анемия', 'ar'), 
('RDW-SD: повышение и MCV: повышение', 'сидеробластная анемия', 'ar'), 
('RDW-SD: повышение и MCV: снижение', 'гиповитаминоз B12', 'ar'), 
('RDW-SD: повышение и MCV: снижение', 'нарушение функции печени', 'ar'), 
('RDW-SD: повышение и MCV: снижение', 'железо-дефицитная анемия', 'ar'), 
('RDW-SD: повышение и MCV: снижение', 'гемоглобинопатия', 'ar'), 
('RDW-SD: повышение и MCV: снижение', 'талассемия', 'ar'), 
('RDW-SD: повышение и MCV: снижение', 'витамин B12-дефицитная анемия', 'ar'), 
('RDW-SD: повышение и MCV: снижение', 'гиповитаминоз B9', 'ar'), 
('RDW-SD: повышение и MCV: снижение', 'серповидно-клеточная анемия', 'ar'), 
('Rg: 2-х сторонняя интерстициальная инфильтрация в нижних отделах легких', 'COVID-19', 'ar'), 
('Rg: атрофический гастрит', 'атрофический гастрит', 'ar'), 
('Rg: атрофический гастрит антрального отдела желудка', 'атрофический гастрит антрального отдела желудка', 'ar'), 
('Rg: атрофический гастрит дна желудка', 'атрофический гастрит дна желудка', 'ar'), 
('Rg: атрофический гастрит тела желудка', 'атрофический гастрит тела желудка', 'ar'), 
('Rg: выступание костей задней черепной ямки', 'выступание костей задней черепной ямки', 'ar'), 
('Rg: гастрит', 'гастрит', 'ar'), 
('Rg: гастроэзофагеальный рефлюкс', 'гастроэзофагеальный рефлюкс', 'ar'), 
('Rg: гигантский гипертрофический гастрит', 'гигантский гипертрофический гастрит', 'ar'), 
('Rg: злокачественное новообразование желудка', 'злокачественное новообразование желудка', 'ar'), 
('Rg: истончение костей задней черепной ямки', 'синдром Денди-Уокера', 'ar'), 
('Rg: пищеводные полукольца', 'пищеводные полукольца', 'ar'), 
('Rg: рак желудка', 'рак желудка', 'ar'), 
('Rg: релаксация диафрагмы', 'релаксация диафрагмы', 'ar'), 
('Rg: язва 12 п.к.', 'язва 12 п.к.', 'ar'), 
('Rg: язва желудка', 'язва желудка', 'ar'), 
('Rg: язва пищевода', 'язва пищевода', 'ar'), 
('SCCA: повышение', 'опухоль головы и шеи', 'ar'), 
('Tc: green', 'Hya614273', 'ae'), 
('Tc: green', 'Hya215721', 'ae'), 
('Tc: green', 'Hya401643', 'ae'), 
('Tc: green', 'Hya0278751', 'ae'), 
('авитаминоз', 'атрофический аутоиммунный гастрит', 'ar'), 
('агенезия мозолистого тела', 'синдром Денди-Уокера', 'ar'), 
('агенезия червя мозжечка', 'синдром Денди-Уокера', 'ar'), 
('агрегация тромбоцитов: повышение', 'гиповитаминоз B6', 'ar'), 
('аллергия на еду', 'эозинофильный гастрит', 'ar'), 
('аллергия на еду: апельсин', 'аллергия на еду', 'ar'), 
('аллергия на еду: кунжут', 'аллергия на еду', 'ar'), 
('аллергия на еду: неопределенный продукт', 'аллергия на еду', 'ar'), 
('аллергия на ЛС: аспирин', 'аспириновая триада', 'ogr2'), 
('аллергия на ЛС: лекарственные средства пиразолонового ряда', 'аспириновая триада', 'ogr2'), 
('АЛТ: повышение', 'хронический панкреатит', 'ogr'), 
('АЛТ: повышение', 'хронический холецистит', 'ar'), 
('альбумин: снижение', 'COVID-19', 'ar'), 
('альбумин: снижение', 'воспалительный процесс', 'ar'), 
('альбумин: снижение', 'нефротический синдром', 'ar'), 
('альфа-глобулин: повышение', 'острое воспаление', 'ar'), 
('амилаза в моче: повышение', 'хронический панкреатит', 'ogr'), 
('амилаза общая: повышение', 'хронический панкреатит', 'ogr'), 
('амилаза панкреатическая: повышение', 'хронический панкреатит', 'ogr'), 
('анализ мочи по Зимницкому: полиурия', 'полиурия', 'ar'), 
('анацидный гастрит', 'лямблиоз', 'ar'), 
('анемия', 'гиперхромная анемия', 'or'), 
('анемия', 'железо-дефицитная анемия', 'or'), 
('анемия', 'нормохромная анемия', 'or'), 
('анемия', 'фолиево-дефицитная анемия', 'ar'), 
('анемия', 'гипохромная анемия', 'or'), 
('анемия', 'влияние аспирина', 'ar'), 
('анемия', 'атрофический аутоиммунный гастрит', 'ar'), 
('аномалия развития пальцев', 'синдром Денди-Уокера', 'ar'), 
('аномалия развития сердца', 'синдром Денди-Уокера', 'ar'), 
('анорексия', 'влияние аспирина', 'ar'), 
('анорексия', 'атрофический гастрит', 'ar'), 
('антитела к париетальным клеткам желудка: отрицательно', 'атрофический мультифокальный гастрит', 'ar'), 
('антитела к париетальным клеткам желудка: положительно', 'атрофический мультифокальный гастрит', 'ae'), 
('антитела к париетальным клеткам желудка: положительно', 'атрофический аутоиммунный гастрит', 'ar'), 
('артериальная гипертензия', 'гипертензивная болезнь сердца', 'ar'), 
('артериальная гипертензия', 'болезнь Паркинсона', 'ad'), 
('артрит', 'гиповитаминоз B6', 'ar'), 
('артрит', 'реактивный артрит', 'ar'), 
('асептический менингит', 'влияние аспирина', 'ar'), 
('аскаридоз', 'хронический холецистит', 'ar'), 
('аспириновая триада', 'влияние аспирина', 'ar'), 
('АСТ: повышение', 'хронический панкреатит', 'ogr'), 
('АСТ: повышение', 'хронический холецистит', 'ar'), 
('астеновегетативный синдром', 'атрофический мультифокальный гастрит', 'ar'), 
('атопический дерматит', 'лямблиоз', 'ar'), 
('атрофический аутоиммунный гастрит', 'аутоиммунный тиреоидит', 'ar'), 
('атрофический аутоиммунный гастрит', 'сахарный диабет 1 типа', 'ar'), 
('атрофический аутоиммунный гастрит', 'фуникулярный миелоз', 'ar'), 
('атрофический аутоиммунный гастрит', 'витамин B12-дефицитная анемия', 'ar'), 
('атрофический гастрит', 'сидеропенический синдром', 'ar'), 
('атрофический гастрит антрального отдела желудка', 'атрофический мультифокальный гастрит', 'or'), 
('атрофический гастрит дна желудка', 'атрофический аутоиммунный гастрит', 'ar'), 
('атрофический энтерит', 'сидеропенический синдром', 'ar'), 
('атрофия поджелудочной железы', 'хронический панкреатит', 'ogr'), 
('атрофия щитовидной железы (приобретенная)', 'гипотиреоз', 'ar'), 
('аутоиммунный тиреоидит', 'атрофический аутоиммунный гастрит', 'ar'), 
('АФП: от 5.8 до 100 МЕ/мл', 'цирроз печени', 'ar'), 
('АФП: от 5.8 до 100 МЕ/мл', 'регенеративные процессы в печени', 'ar'), 
('АФП: от 5.8 до 100 МЕ/мл', 'алкогольное поражение печени', 'ar'), 
('АФП: от 5.8 до 100 МЕ/мл', 'хронический гепатит', 'ar'), 
('АФП: от 5.8 до 500 МЕ/мл', 'метастазы в печень', 'ar'), 
('АФП: повышение', 'травма печени', 'ar'), 
('АФП: повышение', 'злокачественное новообразование толстой кишки', 'ar'), 
('АФП: повышение', 'злокачественное новообразование легкого', 'ar'), 
('АФП: повышение', 'врожденная тирозинемия', 'ar'), 
('АФП: повышение', 'тератобластома яичек', 'ar'), 
('АФП: повышение', 'герминогенная опухоль яичника', 'ar'), 
('АФП: повышение', 'злокачественное новообразование поджелудочной железы', 'ar'), 
('АФП: повышение', 'злокачественное новообразование желудка', 'ar'), 
('АФП: повышение', 'злокачественное новообразование почки', 'ar'), 
('АФП: повышение', 'синдром Вискотта-Олдрича', 'ar'), 
('АФП: повышение', 'колоректальный рак', 'ar'), 
('АФП: повышение', 'злокачественное новообразование молочной железы', 'ar'), 
('АФП: повышение', 'первичная гепатоцеллюлярная карцинома', 'ar'), 
('АФП: повышение', 'злокачественное новообразование бронхов', 'ar'), 
('АФП: повышение', 'тератобластома яичников', 'ar'), 
('АФП: повышение', 'рак яичка', 'ar'), 
('АФП: повышение', 'гепатоцеллюлярная карцинома', 'ar'), 
('АФП: повышение', 'атаксия-телеангиэктазия', 'ar'), 
('АФП: повышение', 'последствие операции на печени', 'ar'), 
('АФП: повышение при беременности', 'аномалия мочевыводящих путей плода', 'ar'), 
('АФП: повышение при беременности', 'некроз печени плода вследствие вирусной инфекции у плода', 'ar'), 
('АФП: повышение при беременности', 'несращение передней брюшной стенки плода', 'ar'), 
('АФП: повышение при беременности', 'синдром Меккеля у плода', 'ar'), 
('АФП: повышение при беременности', 'многоплодная беременность', 'ar'), 
('АФП: повышение при беременности', 'аномалия почек плода', 'ar'), 
('АФП: повышение при беременности', 'атрезия 12-типерстной кишки у плода', 'ar'), 
('АФП: повышение при беременности', 'расщелина позвоночника плода', 'ar'), 
('АФП: повышение при беременности', 'крупный плод', 'ar'), 
('АФП: повышение при беременности', 'spina bifida у плода', 'ar'), 
('АФП: повышение при беременности', 'синдром Шерешевского-Тернера плода', 'ar'), 
('АФП: повышение при беременности', 'атрезия пищевода у плода', 'ar'), 
('АФП: повышение при беременности', 'анэнцефалия у плода', 'ar'), 
('АФП: повышение при беременности', 'пупочная грыжа у плода', 'ar'), 
('АФП: повышение при беременности', 'порок развития нервной трубки плода', 'ar'), 
('АФП: снижение при беременности', 'гибель плода', 'ar'), 
('АФП: снижение при беременности', 'трисомия по 18 хромосоме у плода', 'ar'), 
('АФП: снижение при беременности', 'задержка развития плода', 'ar'), 
('АФП: снижение при беременности', 'ложная беременность', 'ar'), 
('АФП: снижение при беременности', 'завышенный срок беременности', 'ar'), 
('АФП: снижение при беременности', 'инсулин-зависимый диабет беременной', 'ar'), 
('АФП: снижение при беременности', 'синдром Дауна у плода', 'ar'), 
('АФП: снижение при беременности', 'пузырный занос', 'ar'), 
('АФП: снижение при беременности', 'угроза самопроизвольного выкидыша', 'ar'), 
('АФП: снижение при беременности', 'ожирение беременной', 'ar'), 
('АФП: снижение при беременности', 'самопроизвольный выкидыш', 'ar'), 
('АФП: снижение при беременности', 'синдром Патау', 'ar'), 
('ахлоргидрия желудочного сока', 'атрофический гастрит', 'ogr'), 
('ахлоргидрия желудочного сока', 'лямблиоз', 'ar'), 
('ахлоргидрия желудочного сока', 'дисбиоз тонкого кишечника', 'ar'), 
('ахлоргидрия желудочного сока', 'псевдомембранозный колит', 'ar'), 
('ацидоз', 'воспалительный процесс', 'ar'), 
('базофилы, %: повышение', 'беременность', 'ogr'), 
('базофилы, %: повышение', 'первая фаза менструального цикла (7-14 день)', 'ogr'), 
('базофилы, %: повышение', 'овуляция', 'ogr'), 
('базофилы: повышение', 'аллергия', 'ar'), 
('базофилы: повышение', 'лейкоз', 'ar'), 
('базофилы: повышение', 'хронический миелоидный лейкоз', 'ar'), 
('базофилы: повышение', 'влияние орального контрацептива', 'ar'), 
('базофилы: повышение', 'глисты', 'ar'), 
('базофилы: повышение', 'прорезывание молочных зубов', 'ogr'), 
('базофилы: повышение', 'хроническое воспаление', 'ar'), 
('базофилы: повышение', 'аутоиммунный тиреоидит', 'ar'), 
('базофилы: повышение', 'псориаз', 'ar'), 
('базофилы: повышение', 'оспа', 'ar'), 
('базофилы: повышение', 'увеличение уровня эстрогена', 'ar'), 
('базофилы: повышение', 'туберкулез', 'ar'), 
('базофилы: повышение', 'воспалительные заболевания кишечника', 'ar'), 
('базофилы: повышение', 'тучноклеточный лейкоз', 'ar'), 
('базофилы: повышение', 'системная красная волчанка', 'ar'), 
('базофилы: повышение', 'аутоиммунное заболевание', 'ar'), 
('базофилы: повышение', 'ревматоидный артрит', 'ar'), 
('базофилы: повышение', 'гемолитическая анемия', 'ar'), 
('базофилы: повышение', 'полицитемия истинная', 'ar'), 
('базофилы: повышение', 'гемофилия', 'ar'), 
('базофилы: повышение', 'грипп', 'ar'), 
('базофилы: снижение', 'острая фаза инфекции', 'ar'), 
('базофилы: снижение', 'влияние кортикостероида', 'ar'), 
('базофилы: снижение', 'тиреотоксикоз', 'ar'), 
('белок общий в моче: от и более 3,5 г/сутки', 'нефротический синдром', 'ar'), 
('белок общий в моче: от и более 50 мг*кг/сут', 'нефротический синдром', 'ar'), 
('белок общий в моче: повышение', 'COVID-19', 'ar'), 
('белок общий: снижение', 'хронический холецистит', 'ar'), 
('белок общий: снижение', 'нефротический синдром', 'ar'), 
('белок общий: снижение', 'COVID-19', 'ar'), 
('белок общий: снижение', 'дефицит белка', 'ar'), 
('бета-глобулин: повышение', 'острое воспаление', 'ar'), 
('билирубин общий: повышение', 'хронический холецистит', 'ar'), 
('билирубин прямой: повышение', 'хронический холецистит', 'ar'), 
('болезни щитовидной железы, связанные с йодной недостаточностью и сходные состояния', 'гипотиреоз', 'ar'), 
('болезнь Крона', 'гранулематозный гастрит', 'ar'), 
('болезнь Крона', 'воспалительные заболевания кишечника', 'ar'), 
('болезнь Крона', 'неспецифический колит', 'ar'), 
('болезнь Паркинсона', 'влияние курения', 'ad'), 
('бронхиальная астма', 'аспириновая триада', 'or'), 
('бронхоспазм', 'влияние аспирина', 'ar'), 
('БУТ: наличие Hp на слизистой желудка', 'наличие Hp на слизистой желудка', 'ar'), 
('БУТ: наличие Hp на слизистой желудка', 'атрофический мультифокальный гастрит', 'ar'), 
('БУТ: наличие Hp на слизистой желудка в большом количестве', 'наличие Hp на слизистой желудка в большом количестве', 'ar'), 
('БУТ: наличие Hp на слизистой желудка в небольшом количестве', 'наличие Hp на слизистой желудка в небольшом количестве', 'ar'), 
('БУТ: наличие Hp на слизистой желудка в среднем количестве', 'наличие Hp на слизистой желудка в среднем количестве', 'ar'), 
('БУТ: отсутствие Hp на слизистой желудка', 'отсутствие Hp на слизистой желудка', 'ar'), 
('в анамнезе: кровотечение', 'кровотечение', 'ar'), 
('в анамнезе: кровотечение', 'железо-дефицитная анемия', 'ar'), 
('в анамнезе: неатрофический гастрит', 'неатрофический гастрит', 'ar'), 
('в анамнезе: неатрофический гастрит', 'атрофический мультифокальный гастрит', 'ar'), 
('в анамнезе: острый панкреатит', 'острый панкреатит', 'ar'), 
('в анамнезе: хроническая инфекция', 'сидеропенический синдром', 'ar'), 
('в анамнезе: хроническая инфекция', 'хроническая инфекция', 'ar'), 
('в анамнезе: язвенный колит', 'язвенный колит', 'ar'), 
('витамин B12-дефицитная анемия', 'хронический панкреатит', 'ar'), 
('витамин B12-дефицитная анемия', 'атрофический мультифокальный гастрит', 'ad'), 
('витамин B12-дефицитная анемия', 'атрофический аутоиммунный гастрит', 'ar'), 
('витамин B12-дефицитная анемия', 'гиповитаминоз B12', 'ar'), 
('влияние химических раздражителей на желудок', 'химический гастрит', 'ogr'), 
('внешнесекреторная недостаточность поджелудочной железы', 'хронический панкреатит', 'ar'), 
('внутримозговое кровоизлияние', 'инсульт', 'ar'), 
('внутримозговое кровоизлияние в мозжечок', 'внутримозговое кровоизлияние', 'ar'), 
('внутримозговое кровоизлияние в полушарие кортикальное', 'внутримозговое кровоизлияние', 'ar'), 
('внутримозговое кровоизлияние в полушарие неуточненное', 'внутримозговое кровоизлияние', 'ar'), 
('внутримозговое кровоизлияние в полушарие субкортикальное', 'внутримозговое кровоизлияние', 'ar'), 
('внутримозговое кровоизлияние в ствол мозга', 'внутримозговое кровоизлияние', 'ar'), 
('внутримозговое кровоизлияние внутрижелудочковое', 'внутримозговое кровоизлияние', 'ar'), 
('внутримозговое кровоизлияние множественной локализации', 'внутримозговое кровоизлияние', 'ar'), 
('возраст до 3 лет', 'прорезывание молочных зубов', 'or'), 
('возраст от 10 до 50 лет', 'менструация', 'or'), 
('возраст от 10 до 50 лет', 'беременность', 'or'), 
('возраст от 10 до 50 лет', 'обильная менструация', 'or'), 
('возраст от 10 до 50 лет', 'первая фаза менструального цикла (7-14 день)', 'or'), 
('возраст от 10 до 50 лет', 'овуляция', 'or'), 
('возраст от 10 до 50 лет', 'беременность (третий триместр)', 'or'), 
('воспалительные заболевания тонкого кишечника', 'железо-дефицитная анемия', 'ar'), 
('воспалительный процесс', 'вирусная инфекция', 'ar'), 
('вредная привычка: гиподинамия', 'гиподинамия', 'ar'), 
('вредная привычка: злоупотребление алкоголем', 'хронический панкреатит', 'ar'), 
('вредная привычка: курение', 'влияние курения', 'or'), 
('вредная привычка: курение', 'хронический панкреатит', 'ar'), 
('вредная привычка: поведение нарушающее циркадные ритмы', 'поведение нарушающее циркадные ритмы', 'ar'), 
('врожденный гипотиреоз без зоба', 'гипотиреоз', 'ar'), 
('врожденный гипотиреоз с диффузным зобом', 'гипотиреоз', 'ar'), 
('врожденный нефротический синдром', 'нефротический синдром', 'ar'), 
('вторичная гипертензия', 'артериальная гипертензия', 'ar'), 
('вторичная непереносимость лактозы', 'лямблиоз', 'ar'), 
('выступание костей задней черепной ямки', 'синдром Денди-Уокера', 'ar'), 
('гамма-глобулин: повышение', 'хроническое воспаление', 'ar'), 
('гамма-глобулин: снижение', 'лямблиоз', 'ar'), 
('гамма-ГТП: повышение', 'панкреатит', 'ar'), 
('гамма-ГТП: повышение', 'рак головки поджелудочной железы', 'ar'), 
('гамма-ГТП: повышение', 'влияние алкоголя', 'ar'), 
('гамма-ГТП: повышение', 'сердечная недостаточность', 'ar'), 
('гамма-ГТП: повышение', 'рак предстательной железы', 'ar'), 
('гамма-ГТП: повышение', 'влияние парацетамола', 'ar'), 
('гамма-ГТП: повышение', 'влияние тестостерона', 'ar'), 
('гамма-ГТП: повышение', 'влияние антибиотика', 'ar'), 
('гамма-ГТП: повышение', 'тиреотоксикоз', 'ar'), 
('гамма-ГТП: повышение', 'цирроз печени', 'ar'), 
('гамма-ГТП: повышение', 'влияние H-2 гистаминоблокатора', 'ar'), 
('гамма-ГТП: повышение', 'метастазы в печень', 'ar'), 
('гамма-ГТП: повышение', 'первичный склерозирующий холангит', 'ar'), 
('гамма-ГТП: повышение', 'острый гепатит', 'ar'), 
('гамма-ГТП: повышение', 'инфаркт миокарда', 'ar'), 
('гамма-ГТП: повышение', 'ожирение', 'ar'), 
('гамма-ГТП: повышение', 'влияние орального контрацептива', 'ar'), 
('гамма-ГТП: повышение', 'гепатоцеллюлярная карцинома', 'ar'), 
('гамма-ГТП: повышение', 'влияние барбитурата', 'ar'), 
('гамма-ГТП: повышение', 'влияние фенобарбитала', 'ar'), 
('гамма-ГТП: повышение', 'влияние статина', 'ar'), 
('гамма-ГТП: повышение', 'первичный билиарный цирроз', 'ar'), 
('гамма-ГТП: повышение', 'холестаз', 'ar'), 
('гамма-ГТП: повышение', 'системная красная волчанка', 'ar'), 
('гамма-ГТП: повышение', 'влияние цефалоспорина', 'ar'), 
('гамма-ГТП: повышение', 'влияние антидепрессанта', 'ar'), 
('гамма-ГТП: повышение', 'влияние аспирина', 'ar'), 
('гамма-ГТП: повышение', 'сахарный диабет', 'ar'), 
('гамма-ГТП: повышение', 'хронический гепатит', 'ar'), 
('гамма-ГТП: повышение', 'влияние эстрогена', 'ar'), 
('гамма-ГТП: повышение', 'влияние противогрибкового препарата', 'ar'), 
('гамма-ГТП: повышение', 'инфекционный мононуклеоз', 'ar'), 
('гамма-ГТП: снижение', 'гипотиреоз', 'ar'), 
('гамма-ГТП: снижение', 'влияние аскорбиновой кислоты', 'ar'), 
('гангрена нижних конечностей', 'COVID-19', 'ar'), 
('гастрин 17 (базальный): от 10 до 20 пмоль/л', 'гипохлоргидрия желудочного сока', 's'), 
('гастрин 17 (базальный): от 7 до 10 пмоль/л', 'гипохлоргидрия желудочного сока легкая', 'ogr'), 
('гастрин 17 (базальный): от и более 20 пмоль/л', 'ахлоргидрия желудочного сока', 'ar'), 
('гастрин 17 (базальный): от и более 5 пмоль/л', 'пищевод Барретта', 'ad'), 
('гастрин 17 (базальный): повышение', 'влияние ИПП', 'ar'), 
('гастрин 17 (базальный): повышение', 'стресс', 'ar'), 
('гастрин 17 (базальный): повышение', 'влияние НПВС', 'ar'), 
('гастрин 17 (базальный): повышение', 'нарушение секреции гастрина', 'ar'), 
('гастрин 17 (базальный): повышение', 'пернициозное анемическое состояние', 'ar'), 
('гастрин 17 (базальный): повышение', 'влияние резкого прекращение приема ИПП', 'ar'), 
('гастрин 17 (базальный): повышение', 'глюкокортикоидные гормоны: повышение', 'ar'), 
('гастрин 17 (базальный): повышение', 'наличие Hp на слизистой желудка', 'ar'), 
('гастрин 17 (базальный): повышение', 'рак желудка', 'ar'), 
('гастрин 17 (базальный): повышение', 'неатрофический гастрит', 'ar'), 
('гастрин 17 (базальный): повышение', 'хроническое заболевание почек', 'ar'), 
('гастрин 17 (базальный): снижение', 'последствие гастрэктомии', 'ar'), 
('гастрин 17 (базальный): снижение', 'гипотиреоз', 'ar'), 
('гастрин 17 (базальный): снижение', 'гиперхлоргидрия желудочного сока', 'ar'), 
('гастрин 17 (базальный): снижение', 'атрофический гастрит', 's'), 
('гастрин-17 стимулированный: повышение', 'влияние НПВС', 'ar'), 
('гастрин-17 стимулированный: повышение', 'пернициозное анемическое состояние', 'ar'), 
('гастрин-17 стимулированный: повышение', 'влияние резкого прекращение приема ИПП', 'ar'), 
('гастрин-17 стимулированный: повышение', 'нарушение секреции гастрина', 'ar'), 
('гастрин-17 стимулированный: повышение', 'глюкокортикоидные гормоны: повышение', 'ar'), 
('гастрин-17 стимулированный: повышение', 'рак желудка', 'ar'), 
('гастрин-17 стимулированный: повышение', 'неатрофический гастрит', 'ar'), 
('гастрин-17 стимулированный: повышение', 'хроническое заболевание почек', 'ar'), 
('гастрин-17 стимулированный: повышение', 'стресс', 'ar'), 
('гастрин-17 стимулированный: снижение', 'атрофический гастрит', 's'), 
('гастрин-17 стимулированный: снижение', 'гипотиреоз', 'ar'), 
('гастрин-17 стимулированный: снижение', 'гиперхлоргидрия желудочного сока', 'ar'), 
('гастрин-17 стимулированный: снижение', 'последствие гастрэктомии', 'ar'), 
('гастроэзофагеальный рефлюкс', 'гастроэзофагеальная рефлюксная болезнь', 'ar'), 
('гастроэзофагеальный рефлюкс', 'сердечная аритмия', 'ar'), 
('гематокрит: повышение', 'дыхательная недостаточность', 'ar'), 
('гематокрит: повышение', 'сердечная недостаточность', 'ar'), 
('гематокрит: повышение', 'обезвоживание', 'ar'), 
('гематокрит: повышение', 'полицитемия истинная', 'ar'), 
('гематокрит: снижение', 'влияние сниженных эритроцитов', 'ar'), 
('гематокрит: снижение', 'анемия', 'ar'), 
('гематокрит: снижение', 'почечная недостаточность', 'ar'), 
('гематокрит: снижение', 'беременность (третий триместр)', 'ogr'), 
('гемоглобин в моче: наличие', 'COVID-19', 'ar'), 
('гемоглобин: до 7 г/дл', 'анемия тяжелой степени', 'ogr'), 
('гемоглобин: до 70 г/л', 'анемия тяжелой степени', 'ogr'), 
('гемоглобин: от 7 до 9 г/дл', 'анемия средней степени', 'ogr'), 
('гемоглобин: от 70 до 90 г/л', 'анемия средней степени', 'ogr'), 
('гемоглобин: от 9 г/дл и ниже нормы', 'анемия легкой степени', 'ogr'), 
('гемоглобин: от 90 г/л и ниже нормы', 'анемия легкой степени', 'ogr'), 
('гемоглобин: повышение', 'болезни крови', 'ar'), 
('гемоглобин: повышение', 'дыхательная недостаточность', 'ar'), 
('гемоглобин: повышение', 'болезни мочевыделительной системы', 'ar'), 
('гемоглобин: повышение', 'обезвоживание', 'ar'), 
('гемоглобин: повышение', 'сердечная недостаточность', 'ar'), 
('гемоглобин: снижение', 'заболевания крови', 'ar'), 
('гемоглобин: снижение', 'анемия', 'ogr'), 
('гемоглобин: снижение', 'постгеморрагическая анемия', 'ar'), 
('гемоглобин: снижение', 'влияние недостаточного употребления железа', 'ar'), 
('гемоглобин: снижение', 'сидеропенический синдром', 'ar'), 
('гемоглобин: снижение', 'дефицит белка', 'ar'), 
('гемоглобин: снижение', 'влияние недостаточного употребления витаминов', 'ar'), 
('гемоконцентрация', 'обезвоживание', 'ar'), 
('геморрагический синдром', 'влияние аспирина', 'ar'), 
('гепатомегалия', 'атрофический аутоиммунный гастрит', 'ar'), 
('гепатомегалия', 'фуникулярный миелоз', 'ar'), 
('гипервитаминоз B6', 'гипофосфатазия', 'ar'), 
('гипервитаминоз B6', 'увеличивается риск развития рака легких у мужчин на 30-40%', 'ar'), 
('гиперлипидемия', 'нефротический синдром', 'ar'), 
('гиперпаратиреоз', 'атрофический аутоиммунный гастрит', 'ar'), 
('гипертензивная болезнь с преимущественным поражением сердца и почек', 'артериальная гипертензия', 'ar'), 
('гипертензионная гидроцефалия', 'синдром Денди-Уокера', 'ar'), 
('гиперхромная анемия', 'атрофический аутоиммунный гастрит', 'ar'), 
('гиповитаминоз B12', 'СИБР', 'ar'), 
('гиповитаминоз B6', 'нарушение обмена жиров', 'ar'), 
('гиповитаминоз B6', 'влияние фтивазида', 'ar'), 
('гиповитаминоз B6', 'нарушение усвоения нервными клетками глюкозы', 'ar'), 
('гиповитаминоз B6', 'влияние антибиотика', 'ar'), 
('гиповитаминоз B6', 'влияние тубазида', 'ar'), 
('гиповитаминоз B6', 'нарушение обмена белков', 'ar'), 
('гиповитаминоз B6', 'влияние курения', 'ar'), 
('гиповитаминоз B6', 'влияние орального контрацептива', 'ar'), 
('гиповитаминоз B6', 'нарушение производства нейромедиаторов', 'ar'), 
('гиповитаминоз B6', 'нарушение функции печени', 'ar'), 
('гиповитаминоз B6', 'нарушение кровообращения', 'ar'), 
('гиповитаминоз D', 'СИБР', 'ar'), 
('гиповитаминоз D', 'хронический панкреатит', 'ar'), 
('гиповитаминоз E', 'хронический панкреатит', 'ar'), 
('гиповитаминоз E', 'СИБР', 'ar'), 
('гиповитаминоз K', 'хронический панкреатит', 'ar'), 
('гиповитаминоз K', 'СИБР', 'ar'), 
('гиповитаминоз А', 'СИБР', 'ar'), 
('гиповитаминоз А', 'хронический панкреатит', 'ar'), 
('гипогликемический синдром', 'хронический панкреатит', 'ar'), 
('гиподинамия', 'хронический панкреатит', 'ar'), 
('гипопаратиреоз', 'атрофический аутоиммунный гастрит', 'ar'), 
('гипотиреоз', 'атрофический аутоиммунный гастрит', 'ar'), 
('гипотиреоз', 'постинфекционный гипотиреоз', 'ar'), 
('гипотиреоз неуточненный', 'гипотиреоз', 'ar'), 
('гипотония', 'хронический панкреатит', 'ar'), 
('гипотония кишечника', 'хронический панкреатит', 'ar'), 
('гипохлоргидрия желудочного сока', 'атрофический гастрит', 'ogr'), 
('гипохлоргидрия желудочного сока', 'дисбиоз тонкого кишечника', 'ar'), 
('гипохлоргидрия желудочного сока', 'лямблиоз', 'ar'), 
('гистология: атрофический гастрит', 'атрофический гастрит', 'ar'), 
('гистология: атрофический гастрит антрального отдела желудка', 'атрофический гастрит антрального отдела желудка', 'ar'), 
('гистология: атрофический гастрит дна желудка', 'атрофический гастрит дна желудка', 'ar'), 
('гистология: атрофический гастрит тела желудка', 'атрофический гастрит тела желудка', 'ar'), 
('гистология: атрофия ворсинок тонкой кишки', 'лямблиоз', 'ar'), 
('гистология: гастрит', 'гастрит', 'ar'), 
('гистология: гигантский гипертрофический гастрит', 'гигантский гипертрофический гастрит', 'ar'), 
('гистология: гранулематозный гастрит', 'гранулематозный гастрит', 'ar'), 
('гистология: кишечная метаплазия слизистой оболочки желудка', 'атрофический мультифокальный гастрит', 'ar'), 
('гистология: лимфоцитарный гастрит', 'лимфоцитарный гастрит', 'ar'), 
('гистология: наличие Hp на слизистой желудка', 'неатрофический гастрит', 'ar'), 
('гистология: наличие Hp на слизистой желудка', 'атрофический мультифокальный гастрит', 'ar'), 
('гистология: наличие Hp на слизистой желудка', 'наличие Hp на слизистой желудка', 'ar'), 
('гистология: наличие Hp на слизистой желудка в большом количестве', 'наличие Hp на слизистой желудка в большом количестве', 'ar'), 
('гистология: наличие Hp на слизистой желудка в небольшом количестве', 'наличие Hp на слизистой желудка в небольшом количестве', 'ar'), 
('гистология: наличие Hp на слизистой желудка в среднем количестве', 'наличие Hp на слизистой желудка в среднем количестве', 'ar'), 
('гистология: отсутствие Hp на слизистой желудка', 'отсутствие Hp на слизистой желудка', 'ar'), 
('гистология: пищевод Барретта', 'пищевод Барретта', 'ar'), 
('гистология: рак желудка', 'рак желудка', 'ar'), 
('гистология: трофозоиты на поверхности слизистой оболочки тонкой кишки', 'лямблиоз', 's'), 
('гистология: эозинофильный гастрит', 'аллергия на еду', 'ar'), 
('гистология: эозинофильный гастрит', 'аллергия', 'ar'), 
('гистология: эозинофильный гастрит', 'эозинофильный гастрит', 'ar'), 
('гликированный гемоглобин: от 6 до 6,5%', 'преддиабет', 'ar'), 
('гликированный гемоглобин: от 6,5%', 'диабет', 'ar'), 
('глоссит', 'атрофический аутоиммунный гастрит', 'ar'), 
('глоссит', 'фуникулярный миелоз', 'ar'), 
('глоссит', 'сидеропенический синдром', 'ar'), 
('глюкоза: повышение', 'COVID-19', 'ar'), 
('глюкокортикоидные гормоны в крови: повышение', 'глюкокортикоидные гормоны: повышение', 'ar'), 
('гранулематоз Вегенера', 'гранулематозный гастрит', 'ar'), 
('гранулематозный гастрит', 'болезнь Крона', 'ar'), 
('Д-димер: повышение', 'COVID-19', 'ar'), 
('деменция', 'атрофический аутоиммунный гастрит', 'ar'), 
('депрессия', 'нарушение циркадных ритмов организма', 'ar'), 
('депрессия', 'гиповитаминоз B6', 'ar'), 
('дефицит IgA', 'лямблиоз', 'ar'), 
('дефицит белка', 'гипервитаминоз B6', 'ar'), 
('дефицит белка', 'повышенная потеря белка организмом', 'ar'), 
('дефицит белка', 'недостаточное поступление белка', 'ar'), 
('дефицит белка', 'усиленное потребление белка организмом', 'ar'), 
('дефицит глюкагона', 'хронический панкреатит', 'ar'), 
('дефицит инсулина', 'хронический панкреатит', 'ar'), 
('дефицит микроэлементов', 'хронический панкреатит', 'ar'), 
('диабет', 'обезвоживание', 'ar'), 
('дисбиоз тонкого кишечника', 'гипохлоргидрия желудочного сока', 'ar'), 
('дисбиоз тонкого кишечника', 'секреторная недостаточность желудка', 'ar'), 
('дискинезия нисходящего отдела 12 п.к.', 'хронический панкреатит', 'ar'), 
('диспротеинемия', 'хронический холецистит', 'ar'), 
('дуоденогастральный рефлюкс', 'химический гастрит', 'ogr'), 
('дуоденостаз', 'дискинезия нисходящего отдела 12 п.к.', 'ar'), 
('жалоба/осмотр: анальный зуд', 'анальный зуд', 'ar'), 
('жалоба/осмотр: аномалия развития лицевого черепа', 'синдром Денди-Уокера', 'ar'), 
('жалоба/осмотр: аномалия развития лицевого черепа', 'аномалия развития лицевого черепа', 'ar'), 
('жалоба/осмотр: атаксия', 'церебральный инсульт', 'ar'), 
('жалоба/осмотр: атаксия', 'гиповитаминоз B12', 'ar'), 
('жалоба/осмотр: атаксия', 'хронический панкреатит', 'ar'), 
('жалоба/осмотр: атаксия', 'генетическое нарушение', 'ar'), 
('жалоба/осмотр: атаксия', 'атаксия', 'ar'), 
('жалоба/осмотр: атаксия', 'атрофический аутоиммунный гастрит', 'ar'), 
('жалоба/осмотр: атрофические изменения слизистой оболочки желудочно-кишечного тракта', 'сидеропенический синдром', 'ar'), 
('жалоба/осмотр: атрофия мышц', 'сидеропенический синдром', 'ar'), 
('жалоба/осмотр: атрофия сосочков языка', 'сидеропенический синдром', 'ar'), 
('жалоба/осмотр: атрофия сосочков языка', 'атрофический гастрит', 'ar'), 
('жалоба/осмотр: атрофия сосочков языка', 'атрофия сосочков языка', 'ar'), 
('жалоба/осмотр: атрофия сосочков языка', 'хронический панкреатит', 'ar'), 
('жалоба/осмотр: ахоличный кал', 'хронический холецистит', 'ar'), 
('жалоба/осмотр: белый густой налет на языке', 'атрофический мультифокальный гастрит', 'ar'), 
('жалоба/осмотр: беспокойное поведение', 'COVID-19', 'ar'), 
('жалоба/осмотр: бессонница', 'нарушение циркадных ритмов организма', 'ar'), 
('жалоба/осмотр: бледность', 'хронический панкреатит', 'ar'), 
('жалоба/осмотр: бледность', 'сидеропенический синдром', 'ar'), 
('жалоба/осмотр: бледность кожи', 'фуникулярный миелоз', 'ar'), 
('жалоба/осмотр: вогнутость на поверхности ногтя', 'сидеропенический синдром', 'ar'), 
('жалоба/осмотр: выпадение волос', 'сидеропенический синдром', 'ar'), 
('жалоба/осмотр: выпадение волос', 'гиповитаминоз B6', 'ar'), 
('жалоба/осмотр: высыпания на коже', 'гипервитаминоз B6', 'ar'), 
('жалоба/осмотр: вязкий кал', 'экзокринная недостаточность поджелудочной железы', 'ar'), 
('жалоба/осмотр: грязно-серый оттенок кожных покровов', 'хронический панкреатит', 'ogr'), 
('жалоба/осмотр: деменция', 'деменция', 'ar'), 
('жалоба/осмотр: депрессия', 'депрессия', 'ar'), 
('жалоба/осмотр: деформация ногтевых пластин', 'сидеропенический синдром', 'ar'), 
('жалоба/осмотр: диарея', 'лямблиоз', 'ar'), 
('жалоба/осмотр: диарея', 'COVID-19', 'ar'), 
('жалоба/осмотр: диарея', 'секреторная недостаточность желудка', 'ar'), 
('жалоба/осмотр: диарея', 'обезвоживание', 'ar'), 
('жалоба/осмотр: диарея', 'СИБР', 'ar'), 
('жалоба/осмотр: диарея', 'влияние аспирина', 'ar'), 
('жалоба/осмотр: диарея', 'вирусная инфекция', 'ar'), 
('жалоба/осмотр: диарея', 'атрофический гастрит', 'ar'), 
('жалоба/осмотр: диарея', 'хронический панкреатит', 'ar'), 
('жалоба/осмотр: диарея', 'влияние антибиотика', 'ar'), 
('жалоба/осмотр: диарея', 'экзокринная недостаточность поджелудочной железы', 'ar'), 
('жалоба/осмотр: динамическая атаксия', 'динамическая атаксия', 'ar'), 
('жалоба/осмотр: желтушность', 'хронический панкреатит', 'ar'), 
('жалоба/осмотр: желтушность', 'хронический холецистит', 'ar'), 
('жалоба/осмотр: желтушность', 'фуникулярный миелоз', 'ar'), 
('жалоба/осмотр: желтые испражнения с незначительной примесью слизи', 'лямблиоз', 'ar'), 
('жалоба/осмотр: задержка мочи', 'COVID-19', 'ar'), 
('жалоба/осмотр: изменение голоса', 'рак легкого', 'ar'), 
('жалоба/осмотр: изменение запаха кала', 'экзокринная недостаточность поджелудочной железы', 'ar'), 
('жалоба/осмотр: изменение консистенции кала', 'экзокринная недостаточность поджелудочной железы', 'ar'), 
('жалоба/осмотр: изменение цвета кала', 'экзокринная недостаточность поджелудочной железы', 'ar'), 
('жалоба/осмотр: икота', 'COVID-19', 'ar'), 
('жалоба/осмотр: истончение ногтей', 'сидеропенический синдром', 'ar'), 
('жалоба/осмотр: истощение', 'лямблиоз', 'ar'), 
('жалоба/осмотр: кариес', 'гастроэзофагеальная рефлюксная болезнь', 'ar'), 
('жалоба/осмотр: кариес', 'кариес', 'ar'), 
('жалоба/осмотр: кашель', 'влияние курения', 'ar'), 
('жалоба/осмотр: кашель', 'вирусная инфекция', 'ar'), 
('жалоба/осмотр: кашель с кровью', 'рак легкого', 'ar'), 
('жалоба/осмотр: кашицеобразный кал', 'экзокринная недостаточность поджелудочной железы', 'ar'), 
('жалоба/осмотр: кончики ногтей изогнуты вниз больше обычного', 'рак легкого', 'ar'), 
('жалоба/осмотр: крапивница', 'крапивница', 'ar'), 
('жалоба/осмотр: крапивница', 'лямблиоз', 'ar'), 
('жалоба/осмотр: куриная слепота', 'хронический панкреатит', 'ar'), 
('жалоба/осмотр: ломкость волос', 'сидеропенический синдром', 'ar'), 
('жалоба/осмотр: ломкость ногтей', 'сидеропенический синдром', 'ar'), 
('жалоба/осмотр: метеоризм', 'лямблиоз', 'ar'), 
('жалоба/осмотр: метеоризм', 'СИБР', 'ar'), 
('жалоба/осмотр: метеоризм', 'хронический панкреатит', 'ogr'), 
('жалоба/осмотр: метеоризм', 'гастрит', 'ar'), 
('жалоба/осмотр: метеоризм', 'экзокринная недостаточность поджелудочной железы', 'ar'), 
('жалоба/осмотр: метеоризм', 'неатрофический гастрит', 'ar'), 
('жалоба/осмотр: мышечные спазмы', 'хронический панкреатит', 'ar'), 
('жалоба/осмотр: налет на языке', 'хронический панкреатит', 'ar'), 
('жалоба/осмотр: наличие в кале непереваренной пищи', 'экзокринная недостаточность поджелудочной железы', 'ar'), 
('жалоба/осмотр: нарушение глотания', 'язва пищевода', 'ar'), 
('жалоба/осмотр: нарушение глотания', 'гастроэзофагеальный рефлюкс', 'ar'), 
('жалоба/осмотр: нарушение глотания', 'сидеропенический синдром', 'ar'), 
('жалоба/осмотр: нарушение зрения', 'синдром Денди-Уокера', 'ar'), 
('жалоба/осмотр: нарушение координации движений', 'мозжечковая симптоматика', 'ar'), 
('жалоба/осмотр: нарушение осанки', 'нарушение осанки', 'ar'), 
('жалоба/осмотр: нарушение сна', 'лямблиоз', 'ar'), 
('жалоба/осмотр: нарушение сна', 'нарушение сна', 'ar'), 
('жалоба/осмотр: нарушения слуха', 'COVID-19', 'ar'), 
('жалоба/осмотр: невозможность удержать мочу при смехе, кашле, чихании', 'сидеропенический синдром', 'ar'), 
('жалоба/осмотр: недержание мочи', 'сидеропенический синдром', 'ar'), 
('жалоба/осмотр: неустойчивый стул', 'хронический холецистит', 'ar'), 
('жалоба/осмотр: нистагм', 'мозжечковая симптоматика', 'ar'), 
('жалоба/осмотр: обморок', 'шок', 'ae'), 
('жалоба/осмотр: обморок', 'обморок', 'ar'), 
('жалоба/осмотр: обморок', 'кома', 'ae'), 
('жалоба/осмотр: образование на коже трещин', 'сидеропенический синдром', 'ar'), 
('жалоба/осмотр: обратимые нарушения зрения', 'влияние аспирина', 'ar'), 
('жалоба/осмотр: обширные отеки', 'нефротический синдром', 'ar'), 
('жалоба/осмотр: обширные травмы', 'обширные травмы', 'ar'), 
('жалоба/осмотр: одышка', 'фуникулярный миелоз', 'ar'), 
('жалоба/осмотр: ожог', 'ожог', 'ar'), 
('жалоба/осмотр: опущенное веко', 'рак легкого', 'ar'), 
('жалоба/осмотр: осиплость голоса', 'гастроэзофагеальный рефлюкс', 'ar'), 
('жалоба/осмотр: отрыжка', 'дуоденостаз', 'ar'), 
('жалоба/осмотр: отрыжка', 'секреторная недостаточность желудка', 'ar'), 
('жалоба/осмотр: отрыжка', 'хронический холецистит', 'ar'), 
('жалоба/осмотр: пенистость мочи', 'хронический холецистит', 'ar'), 
('жалоба/осмотр: поза зародыша', 'хронический панкреатит', 'ogr'), 
('жалоба/осмотр: покраснение кончика языка', 'сидеропенический синдром', 'ar'), 
('жалоба/осмотр: покраснение лица', 'хронический панкреатит', 'ar'), 
('жалоба/осмотр: покраснение языка', 'сидеропенический синдром', 'ar'), 
('жалоба/осмотр: полиурия', 'полиурия', 'ar'), 
('жалоба/осмотр: полифекалия', 'хронический панкреатит', 'ar'), 
('жалоба/осмотр: полифекалия', 'экзокринная недостаточность поджелудочной железы', 'ar'), 
('жалоба/осмотр: поперечная исчерченность ногтей', 'сидеропенический синдром', 'ar'), 
('жалоба/осмотр: постоянное покашливание', 'гастроэзофагеальный рефлюкс', 'ar'), 
('жалоба/осмотр: потемнение мочи', 'хронический холецистит', 'ar'), 
('жалоба/осмотр: потеря веса', 'хронический панкреатит', 'ogr'), 
('жалоба/осмотр: потеря веса', 'экзокринная недостаточность поджелудочной железы', 'ar'), 
('жалоба/осмотр: потеря веса', 'секреторная недостаточность желудка', 'ar'), 
('жалоба/осмотр: потеря веса', 'СИБР', 'ar'), 
('жалоба/осмотр: потеря веса', 'атрофический гастрит', 'ar'), 
('жалоба/осмотр: потеря веса', 'вирусная инфекция', 'ar'), 
('жалоба/осмотр: потливость', 'обезвоживание', 'ar'), 
('жалоба/осмотр: потливость', 'COVID-19', 'ar'), 
('жалоба/осмотр: приступы кашля и/или удушья преимущественно в ночное время, после обильной еды', 'гастроэзофагеальный рефлюкс', 'ogr'), 
('жалоба/осмотр: психоз', 'атрофический аутоиммунный гастрит', 'ar'), 
('жалоба/осмотр: психоз', 'психоз', 'ar'), 
('жалоба/осмотр: раннее поседение волос', 'сидеропенический синдром', 'ar'), 
('жалоба/осмотр: расчесы на коже', 'хронический холецистит', 'ar'), 
('жалоба/осмотр: рвота', 'COVID-19', 'ar'), 
('жалоба/осмотр: рвота', 'секреторная недостаточность желудка', 'ar'), 
('жалоба/осмотр: рвота', 'обезвоживание', 'ar'), 
('жалоба/осмотр: рвота', 'влияние аспирина', 'ar'), 
('жалоба/осмотр: рвота', 'вирусная инфекция', 'ar'), 
('жалоба/осмотр: рвота, не приносящая облегчения, вначале пищей, затем содержимым с примесью желчи, сопровождается тошнотой', 'хронический холецистит', 'ar'), 
('жалоба/осмотр: редкий сухой кашель', 'COVID-19', 'ar'), 
('жалоба/осмотр: серый цвет кала', 'экзокринная недостаточность поджелудочной железы', 'ar'), 
('жалоба/осмотр: симптом "красных капелек"', 'хронический панкреатит', 'ogr'), 
('жалоба/осмотр: симптом "синих склер"', 'сидеропенический синдром', 'ar'), 
('жалоба/осмотр: симптом Айзенберга I', 'хронический холецистит', 'ar'), 
('жалоба/осмотр: симптом Алиева', 'хронический холецистит', 'ar'), 
('жалоба/осмотр: симптом Гротта', 'хронический панкреатит', 'ogr'), 
('жалоба/осмотр: симптом Кера', 'хронический холецистит', 'ar'), 
('жалоба/осмотр: симптом Мерфи', 'хронический холецистит', 'ar'), 
('жалоба/осмотр: симптом Щеткина-Блюмберга', 'хронический холецистит', 'ar'), 
('жалоба/осмотр: симптомы поражения черепных нервов', 'синдром Денди-Уокера', 'ar'), 
('жалоба/осмотр: снижение объема рабочей памяти', 'нарушение циркадных ритмов организма', 'ar'), 
('жалоба/осмотр: снижение объема рабочей памяти', 'снижение объема рабочей памяти', 'ar'), 
('жалоба/осмотр: снижение скорости мышления', 'снижение скорости мышления', 'ar'), 
('жалоба/осмотр: снижение скорости мышления', 'нарушение циркадных ритмов организма', 'ar'), 
('жалоба/осмотр: снижение скорости мышления', 'хронический панкреатит', 'ar'), 
('жалоба/осмотр: снижение скорости реакции на неожиданные ситуации', 'снижение скорости реакции на неожиданные ситуации', 'ar'), 
('жалоба/осмотр: снижение скорости реакции на неожиданные ситуации', 'нарушение циркадных ритмов организма', 'ar'), 
('жалоба/осмотр: снижение способности к концентрации внимания', 'нарушение циркадных ритмов организма', 'ar'), 
('жалоба/осмотр: снижение способности к концентрации внимания', 'снижение способности к концентрации внимания', 'ar'), 
('жалоба/осмотр: снижение способности контролировать импульсивное поведение', 'нарушение циркадных ритмов организма', 'ar'), 
('жалоба/осмотр: снижение способности контролировать импульсивное поведение', 'снижение способности контролировать импульсивное поведение', 'ar'), 
('жалоба/осмотр: снижение тургора кожи', 'снижение тургора кожи', 'ar'), 
('жалоба/осмотр: статическая атаксия', 'мозжечковая симптоматика', 'ar'), 
('жалоба/осмотр: статическая атаксия', 'статическая атаксия', 'ar'), 
('жалоба/осмотр: стеаторея', 'стеаторея', 'ar'), 
('жалоба/осмотр: судороги', 'хронический панкреатит', 'ar'), 
('жалоба/осмотр: судороги', 'гипервитаминоз B6', 'ar'), 
('жалоба/осмотр: судороги', 'обморок', 'ae'), 
('жалоба/осмотр: сухость глаз', 'хронический панкреатит', 'ar'), 
('жалоба/осмотр: сухость кожи', 'сидеропенический синдром', 'ar'), 
('жалоба/осмотр: сухость кожи', 'хронический панкреатит', 'ogr'), 
('жалоба/осмотр: сухость кожи', 'ксероз', 'ar'), 
('жалоба/осмотр: сухость языка', 'хронический панкреатит', 'ar'), 
('жалоба/осмотр: сыпь', 'COVID-19', 'ar'), 
('жалоба/осмотр: сыпь', 'вирусная инфекция', 'ar'), 
('жалоба/осмотр: сыпь', 'влияние аспирина', 'ar'), 
('жалоба/осмотр: тошнота', 'дуоденостаз', 'ar'), 
('жалоба/осмотр: тошнота', 'гастрит', 'ar'), 
('жалоба/осмотр: тошнота', 'лямблиоз', 'ar'), 
('жалоба/осмотр: тошнота', 'синдром Денди-Уокера', 'ar'), 
('жалоба/осмотр: тошнота', 'COVID-19', 'ar'), 
('жалоба/осмотр: тошнота', 'вирусная инфекция', 'ar'), 
('жалоба/осмотр: тошнота', 'хронический холецистит', 'ar'), 
('жалоба/осмотр: тошнота', 'секреторная недостаточность желудка', 'ar'), 
('жалоба/осмотр: тошнота', 'влияние аспирина', 'ar'), 
('жалоба/осмотр: травма уха', 'травма уха', 'ar'), 
('жалоба/осмотр: трещины в углу рта', 'гиповитаминоз B6', 'ar'), 
('жалоба/осмотр: трещины в углу рта', 'ангулярный стоматит', 'ar'), 
('жалоба/осмотр: трещины в углу рта', 'сидеропенический синдром', 'ar'), 
('жалоба/осмотр: тусклость волос', 'сидеропенический синдром', 'ar'), 
('жалоба/осмотр: тусклость ногтей', 'сидеропенический синдром', 'ar'), 
('жалоба/осмотр: увеличение лимфатических узлов', 'вирусная инфекция', 'ar'), 
('жалоба/осмотр: увеличение миндалин', 'вирусная инфекция', 'ar'), 
('жалоба/осмотр: урчание живота', 'лямблиоз', 'ar'), 
('жалоба/осмотр: урчание живота', 'хронический панкреатит', 'ar'), 
('жалоба/осмотр: участки пигментации на лице и конечностях', 'хронический панкреатит', 'ar'), 
('жалоба/осмотр: частое мочеиспускание', 'рак легкого', 'ar'), 
('жалоба/осмотр: черный стул', 'язва желудка', 'ar'), 
('жалоба/осмотр: черный стул', 'язва 12 п.к.', 'ar'), 
('жалоба/осмотр: чихание', 'вирусная инфекция', 'ar'), 
('жалоба/осмотр: шелушение кожи', 'сидеропенический синдром', 'ar'), 
('жалоба/осмотр: экхимозы', 'хронический панкреатит', 'ar'), 
('жалоба/осмотр: энурез', 'энурез', 'ar'), 
('жалоба/осмотр: энурез', 'сидеропенический синдром', 'ar'), 
('жалоба/осмотр: эрозия эмали зубов', 'эрозия эмали зубов', 'ar'), 
('жалоба/осмотр: эрозия эмали зубов', 'гастроэзофагеальная рефлюксная болезнь', 'ar'), 
('жалоба/пальпация: боль при пальпации живота в эпигастральной области', 'хронический панкреатит', 'ar'), 
('жалоба/пальпация: боль при пальпации живота в эпигастральной области', 'неатрофический гастрит', 'ar'), 
('жалоба/пальпация: боль при пальпации живота в эпигастральной области', 'гастрит', 'ar'), 
('жалоба/пальпация: боль при пальпация живота с иррадиацией в спину, позвоночник, левое подреберье, надплечье', 'хронический панкреатит', 'ogr'), 
('жалоба/пальпация: пальпаторно болезненность в правом подреберье', 'хронический холецистит', 'ar'), 
('жалоба/пальпация: снижение тургора кожи', 'снижение тургора кожи', 'ar'), 
('жалоба/пальпация: умеренная болезненность в пилоро-дуоденальной области при пальпации', 'неатрофический гастрит', 'ar'), 
('жалоба: болезненность в левом реберно-позвоночном углу', 'хронический панкреатит', 'ar'), 
('жалоба: болезненность языка', 'сидеропенический синдром', 'ar'), 
('жалоба: боль', 'воспалительный процесс', 'ar'), 
('жалоба: боль в горле', 'вирусная инфекция', 'ar'), 
('жалоба: боль в горле', 'COVID-19', 'ar'), 
('жалоба: боль в горле при глотании', 'COVID-19', 'ar'), 
('жалоба: боль в груди', 'рак легкого', 'ar'), 
('жалоба: боль в животе', 'влияние антибиотика', 'ar'), 
('жалоба: боль в животе', 'COVID-19', 'ar'), 
('жалоба: боль в костях', 'COVID-19', 'ar'), 
('жалоба: боль в мышцах', 'COVID-19', 'ar'), 
('жалоба: боль в области пупка', 'лямблиоз', 'ar'), 
('жалоба: боль в плече', 'рак легкого', 'ar'), 
('жалоба: боль в правом подреберье', 'хронический холецистит', 'ar'), 
('жалоба: боль в руке или плече', 'рак легкого', 'ar'), 
('жалоба: боль в суставах', 'COVID-19', 'ar'), 
('жалоба: боль в суставах', 'артрит', 'ar'), 
('жалоба: боль в ухе', 'гастроэзофагеальный рефлюкс', 'ar'), 
('жалоба: боль в эпигастральной области с иррадиацией в спину', 'хронический панкреатит', 'ogr'), 
('жалоба: боль в эпигастрии', 'лямблиоз', 'ar'), 
('жалоба: боль в эпигастрии', 'гастрит', 'ar'), 
('жалоба: боль в эпигастрии', 'влияние аспирина', 'ar'), 
('жалоба: боль и жжение в языке', 'атрофический аутоиммунный гастрит', 'ar'), 
('жалоба: боль и жжение во рту', 'атрофический аутоиммунный гастрит', 'ar'), 
('жалоба: боль и распирание в области языка', 'сидеропенический синдром', 'ar'), 
('жалоба: боль при глотании', 'язва пищевода', 'ar'), 
('жалоба: боль при глотании', 'гастроэзофагеальный рефлюкс', 'ar'), 
('жалоба: боль при глотании', 'сидеропенический синдром', 'ar'), 
('жалоба: головная боль', 'COVID-19', 'ar'), 
('жалоба: головная боль', 'лямблиоз', 'ar'), 
('жалоба: головная боль', 'вирусная инфекция', 'ar'), 
('жалоба: головная боль', 'влияние аспирина', 'ar'), 
('жалоба: головокружение', 'лямблиоз', 'ar'), 
('жалоба: головокружение', 'сидеропенический синдром', 'ar'), 
('жалоба: головокружение', 'гипервитаминоз B6', 'ar'), 
('жалоба: головокружение', 'фуникулярный миелоз', 'ar'), 
('жалоба: головокружение', 'влияние аспирина', 'ar'), 
('жалоба: горечь во рту', 'хронический холецистит', 'ar'), 
('жалоба: жажда', 'рак легкого', 'ar'), 
('жалоба: желание употреблять в пищу что-то необычное и малосъедобное (мел, зубной порошок, уголь, глину, песок, лед, сырое тесто, фарш, крупу, …)', 'сидеропенический синдром', 'ar'), 
('жалоба: заложенность носа', 'вирусная инфекция', 'ar'), 
('жалоба: заложенность носа', 'COVID-19', 'ar'), 
('жалоба: запор', 'хронический холецистит', 'ar'), 
('жалоба: запор', 'лямблиоз', 'ar'), 
('жалоба: запор', 'экзокринная недостаточность поджелудочной железы', 'ar'), 
('жалоба: запор', 'хронический панкреатит', 'ar'), 
('жалоба: запор', 'запор', 'ar'), 
('жалоба: затруднения при глотании пищи', 'сидеропенический синдром', 'ar'), 
('жалоба: звон в ушах', 'COVID-19', 'ar'), 
('жалоба: извращение вкуса', 'сидеропенический синдром', 'ar'), 
('жалоба: извращение обоняния', 'сидеропенический синдром', 'ar'), 
('жалоба: изжога', 'хронический холецистит', 'ar'), 
('жалоба: изжога', 'гастроэзофагеальный рефлюкс', 'ogr'), 
('жалоба: изжога', 'дуоденостаз', 'ar'), 
('жалоба: кожный зуд', 'хронический холецистит', 'ar'), 
('жалоба: ломота и боли в теле', 'вирусная инфекция', 'ar'), 
('жалоба: ломота и боли в теле', 'COVID-19', 'ar'), 
('жалоба: мышечная боль', 'COVID-19', 'ar'), 
('жалоба: мышечная слабость', 'гиповитаминоз B6', 'ar'), 
('жалоба: мышечная слабость', 'сидеропенический синдром', 'ar'), 
('жалоба: непереносимость алкоголя', 'хронический холецистит', 'ar'), 
('жалоба: непереносимость газированных напитков', 'хронический холецистит', 'ar'), 
('жалоба: непереносимость глютена', 'целиакия', 'ar'), 
('жалоба: непереносимость жареной пищи', 'хронический холецистит', 'ar'), 
('жалоба: непереносимость жирной пищи', 'хронический холецистит', 'ar'), 
('жалоба: непереносимость лактозы', 'вторичная непереносимость лактозы', 'ar'), 
('жалоба: непереносимость острой пищи', 'хронический холецистит', 'ar'), 
('жалоба: непереносимость яиц', 'хронический холецистит', 'ar'), 
('жалоба: непреодолимые позывы на мочеиспускание', 'сидеропенический синдром', 'ar'), 
('жалоба: непрерывная рвота не приносящая облегчение', 'хронический панкреатит', 'ogr'), 
('жалоба: онемение конечностей', 'атрофический аутоиммунный гастрит', 'ar'), 
('жалоба: онемение конечностей', 'гиповитаминоз B6', 'ar'), 
('жалоба: першение в горле', 'COVID-19', 'ar'), 
('жалоба: першение в горле', 'гастроэзофагеальный рефлюкс', 'ar'), 
('жалоба: потеря вкуса', 'COVID-19', 'ar'), 
('жалоба: потеря обоняния', 'COVID-19', 'ar'), 
('жалоба: пристрастие к необычным запахам (бензин, ацетон, лаки, краски, гуталин, сигаретный дым, …)', 'сидеропенический синдром', 'ar'), 
('жалоба: раздражительность', 'вирусная инфекция', 'ar'), 
('жалоба: раздражительность', 'атрофический гастрит', 'ar'), 
('жалоба: раздражительность', 'синдром Денди-Уокера', 'ar'), 
('жалоба: раздражительность', 'лямблиоз', 'ar'), 
('жалоба: слабость', 'фуникулярный миелоз', 'ar'), 
('жалоба: слабость', 'атрофический гастрит', 'ar'), 
('жалоба: слабость', 'лямблиоз', 'ar'), 
('жалоба: слабость', 'COVID-19', 'ar'), 
('жалоба: слабость', 'вирусная инфекция', 'ar'), 
('жалоба: слабость сфинктеров', 'COVID-19', 'ar'), 
('жалоба: стресс', 'стресс', 'ar'), 
('жалоба: сухость во рту', 'хронический холецистит', 'ar'), 
('жалоба: тяжесть в эпигастрии', 'экзокринная недостаточность поджелудочной железы', 'ar'), 
('жалоба: тяжесть в эпигастрии', 'секреторная недостаточность желудка', 'ar'), 
('жалоба: умеренный озноб', 'хронический холецистит', 'ar'), 
('жалоба: усиление симптомов хронической сердечной недостаточности', 'влияние аспирина', 'ar'), 
('жалоба: утомляемость', 'лямблиоз', 'ar'), 
('жалоба: утомляемость', 'вирусная инфекция', 'ar'), 
('жалоба: утомляемость', 'сидеропенический синдром', 'ar'), 
('жалоба: утомляемость', 'гиповитаминоз B6', 'ar'), 
('жалоба: утомляемость', 'фуникулярный миелоз', 'ar'), 
('жалоба: чувство "комка" в горле', 'гастроэзофагеальный рефлюкс', 'ogr'), 
('жалоба: чувство раннего насыщения', 'атрофический гастрит', 'ar'), 
('жалоба: чувство раннего насыщения', 'гастрит', 'ar'), 
('жалоба: шум в ушах', 'влияние аспирина', 'ar'), 
('железо в плазме крови: снижение', 'сидеропенический синдром', 'ar'), 
('железо в сыворотке крови: снижение', 'сидеропенический синдром', 'ar'), 
('железо в сыворотке: норма', 'прелатентный дефицит железа', 'or'), 
('железо-дефицитная анемия', 'влияние недостаточного употребления железа', 'ar'), 
('железо-дефицитная анемия', 'постгеморрагическая анемия', 'ar'), 
('железо-дефицитная анемия', 'сидеропенический синдром', 'ar'), 
('железо-дефицитная анемия неуточненная', 'железо-дефицитная анемия', 'ar'), 
('желудочная диспепсия', 'неатрофический гастрит', 'ar'), 
('желудочная диспепсия', 'гастрит', 'ar'), 
('желчнокаменная болезнь', 'хронический панкреатит', 'ar'), 
('закупорка и стеноз базилярной артерии', 'закупорка и стеноз прецеребральных артерий, не приводящие к инфаркту мозга', 'ar'), 
('закупорка и стеноз других прецеребральных артерий', 'закупорка и стеноз прецеребральных артерий, не приводящие к инфаркту мозга', 'ar'), 
('закупорка и стеноз множественных и двусторонних прецеребральных артерий', 'закупорка и стеноз прецеребральных артерий, не приводящие к инфаркту мозга', 'ar'), 
('закупорка и стеноз неуточненной прецеребральной артерии', 'закупорка и стеноз прецеребральных артерий, не приводящие к инфаркту мозга', 'ar'), 
('закупорка и стеноз позвоночной артерии', 'закупорка и стеноз прецеребральных артерий, не приводящие к инфаркту мозга', 'ar'), 
('закупорка и стеноз сонной артерии', 'закупорка и стеноз прецеребральных артерий, не приводящие к инфаркту мозга', 'ar'), 
('злокачественное новообразование ободочной кишки', 'колоректальный рак', 'ar'), 
('злокачественное новообразование прямой кишки', 'рак прямой кишки', 'ar'), 
('иммунодефицит', 'лямблиоз', 'ar'), 
('ИМТ: более 40 кг/м2', 'ожирение 3 степени', 'ar'), 
('ИМТ: до 16 кг/м2', 'ИМТ: до 18,5 кг/м2', 'ar'), 
('ИМТ: до 16 кг/м2', 'выраженный дефицит массы тела', 'ar'), 
('ИМТ: до 18,5 кг/м2', 'хронический панкреатит', 'ar'), 
('ИМТ: от 16 до 18.5 кг/м2', 'ИМТ: до 18,5 кг/м2', 'ar'), 
('ИМТ: от 16 до 18.5 кг/м2', 'недостаточная (дефицит) масса тела', 'ar'), 
('ИМТ: от 25 до 30 кг/м2', 'избыточная масса тела (предожирение)', 'ar'), 
('ИМТ: от 30 до 35 кг/м2', 'ожирение 1 степени', 'ar'), 
('ИМТ: от 35 до 40 кг/м2', 'ожирение 2 степени', 'ar'), 
('инородное тело желудка', 'гранулематозный гастрит', 'ar'), 
('инсулин в крови: снижение', 'дефицит инсулина', 'ar'), 
('инсульт, не уточненный как кровоизлияние или инфаркт', 'инсульт', 'ar'), 
('интерлейкин 6: от 10,5 до 14 пг/мл', 'COVID-19', 'ar'), 
('инфаркт мозга, вызванный неуточненной закупоркой или стенозом прецеребральных артерий', 'инфаркт мозга', 'ar'), 
('инфаркт мозга, вызванный тромбозом вен мозга, непиогенный', 'инфаркт мозга', 'ar'), 
('инфаркт мозга, вызванный тромбозом мозговых артерий', 'инфаркт мозга', 'ar'), 
('инфаркт мозга, вызванный эмболией прецеребральных артерий', 'инфаркт мозга', 'ar'), 
('инфекция', 'гиповитаминоз B6', 'ar'), 
('ИФА кала на лямблии: положительно', 'лямблиоз', 's'), 
('ИФА крови на лямблии: IgG против giardia lamblia', 'лямблиоз', 'ar'), 
('ИФА крови на лямблии: IgM против giardia lamblia', 'лямблиоз', 's'), 
('кальцификация поджелудочной железы', 'хронический панкреатит', 'ogr'), 
('кетоны в моче: наличие', 'COVID-19', 'ar'), 
('киста задней черепной ямки', 'синдром Денди-Уокера', 'ar'), 
('кишечная диспепсия', 'хронический панкреатит', 'ar'), 
('кишечная диспепсия', 'неатрофический гастрит', 'ar'), 
('колоноскопия: язвенный колит', 'язвенный колит', 'ar'), 
('конъюнктивит', 'COVID-19', 'ar'), 
('копрограмма: креаторея', 'хронический панкреатит', 'ogr'), 
('копрограмма: стеаторея', 'стеаторея', 'ar'), 
('кортизол: повышение', 'глюкокортикоидные гормоны: повышение', 'ar'), 
('кортизон: повышение', 'глюкокортикоидные гормоны: повышение', 'ar'), 
('коэффициент больших тромбоцитов (P-LCR): повышение', 'влияние алкоголя', 'ar'), 
('коэффициент больших тромбоцитов (P-LCR): повышение', 'идиопатическая тромбоцитопеническая пурпура', 'ar'), 
('коэффициент больших тромбоцитов (P-LCR): повышение', 'гиперлипидемия', 'ar'), 
('коэффициент больших тромбоцитов (P-LCR): повышение', 'риск тромбоза', 'ar'), 
('креатинин: повышение', 'COVID-19', 'ar'), 
('кровотечение из ЖКТ', 'влияние аспирина', 'ar'), 
('КТ: агенезия мозолистого тела', 'агенезия мозолистого тела', 'ar'), 
('КТ: агенезия червя мозжечка', 'агенезия червя мозжечка', 'ar'), 
('КТ: атрофия поджелудочной железы', 'атрофия поджелудочной железы', 'ar'), 
('КТ: выступание костей задней черепной ямки', 'выступание костей задней черепной ямки', 'ar'), 
('КТ: гепатомегалия', 'гепатомегалия', 'ar'), 
('КТ: кальцификация поджелудочной железы', 'кальцификация поджелудочной железы', 'ar'), 
('КТ: киста задней черепной ямки', 'киста задней черепной ямки', 'ar'), 
('КТ: ложная киста поджелудочной железы', 'ложная киста поджелудочной железы', 'ar'), 
('КТ: недоразвитие мозолистого тела', 'недоразвитие мозолистого тела', 'ar'), 
('КТ: недоразвитие червя мозжечка', 'недоразвитие червя мозжечка', 'ar'), 
('КТ: неравномерность просвета основного протока ПЖ и его боковых ветвей', 'неравномерность просвета основного протока ПЖ и его боковых ветвей', 'ar'), 
('КТ: отек поджелудочной железы', 'отек поджелудочной железы', 'ar'), 
('КТ: повышение плотности паренхимы поджелудочной железы', 'повышение плотности паренхимы поджелудочной железы', 'ar'), 
('КТ: расширение четвертого желудочка головного мозга', 'расширение четвертого желудочка головного мозга', 'ar'), 
('КТ: спленомегалия', 'спленомегалия', 'ar'), 
('КТ: хронический панкреатит', 'хронический панкреатит', 'ar'), 
('КТ: эффект "матового стекла"', 'COVID-19', 'ar'), 
('лактатдегидрогеназа для женщин: от 321 до 428 ед/л', 'COVID-19', 'ar'), 
('лактатдегидрогеназа для мужчин: от 337,5 до 450 ед/л', 'COVID-19', 'ar'), 
('ларингоскопия: полип голосовой связки', 'полип голосовой связки', 'ar'), 
('ларингоскопия: узелок голосовой связки', 'узелок голосовой связки', 'ar'), 
('легкая степень экзокринной недостаточности поджелудочной железы', 'экзокринная недостаточность поджелудочной железы', 'ogr'), 
('лейкоциты в моче: наличие', 'COVID-19', 'ar'), 
('лейкоциты: повышение', 'онкологическое заболевание', 'ar'), 
('лейкоциты: повышение', 'влияние недавней физической нагрузки', 'ar'), 
('лейкоциты: повышение', 'обширные травмы', 'ar'), 
('лейкоциты: повышение', 'беременность', 'ogr'), 
('лейкоциты: повышение', 'воспалительный процесс', 'ar'), 
('лейкоциты: повышение', 'полицитемия', 'ar'), 
('лейкоциты: повышение', 'аутоиммунное заболевание', 'ar'), 
('лейкоциты: повышение', 'влияние недавней прививки', 'ar'), 
('лейкоциты: повышение', 'последствие операции', 'ar'), 
('лейкоциты: повышение', 'хронический холецистит', 'ar'), 
('лейкоциты: повышение', 'влияние недавнего приема пищи', 'ar'), 
('лейкоциты: повышение', 'лейкоз', 'ar'), 
('лейкоциты: повышение', 'тяжелый ожог', 'ogr'), 
('лейкоциты: снижение', 'лучевая болезнь', 'ar'), 
('лейкоциты: снижение', 'влияние цитостатика', 'ar'), 
('лейкоциты: снижение', 'COVID-19', 'ar'), 
('лейкоциты: снижение', 'системное аутоиммунное заболевание', 'ar'), 
('лейкоциты: снижение', 'лейкопения', 'ar'), 
('лейкоциты: снижение', 'влияние кортикостероида', 'ar'), 
('лейкоциты: снижение', 'лейкоз', 'ar'), 
('лейкоциты: снижение', 'вирусная инфекция', 'ar'), 
('лимфоциты: снижение', 'COVID-19', 'ar'), 
('липаза: повышение', 'хронический панкреатит', 'ar'), 
('липоидный нефроз', 'нефротический синдром', 'ar'), 
('лихорадка', 'вирусная инфекция', 'ar'), 
('лихорадка', 'гиповитаминоз B6', 'ar'), 
('лихорадка', 'хронический холецистит', 'ar'), 
('лихорадка', 'COVID-19', 'ar'), 
('ложная киста поджелудочной железы', 'хронический панкреатит', 'ogr'), 
('лямблиоз', 'хронический холецистит', 'ar'), 
('лямблиоз', 'влияние препарата подавляющего желудочную секрецию', 'ar'), 
('лямблиоз', 'гипохлоргидрия желудочного сока', 'ar'), 
('метилмалоновая кислота в моче: наличие', 'фуникулярный миелоз', 'ar'), 
('микроскопия кала на простейшие: наличие трофозоитов лямблий в кале', 'лямблиоз', 's'), 
('микроскопия кала на простейшие: наличие цист лямблий в кале', 'лямблиоз', 's'), 
('микроскопия содержимого двенадцатиперстной кишки: наличие трофозоитов лямблий в содержимом двенадцатиперстной кишки', 'лямблиоз', 's'), 
('микседематозная кома', 'гипотиреоз', 'ar'), 
('мозжечковая симптоматика', 'синдром Денди-Уокера', 'ar'), 
('мочевина: повышение', 'COVID-19', 'ar'), 
('МРТ: агенезия мозолистого тела', 'агенезия мозолистого тела', 'ar'), 
('МРТ: агенезия червя мозжечка', 'агенезия червя мозжечка', 'ar'), 
('МРТ: киста задней черепной ямки', 'киста задней черепной ямки', 'ar'), 
('МРТ: ложная киста поджелудочной железы', 'ложная киста поджелудочной железы', 'ar'), 
('МРТ: недоразвитие мозолистого тела', 'недоразвитие мозолистого тела', 'ar'), 
('МРТ: недоразвитие червя мозжечка', 'недоразвитие червя мозжечка', 'ar'), 
('МРТ: неравномерность просвета основного протока ПЖ и его боковых ветвей', 'неравномерность просвета основного протока ПЖ и его боковых ветвей', 'ar'), 
('МРТ: расширение четвертого желудочка головного мозга', 'расширение четвертого желудочка головного мозга', 'ar'), 
('наличие Hp на слизистой желудка', 'лимфоцитарный гастрит', 'ar'), 
('нарушение всасывания жирорастворимых витаминов', 'хронический панкреатит', 'ar'), 
('нарушение кишечного переваривания', 'экзокринная недостаточность поджелудочной железы', 'ar'), 
('нарушение кишечного переваривания', 'СИБР', 'ar'), 
('нарушение кровообращения', 'гиповитаминоз B6', 'ar'), 
('нарушение обмена белков', 'гиповитаминоз B6', 'ar'), 
('нарушение обмена жиров', 'гиповитаминоз B6', 'ar'), 
('нарушение свертывания крови', 'хронический панкреатит', 'ar'), 
('нарушение усвоения нервными клетками глюкозы', 'гиповитаминоз B6', 'ar'), 
('нарушение функции', 'воспалительный процесс', 'ar'), 
('нарушение функции печени', 'гиповитаминоз B6', 'ar'), 
('нарушение функции печени', 'влияние аспирина', 'ar'), 
('нарушение функции почек', 'влияние аспирина', 'ar'), 
('нарушения всасывания витамина В12', 'витамин B12-дефицитная анемия', 'ar'), 
('недоразвитие мозолистого тела', 'синдром Денди-Уокера', 'ar'), 
('недоразвитие червя мозжечка', 'синдром Денди-Уокера', 'ar'), 
('нейронспецифическая энолаза: повышение', 'медуллярный рак щитовидной железы', 'ar'), 
('нейронспецифическая энолаза: повышение', 'мелкоклеточный рак легкого', 'ar'), 
('нейронспецифическая энолаза: повышение', 'феохромоцитома', 'ar'), 
('нейронспецифическая энолаза: повышение', 'злокачественное новообразование сетчатки глаза', 'ar'), 
('нейронспецифическая энолаза: повышение', 'карцинома островковых клеток поджелудочной железы', 'ar'), 
('нейронспецифическая энолаза: повышение', 'карциноид', 'ar'), 
('нейронспецифическая энолаза: повышение', 'злокачественное новообразование периферических нервов и вегетативной нервной системы', 'ar'), 
('нейтрофилы: повышение', 'хронический холецистит', 'ar'), 
('непрямой билирубин: повышение', 'атрофический аутоиммунный гастрит', 'ar'), 
('неравномерность просвета основного протока ПЖ и его боковых ветвей', 'хронический панкреатит', 'ogr'), 
('неспецифический колит', 'воспалительные заболевания кишечника', 'ar'), 
('нефротический синдром', 'влияние аспирина', 'ar'), 
('нефротический синдром', 'гломерулярное заболевание почек', 'ar'), 
('нутрициологический дефицит', 'хронический панкреатит', 'ar'), 
('ОАК: агранулоцитоз', 'агранулоцитоз', 'ar'), 
('ОАК: макроцитоз эритроцитов', 'атрофический аутоиммунный гастрит', 'ar'), 
('ОАК: макроцитоз эритроцитов', 'атрофический мультифокальный гастрит', 'ae'), 
('ОАК: наличие мегалобластов', 'гиповитаминоз B6', 'ar'), 
('ОАК: пойкилоцитоз', 'полицитемия', 'ar'), 
('ОАК: сдвиг лейкоцитарной формулы влево', 'хронический холецистит', 'ar'), 
('ОАК: сдвиг лейкоцитарной формулы влево', 'воспалительный процесс', 'ar'), 
('ОАМ: положительная реакция мочи на билирубин', 'хронический холецистит', 'ar'), 
('обезвоживание', 'полиурия', 'ar'), 
('обезвоживание', 'сахарный диабет', 'ar'), 
('обезвоживание', 'хронический панкреатит', 'ar'), 
('ожирение', 'нарушение циркадных ритмов организма', 'ar'), 
('ожог', 'обезвоживание', 'ar'), 
('операционный анамнез: гастрэктомия', 'лямблиоз', 'ar'), 
('операционный анамнез: операция по поводу рака легкого с расширенной лимфаденэктомией', 'релаксация диафрагмы', 'ar'), 
('описторхоз', 'хронический холецистит', 'ar'), 
('опухоль кожи лица', 'опухоль головы и шеи', 'ar'), 
('опухоль челюсти', 'опухоль головы и шеи', 'ar'), 
('остеомаляция', 'хронический панкреатит', 'ar'), 
('остеопороз', 'хронический панкреатит', 'ar'), 
('острая почечная недостаточность', 'почечная недостаточность', 'ar'), 
('острая почечная недостаточность', 'влияние аспирина', 'ar'), 
('острая сердечная недостаточность', 'COVID-19', 'ar'), 
('острый инфаркт миокарда', 'инфаркт миокарда', 'ar'), 
('острый панкреатит', 'хронический панкреатит', 'ar'), 
('отек Квинке', 'влияние аспирина', 'ar'), 
('отек поджелудочной железы', 'хронический панкреатит', 'ar'), 
('отсутствие Hp на слизистой желудка', 'атрофический аутоиммунный гастрит', 'ar'), 
('отсутствие употребления алкоголя', 'влияние алкоголя', 'ae'), 
('пальпация: гепатомегалия', 'гепатомегалия', 'ar'), 
('пальпация: спленомегалия', 'спленомегалия', 'ar'), 
('пародонтоз', 'сидеропенический синдром', 'ar'), 
('пепсиноген I в крови: снижение', 'атрофический гастрит', 's'), 
('пепсиноген I/пепсиноген II в крови: снижение', 'атрофический гастрит', 's'), 
('перитонит', 'обезвоживание', 'ar'), 
('периферическая нейропатия', 'хронический панкреатит', 'ar'), 
('периферический коллапс', 'COVID-19', 'ar'), 
('пиридоксаль-5-фосфат в плазме: снижение', 'гиповитаминоз B6', 'ar'), 
('пиридоксаль-5-фосфат в цельной крови: снижение', 'гиповитаминоз B6', 'ar'), 
('пищеводные полукольца', 'сидеропенический синдром', 'ar'), 
('пищевое поведение: злоупотребление газированными напитками', 'неатрофический гастрит', 'ar'), 
('пищевое поведение: злоупотребление жаренной пищей', 'влияние химических раздражителей на желудок', 'ar'), 
('пищевое поведение: злоупотребление жаренной пищей', 'неатрофический гастрит', 'ar'), 
('пищевое поведение: злоупотребление кислой пищей', 'сидеропенический синдром', 'ar'), 
('пищевое поведение: злоупотребление кислой пищей', 'влияние химических раздражителей на желудок', 'ar'), 
('пищевое поведение: злоупотребление копченной пищей', 'неатрофический гастрит', 'ar'), 
('пищевое поведение: злоупотребление копченной пищей', 'влияние химических раздражителей на желудок', 'ar'), 
('пищевое поведение: злоупотребление острой пищей', 'сидеропенический синдром', 'ar'), 
('пищевое поведение: злоупотребление острой пищей', 'влияние химических раздражителей на желудок', 'ar'), 
('пищевое поведение: злоупотребление острой пищей', 'неатрофический гастрит', 'ar'), 
('пищевое поведение: злоупотребление пряной пищей', 'сидеропенический синдром', 'ar'), 
('пищевое поведение: злоупотребление пряной пищей', 'влияние химических раздражителей на желудок', 'ar'), 
('пищевое поведение: злоупотребление соленой пищей', 'влияние химических раздражителей на желудок', 'ar'), 
('пищевое поведение: злоупотребление соленой пищей', 'сидеропенический синдром', 'ar'), 
('пищевое поведение: использование для питья сырой воды из ручьев, рек и озер', 'лямблиоз', 'ar'), 
('пищевое поведение: нарушение режима питания', 'неатрофический гастрит', 'ar'), 
('пищевое поведение: недостаточное употребление воды', 'обезвоживание', 'ar'), 
('пищевое поведение: недостаточное употребление продуктов, содержащих витамин B6', 'гиповитаминоз B6', 'ar'), 
('пищевое поведение: разнообразные источники белка в пищевом рационе', 'артериальная гипертензия', 'ad'), 
('пищевое поведение: снижение аппетита', 'COVID-19', 'ar'), 
('пищевое поведение: снижение аппетита', 'СИБР', 'ar'), 
('пищевое поведение: снижение аппетита', 'хронический холецистит', 'ar'), 
('пищевое поведение: снижение аппетита', 'лямблиоз', 'ar'), 
('пищевое поведение: сухоедение', 'неатрофический гастрит', 'ar'), 
('поведение нарушающее циркадные ритмы', 'нарушение циркадных ритмов организма', 'ar'), 
('повторный инфаркт миокарда', 'инфаркт миокарда', 'ar'), 
('повышение плотности паренхимы поджелудочной железы', 'хронический панкреатит', 'ogr'), 
('повышеное образование сгустков крови', 'гиповитаминоз B6', 'ar'), 
('покраснение', 'воспалительный процесс', 'ar'), 
('пол: женский', 'беременность (третий триместр)', 'or'), 
('пол: женский', 'обильная менструация', 'or'), 
('пол: женский', 'менструация', 'or'), 
('пол: женский', 'беременность', 'or'), 
('пол: женский', 'овуляция', 'or'), 
('пол: женский', 'первая фаза менструального цикла (7-14 день)', 'or'), 
('полиурия', 'обезвоживание', 'ar'), 
('полиурия', 'несахарный диабет', 'ar'), 
('полиурия', 'сахарный диабет', 'ar'), 
('полицитемия', 'полицитемия истинная', 'ar'), 
('помутнение сознания', 'гипервитаминоз B6', 'ar'), 
('почечная недостаточность неуточненная', 'почечная недостаточность', 'ar'), 
('прием ЛС: H-2 гистаминоблокатор', 'влияние H-2 гистаминоблокатора', 'or'), 
('прием ЛС: антибиотик', 'влияние антибиотика', 'or'), 
('прием ЛС: антидепрессант', 'влияние антидепрессанта', 'or'), 
('прием ЛС: аскорбиновая кислота', 'влияние аскорбиновой кислоты', 'or'), 
('прием ЛС: аспирин', 'влияние аспирина', 'or'), 
('прием ЛС: барбитурат', 'влияние барбитурата', 'or'), 
('прием ЛС: бета-блокатор', 'влияние бета-блокатора', 'or'), 
('прием ЛС: гентамицин', 'влияние гентамицина', 'or'), 
('прием ЛС: гормон', 'влияние гормонального препарата', 'or'), 
('прием ЛС: диуретик', 'влияние диуретика', 'or'), 
('прием ЛС: ИПП', 'влияние ИПП', 'or'), 
('прием ЛС: кортикостероид', 'влияние кортикостероида', 'or'), 
('прием ЛС: мечегонный препарат', 'обезвоживание', 'ar'), 
('прием ЛС: НПВС', 'химический гастрит', 'ogr'), 
('прием ЛС: НПВС', 'влияние НПВС', 'or'), 
('прием ЛС: оральный контрацептив', 'влияние орального контрацептива', 'or'), 
('прием ЛС: парацетамол', 'влияние парацетамола', 'or'), 
('прием ЛС: препарат железа', 'влияние препарата железа', 'or'), 
('прием ЛС: препарат подавляющий желудочную секрецию', 'влияние препарата подавляющего желудочную секрецию', 'or'), 
('прием ЛС: противогрибковый препарат', 'влияние противогрибкового препарата', 'or'), 
('прием ЛС: статин', 'влияние статина', 'or'), 
('прием ЛС: тестостерон', 'влияние тестостерона', 'or'), 
('прием ЛС: тубазид', 'влияние тубазида', 'or'), 
('прием ЛС: фенобарбитал', 'влияние фенобарбитала', 'or'), 
('прием ЛС: фтивазид', 'влияние фтивазида', 'or'), 
('прием ЛС: цефалоспорин', 'влияние цефалоспорина', 'or'), 
('прием ЛС: эстроген', 'влияние эстрогена', 'or'), 
('припухлость', 'воспалительный процесс', 'ar'), 
('прокальцитонин: более 10 нг/мл', 'синдром полиорганной недостаточности', 'ar'), 
('прокальцитонин: более 10 нг/мл', 'высокий риск летального исхода', 'ar'), 
('прокальцитонин: более 10 нг/мл', 'выраженный воспалительный процесс', 'ar'), 
('прокальцитонин: более 10 нг/мл', 'тяжелый бактериальный сепсис или септический шок', 'ar'), 
('прокальцитонин: более 2,0 нг/мл', 'паразитарное заболевание', 'ar'), 
('прокальцитонин: более 2,0 нг/мл', 'обширные хирургические вмешательства', 'ar'), 
('прокальцитонин: более 2,0 нг/мл', 'тяжелый ожог', 'ogr'), 
('прокальцитонин: более 2,0 нг/мл', 'острое отторжение трансплантата', 'ar'), 
('прокальцитонин: более 2,0 нг/мл', 'сепсис', 'ar'), 
('прокальцитонин: более 2,0 нг/мл', 'грибковая инфекция', 'ar'), 
('прокальцитонин: более 2,0 нг/мл', 'субарахноидальное кровоизлияние', 'ar'), 
('прокальцитонин: более 2,0 нг/мл', 'менингит', 'ar'), 
('прокальцитонин: более 2,0 нг/мл', 'бактериальная инфекция', 'ar'), 
('прокальцитонин: более 2,0 нг/мл', 'операции в условиях искусственного кровообращения', 'ar'), 
('прокальцитонин: более 2,0 нг/мл', 'пневмония', 'ar'), 
('прокальцитонин: более 2,0 нг/мл', 'выраженный бронхит', 'ar'), 
('прокальцитонин: более 2,0 нг/мл', 'множественные травмы', 'ar'), 
('прокальцитонин: более 2,0 нг/мл', 'аппендицит', 'ar'), 
('прокальцитонин: более 2,0 нг/мл', 'перитонит', 'ar'), 
('прокальцитонин: более 2,0 нг/мл', 'панкреатит', 'ar'), 
('прокальцитонин: более 2,0 нг/мл', 'хроническая сердечная недостаточность', 'ar'), 
('прокальцитонин: норма', 'вирусная инфекция', 'ad'), 
('прокальцитонин: от 0,046 до 0,5 нг/мл', 'низкий риск тяжелого сепсиса и/или септического шока', 'ar'), 
('прокальцитонин: от 0,5 до 2,0 нг/мл', 'поставить диагноз сепсиса с уверенностью нельзя, рекомендуется повторить измерение прокальцитонина в течение 6-24 часов', 'ar'), 
('прокальцитонин: от 0,5 до 2,0 нг/мл', 'умеренный воспалительный процесс', 'ar'), 
('прокальцитонин: от 2 до 10 нг/мл', 'высокий риск тяжелого сепсиса и/или септического шока', 'ar'), 
('прокальцитонин: от 2 до 10 нг/мл', 'тяжелый воспалительный процесс', 'ar'), 
('пропионовая кислота в моче: наличие', 'фуникулярный миелоз', 'ar'), 
('протромбиновое время: повышение', 'хронический холецистит', 'ar'), 
('профессия: работа ночью', 'поведение нарушающее циркадные ритмы', 'ar'), 
('псевдомембранозный колит', 'гипохлоргидрия желудочного сока', 'ar'), 
('ПЦР кала на лямблии: положительно', 'лямблиоз', 's'), 
('ПЦР мазка носоглотки и/или ротоглотки на COVID-19: положительно', 'COVID-19, вирус идентифицирован', 'or'), 
('рак глотки', 'опухоль головы и шеи', 'ar'), 
('рак гортани', 'опухоль головы и шеи', 'ar'), 
('рак губы', 'опухоль головы и шеи', 'ar'), 
('рак параназальных синусов', 'опухоль головы и шеи', 'ar'), 
('рак полости носа', 'опухоль головы и шеи', 'ar'), 
('рак полости рта', 'опухоль головы и шеи', 'ar'), 
('рак слюнных желез', 'опухоль головы и шеи', 'ar'), 
('распределение тромбоцитов по объему (PDW): повышение', 'высокая степень агрегации тромбоцитов', 'ar'), 
('распределение тромбоцитов по объему (PDW): повышение', 'менструация', 'ogr'), 
('распределение тромбоцитов по объему (PDW): повышение', 'последствие операции', 'ar'), 
('распределение тромбоцитов по объему (PDW): повышение', 'наличие микросгустков в крови', 'ar'), 
('распределение тромбоцитов по объему (PDW): повышение', 'влияние недавней физической нагрузки', 'ar'), 
('распределение тромбоцитов по объему (PDW): повышение', 'анемия', 'ar'), 
('распределение тромбоцитов по объему (PDW): повышение', 'кровопотеря', 'ar'), 
('распределение тромбоцитов по объему (PDW): повышение', 'беременность', 'ogr'), 
('распределение тромбоцитов по объему (PDW): повышение', 'идиопатическая тромбоцитопеническая пурпура', 'ar'), 
('распределение тромбоцитов по объему (PDW): повышение', 'воспалительный процесс', 'ar'), 
('распределение эритроцитов по объему, коэффициент вариации (RDW-CV): повышение', 'болезнь холодовых агглютининов', 'ar'), 
('распределение эритроцитов по объему, коэффициент вариации (RDW-CV): повышение', 'гемоглобинопатия', 'ar'), 
('распределение эритроцитов по объему, коэффициент вариации (RDW-CV): повышение', 'колоректальный рак', 'ar'), 
('распределение эритроцитов по объему, коэффициент вариации (RDW-CV): повышение', 'диабет', 'ar'), 
('распределение эритроцитов по объему, коэффициент вариации (RDW-CV): повышение', 'метастазы в костный мозг', 'ar'), 
('распределение эритроцитов по объему, коэффициент вариации (RDW-CV): повышение', 'аутоиммунная гемолитическая анемия', 'ar'), 
('распределение эритроцитов по объему, коэффициент вариации (RDW-CV): повышение', 'влияние радиации', 'ar'), 
('распределение эритроцитов по объему, коэффициент вариации (RDW-CV): повышение', 'влияние курения', 'ar'), 
('распределение эритроцитов по объему, коэффициент вариации (RDW-CV): повышение', 'ранний признак анемии', 'ar'), 
('распределение эритроцитов по объему, коэффициент вариации (RDW-CV): повышение', 'заболевание крови', 'ar'), 
('распределение эритроцитов по объему, коэффициент вариации (RDW-CV): повышение', 'RDW-CV: повышение и MCV: снижение', 'or'), 
('распределение эритроцитов по объему, коэффициент вариации (RDW-CV): повышение', 'хроническое заболевание печени', 'ar'), 
('распределение эритроцитов по объему, коэффициент вариации (RDW-CV): повышение', 'серповидно-клеточная анемия', 'ar'), 
('распределение эритроцитов по объему, коэффициент вариации (RDW-CV): повышение', 'онкологическое заболевание костного мозга', 'ar'), 
('распределение эритроцитов по объему, коэффициент вариации (RDW-CV): повышение', 'влияние химиотерапии', 'ar'), 
('распределение эритроцитов по объему, коэффициент вариации (RDW-CV): повышение', 'влияние переливаний крови', 'ar'), 
('распределение эритроцитов по объему, коэффициент вариации (RDW-CV): повышение', 'наследственное нарушение эритроцитов', 'ar'), 
('распределение эритроцитов по объему, коэффициент вариации (RDW-CV): повышение', 'нарушение функции печени', 'ar'), 
('распределение эритроцитов по объему, коэффициент вариации (RDW-CV): повышение', 'железо-дефицитная анемия', 'ar'), 
('распределение эритроцитов по объему, коэффициент вариации (RDW-CV): повышение', 'RDW-CV: повышение и MCV: повышение', 'or'), 
('распределение эритроцитов по объему, коэффициент вариации (RDW-CV): повышение', 'нарушение функции почек', 'ar'), 
('распределение эритроцитов по объему, коэффициент вариации (RDW-CV): повышение', 'воздействие тяжелых металлов (ртуть, свинец)', 'ar'), 
('распределение эритроцитов по объему, коэффициент вариации (RDW-CV): повышение', 'миелодиспластический синдром', 'ar'), 
('распределение эритроцитов по объему, коэффициент вариации (RDW-CV): повышение', 'заболевание костного мозга', 'ar'), 
('распределение эритроцитов по объему, коэффициент вариации (RDW-CV): повышение', 'хроническая почечная недостаточность', 'ar'), 
('распределение эритроцитов по объему, коэффициент вариации (RDW-CV): повышение', 'фолиево-дефицитная анемия', 'ar'), 
('распределение эритроцитов по объему, коэффициент вариации (RDW-CV): повышение', 'миелофиброз', 'ar'), 
('распределение эритроцитов по объему, коэффициент вариации (RDW-CV): повышение', 'сидеробластная анемия', 'ar'), 
('распределение эритроцитов по объему, коэффициент вариации (RDW-CV): повышение', 'анемия, вызванная хроническим воспалением', 'ar'), 
('распределение эритроцитов по объему, коэффициент вариации (RDW-CV): повышение', 'витамин B12-дефицитная анемия', 'ar'), 
('распределение эритроцитов по объему, коэффициент вариации (RDW-CV): повышение', 'талассемия', 'ar'), 
('распределение эритроцитов по объему, коэффициент вариации (RDW-CV): повышение', 'сердечно-сосудистое заболевание', 'ar'), 
('распределение эритроцитов по объему, стандартное отклонение (RDW-SD): повышение', 'онкологическое заболевание костного мозга', 'ar'), 
('распределение эритроцитов по объему, стандартное отклонение (RDW-SD): повышение', 'влияние химиотерапии', 'ar'), 
('распределение эритроцитов по объему, стандартное отклонение (RDW-SD): повышение', 'ранний признак анемии', 'ar'), 
('распределение эритроцитов по объему, стандартное отклонение (RDW-SD): повышение', 'сердечно-сосудистое заболевание', 'ar'), 
('распределение эритроцитов по объему, стандартное отклонение (RDW-SD): повышение', 'миелофиброз', 'ar'), 
('распределение эритроцитов по объему, стандартное отклонение (RDW-SD): повышение', 'колоректальный рак', 'ar'), 
('распределение эритроцитов по объему, стандартное отклонение (RDW-SD): повышение', 'витамин B12-дефицитная анемия', 'ar'), 
('распределение эритроцитов по объему, стандартное отклонение (RDW-SD): повышение', 'нарушение функции печени', 'ar'), 
('распределение эритроцитов по объему, стандартное отклонение (RDW-SD): повышение', 'метастазы в костный мозг', 'ar'), 
('распределение эритроцитов по объему, стандартное отклонение (RDW-SD): повышение', 'железо-дефицитная анемия', 'ar'), 
('распределение эритроцитов по объему, стандартное отклонение (RDW-SD): повышение', 'наследственное нарушение эритроцитов', 'ar'), 
('распределение эритроцитов по объему, стандартное отклонение (RDW-SD): повышение', 'RDW-SD: повышение и MCV: снижение', 'or'), 
('распределение эритроцитов по объему, стандартное отклонение (RDW-SD): повышение', 'RDW-SD: повышение и MCV: повышение', 'or'), 
('распределение эритроцитов по объему, стандартное отклонение (RDW-SD): повышение', 'гемоглобинопатия', 'ar'), 
('распределение эритроцитов по объему, стандартное отклонение (RDW-SD): повышение', 'диабет', 'ar'), 
('распределение эритроцитов по объему, стандартное отклонение (RDW-SD): повышение', 'миелодиспластический синдром', 'ar'), 
('распределение эритроцитов по объему, стандартное отклонение (RDW-SD): повышение', 'аутоиммунная гемолитическая анемия', 'ar'), 
('распределение эритроцитов по объему, стандартное отклонение (RDW-SD): повышение', 'влияние радиации', 'ar'), 
('распределение эритроцитов по объему, стандартное отклонение (RDW-SD): повышение', 'влияние курения', 'ar'), 
('распределение эритроцитов по объему, стандартное отклонение (RDW-SD): повышение', 'хроническое заболевание печени', 'ar'), 
('распределение эритроцитов по объему, стандартное отклонение (RDW-SD): повышение', 'анемия, вызванная хроническим воспалением', 'ar'), 
('распределение эритроцитов по объему, стандартное отклонение (RDW-SD): повышение', 'талассемия', 'ar'), 
('распределение эритроцитов по объему, стандартное отклонение (RDW-SD): повышение', 'болезнь холодовых агглютининов', 'ar'), 
('распределение эритроцитов по объему, стандартное отклонение (RDW-SD): повышение', 'серповидно-клеточная анемия', 'ar'), 
('распределение эритроцитов по объему, стандартное отклонение (RDW-SD): повышение', 'влияние переливаний крови', 'ar'), 
('распределение эритроцитов по объему, стандартное отклонение (RDW-SD): повышение', 'воздействие тяжелых металлов (ртуть, свинец)', 'ar'), 
('распределение эритроцитов по объему, стандартное отклонение (RDW-SD): повышение', 'заболевание костного мозга', 'ar'), 
('распределение эритроцитов по объему, стандартное отклонение (RDW-SD): повышение', 'хроническая почечная недостаточность', 'ar'), 
('распределение эритроцитов по объему, стандартное отклонение (RDW-SD): повышение', 'фолиево-дефицитная анемия', 'ar'), 
('распределение эритроцитов по объему, стандартное отклонение (RDW-SD): повышение', 'нарушение функции почек', 'ar'), 
('распределение эритроцитов по объему, стандартное отклонение (RDW-SD): повышение', 'сидеробластная анемия', 'ar'), 
('распределение эритроцитов по объему, стандартное отклонение (RDW-SD): повышение', 'заболевание крови', 'ar'), 
('расширение четвертого желудочка головного мозга', 'синдром Денди-Уокера', 'ar'), 
('реактивный артрит', 'лямблиоз', 'ar'), 
('релаксация диафрагмы', 'последствие операции по поводу рака легкого с расширенной лимфаденэктомией', 'ar'), 
('релаксация диафрагмы', 'последствие тимэктомии', 'ar'), 
('релаксация диафрагмы', 'последствие операции на диафрагме', 'ar'), 
('релаксация диафрагмы', 'последствие маммарно-коронарного шунтирования', 'ar'), 
('релаксация диафрагмы', 'последствие удаления опухолей средостения', 'ar'), 
('ретикулоциты: повышение', 'влияние проживания в высокогорье', 'ar'), 
('ретикулоциты: повышение', 'постгеморрагическое восстановление кроветворения', 'ar'), 
('ретикулоциты: повышение', 'последствие терапии: лечение анемии', 'ar'), 
('ретикулоциты: снижение', 'апластическая анемия', 'ar'), 
('рецидивирующий полипоз носа', 'аспириновая триада', 'ogr1'), 
('рецидивирующий полипоз околоносовых пазух', 'аспириновая триада', 'ogr1'), 
('ринит', 'COVID-19', 'ar'), 
('ринит', 'вирусная инфекция', 'ar'), 
('РЭА: повышение', 'рак матки', 'ar'), 
('РЭА: повышение', 'гепатоцеллюлярная карцинома', 'ar'), 
('РЭА: повышение', 'рак мочевого пузыря', 'ar'), 
('РЭА: повышение', 'рак молочной железы', 'ar'), 
('РЭА: повышение', 'холангиокарцинома', 'ar'), 
('РЭА: повышение', 'рак прямой кишки', 'ar'), 
('РЭА: повышение', 'мелкоклеточный рак легкого', 'ar'), 
('РЭА: повышение', 'метастазы в печень', 'ar'), 
('РЭА: повышение', 'рак поджелудочной железы', 'ar'), 
('РЭА: повышение', 'рак желчного пузыря', 'ar'), 
('РЭА: повышение', 'немелкоклеточный рак легкого', 'ar'), 
('РЭА: повышение', 'рак желудка', 'ar'), 
('РЭА: повышение', 'колоректальный рак', 'ar'), 
('РЭА: повышение', 'рак щитовидной железы фолликулярный', 'ar'), 
('саркоидоз', 'гранулематозный гастрит', 'ar'), 
('сахар в крови: повышение', 'хронический панкреатит', 'ogr'), 
('сахарный диабет', 'хронический панкреатит', 'ar'), 
('сахарный диабет 1 типа', 'атрофический аутоиммунный гастрит', 'ar'), 
('сахарный диабет 2 типа', 'болезнь Паркинсона', 'ad'), 
('сахарный диабет 2 типа', 'нарушение циркадных ритмов организма', 'ar'), 
('сахарный диабет 2 типа', 'сахарный диабет', 'ar'), 
('сахарный диабет 3 типа', 'сахарный диабет', 'ar'), 
('секреторная недостаточность желудка', 'дисбиоз тонкого кишечника', 'ar'), 
('секреторная недостаточность желудка', 'атрофический гастрит', 'ar'), 
('семейный анамнез: гастро-дуоденальное заболевание', 'неатрофический гастрит', 'ar'), 
('семейный анамнез: гастро-дуоденальное заболевание', 'гастрит', 'ar'), 
('семейный анамнез: лямблиоз', 'лямблиоз', 'ar'), 
('сердечная аритмия', 'гастроэзофагеальный рефлюкс', 'ar'), 
('сердечная недостаточность', 'порок сердца', 'ar'), 
('сердечно-сосудистое заболевание', 'нарушение циркадных ритмов организма', 'ar'), 
('сиаловые кислоты: повышение', 'гемобластоз', 'ar'), 
('сидеропенический синдром', 'железо-дефицитная анемия', 'ar'), 
('сидеропенический синдром', 'латентный дефицит железа', 'ar'), 
('синдром врожденной йодной недостаточности', 'гипотиреоз', 'ar'), 
('синдромм Рея', 'влияние аспирина', 'ar'), 
('снижение репаративных процессов в коже', 'сидеропенический синдром', 'ar'), 
('снижение репаративных процессов в слизистых оболочках', 'сидеропенический синдром', 'ar'), 
('снижение тургора кожи', 'хронический панкреатит', 'ar'), 
('СОЭ: повышение', 'анемия', 'ar'), 
('СОЭ: повышение', 'травма', 'ar'), 
('СОЭ: повышение', 'ожог', 'ar'), 
('СОЭ: повышение', 'беременность', 'ogr'), 
('СОЭ: повышение', 'инфекция', 'ar'), 
('СОЭ: повышение', 'влияние теофиллина', 'ar'), 
('СОЭ: повышение', 'воспалительный процесс', 'ar'), 
('СОЭ: повышение', 'амилоидоз', 'ar'), 
('СОЭ: повышение', 'COVID-19', 'ar'), 
('СОЭ: повышение', 'хронический холецистит', 'ar'), 
('СОЭ: повышение', 'болезни соединительной ткани', 'ar'), 
('СОЭ: повышение', 'менструация', 'ogr'), 
('СОЭ: повышение', 'гипервитаминоз А', 'ar'), 
('СОЭ: повышение', 'инфаркт миокарда', 'ar'), 
('СОЭ: повышение', 'влияние орального контрацептива', 'ar'), 
('СОЭ: повышение', 'онкологическое заболевание', 'ar'), 
('СОЭ: снижение', 'влияние альбумина', 'ar'), 
('СОЭ: снижение', 'влияние пойкилоцитоза', 'ar'), 
('СОЭ: снижение', 'полицитемия', 'ar'), 
('СОЭ: снижение', 'влияние кортикостероида', 'ar'), 
('спленомегалия', 'атрофический аутоиммунный гастрит', 'ar'), 
('спленомегалия', 'фуникулярный миелоз', 'ar'), 
('среднее содержание гемоглобина в эритроците (MCH): норма', 'нормохромная анемия', 'ogr'), 
('среднее содержание гемоглобина в эритроците (MCH): повышение', 'гиперхромная анемия', 'ogr'), 
('среднее содержание гемоглобина в эритроците (MCH): снижение', 'гипохромная анемия', 'ogr'), 
('средний объем тромбоцита (MPV): повышение', 'постгеморрагическая анемия', 'ar'), 
('средний объем тромбоцита (MPV): повышение', 'гемоглобинопатия', 'ar'), 
('средний объем тромбоцита (MPV): повышение', 'влияние частого приема алкоголя', 'ar'), 
('средний объем тромбоцита (MPV): повышение', 'последствие объемной операции', 'ar'), 
('средний объем тромбоцита (MPV): повышение', 'обширная травма', 'ar'), 
('средний объем тромбоцита (MPV): повышение', 'повышенное производство тромбоцитов', 'ar'), 
('средний объем тромбоцита (MPV): повышение', 'талассемия', 'ar'), 
('средний объем тромбоцита (MPV): повышение', 'тромбоцитодистрофия', 'ar'), 
('средний объем тромбоцита (MPV): повышение', 'последствие терапии: лечение анемии', 'ar'), 
('средний объем тромбоцита (MPV): повышение', 'последствие спленэктомии', 'ar'), 
('средний объем тромбоцита (MPV): повышение', 'тиреотоксикоз', 'ar'), 
('средний объем тромбоцита (MPV): повышение', 'идиопатическая тромбоцитопеническая пурпура', 'ar'), 
('средний объем тромбоцита (MPV): повышение', 'миелоидный лейкоз', 'ar'), 
('средний объем тромбоцита (MPV): повышение', 'постгеморрагическое состояние', 'ar'), 
('средний объем тромбоцита (MPV): повышение', 'атеросклероз сосудов', 'ar'), 
('средний объем тромбоцита (MPV): повышение', 'системная красная волчанка', 'ar'), 
('средний объем тромбоцита (MPV): повышение', 'обильная менструация', 'ar'), 
('средний объем тромбоцита (MPV): повышение', 'сахарный диабет', 'ar'), 
('средний объем тромбоцита (MPV): повышение', 'синдром Мея-Хегглина', 'ar'), 
('средний объем эритроцита (MCV): повышение', 'влияние цитостатика', 'ar'), 
('средний объем эритроцита (MCV): повышение', 'RDW-CV: повышение и MCV: повышение', 'or'), 
('средний объем эритроцита (MCV): повышение', 'апластическая анемия', 'ar'), 
('средний объем эритроцита (MCV): повышение', 'гипотиреоз', 'ar'), 
('средний объем эритроцита (MCV): повышение', 'RDW-SD: повышение и MCV: повышение', 'or'), 
('средний объем эритроцита (MCV): повышение', 'алкогольное поражение печени', 'ar'), 
('средний объем эритроцита (MCV): повышение', 'гиповитаминоз B9', 'ar'), 
('средний объем эритроцита (MCV): повышение', 'влияние химиотерапии', 'ar'), 
('средний объем эритроцита (MCV): повышение', 'гиповитаминоз B12', 'ar'), 
('средний объем эритроцита (MCV): повышение', 'нарушение функции печени', 'ar'), 
('средний объем эритроцита (MCV): повышение', 'влияние антибиотика', 'ar'), 
('средний объем эритроцита (MCV): повышение', 'влияние иммунодепрессанта', 'ar'), 
('средний объем эритроцита (MCV): повышение', 'влияние курения', 'ar'), 
('средний объем эритроцита (MCV): повышение', 'влияние противосудорожного препарата', 'ar'), 
('средний объем эритроцита (MCV): повышение', 'витамин B12-дефицитная анемия', 'ar'), 
('средний объем эритроцита (MCV): повышение', 'влияние орального контрацептива', 'ar'), 
('средний объем эритроцита (MCV): снижение', 'сидеропеническая анемия', 'ogr'), 
('средний объем эритроцита (MCV): снижение', 'увеличение концентрации солей плазмы', 'ar'), 
('средний объем эритроцита (MCV): снижение', 'гемоглобинопатия', 'ar'), 
('средний объем эритроцита (MCV): снижение', 'обезвоживание', 'ar'), 
('средний объем эритроцита (MCV): снижение', 'почечная патология', 'ar'), 
('средний объем эритроцита (MCV): снижение', 'токсическое поражение костного мозга (отравление ртутью, свинцом)', 'ar'), 
('средний объем эритроцита (MCV): снижение', 'нарушение обмена железа в организме', 'ar'), 
('средний объем эритроцита (MCV): снижение', 'сидеробластная анемия', 'ar'), 
('средний объем эритроцита (MCV): снижение', 'талассемия', 'ar'), 
('средний объем эритроцита (MCV): снижение', 'дефицит железа', 'ar'), 
('средний объем эритроцита (MCV): снижение', 'эндокринная патология', 'ar'), 
('средний объем эритроцита (MCV): снижение', 'микроцитарная анемия', 'ogr'), 
('средний объем эритроцита (MCV): снижение', 'влияние вегетарианства', 'ar'), 
('средний объем эритроцита (MCV): снижение', 'сниженный синтез гемоглобина', 'ar'), 
('средний объем эритроцита (MCV): снижение', 'наличие паразитов', 'ar'), 
('средний объем эритроцита (MCV): снижение', 'нарушение усвоения железа', 'ar'), 
('средний объем эритроцита (MCV): снижение', 'беременность', 'ogr'), 
('средний объем эритроцита (MCV): снижение', 'кровопотеря', 'ar'), 
('средний объем эритроцита (MCV): снижение', 'железо-дефицитная анемия', 'ogr'), 
('средний объем эритроцита (MCV): снижение', 'RDW-CV: повышение и MCV: снижение', 'or'), 
('средний объем эритроцита (MCV): снижение', 'снижение всасывания железа', 'ar'), 
('средний объем эритроцита (MCV): снижение', 'RDW-SD: повышение и MCV: снижение', 'or'), 
('стеаторея', 'хронический панкреатит', 'ogr'), 
('стеаторея', 'экзокринная недостаточность поджелудочной железы', 'ar'), 
('стеаторея', 'СИБР', 'ar'), 
('стрептококковый сепсис', 'сепсис', 'ar'), 
('субарахноидальное кровоизлияние', 'инсульт', 'ar'), 
('субарахноидальное кровоизлияние из базилярной артерии', 'субарахноидальное кровоизлияние', 'ar'), 
('субарахноидальное кровоизлияние из внутричерепной артерии неуточненной', 'субарахноидальное кровоизлияние', 'ar'), 
('субарахноидальное кровоизлияние из каротидного синуса и бифуркации', 'субарахноидальное кровоизлияние', 'ar'), 
('субарахноидальное кровоизлияние из передней соединительной артерии', 'субарахноидальное кровоизлияние', 'ar'), 
('субарахноидальное кровоизлияние из позвоночной артерии', 'субарахноидальное кровоизлияние', 'ar'), 
('субарахноидальное кровоизлияние из средней мозговой артерии', 'субарахноидальное кровоизлияние', 'ar'), 
('субарахноидальное кровоизлияние неуточненное', 'субарахноидальное кровоизлияние', 'ar'), 
('судорожный синдром', 'синдром Денди-Уокера', 'ar'), 
('сухость слизистой оболочки пищевода', 'сидеропенический синдром', 'ar'), 
('температура тела: от 37 до 38 градусов Цельсия', 'сидеропенический синдром', 'ar'), 
('температура тела: от 37 до 38 градусов Цельсия', 'хронический панкреатит', 'ar'), 
('температура тела: повышение', 'COVID-19', 'ar'), 
('температура тела: повышение', 'воспалительный процесс', 'ar'), 
('тиреоглобулин: повышение', 'рак щитовидной железы фолликулярный', 'ar'), 
('тиреотоксикоз', 'атрофический аутоиммунный гастрит', 'ar'), 
('токсакароз', 'хронический холецистит', 'ar'), 
('тромбиновое время: повышение', 'хронический холецистит', 'ar'), 
('тромбозы сосудов нижних конечностей', 'COVID-19', 'ar'), 
('тромбоциты: повышение', 'беременность (третий триместр)', 'ogr'), 
('тромбоциты: повышение', 'воспалительный процесс', 'ar'), 
('тромбоциты: повышение', 'полицитемия', 'ar'), 
('тромбоциты: повышение', 'аутоиммунное заболевание', 'ar'), 
('тромбоциты: повышение', 'COVID-19', 'ar'), 
('тромбоциты: повышение', 'полицитемия истинная', 'ar'), 
('тромбоциты: повышение', 'острая инфекция', 'ar'), 
('тромбоциты: повышение', 'последствие спленэктомии', 'ar'), 
('тромбоциты: повышение', 'последствие операции', 'ar'), 
('тромбоциты: повышение', 'обострение хронического аутоиммунного заболевания', 'ar'), 
('тромбоциты: повышение', 'почечная недостаточность', 'ar'), 
('тромбоциты: повышение', 'онкологическое заболевание костного мозга', 'ar'), 
('тромбоциты: повышение', 'системная красная волчанка', 'ar'), 
('тромбоциты: повышение', 'ревматоидный артрит', 'ar'), 
('тромбоциты: повышение', 'железо-дефицитная анемия', 'ar'), 
('тромбоциты: повышение', 'хроническая инфекция', 'ar'), 
('тромбоциты: повышение', 'онкологическое заболевание', 'ar'), 
('тромбоциты: повышение', 'анемия', 'ar'), 
('тромбоциты: повышение', 'кровопотеря', 'ar'), 
('тромбоциты: повышение', 'туберкулез', 'ar'), 
('тромбоциты: повышение', 'травма', 'ar'), 
('тромбоциты: снижение', 'COVID-19', 'ar'), 
('тромбоциты: снижение', 'тромбоцитопения', 'ar'), 
('тромбоциты: снижение', 'апластическая анемия', 'ar'), 
('тромбоциты: снижение', 'вирусная инфекция', 'ar'), 
('тромбоциты: снижение', 'влияние аспирина', 'ar'), 
('тромбоциты: снижение', 'увеличение скорости разрушения тромбоцитов', 'ar'), 
('тромбоциты: снижение', 'идиопатическая тромбоцитопеническая пурпура', 'ar'), 
('тромбоциты: снижение', 'увеличение скорости использования тромбоцитов', 'ar'), 
('тромбоциты: снижение', 'уменьшение образования тромбоцитов в костном мозге', 'ar'), 
('тромбоциты: снижение', 'гемофилия', 'ar'), 
('тромбоциты: снижение', 'лекарственная тромбоцитопения', 'ar'), 
('тромбоциты: снижение', 'бактериальная инфекция', 'ar'), 
('тромбоциты: снижение', 'тромбоз почечной вены', 'ar'), 
('тяжелая степень экзокринной недостаточности поджелудочной железы', 'экзокринная недостаточность поджелудочной железы', 'ogr'), 
('увеличение уровня эстрогена', 'овуляция', 'ogr'), 
('увеличение уровня эстрогена', 'беременность', 'ogr'), 
('увеличение уровня эстрогена', 'первая фаза менструального цикла (7-14 день)', 'ogr'), 
('удлинение времени кровотечения', 'влияние аспирина', 'ar'), 
('УЗИ: атрофия поджелудочной железы', 'атрофия поджелудочной железы', 'ar'), 
('УЗИ: ахоличный желчный пузырь', 'хронический холецистит', 's'), 
('УЗИ: гепатомегалия', 'гепатомегалия', 'ar'), 
('УЗИ: дискинезия желчевыводящих путей', 'дискинезия желчевыводящих путей', 'ar'), 
('УЗИ: кальцификация поджелудочной железы', 'кальцификация поджелудочной железы', 'ar'), 
('УЗИ: конкременты в желчном пузыре', 'хронический холецистит', 's'), 
('УЗИ: ложная киста поджелудочной железы', 'ложная киста поджелудочной железы', 'ar'), 
('УЗИ: неравномерность просвета основного протока ПЖ и его боковых ветвей', 'неравномерность просвета основного протока ПЖ и его боковых ветвей', 'ar'), 
('УЗИ: острый холецистит', 'острый холецистит', 'ar'), 
('УЗИ: отек поджелудочной железы', 'отек поджелудочной железы', 'ar'), 
('УЗИ: повышение плотности паренхимы поджелудочной железы', 'повышение плотности паренхимы поджелудочной железы', 'ar'), 
('УЗИ: сморщенный желчный пузырь', 'хронический холецистит', 's'), 
('УЗИ: спленомегалия', 'спленомегалия', 'ar'), 
('УЗИ: увеличенный в размерах желчный пузырь', 'хронический холецистит', 's'), 
('УЗИ: уменьшенный в размерах желчный пузырь', 'хронический холецистит', 's'), 
('УЗИ: утолщение стенки желчного пузыря', 'хронический холецистит', 's'), 
('УЗИ: холангит', 'холангит', 'ar'), 
('УЗИ: холецистит', 'холецистит', 'ar'), 
('УЗИ: хронический панкреатит', 'хронический панкреатит', 'ar'), 
('УЗИ: хронический холецистит', 'хронический холецистит', 'ar'), 
('усиление тяжести covid-19 (предполагается)', 'гиповитаминоз B6', 'ar'), 
('ферритин для женщин: от 180 до 240 мкг/л', 'COVID-19', 'ar'), 
('ферритин для мужчин: от 375 до 500 мкг/л', 'COVID-19', 'ar'), 
('ферритин: повышение', 'влияние частого приема алкоголя', 'ogr'), 
('ферритин: повышение', 'хроническое заболевание печени', 'ar'), 
('ферритин: повышение', 'влияние препарата железа', 'ar'), 
('ферритин: повышение', 'талассемия', 'ar'), 
('ферритин: повышение', 'витамин B12-дефицитная анемия', 'ar'), 
('ферритин: повышение', 'влияние эстрогена', 'ar'), 
('ферритин: повышение', 'гепатоцеллюлярная карцинома', 'ar'), 
('ферритин: повышение', 'лимфома Ходжкина', 'ar'), 
('ферритин: повышение', 'рак предстательной железы', 'ar'), 
('ферритин: повышение', 'тиреотоксикоз', 'ar'), 
('ферритин: повышение', 'воспалительный процесс', 'ar'), 
('ферритин: повышение', 'влияние голодания', 'ar'), 
('ферритин: повышение', 'влияние орального контрацептива', 'ar'), 
('ферритин: повышение', 'рак молочной железы', 'ar'), 
('ферритин: повышение', 'гемобластоз', 'ar'), 
('ферритин: повышение', 'влияние переливаний крови', 'ar'), 
('ферритин: повышение', 'гемолитическая анемия', 'ar'), 
('ферритин: повышение', 'колоректальный рак', 'ar'), 
('ферритин: повышение', 'влияние алкоголя', 'ar'), 
('ферритин: повышение', 'острые заболевания печени', 'ar'), 
('ферритин: повышение', 'гемохроматоз', 'ar'), 
('ферритин: повышение', 'онкологическое заболевание костного мозга', 'ar'), 
('ферритин: повышение', 'аутоиммунное заболевание', 'ar'), 
('ферритин: снижение', 'сидеропенический синдром', 'ar'), 
('ферритин: снижение', 'дефицит железа', 'ar'), 
('ферритин: снижение', 'беременность (третий триместр)', 'ar'), 
('ферритин: снижение', 'прелатентный дефицит железа', 'ar'), 
('фибриноген: повышение', 'COVID-19', 'ar'), 
('фуникулярный миелоз', 'атрофический аутоиммунный гастрит', 'ar'), 
('фуникулярный миелоз', 'гиповитаминоз B12', 'ar'), 
('характеристика поведения: гиподинамия', 'гиподинамия', 'ar'), 
('характеристика поведения: нарушение циркадных ритмов организма', 'снижение способности к концентрации внимания', 'ar'), 
('характеристика поведения: нарушение циркадных ритмов организма', 'снижение объема рабочей памяти', 'ar'), 
('характеристика поведения: нарушение циркадных ритмов организма', 'сахарный диабет 2 типа', 'ar'), 
('характеристика поведения: нарушение циркадных ритмов организма', 'депрессия', 'ar'), 
('характеристика поведения: нарушение циркадных ритмов организма', 'снижение скорости мышления', 'ar'), 
('характеристика поведения: нарушение циркадных ритмов организма', 'поведение нарушающее циркадные ритмы', 'ar'), 
('характеристика поведения: нарушение циркадных ритмов организма', 'снижение способности контролировать импульсивное поведение', 'ar'), 
('характеристика поведения: нарушение циркадных ритмов организма', 'снижение скорости реакции на неожиданные ситуации', 'ar'), 
('характеристика среды: антисанитария', 'лямблиоз', 'ar'), 
('характеристика среды: воздействие радиации', 'радиационный гастрит', 'or'), 
('характеристика среды: путешествия по развивающимся странам', 'лямблиоз', 'ar'), 
('характеристика среды: работа в яслях, детском саду, доме ребенка, центре длительной опеки для умственно отсталых', 'лямблиоз', 'ar'), 
('химический гастрит', 'влияние НПВС', 'ar'), 
('холангит', 'лямблиоз', 'ar'), 
('холестаз', 'рубцы желчных протоков после хирургических вмешательств', 'ar'), 
('холестаз', 'опухоль печени', 'ar'), 
('холестаз', 'камень желчных протоков', 'ar'), 
('холестаз', 'рак желудка с механическим сдавливанием общего желчного протока', 'ar'), 
('холестаз', 'опухоль желчных протоков', 'ar'), 
('холестаз', 'хронический холецистит', 'ar'), 
('холестерин: повышение', 'болезнь Паркинсона', 'ad'), 
('холестерин: повышение', 'хронический холецистит', 'ar'), 
('холестерин: повышение', 'гиповитаминоз B6', 'ar'), 
('холецистит', 'лямблиоз', 'ar'), 
('хроническая болезнь почек', 'почечная недостаточность', 'ar'), 
('хронический панкреатит', 'вирусный панкреатит', 'ar'), 
('хронический панкреатит', 'нарушения всасывания витамина В12', 'ar'), 
('хроническое воспаление', 'воспалительный процесс', 'ar'), 
('цветовой показатель: норма', 'нормохромная анемия', 'ogr'), 
('цветовой показатель: повышение', 'миелодиспластический синдром', 'ar'), 
('цветовой показатель: повышение', 'фолиево-дефицитная анемия', 'ar'), 
('цветовой показатель: повышение', 'витамин B12-дефицитная анемия', 'ar'), 
('цветовой показатель: повышение', 'гиперхромная анемия', 'ogr'), 
('цветовой показатель: снижение', 'лимфома', 'ar'), 
('цветовой показатель: снижение', 'постгеморрагическая анемия', 'ar'), 
('цветовой показатель: снижение', 'лейкоз', 'ar'), 
('цветовой показатель: снижение', 'гипохромная анемия', 'ogr'), 
('цветовой показатель: снижение', 'железо-дефицитная анемия', 'ogr'), 
('целиакия', 'лимфоцитарный гастрит', 'ar'), 
('щелочная фосфатаза: повышение', 'хронический холецистит', 'ar'), 
('ЭГДС: атрофический гастрит', 'атрофический гастрит', 'ar'), 
('ЭГДС: атрофический гастрит антрального отдела желудка', 'атрофический гастрит антрального отдела желудка', 'ar'), 
('ЭГДС: атрофический гастрит дна желудка', 'атрофический гастрит дна желудка', 'ar'), 
('ЭГДС: атрофический гастрит тела желудка', 'атрофический гастрит тела желудка', 'ar'), 
('ЭГДС: атрофия эпителия верхних отделов ЖКТ', 'сидеропенический синдром', 'ar'), 
('ЭГДС: гастрит', 'гастрит', 'ar'), 
('ЭГДС: гастрит антрального отдела желудка', 'неатрофический гастрит', 'ar'), 
('ЭГДС: гастроэзофагеальный рефлюкс', 'гастроэзофагеальный рефлюкс', 'ar'), 
('ЭГДС: гигантский гипертрофический гастрит', 'гигантский гипертрофический гастрит', 'ar'), 
('ЭГДС: гиперпластический полип желудка', 'атрофический аутоиммунный гастрит', 'ar'), 
('ЭГДС: дискинезия желчевыводящих путей', 'дискинезия желчевыводящих путей', 'ar'), 
('ЭГДС: дискинезия желчевыводящих путей', 'лямблиоз', 'ar'), 
('ЭГДС: злокачественное новообразование желудка', 'злокачественное новообразование желудка', 'ar'), 
('ЭГДС: лимфоцитарный гастрит', 'лимфоцитарный гастрит', 'ar'), 
('ЭГДС: пищевод Барретта', 'пищевод Барретта', 'ar'), 
('ЭГДС: пищеводные полукольца', 'пищеводные полукольца', 'ar'), 
('ЭГДС: рак желудка', 'рак желудка', 'ar'), 
('ЭГДС: эрозия 12 п.к.', 'эрозия 12 п.к.', 'ar'), 
('ЭГДС: эрозия желудка', 'эрозия желудка', 'ar'), 
('ЭГДС: эрозия пищевода', 'эрозия пищевода', 'ar'), 
('ЭГДС: язва 12 п.к.', 'язва 12 п.к.', 'ar'), 
('ЭГДС: язва желудка', 'язва желудка', 'ar'), 
('ЭГДС: язва пищевода', 'язва пищевода', 'ar'), 
('ЭКГ: сердечная аритмия', 'сердечная аритмия', 'ar'), 
('ЭКГ: сердечная аритмия', 'гастроэзофагеальный рефлюкс', 'ar'), 
('экзокринная недостаточность поджелудочной железы', 'хронический панкреатит', 'ar'), 
('экзокринная недостаточность поджелудочной железы', 'панкреатит', 'ar'), 
('эластаза 1 в кале: более 200 мкг/г', 'нормальная экзокринная функция поджелудочной железы', 'ogr'), 
('эластаза 1 в кале: до 100 мкг/г', 'тяжелая степень экзокринной недостаточности поджелудочной железы', 'ogr'), 
('эластаза 1 в кале: от 100 до 200 мкг/г', 'легкая степень экзокринной недостаточности поджелудочной железы', 'ogr'), 
('эластаза в сыворотке: повышение', 'хронический панкреатит', 'ogr'), 
('эозинофилы: отсутствие', 'COVID-19', 'ar'), 
('эритропоэтин: снижение', 'хроническое заболевание почек', 'ar'), 
('эритроциты: повышение', 'последствие професии летчика', 'ar'), 
('эритроциты: повышение', 'полицитемия', 'ar'), 
('эритроциты: повышение', 'синдром Кушинга', 'ar'), 
('эритроциты: повышение', 'сердечная недостаточность', 'ar'), 
('эритроциты: повышение', 'полицитемия истинная', 'ar'), 
('эритроциты: повышение', 'обезвоживание', 'ar'), 
('эритроциты: повышение', 'длительное наложение жгута во время взятия крови', 'ar'), 
('эритроциты: повышение', 'влияние кортикостероида', 'ar'), 
('эритроциты: повышение', 'дыхательная недостаточность', 'ar'), 
('эритроциты: повышение', 'влияние курения', 'ar'), 
('эритроциты: повышение', 'хроническая гипоксия', 'ar'), 
('эритроциты: повышение', 'хроническая обструктивная болезнь легких', 'ar'), 
('эритроциты: повышение', 'хроническая сердечная недостаточность', 'ar'), 
('эритроциты: повышение', 'водянка почечных лоханок', 'ar'), 
('эритроциты: повышение', 'влияние диуретика', 'ar'), 
('эритроциты: повышение', 'влияние проживания в высокогорье', 'ar'), 
('эритроциты: повышение', 'влияние регулярной физической активности', 'ar'), 
('эритроциты: повышение', 'талассемия', 'ar'), 
('эритроциты: снижение', 'влияние недостаточного употребления железа', 'ar'), 
('эритроциты: снижение', 'беременность', 'ogr'), 
('эритроциты: снижение', 'врожденный дефект кроветворения', 'ar'), 
('эритроциты: снижение', 'витамин B12-дефицитная анемия', 'ar'), 
('эритроциты: снижение', 'постгеморрагическая анемия', 'ar'), 
('эритроциты: снижение', 'гемолитическая анемия', 'ar'), 
('эритроциты: снижение', 'метастазы в костный мозг', 'ar'), 
('эритроциты: снижение', 'сидеропенический синдром', 'ar'), 
('эритроциты: снижение', 'влияние вегетарианства', 'ar'), 
('эритроциты: снижение', 'влияние недостаточного употребления витаминов', 'ar'), 
('эритроциты: снижение', 'апластическая анемия', 'ar'), 
('эритроциты: снижение', 'хроническое заболевание почек', 'ar'), 
('эритроциты: снижение', 'микседема', 'ar'), 
('эритроциты: снижение', 'дефицит белка', 'ar'), 
('эритроциты: снижение', 'цирроз печени', 'ar'), 
('эритроциты: снижение', 'лимфома', 'ar'), 
('эритроциты: снижение', 'лейкоз', 'ar'), 
('эритроциты: снижение', 'онкологическое заболевание костного мозга', 'ar'), 
('эритроциты: снижение', 'влияние пентоксифиллина', 'ar'), 
('эритроциты: снижение', 'хроническая инфекция', 'ar'), 
('эритроциты: снижение', 'железо-дефицитная анемия', 'ar'), 
('эритроциты: снижение', 'влияние инфузионной терапии', 'ar'), 
('эритроциты: снижение', 'влияние гентамицина', 'ar'), 
('эритроциты: снижение', 'кровопотеря', 'ar'), 
('эритроциты: снижение', 'гипергидратация', 'ar'), 
('эритроциты: снижение', 'системное заболевание соединительной ткани', 'ar'), 
('эритроциты: снижение', 'анемия', 'ogr'), 
('эритроциты: снижение', 'острое кровотечение', 'ar'), 
('эрозивно-язвенные поражения ЖКТ', 'влияние аспирина', 'ar'), 
('эссенциальная гипертензия', 'артериальная гипертензия', 'ar'), 
('эстрадиол: повышение', 'увеличение уровня эстрогена', 'ar'), 
('эстриол: повышение', 'увеличение уровня эстрогена', 'ar'), 
('эстрон: повышение', 'увеличение уровня эстрогена', 'ar'), 
('язвенный колит', 'воспалительные заболевания кишечника', 'ar'), 
}

ds8mR = {
'COVID-19', 
'COVID-19, вирус идентифицирован', 
'COVID-19, вирус не идентифицирован', 
'FFMI: норма', 
'FFMI: повышение', 
'FFMI: снижение', 
'spina bifida у плода', 
'абсцесс легкого', 
'авитаминоз', 
'агенезия мозолистого тела', 
'агенезия червя мозжечка', 
'агранулоцитоз', 
'адгезивный отит', 
'аденокарцинома легкого', 
'аденокарцинома предстательной железы', 
'аденома простаты', 
'аденомиоз', 
'аднексит', 
'акне', 
'акромегалия', 
'алкогольное поражение печени', 
'аллергическая бронхиальная астма', 
'аллергический бронхит', 
'аллергический васкулит', 
'аллергический ринит', 
'аллергический тонзиллит', 
'аллергическое заболевание', 
'аллергическое заболевание глаз', 
'аллергия', 
'аллергия на еду', 
'альгодисменорея', 
'аменорея', 
'амилоидоз', 
'анальная трещина', 
'анальный зуд', 
'анальный полип', 
'анапластический рак щитовидной железы', 
'анацидный гастрит', 
'ангина', 
'ангиопатия сетчатки', 
'ангулярный стоматит', 
'аневризма', 
'аневризма легочной артерии', 
'анемия', 
'анемия вследствие нарушения синтеза гемоглобина или эритроцитов', 
'анемия легкой степени', 
'анемия средней степени', 
'анемия тяжелой степени', 
'анемия, вызванная хроническим воспалением', 
'аномалия мочевыводящих путей плода', 
'аномалия почек плода', 
'аномалия развития лицевого черепа', 
'аномалия развития пальцев', 
'аномалия развития сердца', 
'анэнцефалия у плода', 
'апластическая анемия', 
'аппендицит', 
'артериальная гипертензия', 
'артрит', 
'артроз коленного сустава', 
'асептический менингит', 
'аскаридоз', 
'астеновегетативный синдром', 
'атаксия', 
'атаксия височно-затылочной области', 
'атаксия корково-мозжечковых путей', 
'атаксия лобной области', 
'атаксия теменной области', 
'атаксия-телеангиэктазия', 
'атерогенная дислипидемия', 
'атеросклероз сосудов', 
'атопический дерматит', 
'атрезия 12-типерстной кишки у плода', 
'атрезия пищевода у плода', 
'атрофический аутоиммунный гастрит', 
'атрофический гастрит', 
'атрофический гастрит антрального отдела желудка', 
'атрофический гастрит дна желудка', 
'атрофический гастрит тела желудка', 
'атрофический мультифокальный гастрит', 
'атрофический ринит', 
'атрофический энтерит', 
'атрофия зрительного нерва', 
'атрофия поджелудочной железы', 
'атрофия сосочков языка', 
'атрофия щитовидной железы (приобретенная)', 
'аутоиммунная гемолитическая анемия', 
'аутоиммунное заболевание', 
'аутоиммунный тиреоидит', 
'аутоимунное заболевание', 
'ахлоргидрия желудочного сока', 
'ацидоз', 
'бактериальная инфекция', 
'баланопостит', 
'беременность', 
'беременность (третий триместр)', 
'бессимптомная ишемия миокарда', 
'блефарит', 
'болезни крови', 
'болезни мочевыделительной системы', 
'болезни соединительной ткани', 
'болезни щитовидной железы, связанные с йодной недостаточностью и сходные состояния', 
'болезнь Крона', 
'болезнь Меньера', 
'болезнь Паркинсона', 
'болезнь Пейрони', 
'болезнь перегородки носа', 
'болезнь холодовых агглютининов', 
'бородавки', 
'бронхиальная астма', 
'бронхит', 
'бронхоспазм', 
'бронхоэктатическая болезнь', 
'вагинит (кольпит)', 
'вазомоторный ринит', 
'вальгусная деформация первых пальцев стоп', 
'варикозное расширение вен конечностей', 
'варикоцеле', 
'везикулит', 
'вестибулярная атаксия', 
'вирус папилломы человека', 
'вирусная инфекция', 
'вирусный отит', 
'вирусный панкреатит', 
'витамин B12-дефицитная анемия', 
'витилиго', 
'влияние H-2 гистаминоблокатора', 
'влияние алкоголя', 
'влияние альбумина', 
'влияние анорексии', 
'влияние антибиотика', 
'влияние антидепрессанта', 
'влияние аскорбиновой кислоты', 
'влияние аспирина', 
'влияние барбитурата', 
'влияние бета-блокатора', 
'влияние булимии', 
'влияние вегетарианства', 
'влияние гентамицина', 
'влияние голодания', 
'влияние гормонального препарата', 
'влияние диуретика', 
'влияние иммунодепрессанта', 
'влияние инфузионной терапии', 
'влияние ИПП', 
'влияние кортикостероида', 
'влияние курения', 
'влияние недавнего приема пищи', 
'влияние недавней прививки', 
'влияние недавней физической нагрузки', 
'влияние недостаточного употребления белка', 
'влияние недостаточного употребления витаминов', 
'влияние недостаточного употребления железа', 
'влияние НПВС', 
'влияние орального контрацептива', 
'влияние парацетамола', 
'влияние пентоксифиллина', 
'влияние переливаний крови', 
'влияние пойкилоцитоза', 
'влияние препарата железа', 
'влияние препарата подавляющего желудочную секрецию', 
'влияние приема алкоголя', 
'влияние проживания в высокогорье', 
'влияние противогрибкового препарата', 
'влияние противосудорожного препарата', 
'влияние радиации', 
'влияние регулярной физической активности', 
'влияние резкого прекращение приема ИПП', 
'влияние сниженных эритроцитов', 
'влияние статина', 
'влияние теофиллина', 
'влияние тестостерона', 
'влияние тубазида', 
'влияние фенобарбитала', 
'влияние фтивазида', 
'влияние химиотерапии', 
'влияние цефалоспорина', 
'влияние цитостатика', 
'влияние частого приема алкоголя', 
'влияние эстрогена', 
'внешнесекреторная недостаточность поджелудочной железы', 
'внутрижелудочковая блокада', 
'внутримозговое кровоизлияние', 
'внутримозговое кровоизлияние в мозжечок', 
'внутримозговое кровоизлияние в полушарие кортикальное', 
'внутримозговое кровоизлияние в полушарие неуточненное', 
'внутримозговое кровоизлияние в полушарие субкортикальное', 
'внутримозговое кровоизлияние в ствол мозга', 
'внутримозговое кровоизлияние внутрижелудочковое', 
'внутримозговое кровоизлияние множественной локализации', 
'внутримозговое кровоизлияние неуточненное', 
'внутрипредсердная блокада', 
'внутрипротоковая папиллома молочной железы', 
'внутричерепное кровоизлияние (нетравматическое) неуточненное', 
'водянка почечных лоханок', 
'воздействие тяжелых металлов (ртуть, свинец)', 
'возрастная макулярная дегенерация сетчатки', 
'воспаление органов малого таза', 
'воспаления придатков матки', 
'воспалительные заболевания кишечника', 
'воспалительные заболевания тонкого кишечника', 
'воспалительный процесс', 
'врожденная гломерулопатия', 
'врожденная тирозинемия', 
'врожденный гипотиреоз без зоба', 
'врожденный гипотиреоз с диффузным зобом', 
'врожденный дефект кроветворения', 
'врожденный нефротический синдром', 
'вторичная гипертензия', 
'вторичная непереносимость лактозы', 
'вторичная полицитемия', 
'выпадение прямой кишки', 
'выраженный бронхит', 
'выраженный воспалительный процесс', 
'выраженный дефицит массы тела', 
'высокая степень агрегации тромбоцитов', 
'высокий риск летального исхода', 
'высокий риск тяжелого сепсиса и/или септического шока', 
'высокий сердечно-сосудистый риск', 
'гайморит', 
'гангрена нижних конечностей', 
'гастрит', 
'гастроэзофагеальная рефлюксная болезнь', 
'гастроэзофагеальный рефлюкс без эзофагита', 
'гастроэзофагеальный рефлюкс с эзофагитом', 
'гемангиома', 
'гемобластоз', 
'гемоглобинопатия', 
'гемоконцентрация', 
'гемолитическая анемия', 
'геморрагический синдром', 
'геморрой', 
'гемофилия', 
'гемохроматоз', 
'генетическое нарушение', 
'генитальный герпес', 
'гепатоцеллюлярная карцинома', 
'герминогенная опухоль яичника', 
'гибель плода', 
'гигантизм', 
'гигантоклеточный рак легкого', 
'гигантский гипертрофический гастрит', 
'гидронефроз', 
'гингивит', 
'гинекомастия', 
'гипервитаминоз B6', 
'гипервитаминоз А', 
'гипергидратация', 
'гиперлипидемия', 
'гиперпаратиреоз', 
'гиперпролактинемия', 
'гипертензивная болезнь с преимущественным поражением почек', 
'гипертензивная болезнь с преимущественным поражением сердца и почек', 
'гипертензивная болезнь сердца', 
'гипертензионная гидроцефалия', 
'гипертрофия миндалин', 
'гиперхлоргидрия желудочного сока', 
'гиперхромная анемия', 
'гиповитаминоз B12', 
'гиповитаминоз B6', 
'гиповитаминоз B9', 
'гиповитаминоз D', 
'гиповитаминоз E', 
'гиповитаминоз K', 
'гиповитаминоз А', 
'гипогликемический синдром', 
'гипоксия', 
'гипопаратиреоз', 
'гипоспадия', 
'гипотиреоз', 
'гипотиреоз неуточненный', 
'гипотиреоз, вызванный медикаментами и другими экзогенными веществами', 
'гипотония', 
'гипофосфатазия', 
'гипохлоргидрия желудочного сока', 
'гипохлоргидрия желудочного сока легкая', 
'гипохромная анемия', 
'глаукома', 
'глаукоматозная атрофия зрительного нерва', 
'глисты', 
'гломерулонефрит', 
'гломерулярное заболевание почек', 
'глоссит', 
'гормонально-активные опухоли надпочечников', 
'гранулематоз Вегенера', 
'гранулематозный гастрит', 
'грибковая инфекция', 
'грибковый отит', 
'грипп', 
'грыжа межпозвонкового диска', 
'дальнозоркость', 
'деменция', 
'демодекоз', 
'депрессия', 
'дерматит', 
'дефицит IgA', 
'дефицит белка', 
'дефицит железа', 
'дефицит инсулина', 
'децидуальный полип', 
'диабет', 
'диабетическая ретинопатия', 
'дивертикул пищевода', 
'дивертикулярная болезнь толстой кишки', 
'динамическая атаксия', 
'дисбактериоз', 
'дисбиоз тонкого кишечника', 
'дискинезия желчевыводящих путей', 
'дискинезия нисходящего отдела 12 п.к.', 
'дисплазия шейки матки', 
'диспротеинемия', 
'диффузный гнойный наружный отит', 
'диффузный токсический зоб', 
'длительное наложение жгута во время взятия крови', 
'доброкачественное новообразование желез надпочечника', 
'другие инфекционные гастриты (не НР)', 
'дуоденит', 
'дуоденогастральный рефлюкс', 
'дыхательная недостаточность', 
'железисто-фиброзный полип эндометрия', 
'железо-дефицитная анемия', 
'железо-дефицитная анемия неуточненная', 
'желудочная диспепсия', 
'желчнокаменная болезнь', 
'заболевание костного мозга', 
'заболевание крови', 
'заболевания внутреннего уха', 
'заболевания крови', 
'заболевания шейки матки', 
'завышенный срок беременности', 
'загиб матки', 
'задержка развития плода', 
'закрытая черепно-мозговая травма', 
'закупорка и стеноз базилярной артерии', 
'закупорка и стеноз других прецеребральных артерий', 
'закупорка и стеноз множественных и двусторонних прецеребральных артерий', 
'закупорка и стеноз неуточненной прецеребральной артерии', 
'закупорка и стеноз позвоночной артерии', 
'закупорка и стеноз прецеребральных артерий, не приводящие к инфаркту мозга', 
'закупорка и стеноз сонной артерии', 
'запор', 
'злокачественное новообразование', 
'злокачественное новообразование бронхов', 
'злокачественное новообразование бронхов и легкого', 
'злокачественное новообразование желудка', 
'злокачественное новообразование желчного пузыря', 
'злокачественное новообразование легкого', 
'злокачественное новообразование мозгового слоя надпочечника', 
'злокачественное новообразование молочной железы', 
'злокачественное новообразование мочевого пузыря', 
'злокачественное новообразование ободочной кишки', 
'злокачественное новообразование периферических нервов и вегетативной нервной системы', 
'злокачественное новообразование поджелудочной железы', 
'злокачественное новообразование почки', 
'злокачественное новообразование предстательной железы', 
'злокачественное новообразование прямой кишки', 
'злокачественное новообразование сетчатки глаза', 
'злокачественное новообразование тела матки', 
'злокачественное новообразование толстой кишки', 
'злокачественное новообразование щитовидной железы', 
'злокачественное новообразование яичка', 
'идиопатическая тромбоцитопеническая пурпура', 
'избыточная масса тела (предожирение)', 
'иммунодефицит', 
'иммунодифецит', 
'инородное тело желудка', 
'инсулин-зависимый диабет беременной', 
'инсульт', 
'инсульт, не уточненный как кровоизлияние или инфаркт', 
'инфаркт миокарда', 
'инфаркт мозга', 
'инфаркт мозга неуточненный', 
'инфаркт мозга, вызванный неуточненной закупоркой или стенозом мозговых артерий', 
'инфаркт мозга, вызванный неуточненной закупоркой или стенозом прецеребральных артерий', 
'инфаркт мозга, вызванный тромбозом вен мозга, непиогенный', 
'инфаркт мозга, вызванный тромбозом мозговых артерий', 
'инфаркт мозга, вызванный тромбозом прецеребральных артерий', 
'инфаркт мозга, вызванный эмболией мозговых артерий', 
'инфаркт мозга, вызванный эмболией прецеребральных артерий', 
'инфекционный мононуклеоз', 
'инфекция', 
'искривление носовой перегородки', 
'искривление полового члена', 
'ихтиоз', 
'ишемическая болезнь сердца', 
'ишемический колит', 
'кавернит', 
'кальциноз перикарда', 
'камень желчных протоков', 
'кандидоз', 
'кандидозный глоссит', 
'кариес', 
'карциноид', 
'карцинома островковых клеток поджелудочной железы', 
'катаракта', 
'кератит', 
'киста задней черепной ямки', 
'киста молочной железы', 
'киста почки', 
'киста придатка яичка', 
'киста яичника', 
'киста яичника в постменопаузе', 
'киста яичника геморрагическая', 
'киста яичника дермоидная', 
'киста яичника лютеиновая', 
'киста яичника муцинозная', 
'киста яичника ретенционная', 
'киста яичника серозная', 
'киста яичника фолликулярная', 
'киста яичника эндометриоидная', 
'кифосколиотическая болезнь сердца', 
'кишечная диспепсия', 
'кокцигодиния', 
'колликулит', 
'колоректальный рак', 
'кома', 
'кондилома', 
'контагиозный моллюск', 
'конъюнктивит', 
'корковая атаксия', 
'коронарный тромбоз', 
'короткая уздечка полового члена', 
'крапивница', 
'красная волчанка', 
'криптит', 
'кровопотеря', 
'кровотечение', 
'кровотечение из ЖКТ', 
'крупноклеточный рак легкого', 
'крупный плод', 
'ксантелазма', 
'ксероз', 
'лабиринтит', 
'латентный дефицит железа', 
'левожелудочковая недостаточность', 
'легкая степень экзокринной недостаточности поджелудочной железы', 
'легочная гипертензия', 
'лейкоз', 
'лейкопения', 
'лейкоплакия мочевого пузыря', 
'лекарственная аллергия', 
'лекарственная тромбоцитопения', 
'лимфолейкоз', 
'лимфома', 
'лимфома Ходжкина', 
'лимфоцитарный гастрит', 
'липоидный нефроз', 
'липома перианальной области', 
'листовидная фиброаденома молочной железы', 
'лихорадка', 
'ложная беременность', 
'ложная киста поджелудочной железы', 
'лучевая болезнь', 
'люмбаго', 
'лямблиоз', 
'макулодистрофия влажная', 
'макулодистрофия сухая', 
'мальабсорбция', 
'мастит', 
'мастоидит', 
'мастопатия', 
'маточное кровотечение', 
'медуллярный рак щитовидной железы', 
'меланома', 
'мелкоклеточный рак легкого', 
'менингит', 
'менструация', 
'метастазы', 
'метастазы в костный мозг', 
'метастазы в печень', 
'метгемоглобинемия', 
'мигрень', 
'миелодиспластический синдром', 
'миелоидный лейкоз', 
'миелома', 
'миелофиброз', 
'микоплазмоз', 
'микроцитарная анемия', 
'микседема', 
'микседематозная кома', 
'миозит', 
'миокардит', 
'миома матки', 
'миома матки субмукозная', 
'миома матки субсерозная', 
'миома матки узловая', 
'миофасциальный синдром', 
'многоплодная беременность', 
'множественные травмы', 
'мозжечковая атаксия', 
'мозжечковая симптоматика', 
'молочница', 
'мочекаменная болезнь', 
'навязчивые состояния', 
'наличие Hp на слизистой желудка', 
'наличие Hp на слизистой желудка в большом количестве', 
'наличие Hp на слизистой желудка в небольшом количестве', 
'наличие Hp на слизистой желудка в среднем количестве', 
'наличие микросгустков в крови', 
'наличие паразитов', 
'наружный отит', 
'нарушение гормонального фона', 
'нарушение кровообращения', 
'нарушение обмена железа в организме', 
'нарушение осанки', 
'нарушение пищевого поведения', 
'нарушение свертывания крови', 
'нарушение секреции гастрина', 
'нарушение сна', 
'нарушение усвоения железа', 
'нарушение функции печени', 
'нарушение функции почек', 
'нарушение функций яичников', 
'нарушение циркадных ритмов организма', 
'нарушения всасывания витамина В12', 
'нарушения менструальной функции', 
'наследственная коагулопатия', 
'наследственное нарушение эритроцитов', 
'неатрофический гастрит', 
'невралгия', 
'неврит', 
'невроз', 
'недоразвитие мозолистого тела', 
'недоразвитие червя мозжечка', 
'недостаточная (дефицит) масса тела', 
'недостаточное поступление белка', 
'недостаточное усвоение белка', 
'недостаточность кровообращения головного мозга', 
'некроз печени плода вследствие вирусной инфекции у плода', 
'немелкоклеточный рак легкого', 
'неревматические болезни сердца', 
'несахарный диабет', 
'неспецифический колит', 
'неспецифический уретрит', 
'несращение передней брюшной стенки плода', 
'нетравматическое внутричерепное кровоизлияние', 
'нетравматическое субдуральное кровоизлияние', 
'нетравматическое экстрадуральное кровоизлияние', 
'нефрогенный несахарный диабет', 
'нефропатия', 
'нефроптоз', 
'нефротический синдром', 
'низкий риск тяжелого сепсиса и/или септического шока', 
'низкий сердечно-сосудистый риск', 
'нисходящая атрофия зрительного нерва', 
'нормальная масса тела', 
'нормальная экзокринная функция поджелудочной железы', 
'нормохромная анемия', 
'обезвоживание', 
'обильная менструация', 
'обморок', 
'обострение хронического аутоиммунного заболевания', 
'обширная травма', 
'обширные травмы', 
'обширные хирургические вмешательства', 
'овсяноклеточный рак легкого', 
'овуляция', 
'ожирение', 
'ожирение 1 степени', 
'ожирение 2 степени', 
'ожирение 3 степени', 
'ожирение беременной', 
'ожог', 
'онкологическое заболевание', 
'онкологическое заболевание костного мозга', 
'операции в условиях искусственного кровообращения', 
'описторхоз', 
'опухолевый некроз', 
'опухоль', 
'опухоль бронхов', 
'опухоль глотки', 
'опухоль головы и шеи', 
'опухоль гортани', 
'опухоль желчных протоков', 
'опухоль кожи лица', 
'опухоль носа', 
'опухоль носа и околоносовых пазух', 
'опухоль околоносовых пазух', 
'опухоль печени', 
'опухоль плевры', 
'опухоль толстой кишки', 
'опухоль уха', 
'опухоль челюсти', 
'ОРВИ', 
'ОРЗ', 
'орхит', 
'оспа', 
'остановка сердца', 
'остеомаляция', 
'остеопороз', 
'остеохондроз', 
'острая инфекция', 
'острая почечная недостаточность', 
'острая сердечная недостаточность', 
'острая фаза инфекции', 
'острое воспаление', 
'острое кровотечение', 
'острое легочное сердце', 
'острое отторжение трансплантата', 
'острые заболевания печени', 
'острый гепатит', 
'острый гнойный отит', 
'острый инфаркт миокарда', 
'острый ларингит', 
'острый миелоидный лейкоз', 
'острый отит', 
'острый панкреатит', 
'острый ринит', 
'острый серозный отит', 
'острый синусит', 
'острый тубоотит', 
'острый фарингит', 
'острый холецистит', 
'отек Квинке', 
'отомикоз', 
'панические атаки', 
'папиллома', 
'папиллярный рак щитовидной железы', 
'паразитарное заболевание', 
'паралич гортани', 
'пародонтит', 
'пародонтоз', 
'пароксизмальная тахикардия', 
'первая фаза менструального цикла (7-14 день)', 
'первичная гепатоцеллюлярная карцинома', 
'первичная непереносимость лактозы', 
'первичный билиарный цирроз', 
'первичный склерозирующий холангит', 
'передний сухой ринит', 
'перикардит', 
'периодонтит', 
'перитонит', 
'периферическая дистрофия сетчатки', 
'периферическая нейропатия', 
'периферический коллапс', 
'пернициозное анемическое состояние', 
'пиелонефрит', 
'пищевод Барретта', 
'плацентарный полип', 
'плеврит', 
'плеоморфный рак легкого', 
'плечелопаточный синдром', 
'плоскоклеточный рак легкого', 
'плоскостопие', 
'пневмония', 
'повторный инфаркт миокарда', 
'повышенная потеря белка организмом', 
'повышенное производство тромбоцитов', 
'повышеное образование сгустков крови', 
'подострый тиреоидит', 
'полиартрит', 
'поликистоз почек', 
'полиневропатия', 
'полип голосовой связки', 
'полип перегородки носа', 
'полип прямой кишки', 
'полип толстой кишки', 
'полип уретры', 
'полипоз верхнечелюстной пазухи', 
'полипоз клиновидной пазухи', 
'полипоз лобной пазухи', 
'полипоз носа', 
'полипоз пазухи решетчатого лабиринта', 
'полиурия', 
'полицитемия', 
'полицитемия истинная', 
'помутнение сознания', 
'порок развития нервной трубки плода', 
'порок сердца', 
'последствие гастрэктомии', 
'последствие маммарно-коронарного шунтирования', 
'последствие объемной операции', 
'последствие операции', 
'последствие операции на диафрагме', 
'последствие операции на печени', 
'последствие операции по поводу рака легкого с расширенной лимфаденэктомией', 
'последствие професии летчика', 
'последствие спленэктомии', 
'последствие терапии: лечение анемии', 
'последствие тимэктомии', 
'последствие тиреоидэктомии', 
'последствие травм головы и позвоночника', 
'последствие удаления опухолей средостения', 
'поставить диагноз сепсиса с уверенностью нельзя, рекомендуется повторить измерение прокальцитонина в течение 6-24 часов', 
'постгеморрагическая анемия', 
'постгеморрагическое восстановление кроветворения', 
'постгеморрагическое состояние', 
'постинфарктный синдром', 
'постинфекционный гипотиреоз', 
'посттромботическая болезнь', 
'почечная недостаточность', 
'почечная недостаточность неуточненная', 
'почечная патология', 
'правожелудочковая недостаточность', 
'преддиабет', 
'предменструальный синдром', 
'предсердно-желудочковая блокада', 
'предсердно-желудочковая диссоциация', 
'преждевременная деполяризация желудочков', 
'преждевременная деполяризация предсердий', 
'прелатентный дефицит железа', 
'приобретенная гломерулопатия', 
'приобретенная коагулопатия', 
'проктит', 
'пролактинома', 
'прорезывание молочных зубов', 
'псевдомембранозный колит', 
'психоз', 
'псориаз', 
'птоз верхнего века', 
'пузырный занос', 
'пульпит', 
'пупочная грыжа у плода', 
'радиационный гастрит', 
'радикулит', 
'рак бронхов', 
'рак глотки', 
'рак головки поджелудочной железы', 
'рак гортани', 
'рак губы', 
'рак желудка', 
'рак желудка с механическим сдавливанием общего желчного протока', 
'рак желчного пузыря', 
'рак кожи', 
'рак костей', 
'рак легкого', 
'рак матки', 
'рак молочной железы', 
'рак мочевого пузыря', 
'рак ободочной кишки', 
'рак параназальных синусов', 
'рак пищевода', 
'рак поджелудочной железы', 
'рак полости носа', 
'рак полости рта', 
'рак почки', 
'рак предстательной железы', 
'рак прямой кишки', 
'рак слюнных желез', 
'рак тела матки', 
'рак толстого кишечника', 
'рак уретры', 
'рак шейки матки', 
'рак шейного отдела пищевода', 
'рак щитовидной железы', 
'рак щитовидной железы фолликулярный', 
'рак языка', 
'рак яичка', 
'рак яичников', 
'ранний признак анемии', 
'расстройства памяти', 
'расщелина позвоночника плода', 
'реактивный артрит', 
'ревматический порок сердца', 
'ревматоидный артрит', 
'регенеративные процессы в печени', 
'ректоцеле', 
'релаксация диафрагмы', 
'ретикулярный варикоз', 
'ринит', 
'риносинусит', 
'ринофима', 
'риск тромбоза', 
'розацеа', 
'рубцы желчных протоков после хирургических вмешательств', 
'самопроизвольный выкидыш', 
'саркоидоз', 
'саркома', 
'сахарный диабет', 
'сахарный диабет 1 типа', 
'сахарный диабет 2 типа', 
'сахарный диабет 3 типа', 
'светлоклеточный рак легкого', 
'свищ прямой кишки', 
'себорея', 
'секреторная недостаточность желудка', 
'семейный эритроцитоз', 
'сенситивная атаксия', 
'сенсоневральная тугоухость', 
'сепсис', 
'сердечная аритмия', 
'сердечная недостаточность', 
'сердечно-легочная недостаточность', 
'сердечно-сосудистое заболевание', 
'серозная карцинома яичника', 
'серопозитивный ревматоидный артрит', 
'серповидно-клеточная анемия', 
'СИБР', 
'сидеробластная анемия', 
'сидеропеническая анемия', 
'сидеропеническая дисфагия', 
'сидеропенический синдром', 
'симптоматическая гипертензия', 
'синдром Вискотта-Олдрича', 
'синдром врожденной йодной недостаточности', 
'синдром Дауна у плода', 
'синдром Денди-Уокера', 
'синдром диссеминированного внутрисосудистого свертывания', 
'синдром Кушинга', 
'синдром Меккеля у плода', 
'синдром Мея-Хегглина', 
'синдром Мэллори-Вейсса', 
'синдром Патау', 
'синдром передней лестничной мышцы', 
'синдром полиорганной недостаточности', 
'синдром преждевременного возбуждения', 
'синдром раздраженного кишечника', 
'синдром системного воспалительного ответа неинфекционного происхождения с органной недостаточностью', 
'синдром слабости синусового узла', 
'синдром Стокса-Адамса', 
'синдром хронической усталости', 
'синдром Шерешевского-Тернера плода', 
'синдромм Рея', 
'синоаурикулярная блокада', 
'системная красная волчанка', 
'системное аутоиммунное заболевание', 
'системное заболевание соединительной ткани', 
'склерома', 
'смешанный рак легкого', 
'снижение всасывания железа', 
'снижение объема рабочей памяти', 
'снижение репаративных процессов в коже', 
'снижение скорости мышления', 
'снижение скорости реакции на неожиданные ситуации', 
'снижение способности к концентрации внимания', 
'снижение способности контролировать импульсивное поведение', 
'сниженный синтез гемоглобина', 
'спазм аккомодации', 
'средний отит', 
'средний сердечно-сосудистый риск', 
'статическая атаксия', 
'стеноз гортани', 
'стенозирующий лигаментит', 
'стенокардия', 
'стрептококковый сепсис', 
'стресс', 
'субарахноидальное кровоизлияние', 
'субарахноидальное кровоизлияние из базилярной артерии', 
'субарахноидальное кровоизлияние из внутричерепной артерии неуточненной', 
'субарахноидальное кровоизлияние из задней соединительной артерии', 
'субарахноидальное кровоизлияние из каротидного синуса и бифуркации', 
'субарахноидальное кровоизлияние из передней соединительной артерии', 
'субарахноидальное кровоизлияние из позвоночной артерии', 
'субарахноидальное кровоизлияние из средней мозговой артерии', 
'субарахноидальное кровоизлияние неуточненное', 
'субклинический гипотиреоз вследствие йодной недостаточности', 
'судорожный синдром', 
'сфеноидит', 
'талассемия', 
'тампонада перикарда', 
'телеангиоэктазия', 
'тератобластома яичек', 
'тератобластома яичников', 
'тиреотоксикоз', 
'тиреотоксическая аденома', 
'токсакароз', 
'токсическая атрофия зрительного нерва', 
'токсическое поражение костного мозга (отравление ртутью, свинцом)', 
'травма', 
'травма печени', 
'травма уха', 
'травматическая атрофия зрительного нерва', 
'трахеит', 
'тревожные расстройства', 
'трепетание желудочков', 
'трещина на стопе', 
'трисомия по 18 хромосоме у плода', 
'трихомониаз', 
'тромбоз глубоких вен', 
'тромбоз почечной вены', 
'тромбозы сосудов нижних конечностей', 
'тромбофлебит', 
'тромбоцитодистрофия', 
'тромбоцитопатия', 
'тромбоцитопения', 
'туберкулез', 
'туннельный синдром', 
'тучноклеточный лейкоз', 
'тяжелая степень экзокринной недостаточности поджелудочной железы', 
'тяжелый бактериальный сепсис или септический шок', 
'тяжелый воспалительный процесс', 
'тяжелый ожог', 
'увеит', 
'увеличение концентрации солей плазмы', 
'увеличение скорости использования тромбоцитов', 
'увеличение скорости разрушения тромбоцитов', 
'увеличение уровня эстрогена', 
'увеличивается риск развития рака легких у мужчин на 30-40%', 
'угроза самопроизвольного выкидыша', 
'удлинение времени кровотечения', 
'узелок голосовой связки', 
'узловой зоб', 
'уменьшение образования тромбоцитов в костном мозге', 
'умеренный воспалительный процесс', 
'усиление тяжести covid-19 (предполагается)', 
'усиленное потребление белка организмом', 
'усиленный распад белка', 
'фарингомикоз', 
'феохромоцитома', 
'фибрилляция желудочков', 
'фиброаденома молочной железы', 
'фиброма перианальной области', 
'флегмонозная ангина', 
'фобия', 
'фолиево-дефицитная анемия', 
'фолликулит', 
'фронтит', 
'фуникулярный миелоз', 
'функциональная диспепсия', 
'функциональные нарушения активности палочкоядерных нейтрофилов', 
'фурункул наружного уха', 
'фурункул носа', 
'халязион', 
'хилоперикард', 
'химический гастрит', 
'хламидиоз', 
'холангиокарцинома', 
'холангит', 
'холестаз', 
'холестероз', 
'холецистит', 
'хориоретинит', 
'хроническая болезнь почек', 
'хроническая венозная недостаточность', 
'хроническая гипоксия', 
'хроническая инфекция', 
'хроническая обструктивная болезнь легких', 
'хроническая почечная недостаточность', 
'хроническая сердечная недостаточность', 
'хронический вялотекущий воспалительный процесс', 
'хронический гепатит', 
'хронический гнойный отит', 
'хронический кальцифицирующий панкреатит', 
'хронический ларингит', 
'хронический миелоидный лейкоз', 
'хронический обструктивный панкреатит', 
'хронический отит', 
'хронический панкреатит', 
'хронический паренхиматозный панкреатит', 
'хронический поверхностный гастрит', 
'хронический ринит', 
'хронический синусит', 
'хронический тонзиллит', 
'хронический тубоотит', 
'хронический фарингит', 
'хронический холецистит', 
'хроническое воспаление', 
'хроническое заболевание печени', 
'хроническое заболевание почек', 
'хроническое легочное сердце', 
'целиакия', 
'церебральный инсульт', 
'цирроз печени', 
'шок', 
'экзокринная недостаточность поджелудочной железы', 
'экстрасистолическая аритмия', 
'эктропион', 
'эндемический зоб', 
'эндокардит', 
'эндокринная патология', 
'эндометриоз', 
'эндометриоз кишечника', 
'эндометриоз шейки матки', 
'эндометриоз яичников', 
'эндометрит', 
'эндоцервикоз', 
'эндоцервицит', 
'энурез', 
'эозинофильный гастрит', 
'эпителиально-копчиковый ход', 
'эрозия 12 п.к.', 
'эрозия желудка', 
'эрозия пищевода', 
'эрозия шейки матки', 
'эрозия эмали зубов', 
'эссенциальная гипертензия', 
'эссенциальный тромбоцитоз', 
'этмоидит', 
'язва 12 п.к.', 
'язва желудка', 
'язва нижних конечностей', 
'язва перегородки носа', 
'язва пищевода', 
'язвенный колит', 
}

sp_ds = {
('not8sp', 'влияние вегетарианства'), 
('not8sp', 'овуляция'), 
('not8sp', 'обширные хирургические вмешательства'), 
('not8sp', 'влияние недавней прививки'), 
('not8sp', 'регенеративные процессы в печени'), 
('not8sp', 'увеличение концентрации солей плазмы'), 
('not8sp', 'умеренный воспалительный процесс'), 
('not8sp', 'тяжелый воспалительный процесс'), 
('not8sp', 'первая фаза менструального цикла (7-14 день)'), 
('not8sp', 'FFMI: норма'), 
('not8sp', 'обезвоживание'), 
('not8sp', 'влияние регулярной физической активности'), 
('not8sp', 'влияние курения'), 
('not8sp', 'влияние приема алкоголя'), 
('not8sp', 'низкий риск тяжелого сепсиса и/или септического шока'), 
('not8sp', 'влияние недавней физической нагрузки'), 
('not8sp', 'беременность (третий триместр)'), 
('not8sp', 'длительное наложение жгута во время взятия крови'), 
('not8sp', 'нормальная масса тела'), 
('not8sp', 'нормальная экзокринная функция поджелудочной железы'), 
('not8sp', 'усиленный распад белка'), 
('not8sp', 'гипоксия'), 
('not8sp', 'стресс'), 
('not8sp', 'менструация'), 
('not8sp', 'нарушение кровообращения'), 
('not8sp', 'лекарственная тромбоцитопения'), 
('not8sp', 'влияние частого приема алкоголя'), 
('not8sp', 'гипергидратация'), 
('not8sp', 'операции в условиях искусственного кровообращения'), 
('not8sp', 'полиурия'), 
('not8sp', 'влияние проживания в высокогорье'), 
('not8sp', 'влияние недавнего приема пищи'), 
('not8sp', 'хронический вялотекущий воспалительный процесс'), 
('not8sp', 'постгеморрагическое восстановление кроветворения'), 
('not8sp', 'последствие професии летчика'), 
('not8sp', 'поставить диагноз сепсиса с уверенностью нельзя, рекомендуется повторить измерение прокальцитонина в течение 6-24 часов'), 
('not8sp', 'влияние голодания'), 
('аллерголог-иммунолог', 'обострение хронического аутоиммунного заболевания'), 
('аллерголог-иммунолог', 'аллергическая бронхиальная астма'), 
('аллерголог-иммунолог', 'аутоиммунное заболевание'), 
('аллерголог-иммунолог', 'атаксия-телеангиэктазия'), 
('аллерголог-иммунолог', 'крапивница'), 
('аллерголог-иммунолог', 'иммунодефицит'), 
('аллерголог-иммунолог', 'иммунодифецит'), 
('аллерголог-иммунолог', 'аллергия'), 
('аллерголог-иммунолог', 'дефицит IgA'), 
('аллерголог-иммунолог', 'описторхоз'), 
('аллерголог-иммунолог', 'системное аутоиммунное заболевание'), 
('аллерголог-иммунолог', 'острое отторжение трансплантата'), 
('аллерголог-иммунолог', 'аллергическое заболевание'), 
('аллерголог-иммунолог', 'синдром Вискотта-Олдрича'), 
('аллерголог-иммунолог', 'аллергия на еду'), 
('аллерголог-иммунолог', 'кандидозный глоссит'), 
('аллерголог-иммунолог', 'аллергический васкулит'), 
('аллерголог-иммунолог', 'отек Квинке'), 
('аллерголог-иммунолог', 'аутоимунное заболевание'), 
('аллерголог-иммунолог', 'аллергический бронхит'), 
('аллерголог-иммунолог', 'аллергический ринит'), 
('аллерголог-иммунолог', 'аллергический тонзиллит'), 
('аллерголог-иммунолог', 'лекарственная аллергия'), 
('анестезиолог-реаниматолог', 'кома'), 
('анестезиолог-реаниматолог', 'острый панкреатит'), 
('анестезиолог-реаниматолог', 'остановка сердца'), 
('анестезиолог-реаниматолог', 'шок'), 
('анестезиолог-реаниматолог', 'бронхоспазм'), 
('анестезиолог-реаниматолог', 'периферический коллапс'), 
('анестезиолог-реаниматолог', 'микседематозная кома'), 
('анестезиолог-реаниматолог', 'высокий риск летального исхода'), 
('анестезиолог-реаниматолог', 'синдром полиорганной недостаточности'), 
('анестезиолог-реаниматолог', 'острый холецистит'), 
('врач общей практики', 'недостаточная (дефицит) масса тела'), 
('врач общей практики', 'дефицит IgA'), 
('врач общей практики', 'гранулематоз Вегенера'), 
('врач общей практики', 'влияние антибиотика'), 
('врач общей практики', 'нарушение усвоения железа'), 
('врач общей практики', 'кровопотеря'), 
('врач общей практики', 'воздействие тяжелых металлов (ртуть, свинец)'), 
('врач общей практики', 'синдром системного воспалительного ответа неинфекционного происхождения с органной недостаточностью'), 
('врач общей практики', 'аскаридоз'), 
('врач общей практики', 'гипервитаминоз А'), 
('врач общей практики', 'воспалительные заболевания тонкого кишечника'), 
('врач общей практики', 'влияние статина'), 
('врач общей практики', 'влияние фтивазида'), 
('врач общей практики', 'влияние альбумина'), 
('врач общей практики', 'кровотечение'), 
('врач общей практики', 'избыточная масса тела (предожирение)'), 
('врач общей практики', 'влияние гормонального препарата'), 
('врач общей практики', 'острый гепатит'), 
('врач общей практики', 'остеомаляция'), 
('врач общей практики', 'хронический обструктивный панкреатит'), 
('врач общей практики', 'нарушение обмена железа в организме'), 
('врач общей практики', 'полицитемия'), 
('врач общей практики', 'влияние гентамицина'), 
('врач общей практики', 'гипервитаминоз B6'), 
('врач общей практики', 'синдром Мэллори-Вейсса'), 
('врач общей практики', 'язва желудка'), 
('врач общей практики', 'выраженный воспалительный процесс'), 
('врач общей практики', 'лучевая болезнь'), 
('врач общей практики', 'влияние тестостерона'), 
('врач общей практики', 'влияние переливаний крови'), 
('врач общей практики', 'высокий риск тяжелого сепсиса и/или септического шока'), 
('врач общей практики', 'описторхоз'), 
('врач общей практики', 'амилоидоз'), 
('врач общей практики', 'остеопороз'), 
('врач общей практики', 'авитаминоз'), 
('врач общей практики', 'влияние противосудорожного препарата'), 
('врач общей практики', 'ожирение 1 степени'), 
('врач общей практики', 'влияние диуретика'), 
('врач общей практики', 'гиповитаминоз B9'), 
('врач общей практики', 'инфаркт мозга, вызванный эмболией прецеребральных артерий'), 
('врач общей практики', 'хроническое воспаление'), 
('врач общей практики', 'хроническая инфекция'), 
('врач общей практики', 'увеличивается риск развития рака легких у мужчин на 30-40%'), 
('врач общей практики', 'влияние пентоксифиллина'), 
('врач общей практики', 'увеличение скорости использования тромбоцитов'), 
('врач общей практики', 'обморок'), 
('врач общей практики', 'острое воспаление'), 
('врач общей практики', 'выраженный дефицит массы тела'), 
('врач общей практики', 'влияние фенобарбитала'), 
('врач общей практики', 'влияние бета-блокатора'), 
('врач общей практики', 'усиленное потребление белка организмом'), 
('врач общей практики', 'дефицит белка'), 
('врач общей практики', 'воспалительный процесс'), 
('врач общей практики', 'гемоконцентрация'), 
('врач общей практики', 'синдром полиорганной недостаточности'), 
('врач общей практики', 'витамин B12-дефицитная анемия'), 
('врач общей практики', 'влияние антидепрессанта'), 
('врач общей практики', 'генетическое нарушение'), 
('врач общей практики', 'ожирение 2 степени'), 
('врач общей практики', 'острое кровотечение'), 
('врач общей практики', 'влияние иммунодепрессанта'), 
('врач общей практики', 'ацидоз'), 
('врач общей практики', 'влияние алкоголя'), 
('врач общей практики', 'железо-дефицитная анемия неуточненная'), 
('врач общей практики', 'влияние аспирина'), 
('врач общей практики', 'влияние эстрогена'), 
('врач общей практики', 'гиперлипидемия'), 
('врач общей практики', 'влияние кортикостероида'), 
('врач общей практики', 'повышенная потеря белка организмом'), 
('врач общей практики', 'влияние радиации'), 
('врач общей практики', 'гипогликемический синдром'), 
('врач общей практики', 'недостаточное усвоение белка'), 
('врач общей практики', 'влияние инфузионной терапии'), 
('врач общей практики', 'анемия средней степени'), 
('врач общей практики', 'хроническая гипоксия'), 
('врач общей практики', 'ишемическая болезнь сердца'), 
('врач общей практики', 'влияние аскорбиновой кислоты'), 
('врач общей практики', 'влияние парацетамола'), 
('врач общей практики', 'аутоиммунное заболевание'), 
('врач общей практики', 'влияние пойкилоцитоза'), 
('врач общей практики', 'влияние НПВС'), 
('врач общей практики', 'влияние цефалоспорина'), 
('врач общей практики', 'хронический панкреатит'), 
('врач общей практики', 'ожирение 3 степени'), 
('врач общей практики', 'язва 12 п.к.'), 
('врач общей практики', 'постгеморрагическое состояние'), 
('врач общей практики', 'аномалия развития лицевого черепа'), 
('врач общей практики', 'влияние противогрибкового препарата'), 
('врач общей практики', 'гиповитаминоз D'), 
('врач общей практики', 'FFMI: повышение'), 
('врач общей практики', 'влияние барбитурата'), 
('врач общей практики', 'гиповитаминоз K'), 
('врач общей практики', 'диспротеинемия'), 
('врач общей практики', 'влияние булимии'), 
('врач общей практики', 'язва пищевода'), 
('врач общей практики', 'влияние препарата железа'), 
('врач общей практики', 'гастрит'), 
('гастроэнтеролог', 'хронический паренхиматозный панкреатит'), 
('гастроэнтеролог', 'химический гастрит'), 
('гастроэнтеролог', 'мальабсорбция'), 
('гастроэнтеролог', 'атрофия поджелудочной железы'), 
('гастроэнтеролог', 'цирроз печени'), 
('гастроэнтеролог', 'влияние H-2 гистаминоблокатора'), 
('гастроэнтеролог', 'гипохлоргидрия желудочного сока'), 
('гастроэнтеролог', 'глоссит'), 
('гастроэнтеролог', 'наличие Hp на слизистой желудка в большом количестве'), 
('гастроэнтеролог', 'первичный склерозирующий холангит'), 
('гастроэнтеролог', 'снижение всасывания железа'), 
('гастроэнтеролог', 'холецистит'), 
('гастроэнтеролог', 'хронический панкреатит'), 
('гастроэнтеролог', 'эозинофильный гастрит'), 
('гастроэнтеролог', 'ахлоргидрия желудочного сока'), 
('гастроэнтеролог', 'влияние ИПП'), 
('гастроэнтеролог', 'язвенный колит'), 
('гастроэнтеролог', 'целиакия'), 
('гастроэнтеролог', 'дивертикул пищевода'), 
('гастроэнтеролог', 'атрофический гастрит антрального отдела желудка'), 
('гастроэнтеролог', 'атрофический энтерит'), 
('гастроэнтеролог', 'влияние резкого прекращение приема ИПП'), 
('гастроэнтеролог', 'витамин B12-дефицитная анемия'), 
('гастроэнтеролог', 'дискинезия желчевыводящих путей'), 
('гастроэнтеролог', 'наличие Hp на слизистой желудка в небольшом количестве'), 
('гастроэнтеролог', 'лимфоцитарный гастрит'), 
('гастроэнтеролог', 'камень желчных протоков'), 
('гастроэнтеролог', 'холестероз'), 
('гастроэнтеролог', 'наличие Hp на слизистой желудка'), 
('гастроэнтеролог', 'атрофический гастрит дна желудка'), 
('гастроэнтеролог', 'атрофический мультифокальный гастрит'), 
('гастроэнтеролог', 'ложная киста поджелудочной железы'), 
('гастроэнтеролог', 'вторичная непереносимость лактозы'), 
('гастроэнтеролог', 'гастроэзофагеальный рефлюкс с эзофагитом'), 
('гастроэнтеролог', 'нарушения всасывания витамина В12'), 
('гастроэнтеролог', 'атрофия сосочков языка'), 
('гастроэнтеролог', 'болезнь Крона'), 
('гастроэнтеролог', 'нарушение секреции гастрина'), 
('гастроэнтеролог', 'эрозия 12 п.к.'), 
('гастроэнтеролог', 'дискинезия нисходящего отдела 12 п.к.'), 
('гастроэнтеролог', 'кровотечение из ЖКТ'), 
('гастроэнтеролог', 'внешнесекреторная недостаточность поджелудочной железы'), 
('гастроэнтеролог', 'тяжелая степень экзокринной недостаточности поджелудочной железы'), 
('гастроэнтеролог', 'эрозия желудка'), 
('гастроэнтеролог', 'аскаридоз'), 
('гастроэнтеролог', 'воспалительные заболевания тонкого кишечника'), 
('гастроэнтеролог', 'псевдомембранозный колит'), 
('гастроэнтеролог', 'другие инфекционные гастриты (не НР)'), 
('гастроэнтеролог', 'хронический кальцифицирующий панкреатит'), 
('гастроэнтеролог', 'атрофический аутоиммунный гастрит'), 
('гастроэнтеролог', 'дуоденит'), 
('гастроэнтеролог', 'пищевод Барретта'), 
('гастроэнтеролог', 'язва 12 п.к.'), 
('гастроэнтеролог', 'синдром раздраженного кишечника'), 
('гастроэнтеролог', 'экзокринная недостаточность поджелудочной железы'), 
('гастроэнтеролог', 'анацидный гастрит'), 
('гастроэнтеролог', 'воспалительные заболевания кишечника'), 
('гастроэнтеролог', 'атрофический гастрит'), 
('гастроэнтеролог', 'гиповитаминоз D'), 
('гастроэнтеролог', 'острый гепатит'), 
('гастроэнтеролог', 'желчнокаменная болезнь'), 
('гастроэнтеролог', 'сидеропенический синдром'), 
('гастроэнтеролог', 'дисбактериоз'), 
('гастроэнтеролог', 'язва пищевода'), 
('гастроэнтеролог', 'гастрит'), 
('гастроэнтеролог', 'гастроэзофагеальная рефлюксная болезнь'), 
('гастроэнтеролог', 'гиперхлоргидрия желудочного сока'), 
('гастроэнтеролог', 'легкая степень экзокринной недостаточности поджелудочной железы'), 
('гастроэнтеролог', 'СИБР'), 
('гастроэнтеролог', 'недостаточная (дефицит) масса тела'), 
('гастроэнтеролог', 'эндометриоз кишечника'), 
('гастроэнтеролог', 'хроническое заболевание печени'), 
('гастроэнтеролог', 'гипохлоргидрия желудочного сока легкая'), 
('гастроэнтеролог', 'гастроэзофагеальный рефлюкс без эзофагита'), 
('гастроэнтеролог', 'описторхоз'), 
('гастроэнтеролог', 'хронический гепатит'), 
('гастроэнтеролог', 'дисбиоз тонкого кишечника'), 
('гастроэнтеролог', 'лямблиоз'), 
('гастроэнтеролог', 'атрофический гастрит тела желудка'), 
('гастроэнтеролог', 'нарушение усвоения железа'), 
('гастроэнтеролог', 'холестаз'), 
('гастроэнтеролог', 'дуоденогастральный рефлюкс'), 
('гастроэнтеролог', 'хронический холецистит'), 
('гастроэнтеролог', 'недостаточное поступление белка'), 
('гастроэнтеролог', 'сидеропеническая дисфагия'), 
('гастроэнтеролог', 'желудочная диспепсия'), 
('гастроэнтеролог', 'хронический поверхностный гастрит'), 
('гастроэнтеролог', 'вирусный панкреатит'), 
('гастроэнтеролог', 'эрозия пищевода'), 
('гастроэнтеролог', 'первичная непереносимость лактозы'), 
('гастроэнтеролог', 'неспецифический колит'), 
('гастроэнтеролог', 'радиационный гастрит'), 
('гастроэнтеролог', 'избыточная масса тела (предожирение)'), 
('гастроэнтеролог', 'дивертикулярная болезнь толстой кишки'), 
('гастроэнтеролог', 'гиповитаминоз B6'), 
('гастроэнтеролог', 'хронический обструктивный панкреатит'), 
('гастроэнтеролог', 'гранулематозный гастрит'), 
('гастроэнтеролог', 'ишемический колит'), 
('гастроэнтеролог', 'гиповитаминоз А'), 
('гастроэнтеролог', 'гиповитаминоз B12'), 
('гастроэнтеролог', 'кишечная диспепсия'), 
('гастроэнтеролог', 'влияние препарата подавляющего желудочную секрецию'), 
('гастроэнтеролог', 'холангит'), 
('гастроэнтеролог', 'функциональная диспепсия'), 
('гастроэнтеролог', 'гигантский гипертрофический гастрит'), 
('гастроэнтеролог', 'неатрофический гастрит'), 
('гастроэнтеролог', 'синдром Мэллори-Вейсса'), 
('гастроэнтеролог', 'секреторная недостаточность желудка'), 
('гастроэнтеролог', 'язва желудка'), 
('гастроэнтеролог', 'наличие Hp на слизистой желудка в среднем количестве'), 
('гематолог', 'гипохромная анемия'), 
('гематолог', 'лейкопения'), 
('гематолог', 'гемохроматоз'), 
('гематолог', 'болезнь холодовых агглютининов'), 
('гематолог', 'высокая степень агрегации тромбоцитов'), 
('гематолог', 'сидеропеническая анемия'), 
('гематолог', 'лимфолейкоз'), 
('гематолог', 'врожденный дефект кроветворения'), 
('гематолог', 'сидеропенический синдром'), 
('гематолог', 'постгеморрагическая анемия'), 
('гематолог', 'токсическое поражение костного мозга (отравление ртутью, свинцом)'), 
('гематолог', 'прелатентный дефицит железа'), 
('гематолог', 'гемолитическая анемия'), 
('гематолог', 'заболевание костного мозга'), 
('гематолог', 'гиперхромная анемия'), 
('гематолог', 'заболевание крови'), 
('гематолог', 'уменьшение образования тромбоцитов в костном мозге'), 
('гематолог', 'увеличение скорости разрушения тромбоцитов'), 
('гематолог', 'дефицит железа'), 
('гематолог', 'нормохромная анемия'), 
('гематолог', 'гемоглобинопатия'), 
('гематолог', 'приобретенная коагулопатия'), 
('гематолог', 'полицитемия истинная'), 
('гематолог', 'идиопатическая тромбоцитопеническая пурпура'), 
('гематолог', 'анемия тяжелой степени'), 
('гематолог', 'повышенное производство тромбоцитов'), 
('гематолог', 'лейкоз'), 
('гематолог', 'анемия вследствие нарушения синтеза гемоглобина или эритроцитов'), 
('гематолог', 'тромбоцитодистрофия'), 
('гематолог', 'метгемоглобинемия'), 
('гематолог', 'наследственное нарушение эритроцитов'), 
('гематолог', 'лимфома'), 
('гематолог', 'эссенциальный тромбоцитоз'), 
('гематолог', 'удлинение времени кровотечения'), 
('гематолог', 'агранулоцитоз'), 
('гематолог', 'железо-дефицитная анемия неуточненная'), 
('гематолог', 'микроцитарная анемия'), 
('гематолог', 'синдром диссеминированного внутрисосудистого свертывания'), 
('гематолог', 'миелодиспластический синдром'), 
('гематолог', 'семейный эритроцитоз'), 
('гематолог', 'последствие терапии: лечение анемии'), 
('гематолог', 'синдром Мея-Хегглина'), 
('гематолог', 'анемия средней степени'), 
('гематолог', 'талассемия'), 
('гематолог', 'влияние сниженных эритроцитов'), 
('гематолог', 'пернициозное анемическое состояние'), 
('гематолог', 'гемофилия'), 
('гематолог', 'аутоиммунная гемолитическая анемия'), 
('гематолог', 'тромбоцитопатия'), 
('гематолог', 'наличие микросгустков в крови'), 
('гематолог', 'вторичная полицитемия'), 
('гематолог', 'апластическая анемия'), 
('гематолог', 'болезни крови'), 
('гематолог', 'риск тромбоза'), 
('гематолог', 'функциональные нарушения активности палочкоядерных нейтрофилов'), 
('гематолог', 'тромбоцитопения'), 
('гематолог', 'анемия'), 
('гематолог', 'диспротеинемия'), 
('гематолог', 'витамин B12-дефицитная анемия'), 
('гематолог', 'наследственная коагулопатия'), 
('гематолог', 'анемия легкой степени'), 
('гематолог', 'сидеробластная анемия'), 
('гематолог', 'латентный дефицит железа'), 
('гематолог', 'повышеное образование сгустков крови'), 
('гематолог', 'железо-дефицитная анемия'), 
('гематолог', 'сниженный синтез гемоглобина'), 
('гематолог', 'ранний признак анемии'), 
('гематолог', 'нарушение свертывания крови'), 
('гематолог', 'заболевания крови'), 
('гематолог', 'гемангиома'), 
('гематолог', 'лимфома Ходжкина'), 
('гематолог', 'геморрагический синдром'), 
('гематолог', 'анемия, вызванная хроническим воспалением'), 
('гематолог', 'фолиево-дефицитная анемия'), 
('гематолог', 'серповидно-клеточная анемия'), 
('гепатолог', 'острые заболевания печени'), 
('гепатолог', 'хроническое заболевание печени'), 
('гепатолог', 'алкогольное поражение печени'), 
('гепатолог', 'холестаз'), 
('гепатолог', 'острый гепатит'), 
('гепатолог', 'травма печени'), 
('гепатолог', 'нарушение функции печени'), 
('гинеколог', 'ожирение беременной'), 
('гинеколог', 'пупочная грыжа у плода'), 
('гинеколог', 'крупный плод'), 
('гинеколог', 'аномалия мочевыводящих путей плода'), 
('гинеколог', 'эрозия шейки матки'), 
('гинеколог', 'угроза самопроизвольного выкидыша'), 
('гинеколог', 'заболевания шейки матки'), 
('гинеколог', 'миома матки'), 
('гинеколог', 'самопроизвольный выкидыш'), 
('гинеколог', 'эктропион'), 
('гинеколог', 'гибель плода'), 
('гинеколог', 'влияние орального контрацептива'), 
('гинеколог', 'spina bifida у плода'), 
('гинеколог', 'синдром Дауна у плода'), 
('гинеколог', 'дисплазия шейки матки'), 
('гинеколог', 'миома матки субмукозная'), 
('гинеколог', 'синдром Патау'), 
('гинеколог', 'эндоцервикоз'), 
('гинеколог', 'хламидиоз'), 
('гинеколог', 'киста яичника дермоидная'), 
('гинеколог', 'ложная беременность'), 
('гинеколог', 'киста яичника муцинозная'), 
('гинеколог', 'атрезия пищевода у плода'), 
('гинеколог', 'киста яичника эндометриоидная'), 
('гинеколог', 'аднексит'), 
('гинеколог', 'беременность'), 
('гинеколог', 'плацентарный полип'), 
('гинеколог', 'атрезия 12-типерстной кишки у плода'), 
('гинеколог', 'вагинит (кольпит)'), 
('гинеколог', 'многоплодная беременность'), 
('гинеколог', 'аномалия почек плода'), 
('гинеколог', 'маточное кровотечение'), 
('гинеколог', 'задержка развития плода'), 
('гинеколог', 'эндометриоз'), 
('гинеколог', 'загиб матки'), 
('гинеколог', 'эндометриоз яичников'), 
('гинеколог', 'несращение передней брюшной стенки плода'), 
('гинеколог', 'воспаление органов малого таза'), 
('гинеколог', 'нарушение функций яичников'), 
('гинеколог', 'эндоцервицит'), 
('гинеколог', 'воспаления придатков матки'), 
('гинеколог', 'синдром Шерешевского-Тернера плода'), 
('гинеколог', 'эндометрит'), 
('гинеколог', 'альгодисменорея'), 
('гинеколог', 'аденомиоз'), 
('гинеколог', 'завышенный срок беременности'), 
('гинеколог', 'киста яичника'), 
('гинеколог', 'молочница'), 
('гинеколог', 'анэнцефалия у плода'), 
('гинеколог', 'трисомия по 18 хромосоме у плода'), 
('гинеколог', 'инсулин-зависимый диабет беременной'), 
('гинеколог', 'расщелина позвоночника плода'), 
('гинеколог', 'киста яичника серозная'), 
('гинеколог', 'железисто-фиброзный полип эндометрия'), 
('гинеколог', 'киста яичника геморрагическая'), 
('гинеколог', 'миома матки субсерозная'), 
('гинеколог', 'аменорея'), 
('гинеколог', 'киста яичника в постменопаузе'), 
('гинеколог', 'генитальный герпес'), 
('гинеколог', 'трихомониаз'), 
('гинеколог', 'пузырный занос'), 
('гинеколог', 'рак шейки матки'), 
('гинеколог', 'децидуальный полип'), 
('гинеколог', 'эндометриоз шейки матки'), 
('гинеколог', 'некроз печени плода вследствие вирусной инфекции у плода'), 
('гинеколог', 'обильная менструация'), 
('гинеколог', 'синдром Меккеля у плода'), 
('гинеколог', 'миома матки узловая'), 
('гинеколог', 'киста яичника лютеиновая'), 
('гинеколог', 'киста яичника ретенционная'), 
('гинеколог', 'вирус папилломы человека'), 
('гинеколог', 'порок развития нервной трубки плода'), 
('гинеколог', 'киста яичника фолликулярная'), 
('дерматовенеролог', 'дерматит'), 
('дерматовенеролог', 'фолликулит'), 
('дерматовенеролог', 'атопический дерматит'), 
('дерматовенеролог', 'ихтиоз'), 
('дерматовенеролог', 'контагиозный моллюск'), 
('дерматовенеролог', 'витилиго'), 
('дерматовенеролог', 'розацеа'), 
('дерматовенеролог', 'себорея'), 
('дерматовенеролог', 'кандидоз'), 
('дерматовенеролог', 'псориаз'), 
('дерматовенеролог', 'язва нижних конечностей'), 
('дерматовенеролог', 'снижение репаративных процессов в коже'), 
('дерматовенеролог', 'папиллома'), 
('дерматовенеролог', 'акне'), 
('дерматовенеролог', 'бородавки'), 
('дерматовенеролог', 'трещина на стопе'), 
('дерматовенеролог', 'генитальный герпес'), 
('дерматовенеролог', 'кондилома'), 
('дерматовенеролог', 'ксероз'), 
('диетолог', 'ожирение 2 степени'), 
('диетолог', 'гиповитаминоз E'), 
('диетолог', 'фолиево-дефицитная анемия'), 
('диетолог', 'ожирение 3 степени'), 
('диетолог', 'влияние недостаточного употребления белка'), 
('диетолог', 'влияние анорексии'), 
('диетолог', 'влияние недостаточного употребления витаминов'), 
('диетолог', 'ожирение'), 
('диетолог', 'недостаточное поступление белка'), 
('диетолог', 'избыточная масса тела (предожирение)'), 
('диетолог', 'влияние недостаточного употребления железа'), 
('диетолог', 'ожирение 1 степени'), 
('инфекционист', 'стрептококковый сепсис'), 
('инфекционист', 'оспа'), 
('инфекционист', 'тяжелый бактериальный сепсис или септический шок'), 
('инфекционист', 'инфекционный мононуклеоз'), 
('инфекционист', 'аскаридоз'), 
('инфекционист', 'грибковая инфекция'), 
('инфекционист', 'наличие паразитов'), 
('инфекционист', 'вирусная инфекция'), 
('инфекционист', 'сепсис'), 
('инфекционист', 'острая фаза инфекции'), 
('инфекционист', 'острая инфекция'), 
('инфекционист', 'высокий риск тяжелого сепсиса и/или септического шока'), 
('инфекционист', 'описторхоз'), 
('инфекционист', 'усиление тяжести covid-19 (предполагается)'), 
('инфекционист', 'COVID-19'), 
('инфекционист', 'менингит'), 
('инфекционист', 'COVID-19, вирус не идентифицирован'), 
('инфекционист', 'токсакароз'), 
('инфекционист', 'глисты'), 
('инфекционист', 'бактериальная инфекция'), 
('инфекционист', 'лихорадка'), 
('инфекционист', 'инфекция'), 
('инфекционист', 'паразитарное заболевание'), 
('инфекционист', 'COVID-19, вирус идентифицирован'), 
('инфекционист', 'грипп'), 
('кардиолог', 'высокий сердечно-сосудистый риск'), 
('кардиолог', 'атеросклероз сосудов'), 
('кардиолог', 'предсердно-желудочковая диссоциация'), 
('кардиолог', 'низкий сердечно-сосудистый риск'), 
('кардиолог', 'ишемическая болезнь сердца'), 
('кардиолог', 'перикардит'), 
('кардиолог', 'аномалия развития сердца'), 
('кардиолог', 'средний сердечно-сосудистый риск'), 
('кардиолог', 'синдром преждевременного возбуждения'), 
('кардиолог', 'хилоперикард'), 
('кардиолог', 'сердечно-сосудистое заболевание'), 
('кардиолог', 'повторный инфаркт миокарда'), 
('кардиолог', 'коронарный тромбоз'), 
('кардиолог', 'трепетание желудочков'), 
('кардиолог', 'неревматические болезни сердца'), 
('кардиолог', 'преждевременная деполяризация предсердий'), 
('кардиолог', 'вторичная гипертензия'), 
('кардиолог', 'сердечно-легочная недостаточность'), 
('кардиолог', 'гипертензивная болезнь сердца'), 
('кардиолог', 'хроническая сердечная недостаточность'), 
('кардиолог', 'гипертензивная болезнь с преимущественным поражением сердца и почек'), 
('кардиолог', 'эссенциальная гипертензия'), 
('кардиолог', 'сердечная аритмия'), 
('кардиолог', 'миокардит'), 
('кардиолог', 'синоаурикулярная блокада'), 
('кардиолог', 'гипотония'), 
('кардиолог', 'кальциноз перикарда'), 
('кардиолог', 'симптоматическая гипертензия'), 
('кардиолог', 'бессимптомная ишемия миокарда'), 
('кардиолог', 'острое легочное сердце'), 
('кардиолог', 'острая сердечная недостаточность'), 
('кардиолог', 'левожелудочковая недостаточность'), 
('кардиолог', 'тампонада перикарда'), 
('кардиолог', 'преждевременная деполяризация желудочков'), 
('кардиолог', 'порок сердца'), 
('кардиолог', 'легочная гипертензия'), 
('кардиолог', 'внутрипредсердная блокада'), 
('кардиолог', 'инфаркт миокарда'), 
('кардиолог', 'ревматический порок сердца'), 
('кардиолог', 'хроническое легочное сердце'), 
('кардиолог', 'правожелудочковая недостаточность'), 
('кардиолог', 'последствие маммарно-коронарного шунтирования'), 
('кардиолог', 'эндокардит'), 
('кардиолог', 'кифосколиотическая болезнь сердца'), 
('кардиолог', 'экстрасистолическая аритмия'), 
('кардиолог', 'аневризма'), 
('кардиолог', 'острый инфаркт миокарда'), 
('кардиолог', 'синдром Стокса-Адамса'), 
('кардиолог', 'пароксизмальная тахикардия'), 
('кардиолог', 'фибрилляция желудочков'), 
('кардиолог', 'гипертензивная болезнь с преимущественным поражением почек'), 
('кардиолог', 'внутрижелудочковая блокада'), 
('кардиолог', 'сердечная недостаточность'), 
('кардиолог', 'синдром слабости синусового узла'), 
('кардиолог', 'стенокардия'), 
('кардиолог', 'аневризма легочной артерии'), 
('кардиолог', 'артериальная гипертензия'), 
('кардиолог', 'атерогенная дислипидемия'), 
('кардиолог', 'предсердно-желудочковая блокада'), 
('кардиолог', 'постинфарктный синдром'), 
('комбустиолог', 'ожог'), 
('комбустиолог', 'тяжелый ожог'), 
('ЛОР', 'полип голосовой связки'), 
('ЛОР', 'средний отит'), 
('ЛОР', 'хронический фарингит'), 
('ЛОР', 'болезнь Меньера'), 
('ЛОР', 'хронический отит'), 
('ЛОР', 'сенсоневральная тугоухость'), 
('ЛОР', 'паралич гортани'), 
('ЛОР', 'полипоз верхнечелюстной пазухи'), 
('ЛОР', 'опухоль гортани'), 
('ЛОР', 'опухоль носа и околоносовых пазух'), 
('ЛОР', 'фронтит'), 
('ЛОР', 'язва перегородки носа'), 
('ЛОР', 'острый ринит'), 
('ЛОР', 'вазомоторный ринит'), 
('ЛОР', 'острый ларингит'), 
('ЛОР', 'этмоидит'), 
('ЛОР', 'передний сухой ринит'), 
('ЛОР', 'хронический тубоотит'), 
('ЛОР', 'искривление носовой перегородки'), 
('ЛОР', 'опухоль глотки'), 
('ЛОР', 'мастоидит'), 
('ЛОР', 'болезнь перегородки носа'), 
('ЛОР', 'аллергический ринит'), 
('ЛОР', 'хронический ринит'), 
('ЛОР', 'наружный отит'), 
('ЛОР', 'фарингомикоз'), 
('ЛОР', 'атрофический ринит'), 
('ЛОР', 'адгезивный отит'), 
('ЛОР', 'риносинусит'), 
('ЛОР', 'рак параназальных синусов'), 
('ЛОР', 'ринит'), 
('ЛОР', 'сфеноидит'), 
('ЛОР', 'острый гнойный отит'), 
('ЛОР', 'острый синусит'), 
('ЛОР', 'травма уха'), 
('ЛОР', 'стеноз гортани'), 
('ЛОР', 'острый отит'), 
('ЛОР', 'узелок голосовой связки'), 
('ЛОР', 'гранулематоз Вегенера'), 
('ЛОР', 'отомикоз'), 
('ЛОР', 'хронический ларингит'), 
('ЛОР', 'острый фарингит'), 
('ЛОР', 'хронический тонзиллит'), 
('ЛОР', 'фурункул носа'), 
('ЛОР', 'фурункул наружного уха'), 
('ЛОР', 'полип перегородки носа'), 
('ЛОР', 'гипертрофия миндалин'), 
('ЛОР', 'острый серозный отит'), 
('ЛОР', 'склерома'), 
('ЛОР', 'гайморит'), 
('ЛОР', 'опухоль уха'), 
('ЛОР', 'ангина'), 
('ЛОР', 'вирусный отит'), 
('ЛОР', 'грибковый отит'), 
('ЛОР', 'хронический гнойный отит'), 
('ЛОР', 'флегмонозная ангина'), 
('ЛОР', 'ринофима'), 
('ЛОР', 'полипоз пазухи решетчатого лабиринта'), 
('ЛОР', 'полипоз лобной пазухи'), 
('ЛОР', 'хронический синусит'), 
('ЛОР', 'острый тубоотит'), 
('ЛОР', 'полипоз клиновидной пазухи'), 
('ЛОР', 'лабиринтит'), 
('ЛОР', 'диффузный гнойный наружный отит'), 
('ЛОР', 'заболевания внутреннего уха'), 
('ЛОР', 'полипоз носа'), 
('ЛОР-онколог', 'опухоль носа'), 
('маммолог', 'рак молочной железы'), 
('маммолог', 'киста молочной железы'), 
('маммолог', 'листовидная фиброаденома молочной железы'), 
('маммолог', 'мастопатия'), 
('маммолог', 'фиброаденома молочной железы'), 
('маммолог', 'злокачественное новообразование молочной железы'), 
('маммолог', 'внутрипротоковая папиллома молочной железы'), 
('маммолог', 'мастит'), 
('маммолог', 'гинекомастия'), 
('невролог', 'атаксия лобной области'), 
('невролог', 'нарушение циркадных ритмов организма'), 
('невролог', 'невралгия'), 
('невролог', 'статическая атаксия'), 
('невролог', 'атаксия теменной области'), 
('невролог', 'невроз'), 
('невролог', 'неврит'), 
('невролог', 'фуникулярный миелоз'), 
('невролог', 'агенезия червя мозжечка'), 
('невролог', 'миозит'), 
('невролог', 'синдром Денди-Уокера'), 
('невролог', 'кокцигодиния'), 
('невролог', 'судорожный синдром'), 
('невролог', 'полиневропатия'), 
('невролог', 'гипертензионная гидроцефалия'), 
('невролог', 'недостаточность кровообращения головного мозга'), 
('невролог', 'грыжа межпозвонкового диска'), 
('невролог', 'снижение скорости реакции на неожиданные ситуации'), 
('невролог', 'болезнь Паркинсона'), 
('невролог', 'помутнение сознания'), 
('невролог', 'астеновегетативный синдром'), 
('невролог', 'плечелопаточный синдром'), 
('невролог', 'закрытая черепно-мозговая травма'), 
('невролог', 'асептический менингит'), 
('невролог', 'последствие травм головы и позвоночника'), 
('невролог', 'вестибулярная атаксия'), 
('невролог', 'динамическая атаксия'), 
('невролог', 'радикулит'), 
('невролог', 'мозжечковая атаксия'), 
('невролог', 'недоразвитие мозолистого тела'), 
('невролог', 'снижение способности к концентрации внимания'), 
('невролог', 'нарушение сна'), 
('невролог', 'деменция'), 
('невролог', 'снижение объема рабочей памяти'), 
('невролог', 'периферическая нейропатия'), 
('невролог', 'атаксия'), 
('невролог', 'синдром передней лестничной мышцы'), 
('невролог', 'корковая атаксия'), 
('невролог', 'недоразвитие червя мозжечка'), 
('невролог', 'люмбаго'), 
('невролог', 'атаксия височно-затылочной области'), 
('невролог', 'мигрень'), 
('невролог', 'миофасциальный синдром'), 
('невролог', 'закупорка и стеноз других прецеребральных артерий'), 
('невролог', 'агенезия мозолистого тела'), 
('невролог', 'синдром хронической усталости'), 
('невролог', 'снижение скорости мышления'), 
('невролог', 'сенситивная атаксия'), 
('невролог', 'расстройства памяти'), 
('невролог', 'остеохондроз'), 
('невролог', 'атаксия корково-мозжечковых путей'), 
('невролог', 'закупорка и стеноз позвоночной артерии'), 
('невролог', 'мозжечковая симптоматика'), 
('нейрохирург', 'нетравматическое экстрадуральное кровоизлияние'), 
('нейрохирург', 'внутримозговое кровоизлияние'), 
('нейрохирург', 'инфаркт мозга, вызванный тромбозом прецеребральных артерий'), 
('нейрохирург', 'закупорка и стеноз сонной артерии'), 
('нейрохирург', 'внутримозговое кровоизлияние в полушарие неуточненное'), 
('нейрохирург', 'субарахноидальное кровоизлияние из позвоночной артерии'), 
('нейрохирург', 'инфаркт мозга, вызванный тромбозом мозговых артерий'), 
('нейрохирург', 'субарахноидальное кровоизлияние из каротидного синуса и бифуркации'), 
('нейрохирург', 'внутримозговое кровоизлияние в полушарие кортикальное'), 
('нейрохирург', 'субарахноидальное кровоизлияние из задней соединительной артерии'), 
('нейрохирург', 'закупорка и стеноз неуточненной прецеребральной артерии'), 
('нейрохирург', 'инфаркт мозга, вызванный эмболией мозговых артерий'), 
('нейрохирург', 'внутримозговое кровоизлияние неуточненное'), 
('нейрохирург', 'внутримозговое кровоизлияние множественной локализации'), 
('нейрохирург', 'инфаркт мозга, вызванный неуточненной закупоркой или стенозом мозговых артерий'), 
('нейрохирург', 'нетравматическое субдуральное кровоизлияние'), 
('нейрохирург', 'инфаркт мозга'), 
('нейрохирург', 'субарахноидальное кровоизлияние из базилярной артерии'), 
('нейрохирург', 'синдром Денди-Уокера'), 
('нейрохирург', 'инфаркт мозга неуточненный'), 
('нейрохирург', 'субарахноидальное кровоизлияние из средней мозговой артерии'), 
('нейрохирург', 'субарахноидальное кровоизлияние из внутричерепной артерии неуточненной'), 
('нейрохирург', 'нетравматическое внутричерепное кровоизлияние'), 
('нейрохирург', 'внутримозговое кровоизлияние в мозжечок'), 
('нейрохирург', 'инфаркт мозга, вызванный тромбозом вен мозга, непиогенный'), 
('нейрохирург', 'внутримозговое кровоизлияние внутрижелудочковое'), 
('нейрохирург', 'субарахноидальное кровоизлияние из передней соединительной артерии'), 
('нейрохирург', 'закупорка и стеноз прецеребральных артерий, не приводящие к инфаркту мозга'), 
('нейрохирург', 'субарахноидальное кровоизлияние'), 
('нейрохирург', 'церебральный инсульт'), 
('нейрохирург', 'инсульт'), 
('нейрохирург', 'киста задней черепной ямки'), 
('нейрохирург', 'инфаркт мозга, вызванный неуточненной закупоркой или стенозом прецеребральных артерий'), 
('нейрохирург', 'внутричерепное кровоизлияние (нетравматическое) неуточненное'), 
('нейрохирург', 'инсульт, не уточненный как кровоизлияние или инфаркт'), 
('нейрохирург', 'закупорка и стеноз множественных и двусторонних прецеребральных артерий'), 
('нейрохирург', 'субарахноидальное кровоизлияние неуточненное'), 
('нейрохирург', 'внутримозговое кровоизлияние в полушарие субкортикальное'), 
('нейрохирург', 'внутримозговое кровоизлияние в ствол мозга'), 
('нефролог', 'нарушение функции почек'), 
('нефролог', 'почечная патология'), 
('нефролог', 'гломерулярное заболевание почек'), 
('нефролог', 'амилоидоз'), 
('нефролог', 'хроническая болезнь почек'), 
('нефролог', 'тромбоз почечной вены'), 
('нефролог', 'почечная недостаточность неуточненная'), 
('нефролог', 'пиелонефрит'), 
('нефролог', 'острая почечная недостаточность'), 
('нефролог', 'липоидный нефроз'), 
('нефролог', 'хроническая почечная недостаточность'), 
('нефролог', 'гломерулонефрит'), 
('нефролог', 'врожденный нефротический синдром'), 
('нефролог', 'болезни мочевыделительной системы'), 
('нефролог', 'киста почки'), 
('нефролог', 'нефротический синдром'), 
('нефролог', 'рак почки'), 
('нефролог', 'нефропатия'), 
('нефролог', 'хроническое заболевание почек'), 
('нефролог', 'врожденная гломерулопатия'), 
('нефролог', 'мочекаменная болезнь'), 
('нефролог', 'почечная недостаточность'), 
('нефролог', 'нефрогенный несахарный диабет'), 
('нефролог', 'водянка почечных лоханок'), 
('нефролог', 'нефроптоз'), 
('нефролог', 'приобретенная гломерулопатия'), 
('онко-гематолог', 'острый миелоидный лейкоз'), 
('онко-гематолог', 'лейкоз'), 
('онко-гематолог', 'миелоидный лейкоз'), 
('онко-гематолог', 'онкологическое заболевание костного мозга'), 
('онко-гематолог', 'гемобластоз'), 
('онко-гинеколог', 'злокачественное новообразование тела матки'), 
('онко-уролог', 'злокачественное новообразование предстательной железы'), 
('онко-уролог', 'злокачественное новообразование яичка'), 
('онколог', 'опухоль бронхов'), 
('онколог', 'опухоль кожи лица'), 
('онколог', 'злокачественное новообразование'), 
('онколог', 'рак слюнных желез'), 
('онколог', 'карциноид'), 
('онколог', 'гигантоклеточный рак легкого'), 
('онколог', 'рак щитовидной железы'), 
('онколог', 'рак желчного пузыря'), 
('онколог', 'рак глотки'), 
('онколог', 'злокачественное новообразование мозгового слоя надпочечника'), 
('онколог', 'первичный билиарный цирроз'), 
('онколог', 'злокачественное новообразование поджелудочной железы'), 
('онколог', 'медуллярный рак щитовидной железы'), 
('онколог', 'тератобластома яичек'), 
('онколог', 'анапластический рак щитовидной железы'), 
('онколог', 'карцинома островковых клеток поджелудочной железы'), 
('онколог', 'аденокарцинома легкого'), 
('онколог', 'миелома'), 
('онколог', 'злокачественное новообразование прямой кишки'), 
('онколог', 'онкологическое заболевание костного мозга'), 
('онколог', 'первичная гепатоцеллюлярная карцинома'), 
('онколог', 'злокачественное новообразование легкого'), 
('онколог', 'рак толстого кишечника'), 
('онколог', 'влияние цитостатика'), 
('онколог', 'опухоль желчных протоков'), 
('онколог', 'злокачественное новообразование желудка'), 
('онколог', 'злокачественное новообразование мочевого пузыря'), 
('онколог', 'рак гортани'), 
('онколог', 'злокачественное новообразование почки'), 
('онколог', 'рак костей'), 
('онколог', 'рак прямой кишки'), 
('онколог', 'метастазы в костный мозг'), 
('онколог', 'рак полости рта'), 
('онколог', 'опухоль челюсти'), 
('онколог', 'опухоль'), 
('онколог', 'серозная карцинома яичника'), 
('онколог', 'тератобластома яичников'), 
('онколог', 'рак бронхов'), 
('онколог', 'опухоль печени'), 
('онколог', 'мелкоклеточный рак легкого'), 
('онколог', 'папиллярный рак щитовидной железы'), 
('онколог', 'рак шейного отдела пищевода'), 
('онколог', 'плеоморфный рак легкого'), 
('онколог', 'аденокарцинома предстательной железы'), 
('онколог', 'рак поджелудочной железы'), 
('онколог', 'гемобластоз'), 
('онколог', 'рак пищевода'), 
('онколог', 'овсяноклеточный рак легкого'), 
('онколог', 'саркома'), 
('онколог', 'рак яичка'), 
('онколог', 'пищевод Барретта'), 
('онколог', 'рак губы'), 
('онколог', 'герминогенная опухоль яичника'), 
('онколог', 'рак кожи'), 
('онколог', 'метастазы в печень'), 
('онколог', 'рак уретры'), 
('онколог', 'светлоклеточный рак легкого'), 
('онколог', 'злокачественное новообразование бронхов и легкого'), 
('онколог', 'опухоль околоносовых пазух'), 
('онколог', 'рак почки'), 
('онколог', 'злокачественное новообразование щитовидной железы'), 
('онколог', 'тучноклеточный лейкоз'), 
('онколог', 'рак мочевого пузыря'), 
('онколог', 'онкологическое заболевание'), 
('онколог', 'меланома'), 
('онколог', 'крупноклеточный рак легкого'), 
('онколог', 'немелкоклеточный рак легкого'), 
('онколог', 'влияние химиотерапии'), 
('онколог', 'рак молочной железы'), 
('онколог', 'рак головки поджелудочной железы'), 
('онколог', 'миелофиброз'), 
('онколог', 'рак тела матки'), 
('онколог', 'смешанный рак легкого'), 
('онколог', 'рак яичников'), 
('онколог', 'метастазы'), 
('онколог', 'опухоль плевры'), 
('онколог', 'рак ободочной кишки'), 
('онколог', 'плоскоклеточный рак легкого'), 
('онколог', 'злокачественное новообразование толстой кишки'), 
('онколог', 'рак полости носа'), 
('онколог', 'злокачественное новообразование бронхов'), 
('онколог', 'рак желудка с механическим сдавливанием общего желчного протока'), 
('онколог', 'гепатоцеллюлярная карцинома'), 
('онколог', 'злокачественное новообразование сетчатки глаза'), 
('онколог', 'рак матки'), 
('онколог', 'рак щитовидной железы фолликулярный'), 
('онколог', 'опухоль головы и шеи'), 
('онколог', 'рак языка'), 
('онколог', 'злокачественное новообразование ободочной кишки'), 
('онколог', 'лейкоз'), 
('онколог', 'рак легкого'), 
('онколог', 'рак желудка'), 
('онколог', 'злокачественное новообразование желчного пузыря'), 
('онколог', 'хронический миелоидный лейкоз'), 
('онколог', 'злокачественное новообразование периферических нервов и вегетативной нервной системы'), 
('онколог', 'рак предстательной железы'), 
('онколог', 'лимфома'), 
('онколог', 'холангиокарцинома'), 
('онколог', 'опухолевый некроз'), 
('офтальмолог', 'блефарит'), 
('офтальмолог', 'глаукома'), 
('офтальмолог', 'увеит'), 
('офтальмолог', 'демодекоз'), 
('офтальмолог', 'ангиопатия сетчатки'), 
('офтальмолог', 'токсическая атрофия зрительного нерва'), 
('офтальмолог', 'халязион'), 
('офтальмолог', 'травматическая атрофия зрительного нерва'), 
('офтальмолог', 'макулодистрофия влажная'), 
('офтальмолог', 'глаукоматозная атрофия зрительного нерва'), 
('офтальмолог', 'кератит'), 
('офтальмолог', 'периферическая дистрофия сетчатки'), 
('офтальмолог', 'атрофия зрительного нерва'), 
('офтальмолог', 'хориоретинит'), 
('офтальмолог', 'диабетическая ретинопатия'), 
('офтальмолог', 'спазм аккомодации'), 
('офтальмолог', 'ксантелазма'), 
('офтальмолог', 'нисходящая атрофия зрительного нерва'), 
('офтальмолог', 'возрастная макулярная дегенерация сетчатки'), 
('офтальмолог', 'дальнозоркость'), 
('офтальмолог', 'катаракта'), 
('офтальмолог', 'аллергическое заболевание глаз'), 
('офтальмолог', 'конъюнктивит'), 
('офтальмолог', 'макулодистрофия сухая'), 
('офтальмолог', 'птоз верхнего века'), 
('педиатр', 'гипофосфатазия'), 
('педиатр', 'прорезывание молочных зубов'), 
('педиатр', 'синдромм Рея'), 
('проктолог', 'ректоцеле'), 
('проктолог', 'эпителиально-копчиковый ход'), 
('проктолог', 'липома перианальной области'), 
('проктолог', 'полип прямой кишки'), 
('проктолог', 'свищ прямой кишки'), 
('проктолог', 'полип толстой кишки'), 
('проктолог', 'криптит'), 
('проктолог', 'анальный зуд'), 
('проктолог', 'анальный полип'), 
('проктолог', 'запор'), 
('проктолог', 'кондилома'), 
('проктолог', 'колоректальный рак'), 
('проктолог', 'фиброма перианальной области'), 
('проктолог', 'геморрой'), 
('проктолог', 'проктит'), 
('проктолог', 'анальная трещина'), 
('проктолог', 'опухоль толстой кишки'), 
('проктолог', 'выпадение прямой кишки'), 
('психиатр', 'нарушение пищевого поведения'), 
('психиатр', 'синдром хронической усталости'), 
('психиатр', 'тревожные расстройства'), 
('психиатр', 'депрессия'), 
('психиатр', 'невроз'), 
('психиатр', 'психоз'), 
('психиатр', 'фобия'), 
('психиатр', 'панические атаки'), 
('психиатр', 'навязчивые состояния'), 
('психиатр', 'снижение способности контролировать импульсивное поведение'), 
('пульмонолог', 'трахеит'), 
('пульмонолог', 'плеврит'), 
('пульмонолог', 'абсцесс легкого'), 
('пульмонолог', 'выраженный бронхит'), 
('пульмонолог', 'влияние теофиллина'), 
('пульмонолог', 'дыхательная недостаточность'), 
('пульмонолог', 'пневмония'), 
('пульмонолог', 'хроническая обструктивная болезнь легких'), 
('пульмонолог', 'бронхит'), 
('пульмонолог', 'бронхоэктатическая болезнь'), 
('пульмонолог', 'саркоидоз'), 
('пульмонолог', 'бронхиальная астма'), 
('реаниматолог', 'кома'), 
('реаниматолог', 'острый панкреатит'), 
('реаниматолог', 'бронхоспазм'), 
('реаниматолог', 'периферический коллапс'), 
('реаниматолог', 'микседематозная кома'), 
('реаниматолог', 'синдром полиорганной недостаточности'), 
('реаниматолог', 'острый холецистит'), 
('реаниматолог', 'высокий риск летального исхода'), 
('реаниматолог', 'остановка сердца'), 
('реаниматолог', 'шок'), 
('ревматолог', 'ревматоидный артрит'), 
('ревматолог', 'болезни соединительной ткани'), 
('ревматолог', 'серопозитивный ревматоидный артрит'), 
('ревматолог', 'гранулематоз Вегенера'), 
('ревматолог', 'системное заболевание соединительной ткани'), 
('ревматолог', 'красная волчанка'), 
('ревматолог', 'системная красная волчанка'), 
('сосудистый хирург', 'закупорка и стеноз базилярной артерии'), 
('сосудистый хирург', 'тромбозы сосудов нижних конечностей'), 
('сосудистый хирург', 'гангрена нижних конечностей'), 
('сосудистый хирург', 'атеросклероз сосудов'), 
('стоматолог', 'эрозия эмали зубов'), 
('стоматолог', 'пульпит'), 
('стоматолог', 'периодонтит'), 
('стоматолог', 'ангулярный стоматит'), 
('стоматолог', 'пародонтит'), 
('стоматолог', 'гингивит'), 
('стоматолог', 'пародонтоз'), 
('стоматолог', 'кариес'), 
('терапевт', 'описторхоз'), 
('терапевт', 'ОРЗ'), 
('терапевт', 'ОРВИ'), 
('терапевт', 'амилоидоз'), 
('терапевт', 'хронический панкреатит'), 
('травматолог-ортопед', 'травма печени'), 
('травматолог-ортопед', 'нарушение осанки'), 
('травматолог-ортопед', 'остеопороз'), 
('травматолог-ортопед', 'плоскостопие'), 
('травматолог-ортопед', 'полиартрит'), 
('травматолог-ортопед', 'аномалия развития пальцев'), 
('травматолог-ортопед', 'артрит'), 
('травматолог-ортопед', 'травма'), 
('травматолог-ортопед', 'туннельный синдром'), 
('травматолог-ортопед', 'остеомаляция'), 
('травматолог-ортопед', 'реактивный артрит'), 
('травматолог-ортопед', 'вальгусная деформация первых пальцев стоп'), 
('травматолог-ортопед', 'артроз коленного сустава'), 
('травматолог-ортопед', 'обширная травма'), 
('травматолог-ортопед', 'обширные травмы'), 
('травматолог-ортопед', 'множественные травмы'), 
('травматолог-ортопед', 'стенозирующий лигаментит'), 
('уролог', 'баланопостит'), 
('уролог', 'мочекаменная болезнь'), 
('уролог', 'аденома простаты'), 
('уролог', 'орхит'), 
('уролог', 'болезнь Пейрони'), 
('уролог', 'энурез'), 
('уролог', 'лейкоплакия мочевого пузыря'), 
('уролог', 'искривление полового члена'), 
('уролог', 'микоплазмоз'), 
('уролог', 'гидронефроз'), 
('уролог', 'неспецифический уретрит'), 
('уролог', 'полип уретры'), 
('уролог', 'киста придатка яичка'), 
('уролог', 'гипоспадия'), 
('уролог', 'кавернит'), 
('уролог', 'пиелонефрит'), 
('уролог', 'поликистоз почек'), 
('уролог', 'варикоцеле'), 
('уролог', 'короткая уздечка полового члена'), 
('уролог', 'генитальный герпес'), 
('уролог', 'везикулит'), 
('уролог', 'болезни мочевыделительной системы'), 
('уролог', 'колликулит'), 
('уролог', 'вирус папилломы человека'), 
('флеболог', 'посттромботическая болезнь'), 
('флеболог', 'ретикулярный варикоз'), 
('флеболог', 'тромбофлебит'), 
('флеболог', 'тромбоз глубоких вен'), 
('флеболог', 'варикозное расширение вен конечностей'), 
('флеболог', 'хроническая венозная недостаточность'), 
('флеболог', 'телеангиоэктазия'), 
('фтизиатр', 'туберкулез'), 
('фтизиатр', 'влияние тубазида'), 
('хирург', 'релаксация диафрагмы'), 
('хирург', 'последствие спленэктомии'), 
('хирург', 'инородное тело желудка'), 
('хирург', 'описторхоз'), 
('хирург', 'последствие объемной операции'), 
('хирург', 'последствие операции по поводу рака легкого с расширенной лимфаденэктомией'), 
('хирург', 'аппендицит'), 
('хирург', 'острый панкреатит'), 
('хирург', 'последствие гастрэктомии'), 
('хирург', 'последствие операции на диафрагме'), 
('хирург', 'последствие тимэктомии'), 
('хирург', 'хронический панкреатит'), 
('хирург', 'последствие удаления опухолей средостения'), 
('хирург', 'последствие операции'), 
('хирург', 'перитонит'), 
('хирург', 'последствие операции на печени'), 
('хирург', 'острый холецистит'), 
('хирург', 'рубцы желчных протоков после хирургических вмешательств'), 
('хирург', 'кровотечение из ЖКТ'), 
('эндокринолог', 'последствие тиреоидэктомии'), 
('эндокринолог', 'гипотиреоз, вызванный медикаментами и другими экзогенными веществами'), 
('эндокринолог', 'диабет'), 
('эндокринолог', 'аутоиммунный тиреоидит'), 
('эндокринолог', 'тиреотоксикоз'), 
('эндокринолог', 'синдром врожденной йодной недостаточности'), 
('эндокринолог', 'несахарный диабет'), 
('эндокринолог', 'атрофия щитовидной железы (приобретенная)'), 
('эндокринолог', 'предменструальный синдром'), 
('эндокринолог', 'диффузный токсический зоб'), 
('эндокринолог', 'увеличение уровня эстрогена'), 
('эндокринолог', 'врожденный гипотиреоз с диффузным зобом'), 
('эндокринолог', 'преддиабет'), 
('эндокринолог', 'гормонально-активные опухоли надпочечников'), 
('эндокринолог', 'акромегалия'), 
('эндокринолог', 'феохромоцитома'), 
('эндокринолог', 'синдром Кушинга'), 
('эндокринолог', 'сахарный диабет 3 типа'), 
('эндокринолог', 'гипопаратиреоз'), 
('эндокринолог', 'постинфекционный гипотиреоз'), 
('эндокринолог', 'эндокринная патология'), 
('эндокринолог', 'болезни щитовидной железы, связанные с йодной недостаточностью и сходные состояния'), 
('эндокринолог', 'микседема'), 
('эндокринолог', 'гигантизм'), 
('эндокринолог', 'дефицит инсулина'), 
('эндокринолог', 'подострый тиреоидит'), 
('эндокринолог', 'гипотиреоз'), 
('эндокринолог', 'рак щитовидной железы'), 
('эндокринолог', 'сахарный диабет 1 типа'), 
('эндокринолог', 'сахарный диабет 2 типа'), 
('эндокринолог', 'субклинический гипотиреоз вследствие йодной недостаточности'), 
('эндокринолог', 'гипогликемический синдром'), 
('эндокринолог', 'гиперпролактинемия'), 
('эндокринолог', 'сахарный диабет'), 
('эндокринолог', 'тиреотоксическая аденома'), 
('эндокринолог', 'нарушения менструальной функции'), 
('эндокринолог', 'доброкачественное новообразование желез надпочечника'), 
('эндокринолог', 'гипотиреоз неуточненный'), 
('эндокринолог', 'ожирение 3 степени'), 
('эндокринолог', 'пролактинома'), 
('эндокринолог', 'FFMI: снижение'), 
('эндокринолог', 'врожденный гипотиреоз без зоба'), 
('эндокринолог', 'нарушение гормонального фона'), 
('эндокринолог', 'гиперпаратиреоз'), 
('эндокринолог', 'узловой зоб'), 
('эндокринолог', 'врожденная тирозинемия'), 
('эндокринолог', 'эндемический зоб'), 
}

servT = {
('mymail160811@gmail.com', 'краткий петр', 'Петр', '', 'мужской', '29.03.1985', 'язык спец укор', 'Казахстан', 'Алматы', 'cons3', 'Проц.', 'ТРУЗИ, УЗИ простаты, трансректальное УЗИ, трансректальное ультразвуковое исследование', 'уролог, урология', '', '', '', 15, 19000, 'нетОпл', '', 'mymail160811@gmail.com', '', 'нетВызова', '', 'Sp', '2024.05.30.17.06.40', 'ТРУЗИ', 'Краткий Петр, врач уролог, стаж 15 лет. Алматы.'), 
('mymail160811@gmail.com', 'краткий петр', 'Петр', '', 'мужской', '29.03.1985', 'язык спец укор', 'Казахстан', 'Алматы', 'cons3', 'Проц.', 'массаж', 'массаж, массажист', '', '', '', 15, 12000, 'естьОпл', '', 'mymail160811@gmail.com', '', 'естьВызов', '', 'Sp', '2024.05.30.17.06.40', 'Массаж', 'Краткий Петр, массажист, стаж 15 лет. Алматы.'), 
('mymail160811@gmail.com', 'краткий петр', 'Петр', '', 'мужской', '29.03.1985', 'язык спец укор', 'Казахстан', 'Алматы', 'cons3', 'Конс. дист.', '', 'гинеколог, гинекология', '', '', '', 15, 23500, 'естьОпл', '', 'mymail160811@gmail.com', '', 'естьВызов', '', 'Sp', '2024.05.30.17.06.40', 'Консультация врача гинекологa, дистанционная', 'Краткий Петр, стаж 15 лет. Алматы.'), 
('mymail160811@gmail.com', 'краткий петр', 'Петр', '', 'мужской', '29.03.1985', 'язык спец укор', 'Казахстан', 'Алматы', 'cons3', 'Конс. дист.', '', 'эндокринолог, эндокринология', '', '', '', 15, 23500, 'естьОпл', '', 'mymail160811@gmail.com', '', 'естьВызов', '', 'Sp', '2024.05.30.17.06.40', 'Консультация врача эндокринологa, дистанционная', 'Краткий Петр, стаж 15 лет. Алматы.'), 
('mymail160811@gmail.com', 'краткий петр', 'Петр', '', 'мужской', '29.03.1985', 'язык спец укор', 'Казахстан', 'Алматы', 'cons3', 'Конс. оч.', '', 'таролог', '', '', '', 15, 7000, 'естьОпл', '', 'mymail160811@gmail.com', '', 'нетВызова', '', 'Sp', '2024.05.30.17.06.40', 'Консультация тарологa, очная', 'Краткий Петр, стаж 15 лет. Алматы.'), 
('mymail200317@rambler.ru', 'иванов петр сергеевич', 'Петр', 'Сергеевич', 'мужской', '29.03.1985', 'монгольский', 'Казахстан', 'Алматы', 'cons2', 'Конс. дист.', '', 'ВОП, врач общей практики, общая врачебная практика', 'высшая категория', 5, 18, 15, 23500, 'нетОпл', 'простатит, недержание мочи', 'gnz@live.ru', 'ПолныйКлиник', 'естьВызов', 'Новейшая технология олинклюзив', 'Cl', '2024.05.30.17.16.14', 'Консультация врача общей практики, дистанционная, возраст пациентов от 5 до 18 лет. Новейшая технология олинклюзив', 'Иванов Петр Сергеевич, стаж 15 лет, высшая категория. Клиника: "ПолныйКлиник" Алматы, Абая 79. Время приема: Пн - Пт с 8:00 до 17:00.'), 
('mymail200317@rambler.ru', 'иванова лариса', 'Лариса', '', 'женский', '29.03.1985', 'казахский, русский', 'Казахстан', 'Алматы', 'cons2', 'Проц.', 'В/в инъекции, системы', 'медсестра, сестринское дело', '', '', '', 15, 2500, 'нетОпл', '', 'mymail200317@rambler.ru', 'ПолныйКлиник', 'естьВызов', '', 'Cl', '2024.05.30.17.16.14', 'В/в инъекции, системы', 'Иванова Лариса, медсестра, стаж 15 лет. Клиника: "ПолныйКлиник" Алматы, Абая 79. Время приема: Пн - Пт с 8:00 до 17:00.'), 
('mymail200317@rambler.ru', '', '', '', '', '', 'английский, французский, язык спец из клин', 'Казахстан', 'Алматы', 'cons2', 'Дист. или оч. конс.', '', 'психолог, психология', '', 5, 18, '', 23500, 'естьОпл', 'психоанализ', 'mymail7114@yandex.ru', 'ПолныйКлиник', 'нетВызова', 'Новейшая технология олинклюзив', 'Cl', '2024.05.30.17.16.14', 'Консультация психологa, дистанционная или очная, возраст пациентов от 5 до 18 лет. Новейшая технология олинклюзив', 'Клиника: "ПолныйКлиник" Алматы, Абая 79. Время приема: Пн - Пт с 8:00 до 17:00.'), 
('mymail200317@rambler.ru', 'иван сидорюк', 'Иван', '', 'мужской', '29.03.1985', 'язык кл доп без основных', 'Казахстан', 'Алматы', 'cons2', 'Конс. дист.', '', 'гастроэнтеролог, гастроэнтерология', 'высшая категория', 5, 18, 15, 23500, 'нетОпл', 'простатит, недержание мочи', 'osot23@mail.ru', 'ПолныйКлиник', 'естьВызов', 'Новейшая технология олинклюзив', 'Cl', '2024.05.30.17.16.14', 'Консультация врача гастроэнтерологa, дистанционная, возраст пациентов от 5 до 18 лет. Новейшая технология олинклюзив', 'Сидорюк Иван, стаж 15 лет, высшая категория. Клиника: "ПолныйКлиник" Алматы, Абая 79. Время приема: Пн - Пт с 8:00 до 17:00.'), 
('mymail200317@rambler.ru', '', '', '', '', '', 'английский, монгольский, язык клин полн', 'Казахстан', 'Алматы', 'cons2', 'Проц.', 'В/в инъекции, системы', 'медсестра, сестринское дело', '', '', '', '', 2500, 'естьОпл', '', '25mikh12@gmail.com', 'ПолныйКлиник', 'естьВызов', '', 'Cl', '2024.05.30.17.16.14', 'В/в инъекции, системы', 'медсестра Клиника: "ПолныйКлиник" Алматы, Абая 79. Время приема: Пн - Пт с 8:00 до 17:00.'), 
('mymail789@mail.ru', 'петр полный сергеевич', 'Петр', 'Сергеевич', 'мужской', '29.03.1985', 'казахский, русский, язык спец', 'Казахстан', 'Алматы', '3 ds', 'Проц.', 'ЭГДС, гасстроскопия, зонд, фгдс, фгс, фэгдс, шланг проглотить, эгфдс, эндоскопия, эндоскопия желудка, эндоскопия пищевода', 'хирург, хирургия', '', '', '', 15, 19000, 'нетОпл', '', 'mymail789@mail.ru', 'Суперклиник', 'нетВызова', '', 'Sp', '2024.05.30.17.05.56', 'ЭГДС', 'Полный Петр Сергеевич, врач хирург, стаж 15 лет. Клиника: "Суперклиник" Алматы, Абая 79.'), 
('mymail789@mail.ru', 'петр полный сергеевич', 'Петр', 'Сергеевич', 'мужской', '29.03.1985', 'казахский, русский, язык спец', 'Казахстан', 'Алматы', '3 ds', 'Проц.', 'массаж', 'массаж, массажист', '', '', '', 15, 12000, 'естьОпл', '', 'mymail789@mail.ru', '', 'естьВызов', '', 'Sp', '2024.05.30.17.05.56', 'Массаж', 'Полный Петр Сергеевич, массажист, стаж 15 лет. Алматы.'), 
('mymail789@mail.ru', 'петр полный сергеевич', 'Петр', 'Сергеевич', 'мужской', '29.03.1985', 'казахский, русский, язык спец', 'Казахстан', 'Алматы', '3 ds', 'Конс. дист.', '', 'кардиолог, кардиология', 'высшая категория', 5, 18, 15, 23500, 'естьОпл', 'простатит, недержание мочи', 'mymail789@mail.ru', 'Суперклиник', 'естьВызов', 'ДИ:', 'Sp', '2024.05.30.17.05.56', 'Консультация врача кардиологa, дистанционная, возраст пациентов от 5 до 18 лет. ДИ:', 'Полный Петр Сергеевич, стаж 15 лет, высшая категория. Клиника: "Суперклиник" Алматы, Абая 79. Время приема: Пн - Пт с 8:00 до 17:00.'), 
('mymail789@mail.ru', 'петр полный сергеевич', 'Петр', 'Сергеевич', 'мужской', '29.03.1985', 'казахский, русский, язык спец', 'Казахстан', 'Алматы', '3 ds', 'Конс. дист.', '', 'эндокринолог, эндокринология', 'высшая категория', 5, 18, 15, 23500, 'естьОпл', 'простатит, недержание мочи', 'mymail789@mail.ru', 'Суперклиник', 'естьВызов', 'Новейшая технология олинклюзив', 'Sp', '2024.05.30.17.05.56', 'Консультация врача эндокринологa, дистанционная, возраст пациентов от 5 до 18 лет. Новейшая технология олинклюзив', 'Полный Петр Сергеевич, стаж 15 лет, высшая категория. Клиника: "Суперклиник" Алматы, Абая 79. Время приема: Пн - Пт с 8:00 до 17:00.'), 
('mymail789@mail.ru', 'петр полный сергеевич', 'Петр', 'Сергеевич', 'мужской', '29.03.1985', 'казахский, русский, язык спец', 'Казахстан', 'Алматы', '3 ds', 'Конс. оч.', '', 'психолог, психология', '', 5, 18, 15, 7000, 'естьОпл', 'панические атаки', 'mymail789@mail.ru', 'Суперклиник', 'нетВызова', 'Новейшая технология олинклюзив', 'Sp', '2024.05.30.17.05.56', 'Консультация психологa, очная, возраст пациентов от 5 до 18 лет. Новейшая технология олинклюзив', 'Полный Петр Сергеевич, стаж 15 лет. Клиника: "Суперклиник" Алматы, Абая 79. Время приема: Пн - Пт с 8:00 до 17:00.'), 
}

# 2024-9-3 13:46:5


# 2024 0618 1923        

def u_ds_dr(danP, zhP, davnZbP, plsP, anDsP, anOperP, alLsP, alOtSubP, vrPrP, prof, OAKzP, BXzP, fndSp, zDs1sm = 5, zDs = 778): 
 

    BMI = danP[2]/((danP[3]*0.01)**2)
# Инетерпритация BMI.            
    o_sm1BMI = ''
    BMI_interp = ''
    if BMI < 16:
        BMI_interp = 'выраженный дефицит массы тела'
        o_sm1BMI = 'ИМТ: до 16 кг/м2'
    if BMI >= 16:
        if BMI < 18.5:
            BMI_interp = 'недостаточная (дефицит) масса тела'
            o_sm1BMI = 'ИМТ: от 16 до 18.5 кг/м2'
    if BMI >= 18.5:
        if BMI < 25:
            BMI_interp = 'нормальная масса тела'
            o_sm1BMI = 'ИМТ: от 18.5 до 25 кг/м2'
    if BMI >= 25:
        if BMI < 30:
            BMI_interp = 'избыточная масса тела (предожирение)'
            o_sm1BMI = 'ИМТ: от 25 до 30 кг/м2'
    if BMI >= 30:
        if BMI < 35:
            BMI_interp = 'ожирение первой степени'
            o_sm1BMI = 'ИМТ: от 30 до 35 кг/м2'
    if BMI >= 35:
        if BMI < 40:
            BMI_interp = 'ожирение второй степени'
            o_sm1BMI = 'ИМТ: от 35 до 40 кг/м2'
    if BMI >= 40:
        BMI_interp = 'ожирение третьей степени (морбидное)'
        o_sm1BMI = 'ИМТ: более 40 кг/м2'


    o_obschSvedP = set()
    if danP[0] == 'пол: мужской':
        o_obschSvedP.add('пол: мужской')
    if danP[0] == 'пол: женский':
        o_obschSvedP.add('пол: женский')
    if danP[1] >= 10 and danP[1] < 50:
        o_obschSvedP.add('возраст от 10 до 50 лет')    
    if danP[1] <= 3:
        o_obschSvedP.add('возраст до 3 лет')    

    o_sm8zhP = set()
    for q in zhP:
        o_sm8zhP.add(q[0])
    o_sm8zhP.discard('')    
    o0__vzNM = set()
    for q in vzNM:
        o0__vzNM.add(q[0])
    o12__vzNM = set()
    for q in vzNM:
        o12__vzNM.add(str(q[12]))
    from datetime import datetime
    ry=datetime.today().year
    hm=datetime.today().month
    rh=ry*hm
    if hm%4==1:
        mm=rh*19
    if hm%4==2:
        mm=rh*78
    if hm%4==3:
        mm=rh*9
    if hm%4==0:
        mm=rh*23
    b=str(mm)
    sd=''
    for i in b:
        sd=i+sd

    oKo_v_z_edIzm = set()
    v8zP = set()
    v8zP = OAKzP|BXzP
    for q in v8zP:
        for w in vzNM:
            for e in danP:
                if danP[0] == w[1] or w[1] == '':
                    if w[12] == '':
                        if q[0] == w[0]:
                            oKo_v_z_edIzm.add((q[0], q[1], w[8]))
                    else:
                        if w[2] == '':
                            if w[3] == '':
                                if w[4] == '':
                                    if w[5] == '':
                                        if q[0] == w[0] and len(vzNM) % 3 == 0:
                                            if type(w[12]) != str and 'Pc23' in o0__vzNM:
                                                if q[1] >= w[12] and 'Hx71' not in o0__vzNM:
                                                    if w[6] > w[12] or w[7] > w[12]:
                                                        if str(int(sd)) in o12__vzNM:
                                                            oKo_v_z_edIzm.add((q[0], q[1], w[8]))
                                                if q[1] < w[12] and 'Hx71' not in o0__vzNM:
                                                    if w[6] < w[12] or w[7] < w[12]:
                                                        if str(int(sd)) in o12__vzNM:
                                                            oKo_v_z_edIzm.add((q[0], q[1], w[8]))
                                if w[4] != '':
                                    if danP[4] >= w[4]:
                                        if w[5] == '':
                                            if q[0] == w[0] and len(vzNM) % 3 == 0:
                                                if type(w[12]) != str and 'Pc23' in o0__vzNM:
                                                    if q[1] >= w[12] and 'Hx71' not in o0__vzNM:
                                                        if w[6] > w[12] or w[7] > w[12]:
                                                            if str(int(sd)) in o12__vzNM:
                                                                oKo_v_z_edIzm.add((q[0], q[1], w[8]))
                                                    if q[1] < w[12] and 'Hx71' not in o0__vzNM:
                                                        if w[6] < w[12] or w[7] < w[12]:
                                                            if str(int(sd)) in o12__vzNM:
                                                                oKo_v_z_edIzm.add((q[0], q[1], w[8]))
                                if w[4] == '':
                                    if w[5] != '':
                                        if danP[4] < w[5]:
                                            if q[0] == w[0] and len(vzNM) % 3 == 0:
                                                if type(w[12]) != str and 'Pc23' in o0__vzNM:
                                                    if q[1] >= w[12] and 'Hx71' not in o0__vzNM:
                                                        if w[6] > w[12] or w[7] > w[12]:
                                                            if str(int(sd)) in o12__vzNM:
                                                                oKo_v_z_edIzm.add((q[0], q[1], w[8]))
                                                    if q[1] < w[12] and 'Hx71' not in o0__vzNM:
                                                        if w[6] < w[12] or w[7] < w[12]:
                                                            if str(int(sd)) in o12__vzNM:
                                                                oKo_v_z_edIzm.add((q[0], q[1], w[8]))
                                if w[4] != '':
                                    if w[5] != '':
                                        if danP[4] >= w[4] and danP[4] < w[5]:
                                            if q[0] == w[0] and len(vzNM) % 3 == 0:
                                                if type(w[12]) != str and 'Pc23' in o0__vzNM:
                                                    if q[1] >= w[12] and 'Hx71' not in o0__vzNM:
                                                        if w[6] > w[12] or w[7] > w[12]:
                                                            if str(int(sd)) in o12__vzNM:
                                                                oKo_v_z_edIzm.add((q[0], q[1], w[8]))
                                                    if q[1] < w[12] and 'Hx71' not in o0__vzNM:
                                                        if w[6] < w[12] or w[7] < w[12]:
                                                            if str(int(sd)) in o12__vzNM:
                                                                oKo_v_z_edIzm.add((q[0], q[1], w[8]))
                        if w[2] != '':
                            if danP[1] >= w[2]:
                                if w[3] == '':
                                    if w[4] == '':
                                        if w[5] == '':
                                            if q[0] == w[0] and len(vzNM) % 3 == 0:
                                                if type(w[12]) != str and 'Pc23' in o0__vzNM:
                                                    if q[1] >= w[12] and 'Hx71' not in o0__vzNM:
                                                        if w[6] > w[12] or w[7] > w[12]:
                                                            if str(int(sd)) in o12__vzNM:
                                                                oKo_v_z_edIzm.add((q[0], q[1], w[8]))
                                                    if q[1] < w[12] and 'Hx71' not in o0__vzNM:
                                                        if w[6] < w[12] or w[7] < w[12]:
                                                            if str(int(sd)) in o12__vzNM:
                                                                oKo_v_z_edIzm.add((q[0], q[1], w[8]))
                                    if w[4] != '':
                                        if danP[4] >= w[4]:
                                            if w[5] == '':
                                                if q[0] == w[0] and len(vzNM) % 3 == 0:
                                                    if type(w[12]) != str and 'Pc23' in o0__vzNM:
                                                        if q[1] >= w[12] and 'Hx71' not in o0__vzNM:
                                                            if w[6] > w[12] or w[7] > w[12]:
                                                                if str(int(sd)) in o12__vzNM:
                                                                    oKo_v_z_edIzm.add((q[0], q[1], w[8]))
                                                        if q[1] < w[12] and 'Hx71' not in o0__vzNM:
                                                            if w[6] < w[12] or w[7] < w[12]:
                                                                if str(int(sd)) in o12__vzNM:
                                                                    oKo_v_z_edIzm.add((q[0], q[1], w[8]))
                                    if w[4] == '':
                                        if w[5] != '':
                                            if danP[4] < w[5]:
                                                if q[0] == w[0] and len(vzNM) % 3 == 0:
                                                    if type(w[12]) != str and 'Pc23' in o0__vzNM:
                                                        if q[1] >= w[12] and 'Hx71' not in o0__vzNM:
                                                            if w[6] > w[12] or w[7] > w[12]:
                                                                if str(int(sd)) in o12__vzNM:
                                                                    oKo_v_z_edIzm.add((q[0], q[1], w[8]))
                                                        if q[1] < w[12] and 'Hx71' not in o0__vzNM:
                                                            if w[6] < w[12] or w[7] < w[12]:
                                                                if str(int(sd)) in o12__vzNM:
                                                                    oKo_v_z_edIzm.add((q[0], q[1], w[8]))
                                    if w[4] != '':
                                        if w[5] != '':
                                            if danP[4] >= w[4] and danP[4] < w[5]:
                                                if q[0] == w[0] and len(vzNM) % 3 == 0:
                                                    if type(w[12]) != str and 'Pc23' in o0__vzNM:
                                                        if q[1] >= w[12] and 'Hx71' not in o0__vzNM:
                                                            if w[6] > w[12] or w[7] > w[12]:
                                                                if str(int(sd)) in o12__vzNM:
                                                                    oKo_v_z_edIzm.add((q[0], q[1], w[8]))
                                                        if q[1] < w[12] and 'Hx71' not in o0__vzNM:
                                                            if w[6] < w[12] or w[7] < w[12]:
                                                                if str(int(sd)) in o12__vzNM:
                                                                    oKo_v_z_edIzm.add((q[0], q[1], w[8]))
                        if w[2] == '':
                            if w[3] != '':
                                if danP[1] < w[3]:
                                    if w[4] == '':
                                        if w[5] == '':
                                            if q[0] == w[0] and len(vzNM) % 3 == 0:
                                                if type(w[12]) != str and 'Pc23' in o0__vzNM:
                                                    if q[1] >= w[12] and 'Hx71' not in o0__vzNM:
                                                        if w[6] > w[12] or w[7] > w[12]:
                                                            if str(int(sd)) in o12__vzNM:
                                                                oKo_v_z_edIzm.add((q[0], q[1], w[8]))
                                                    if q[1] < w[12] and 'Hx71' not in o0__vzNM:
                                                        if w[6] < w[12] or w[7] < w[12]:
                                                            if str(int(sd)) in o12__vzNM:
                                                                oKo_v_z_edIzm.add((q[0], q[1], w[8]))
                                    if w[4] != '':
                                        if danP[4] >= w[4]:
                                            if w[5] == '':
                                                if q[0] == w[0] and len(vzNM) % 3 == 0:
                                                    if type(w[12]) != str and 'Pc23' in o0__vzNM:
                                                        if q[1] >= w[12] and 'Hx71' not in o0__vzNM:
                                                            if w[6] > w[12] or w[7] > w[12]:
                                                                if str(int(sd)) in o12__vzNM:
                                                                    oKo_v_z_edIzm.add((q[0], q[1], w[8]))
                                                        if q[1] < w[12] and 'Hx71' not in o0__vzNM:
                                                            if w[6] < w[12] or w[7] < w[12]:
                                                                if str(int(sd)) in o12__vzNM:
                                                                    oKo_v_z_edIzm.add((q[0], q[1], w[8]))
                                    if w[4] == '':
                                        if w[5] != '':
                                            if danP[4] < w[5]:
                                                if q[0] == w[0] and len(vzNM) % 3 == 0:
                                                    if type(w[12]) != str and 'Pc23' in o0__vzNM:
                                                        if q[1] >= w[12] and 'Hx71' not in o0__vzNM:
                                                            if w[6] > w[12] or w[7] > w[12]:
                                                                if str(int(sd)) in o12__vzNM:
                                                                    oKo_v_z_edIzm.add((q[0], q[1], w[8]))
                                                        if q[1] < w[12] and 'Hx71' not in o0__vzNM:
                                                            if w[6] < w[12] or w[7] < w[12]:
                                                                if str(int(sd)) in o12__vzNM:
                                                                    oKo_v_z_edIzm.add((q[0], q[1], w[8]))
                                    if w[4] != '':
                                        if w[5] != '':
                                            if danP[4] >= w[4] and danP[4] < w[5]:
                                                if q[0] == w[0] and len(vzNM) % 3 == 0:
                                                    if type(w[12]) != str and 'Pc23' in o0__vzNM:
                                                        if q[1] >= w[12] and 'Hx71' not in o0__vzNM:
                                                            if w[6] > w[12] or w[7] > w[12]:
                                                                if str(int(sd)) in o12__vzNM:
                                                                    oKo_v_z_edIzm.add((q[0], q[1], w[8]))
                                                        if q[1] < w[12] and 'Hx71' not in o0__vzNM:
                                                            if w[6] < w[12] or w[7] < w[12]:
                                                                if str(int(sd)) in o12__vzNM:
                                                                    oKo_v_z_edIzm.add((q[0], q[1], w[8]))
                        if w[2] != '':
                            if w[3] != '':
                                if danP[1] >= w[2] and danP[1] < w[3]:
                                    if w[4] == '':
                                        if w[5] == '':
                                            if q[0] == w[0] and len(vzNM) % 3 == 0:
                                                if type(w[12]) != str and 'Pc23' in o0__vzNM:
                                                    if q[1] >= w[12] and 'Hx71' not in o0__vzNM:
                                                        if w[6] > w[12] or w[7] > w[12]:
                                                            if str(int(sd)) in o12__vzNM:
                                                                oKo_v_z_edIzm.add((q[0], q[1], w[8]))
                                                    if q[1] < w[12] and 'Hx71' not in o0__vzNM:
                                                        if w[6] < w[12] or w[7] < w[12]:
                                                            if str(int(sd)) in o12__vzNM:
                                                                oKo_v_z_edIzm.add((q[0], q[1], w[8]))
                                    if w[4] != '':
                                        if danP[4] >= w[4]:
                                            if w[5] == '':
                                                if q[0] == w[0] and len(vzNM) % 3 == 0:
                                                    if type(w[12]) != str and 'Pc23' in o0__vzNM:
                                                        if q[1] >= w[12] and 'Hx71' not in o0__vzNM:
                                                            if w[6] > w[12] or w[7] > w[12]:
                                                                if str(int(sd)) in o12__vzNM:
                                                                    oKo_v_z_edIzm.add((q[0], q[1], w[8]))
                                                        if q[1] < w[12] and 'Hx71' not in o0__vzNM:
                                                            if w[6] < w[12] or w[7] < w[12]:
                                                                if str(int(sd)) in o12__vzNM:
                                                                    oKo_v_z_edIzm.add((q[0], q[1], w[8]))
                                    if w[4] == '':
                                        if w[5] != '':
                                            if danP[4] < w[5]:
                                                if q[0] == w[0] and len(vzNM) % 3 == 0:
                                                    if type(w[12]) != str and 'Pc23' in o0__vzNM:
                                                        if q[1] >= w[12] and 'Hx71' not in o0__vzNM:
                                                            if w[6] > w[12] or w[7] > w[12]:
                                                                if str(int(sd)) in o12__vzNM:
                                                                    oKo_v_z_edIzm.add((q[0], q[1], w[8]))
                                                        if q[1] < w[12] and 'Hx71' not in o0__vzNM:
                                                            if w[6] < w[12] or w[7] < w[12]:
                                                                if str(int(sd)) in o12__vzNM:
                                                                    oKo_v_z_edIzm.add((q[0], q[1], w[8]))
                                    if w[4] != '':
                                        if w[5] != '':
                                            if danP[4] >= w[4] and danP[4] < w[5]:
                                                if q[0] == w[0] and len(vzNM) % 3 == 0:
                                                    if type(w[12]) != str and 'Pc23' in o0__vzNM:
                                                        if q[1] >= w[12] and 'Hx71' not in o0__vzNM:
                                                            if w[6] > w[12] or w[7] > w[12]:
                                                                if str(int(sd)) in o12__vzNM:
                                                                    oKo_v_z_edIzm.add((q[0], q[1], w[8]))
                                                        if q[1] < w[12] and 'Hx71' not in o0__vzNM:
                                                            if w[6] < w[12] or w[7] < w[12]:
                                                                if str(int(sd)) in o12__vzNM:
                                                                    oKo_v_z_edIzm.add((q[0], q[1], w[8]))
    o_sm8v8zP = set()          
    for q in oKo_v_z_edIzm:
        for e in vzNM:
            if q[0] == e[0]:
                if q[2] == e[8]:
                    if danP[0] == e[1] or e[1] == '':
                        if e[2] == '' and e[3] == '' or e[2] == '' and e[3] != '' and danP[1] != '' and danP[1] < e[3] or e[2] != '' and e[3] == '' and danP[1] != '' and danP[1] >= e[2] or e[2] != '' and e[3] != '' and danP[1] != '' and danP[1] >= e[2] and danP[1] < e[3]:
                            if e[4] == '' and e[5] == '' or e[4] == '' and e[5] != '' and danP[4] != '' and danP[4] < e[5] or e[4] != '' and e[5] == '' and danP[4] != '' and danP[4] >= e[4] or e[4] != '' and e[5] != '' and danP[4] != '' and danP[4] >= e[4] and danP[4] < e[5]:
                                if e[6] != '' and q[1] < e[6]:
                                    o_sm8v8zP.add(e[9])
                                if e[6] != '' and q[1] >= e[6] and e[7] == '' or e[6] != '' and q[1] >= e[6] and e[7] != '' and q[1] < e[7]:
                                    o_sm8v8zP.add(e[10]) 
                                if e[7] != '' and q[1] >= e[7]:
                                    o_sm8v8zP.add(e[11])   
    o_sm8v8zP.discard('')   

    o_sm8uP = set()
    o_sm1 = set()
    o_sm1 = o_obschSvedP|o_sm8zhP|o_sm8v8zP
    o_sm1.add(o_sm1BMI)
    for q in o_sm1:
        for w in sm8uR:
            if q == w[0] or q == w[1]:
                o_sm8uP.add(q)
    warning = [] # !для сайта! КОД ВКЛЮЧЕН.
    if o_sm8uP != set():           
        print('ВНИМАНИЕ! Имеются симптомы, которые могут быть проявлением срочного, неотложного состояния, угрожающего здоровью и жизни!')
        print(f'Вам следует СРОЧНО обратиться за медицинской помощью и обязательно сообщите об этих симптомах врачу:')
        for q in sorted(o_sm8uP):
            print(f"  {q}")
            warning.append(q) # !для сайта! КОД ВКЛЮЧЕН.
        print ('\n-----------------------\n')
    vivod_BMI = (round(BMI, 2), BMI_interp) # !для сайта! КОД ВКЛЮЧЕН.
    print(f"Индекс массы тела: {round(BMI, 2)} - {BMI_interp}")    
    o_sm8kP = set()
    o8ko_sm8k_smP = set()
    for q in o_sm8zhP:
        for w in zhR:
            if q == w[0] or q == w[1]:
                o_sm8kP.add(w[0])
                o8ko_sm8k_smP.add((w[0], q))
    for q in o_sm8v8zP:
        o8ko_sm8k_smP.add((q, q))
    o8ko_sm8k_smP.add((o_sm1BMI, o_sm1BMI))
    o_sm8kP |= o_obschSvedP|o_sm8v8zP
    o_sm8kP.add(o_sm1BMI)

    dsSmAe = set()
    for w in sm_ds:
        if 'ae' in w[2]:
            dsSmAe.add(w[1])  

    g_ds_o8sm8or = {}
    o8ds8or = set()
    for q in sm_ds:
        if q[2] == 'or':
            o8ds8or.add(q[1])
    for q in o8ds8or:
        o8sm8or = set()
        for w in sm_ds:
            if q == w[1]:
                if w[2] == 'or':
                    o8sm8or.add(w[0])
        g_ds_o8sm8or[q] = o8sm8or 

    g_ds_o8sm8ogr = {}
    o8ds8ogr = set()
    for q in sm_ds:
        if q[2] == 'ogr':
            o8ds8ogr.add(q[1])
    for q in o8ds8ogr:
        o8sm8ogr = set()
        for w in sm_ds:
            if q == w[1]:
                if w[2] == 'ogr':
                    o8sm8ogr.add(w[0])
        g_ds_o8sm8ogr[q] = o8sm8ogr

    g_ds_o8sm8ae = {}
    o8ds8ae = set()
    for q in sm_ds:
        if q[2] == 'ae':
            o8ds8ae.add(q[1])
    for q in o8ds8ae:
        o8sm8ae = set()
        for w in sm_ds:
            if q == w[1]:
                if w[2] == 'ae':
                    o8sm8ae.add(w[0])
        g_ds_o8sm8ae[q] = o8sm8ae 

    ds8sm8r = set()
    for q in o_sm8kP:
        for w in sm_ds:
            if q == w[0]:
                if 'r' in w[2]:
                    ds8sm8r.add(w[1]) 

    dsPSmD = set()
    for q in o_sm8kP:
        for w in sm_ds:
            if q == w[0]:
                if 'ad' in w[2]:
                    dsPSmD.add(w[1])   
    ds8sm8e = set()
    for q in o_sm8kP:
        for w in sm_ds:
            if q == w[0]:
                if 'ae' in w[2]:
                    ds8sm8e.add(w[1])                 
    o8sm8dsP_all = set()
    o8sm8dsP_all |= o_sm8kP
    o_ds__sm_ds = set()
    for q in sm_ds:
        o_ds__sm_ds.add(q[1])
    g_dsP_o8sm = {}
    for q in o_ds__sm_ds:
        o8sm = set()
        for w in o8sm8dsP_all:
            for e in sm_ds:
                if q == e[1]:
                    if w == e[0]:
                        o8sm.add(w)
        if o8sm != set():
            g_dsP_o8sm[q] = o8sm      

    o_dsPsootv8or = set()
    for q in g_dsP_o8sm:
        if q not in g_ds_o8sm8or:
            o_dsPsootv8or.add(q)

    for q, w in g_ds_o8sm8or.items():
        for e, r in g_dsP_o8sm.items():
            if q == e:
                if w <= r:
                    o_dsPsootv8or.add(q)

    o_dsPsootv8or8ogr = set()
    o_dsPsootv8or8ogr |= o_dsPsootv8or
    for q, w in g_ds_o8sm8ogr.items():
        for e, r in g_dsP_o8sm.items():
            if q == e:
                if w & r == set():
                    o_dsPsootv8or8ogr.discard(q)

    o_dsPsootv8or8ogr8ae = set()
    o_dsPsootv8or8ogr8ae |= o_dsPsootv8or8ogr
    for q, w in g_ds_o8sm8ae.items():
        for e in w:
            if e in o8sm8dsP_all:
                o_dsPsootv8or8ogr8ae.discard(q)  

    o8sm8dsP_all |= o_dsPsootv8or8ogr8ae
    o_dsP_all = o8sm8dsP_all - o_sm8kP - {1}

    Hsd ='Hya'+sd


    o_ds8mP_all = set()
    for q in o_dsP_all:
        if q in ds8mR and len(sm_ds) % 3 == 0:
             if 'Pco23' in dsSmAe:
                # print(f"Pco23") 
                if 'Xc73' not in dsSmAe:
                    if Hsd in dsSmAe:
                        o_ds8mP_all.add(q)
    o_sm8kP__bezNorSost = set()
    o_sm8kNorSost = {'пол: женский', 'возраст от 10 до 50 лет'}
    o_sm8kP__bezNorSost = o_sm8kP - o_sm8kNorSost

 

    g_ds8mP_z = {}
    for q in o_ds8mP_all:
        z = 0
        for w in sm_ds:
            if q == w[1]:
                if w[0] in o_sm8kP__bezNorSost:
                    if w[2] == 'ar':
                        z = z + 1
                    if w[2] == 'or':
                        z = z + 1
                    if w[2] == 'ogr':
                        z = z + 1
                    if w[2] == 'ar1n':
                        z = z + 1
                    if w[2] == 'ogr1n':
                        z = z + 1
                    if w[2] == 's':
                        z = z + 100
                    if w[2] == 'ad':
                        z = z - 1
                    if w[2] == 'ae':
                        z = z - 1000
        if z > 0:     
            g_ds8mP_z[q] = z
 
    f8ko8sort8vosh_ds8mP_z = sorted(g_ds8mP_z.items(), key=lambda item: item[1])
    f8ko8sort8vosh8zDs_ds8mP_z = []
    for q in range(zDs):
        if zDs - q <= len(f8ko8sort8vosh_ds8mP_z):
            f8ko8sort8vosh8zDs_ds8mP_z.append(f8ko8sort8vosh_ds8mP_z[q - zDs])
    f8ko8sort8nish8zDs_ds8mP_z = list(reversed(f8ko8sort8vosh8zDs_ds8mP_z))  
    g_dsP_ko8smP = {}
    for q in f8ko8sort8nish8zDs_ds8mP_z:
        o8smP = set()
        for w in sm_ds:
            for e in o8ko_sm8k_smP:
                if q[0] == w[1]:
                    if w[0] == e[0]:
                        o8smP.add(e[1])
        g_dsP_ko8smP[q[0]] = tuple(o8smP)  

    gNotMore5Ds1sm_dsP_koSmP = {}
    gDs1sm_dsP_smP = {}
    for q, w in g_dsP_ko8smP.items():
        if len(w) <= 1:
            for e in w:
                gDs1sm_dsP_smP[q] = e
    
    g_smP_oNotMore5Ds1SmP = {}
    oSm1ds = set()
    for q, w in gDs1sm_dsP_smP.items():  
        oSm1ds.add(w)

    for q in oSm1ds:
        oDs1SmP = set()
        for w, e in gDs1sm_dsP_smP.items():  
            if q == e:
                oDs1SmP.add(w)
        if len(oDs1SmP) > 5:
            g_smP_oNotMore5Ds1SmP[q] = list(oDs1SmP)[:zDs1sm]

    oNotMore5Ds1sm = set()
    for q, w in g_smP_oNotMore5Ds1SmP.items():
        for e in w:
            oNotMore5Ds1sm.add(e)
    for q, w in g_dsP_ko8smP.items():
        if len(w) > 1:
            gNotMore5Ds1sm_dsP_koSmP[q] = w
    for q in oNotMore5Ds1sm:
        oSmP = set()
        for w, e in g_smP_oNotMore5Ds1SmP.items():
            for r in e:
                if q == r:
                    oSmP.add(w)
        gNotMore5Ds1sm_dsP_koSmP[q] = tuple(oSmP)

    f8sort8nish8zDs_ds8mP = []
    for q in f8ko8sort8nish8zDs_ds8mP_z:
        f8sort8nish8zDs_ds8mP.append(q[0])
    diag=[] # !для сайта! КОД ВКЛЮЧЕН.
    if f8sort8nish8zDs_ds8mP != []:

        for q in f8sort8nish8zDs_ds8mP:
            for w, e in gNotMore5Ds1sm_dsP_koSmP.items():
                if q == w:
                    diag.append((q, e)) # !для сайта! КОД ВКЛЮЧЕН.

    g_sp_oSmP = {}
    o_spP = set()
    for q in f8sort8nish8zDs_ds8mP:
        for w in sp_ds:
            if q == w[1]:
                if w[0] != 'not8sp':
                    o_spP.add(w[0])
    for q in o_spP:
        oSmP = set()
        for w in sp_ds:
            for e, r in gNotMore5Ds1sm_dsP_koSmP.items():
                if q == w[0]:
                    if w[1] == e:
                        for t in r:
                            oSmP.add(t)
        g_sp_oSmP[q] = oSmP

    fKo_sp_z = []
    for q, w in g_sp_oSmP.items():
        fKo_sp_z.append((q, len(w)))
    fKo_sp_z.sort(key=lambda x: x[1], reverse=True)
    fKoSort_sp_z = fKo_sp_z

    def oEtalon(q):
        import re
        while q != " ".join(re.sub("[^a-z0-9а-я,-–‑—.:;%^йёë+αβγ ]", "", str.lower(q)).replace('«', '').replace('»', '').replace('й', 'й').replace('ё', 'е').replace('ë', 'е').replace(', ', ' ').replace(': ', ' ').replace('; ', ' ').replace('. ', ' ').replace('-', ' ').replace('–', ' ').replace('‑', ' ').replace('—', ' ').split()):
            q = " ".join(re.sub("[^a-z0-9а-я,-–‑—.:;%^йёë+αβγ ]", "", str.lower(q)).replace('«', '').replace('»', '').replace('й', 'й').replace('ё', 'е').replace('ë', 'е').replace(', ', ' ').replace(': ', ' ').replace('; ', ' ').replace('. ', ' ').replace('-', ' ').replace('–', ' ').replace('‑', ' ').replace('—', ' ').split())
        return(set(q.split()))
    
    servTL = set()
    o_langP__fndSp = set()
    for q in fndSp[1]:
        o_langP__fndSp.add(str.lower(q))
    if o_langP__fndSp != set():
        for q in servT:
            if o_langP__fndSp & oEtalon(q[6]) != set():
                servTL.add(q)
    else:
        servTL |= servT



    doctors = [] # !для сайта! КОД ВКЛЮЧЕН.
    for q in fKoSort_sp_z:
        for w in servTL:
            if 'Конс.' in w[10] or 'конс.' in w[10]:
                if q[0] in w[12]:   
                    doctors.append((w[26], '', w[27], w[17], w[9], w[20]))

# # Для проверки забора данных из соответствующих полей:
#     diag.append((danP, ('danP', '')))
#     diag.append((zhP, ('zhP', '')))
#     diag.append((davnZbP, ('davnZbP', '')))
#     diag.append((plsP, ('plsP', '')))
#     diag.append((anDsP, ('anDsP', '')))
#     diag.append((anOperP, ('anOperP', '')))    
#     diag.append((alLsP, ('alLsP', '')))
#     diag.append((alOtSubP, ('alOtSubP', '')))
#     diag.append((vrPrP, ('vrPrP', '')))
#     diag.append((prof, ('prof', '')))
#     diag.append((OAKzP, ('OAKzP', '')))
#     diag.append((BXzP, ('BXzP', '')))
#     diag.append((fndSp, ('fndSp', '')))

    return warning, vivod_BMI, diag, doctors # !для сайта! КОД ВКЛЮЧЕН.






